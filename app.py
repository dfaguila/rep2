"""
App Streamlit - Generador REP_2 (Costos y Gastos por Familia de Servicios)

Consolida 7 familias de gasto:
  - GRH: Recursos Humanos                (GRH_8 + GRH_11)
  - GCP: Gastos Generales de Personal    (GCP_4 + GCP_5, ya abierta por recurso)
  - GGV: Vehículos y Equipos             (GGV_4 + GGV_5)
  - GGI: Bienes Inmuebles                (GGI_5, ya abierta por recurso y servicio)
  - GGM: Bienes Muebles                  (GGM_1..GGM_5, sin apertura, recursos 2401-2411)
  - OGG: Otros Gastos Generales          (OGG_5, sin apertura, recursos 2501-2550)
  - MEI: Materiales e Insumos            (MEI_1..MEI_4, sin apertura, recursos 4101-4106)

GGM, OGG y MEI no traen columna de servicio: por defecto 100% se asigna al
servicio regulado 1101, con parametrización opcional para destinar % a
servicios no regulados.
"""

import io
import re
import zipfile
import os
import json
import difflib
import statistics
from collections import defaultdict

import openpyxl
import pandas as pd
import streamlit as st
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from openpyxl.chart import LineChart, BarChart, Reference
from openpyxl.formatting.rule import ColorScaleRule

st.set_page_config(page_title="Generador REP_2 - Suralis", layout="wide")

HEADERS = [
    "CÓDIGO EMPRESA", "PERÍODO INFORMACIÓN", "AÑO INFORMADO",
    "CÓDIGO SECTOR DECRETO TARIFARIO", "CÓDIGO RECURSO",
    "CÓDIGO FAMILIA SERVICIOS REGULADOS",
    "% NO ACTIVADO ASIGNADO FAMILIA SERVICIOS", "GASTO ANUAL",
    "% ACTIVADO ASIGNADO FAMILIA SERVICIOS", "MONTO ACTIVADO",
]

SERVICIOS_NO_REGULADOS = {
    2101: "INSPECCIÓN DE INSTALACIÓN DOMICILIARIA",
    2102: "VERIFICACIÓN DE FUNCIONAMIENTO DEL MEDIDOR A TRAVÉS DEL MEDIDOR PATRÓN",
    2103: "REPARACIÓN DE INSTALACIÓN DOMICILIARIA DE AGUA POTABLE",
    2104: "DETECCIÓN DE FUGA INTRADOMICILIARIA DE AGUA POTABLE",
    2105: "CONSTRUCCIÓN DE ARRANQUE",
    2106: "ELIMINACIÓN DE ARRANQUE DE AGUA POTABLE A SOLICITUD DEL CLIENTE",
    2107: "CAMBIO DE MEDIDOR",
    2108: "VENTA DE MEDIDOR",
    2109: "CAMBIO DE UBICACIÓN DE MEDIDOR DE AGUA POTABLE",
    2110: "CONEXIÓN DE REDES A EJECUTAR CON PERSONAL DE LA EMPRESA",
    2111: "DESOBSTRUCCIÓN DE UNIÓN DOMICILIARIA",
    2112: "LIMPIEZA DE FOSA SÉPTICA",
    2113: "CONSTRUCCIÓN DE UD",
    2114: "REPARACIÓN DE CÁMARAS DE INSPECCIÓN DOMICILIARIAS",
    2115: "CONSTRUCCIÓN DE CÁMARAS DE INSPECCIÓN DOMICILIARIAS",
    2116: "EMPALMES DE COLECTORES LOTEOS NUEVOS REALIZADOS POR PERSONAL DE LA EMPRESA",
    2117: "OTRAS PRESTACIONES EXPRESAMENTE VINCULADAS AL INCISO 2° DEL ARTÍCULO 21 DE LA LEY DE TARIFAS",
    2201: "VENTA DE AGUA Y SERVICIOS DE ALCANTARILLADO",
    2202: "VENTA DE SUBPRODUCTOS",
    2203: "TRATAMIENTO DE RILES",
    2204: "SERVICIOS DE ASESORÍA Y GESTIÓN NO APR (PARA TERCEROS, RELACIONADOS O NO)",
    2205: "SERVICIOS DE ASESORÍA Y GESTIÓN DE APR",
    2206: "SERVICIOS DE INGENIERÍA",
    2207: "SERVICIOS DE CONSTRUCCIÓN E INSTALACIÓN",
    2208: "SERVICIOS DE ANÁLISIS DE LABORATORIO",
    2209: "SERVICIOS DE REPARACIONES, OPERACIONES Y MANTENIMIENTOS",
    2210: "SERVICIOS DE INSPECCIONES Y CERTIFICACIONES",
    2211: "VENTA DE EQUIPOS Y PIEZAS",
    2212: "ARRIENDO DE ACTIVOS",
    2213: "MONITOREOS AMBIENTALES",
    2214: "OTROS SERVICIOS NO REGULADOS",
}

# Recursos parametrizables por familia plana (para el selector de la UI)

# ============================================================================
# Diccionario de códigos de recurso (para el Panel de Visualización)
# ============================================================================
RECURSO_NOMBRE = {
    1101: "Remuneraciones",
    1102: "Honorarios",
    1103: "Horas Extras",
    1201: "Indemnizaciones",
    1202: "Seguro de Cesantía",
    1203: "Seguro de Accidentes",
    1204: "Seguro de Invalidez y Sobrevivencia",
    1205: "Otros Beneficios Adicionales",
    2101: "Alimentacíón",
    2102: "Capacitación",
    2103: "Pasajes",
    2104: "Alojamientos",
    2105: "Viáticos",
    2106: "Accesorios de personal",
    2201: "Arriendo de vehículos y maquinarias",
    2202: "Combustible",
    2203: "Permisos de circulación",
    2204: "Revisión técnica",
    2205: "Seguros",
    2206: "Mantención Preventiva (no incluye combustible)",
    2207: "Recobros y Mantención Correctiva",
    2208: "Peajes",
    2209: "Tags",
    2210: "Implementación tag",
    2211: "Gasto identificación de vehículos",
    2301: "Arriendo de inmuebles",
    2302: "Consumos básicos",
    2303: "Servicio de aseo",
    2304: "Materiales de aseo",
    2305: "Mantención de inmuebles",
    2306: "Mantención de extintores",
    2307: "Mantención de areas verdes",
    2308: "Vigilancia presencial",
    2309: "Vigilancia a distancia",
    2310: "Contribuciones",
    2311: "Comisiones de Corretaje de Inmuebles por compra o arriendo",
    2401: "Arriendo de equipos informáticos",
    2402: "Servicios informáticos",
    2403: "Telefonía Fija",
    2404: "Enlasces de Internet Fijos",
    2405: "Enlaces de datos entre inmuebles (incluye enlaces satelitales)",
    2406: "Telefonía y Banda Ancha Móvil",
    2407: "Tarjetas SIM para aplicaciones M2M",
    2408: "Radio trunking de voz y datos",
    2409: "Telefonía Satelital",
    2410: "Materiales e insumos de oficina, computacionales y bodega",
    2411: "Materiales e insumos de laboratorio",
    2501: "Dietas del Directorio",
    2502: "Gastos de Representación Directorio",
    2503: "Patentes Comerciales",
    2504: "Servicios de Imprenta, Fotocopiado y Reproducción",
    2505: "Trámites y gastos notariales",
    2506: "Actuaciones Judiciales",
    2507: "Inscripciones",
    2508: "Peritajes",
    2509: "Tasaciones",
    2510: "Enlaces Satelitales",
    2511: "Líneas Transmisión de Datos",
    2512: "Enlaces de Internet",
    2513: "Impuestos por uso de Canales de Radiofrecuencia",
    2514: "Fletes",
    2515: "Transporte de Correspondencia (incluye servicios postales y mensajería)",
    2516: "Transporte de Muestras de Laboratorio",
    2517: "Seguros de infraestructura de capacidad",
    2518: "Seguros de inmuebles",
    2519: "Seguros de redes",
    2520: "Autoseguro",
    2521: "Seguros de Responsabilidad Civil",
    2522: "Seguros Menores",
    2523: "Deducibles pagados",
    2524: "Publicidad y Avisos (Radio, TV, Diarios u otros medios)",
    2525: "Diseño gráfico",
    2526: "Rotulaciones gráficas",
    2527: "Materiales de difusión",
    2528: "Auspicios y aportes",
    2529: "Donaciones",
    2530: "Eventos comunitarios",
    2531: "Eventos corporativos (juntas de accionistas, cenas fin de año, etc.)",
    2532: "Campañas de educación",
    2533: "Materiales de campañas",
    2534: "Derechos de asociaciones y canalistas",
    2535: "Derechos de afiliaciones",
    2536: "Derechos SERVIU o Vialidad",
    2537: "Permisos municipales",
    2538: "Canon anual por activos en comodato",
    2539: "Fondo Fijo Rotativo",
    2540: "Garantías a favor de SISS",
    2541: "Gastos financieros asociados a Garantías SERVIU o Vialidad",
    2542: "Multas",
    2543: "Indemnizaciones a terceros",
    2544: "Suscripciones",
    2545: "Impuestos pagados",
    2546: "Servicios Bancarios",
    2547: "Operaciones Financieras",
    2548: "Castigo incobrables",
    2549: "Otros Gastos Generales",
    3101: "Lectura de medidores",
    3102: "Reparto de boletas y otros documentos",
    3103: "Suministro e impresión de boletas y otros documentos",
    3104: "Servicios de recaudación en cajas externas",
    3105: "Servicios de recaudación en cajas propias",
    3106: "Servicios de atención telefónica o distante",
    3107: "Servicios de inspección comercial",
    3108: "Servicios de cobranza prejudicial",
    3109: "Servicios de gestión",
    3110: "Servicios de transporte de personas (buses de acercamiento, radiotaxis, etc.)",
    3111: "Servicios de almacenamiento y bodegaje",
    3112: "Servicios de traslados de mercancías",
    3113: "Servicios de Procesamiento, Archivo y Digitación de Datos",
    3114: "Auditorías a los Estados Financieros",
    3115: "Clasificación de Riesgo",
    3116: "Administración del Registro de Accionistas",
    3117: "Asesorías Tributarias y Contables",
    3118: "Gestión de Recursos Hídricos",
    3119: "Administración del Rol Privado",
    3120: "Selección de Personal",
    3121: "Auditorías Sistemas de Calidad",
    3122: "Asesorías en Servicio al Cliente",
    3123: "Planes de Desarrollo",
    3124: "Estudios Tarifarios",
    3125: "Comisiones de Expertos en Procesos Tarifarios",
    3126: "Defensa por acciones de responsabilidad civil",
    3127: "Defensa de derechos sobre inmuebles",
    3128: "Defensa en juicios laborales",
    3129: "Reclamaciones tributarias",
    3130: "Asesoría y defensa en procesos penales",
    3131: "Laboral Permanente y Negociación Colectiva",
    3132: "Informes Legales o en Derecho",
    3133: "Otros Servicios No Operacionales",
    4101: "Productos químicos",
    4102: "Energía Eléctrica",
    4103: "Materiales y repuestos",
    4104: "Compra de agua cruda",
    4105: "Compra de agua potable",
    4106: "Arriendo de derechos de agua",
    5101: "Servicios de control de calidad de agua potable",
    5102: "Servicios de control de calidad de agua servidas",
    5103: "Servicios de interconexión AP",
    5104: "Servicios de interconexión AS",
    5105: "Servicios de operación de redes y conexiones",
    5106: "Servicios de operación de infraestructura",
    5107: "Servicios de transporte y disposición de lodos",
    5108: "Servicios de control y monitoreo ambiental",
    5109: "Servicios de mantención de infraestructura",
    5110: "Servicios de mantención de redes y conexiones",
    5111: "Servicios de mantención de recintos",
    5112: "Servicios de mantención de servidumbres",
    5113: "Concesiones marítimas",
    5114: "Otros Servicios Operacionales",
    6101: "Servicios de Terceros asociados a Control Directo de Riles",
    6201: "Servicios de Terceros asociados a Mantención de Grifos",
    6301: "Servicios de Terceros asociados a Corte y Reposición",
    6401: "Servicios de Terceros asociados a Revisión de Proyectos de Construcción",
    6501: "Servicios de Terceros asociados a Verificación de Medidores",
    6601: "Servicios de Terceros asociados a Otras Prestaciones Asociadas",
    7101: "Servicios de Terceros asociados a Estudios preliminares,  Hidrológicos e Hidrogeológicos",
    7102: "Servicios de Terceros asociados a Diseños de obras",
    7103: "Servicios de Terceros asociados a Impacto Ambiental",
    7201: "Servicios de Terceros asociados a Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)",
    7301: "Servicios de Terceros asociados a Inspección Técnica de Obras",
    7401: "Servicios de Terceros asociados a Inversiones en Telemetría y Telecontrol",
    7402: "Servicios de Terceros asociados a Inversiones en Comunicaciones y Sistemas de Información",
    7501: "Servicios de Terceros asociados a Adquisición Inmuebles",
    7502: "Servicios de Terceros asociados a Adquisición Vehículos y equipos",
    7503: "Servicios de Terceros asociados a Adquisición Bienes Muebles e Insumos",
}

# Prefijo (2 primeros dígitos del código de recurso) -> familia de gasto
FAMILIA_PREFIJOS = {
    11: 'GRH', 12: 'GRH',
    21: 'GCP',
    22: 'GGV',
    23: 'GGI',
    24: 'GGM',
    25: 'OGG',
    31: 'ST',
    41: 'MEI',
    51: 'ST',
    61: 'GPA', 62: 'GPA', 63: 'GPA', 64: 'GPA', 65: 'GPA', 66: 'GPA',
    71: 'STI', 72: 'STI', 73: 'STI', 74: 'STI', 75: 'STI',
}
FAMILIA_NOMBRE = {
    'GRH': 'GRH - Recursos Humanos',
    'GCP': 'GCP - Gastos Generales de Personal',
    'GGV': 'GGV - Vehículos y Equipos',
    'GGI': 'GGI - Bienes Inmuebles',
    'GGM': 'GGM - Bienes Muebles',
    'OGG': 'OGG - Otros Gastos Generales',
    'ST': 'ST - Servicios Tercerizados',
    'MEI': 'MEI - Materiales e Insumos',
    'GPA': 'GPA - Gasto Prestaciones Asociadas',
    'STI': 'STI - Servicios de Terceros (Estudios e Inversión)',
}

def familia_de_recurso(cod):
    prefijo = cod // 100
    key = FAMILIA_PREFIJOS.get(prefijo, 'OTR')
    return FAMILIA_NOMBRE.get(key, 'Otro (' + str(cod) + ')')

def nombre_recurso(cod):
    return RECURSO_NOMBRE.get(cod, f"Recurso {cod}")


RECURSOS_GGM = list(range(2401, 2412))
RECURSOS_OGG = list(range(2501, 2551))
RECURSOS_MEI = [4101, 4102, 4103, 4104, 4105, 4106]

# ============================================================================
# Familia ST - Servicios Tercerizados
# ============================================================================
# 29 tablas independientes. Cada tabla puede tener uno o más códigos de
# recurso (algunas comparten código con otras, ej. ST_22..ST_30 -> 5110;
# ST_16/ST_17 -> 5106 y 5109 a la vez). Por eso NO se asume un código fijo
# por archivo: se lee la columna "CÓDIGO RECURSO" de cada fila (igual que
# GGI_5 / GCP_5), detectando las columnas por NOMBRE de encabezado ya que
# no se conoce de antemano la estructura exacta de las 29 tablas.
ST_TABLE_TO_CODES = {
    "ST_3": [3101], "ST_4": [3102], "ST_5": [3103], "ST_6": [3104],
    "ST_7": [3105], "ST_8": [3106], "ST_9": [3107], "ST_10": [3108],
    "ST_11": list(range(3114, 3133)),  # 3114 a 3132
    "ST_12": [5101], "ST_13": [5102], "ST_14": [5105], "ST_15": [5106],
    "ST_16": [5106, 5109], "ST_17": [5106, 5109], "ST_18": [5107],
    "ST_19": [5114], "ST_20": [5108], "ST_21": [5109], "ST_22": [5110],
    "ST_23": [5110], "ST_25": [5110], "ST_27": [5110], "ST_28": [5110],
    "ST_29": [5110], "ST_30": [5110], "ST_31": [5111], "ST_32": [5112],
    "ST_33": [5114], "ST_34": [3133],
}
ST_TABLES = list(ST_TABLE_TO_CODES.keys())
RECURSOS_ST = sorted({c for codes in ST_TABLE_TO_CODES.values() for c in codes})

# ============================================================================
# Familia GPA - Gasto Prestaciones Asociadas
# ============================================================================
# 6 tablas (GPA_1 a GPA_6). Cada una trae su propio "CÓDIGO RECURSO" fijo
# (6101, 6201, 6301, 6401, 6501, 6601) y, a diferencia de GGM/OGG/MEI/ST, el
# 100% de su gasto se asigna SIEMPRE a un servicio regulado específico por
# tabla (no 1101, y sin necesidad de parametrizar servicios no regulados).
GPA_TABLE_TO_SERVICIO = {
    "GPA_1": 1201, "GPA_2": 1202, "GPA_3": 1203,
    "GPA_4": 1204, "GPA_5": 1205, "GPA_6": 1201,
}
GPA_TABLES = list(GPA_TABLE_TO_SERVICIO.keys())


def _normalizar_encabezado(texto):
    if texto is None:
        return ""
    txt = str(texto).upper().strip()
    for a, b in {"Á": "A", "É": "E", "Í": "I", "Ó": "O", "Ú": "U", "Ñ": "N"}.items():
        txt = txt.replace(a, b)
    return " ".join(txt.split())


def _match_aproximado(palabra, texto_colapsado, umbral=0.82):
    """True si `palabra` aparece exacta en texto_colapsado, o si existe un
    tramo de largo similar muy parecido (tolera 1-2 errores de tipeo).
    Palabras muy cortas (<4 letras) solo se buscan de forma exacta, para
    evitar falsos positivos."""
    if palabra in texto_colapsado:
        return True
    if len(palabra) < 4:
        return False
    n = len(palabra)
    mejor = 0.0
    for i in range(0, max(1, len(texto_colapsado) - n + 1) + 3):
        trozo = texto_colapsado[i:i + n]
        if not trozo:
            continue
        score = difflib.SequenceMatcher(None, palabra, trozo).ratio()
        if score > mejor:
            mejor = score
    return mejor >= umbral


def _encontrar_columna(header_row, grupos_palabras_clave, evitar=None):
    """Busca una columna cuyo encabezado contenga TODAS las palabras clave
    de algún grupo (en cualquier orden, con o sin espacios entre ellas,
    tolerando errores de tipeo menores como "Recuso" en vez de "Recurso",
    o "CódigoRecurso" pegado sin espacio). Prueba los grupos en orden de
    prioridad y devuelve el índice del primer match. `evitar`: si el
    encabezado contiene alguna de estas palabras, se descarta la columna
    (para no confundir columnas parecidas, ej. no tomar "% Activado" como
    la columna de gasto)."""
    evitar = evitar or []
    normalizados = [_normalizar_encabezado(h) for h in header_row]
    colapsados = [n.replace(" ", "") for n in normalizados]

    for palabras in grupos_palabras_clave:
        for i, (norm, cole) in enumerate(zip(normalizados, colapsados)):
            if any(ev in norm for ev in evitar):
                continue
            if all(_match_aproximado(p, cole) for p in palabras):
                return i
    return None


def _a_numero(valor):
    """Convierte celdas numéricas 'sucias' (texto, '-', separadores de miles
    en formato chileno o estadounidense, vacías) a float de forma segura.
    Las tablas ST son cargadas por el usuario y pueden traer celdas con
    formato de texto en vez de número (ej. '-' como marcador de cero,
    '1.234.567' con puntos de miles, o '1.234,56' con coma decimal)."""
    if valor is None:
        return 0.0
    if isinstance(valor, (int, float)):
        return float(valor)
    if isinstance(valor, str):
        v = valor.strip()
        if v in ("", "-", "—", "–", "N/A", "n/a", "s/i", "S/I"):
            return 0.0
        v = v.replace("$", "").replace(" ", "")
        n_puntos = v.count(".")
        n_comas = v.count(",")
        if n_puntos >= 1 and n_comas >= 1:
            if v.rfind(",") > v.rfind("."):
                v = v.replace(".", "").replace(",", ".")  # chileno: 1.234.567,89
            else:
                v = v.replace(",", "")  # estadounidense: 1,234,567.89
        elif n_comas >= 1:
            if n_comas >= 2:
                v = v.replace(",", "")
            else:
                entero, dec = v.split(",")
                if entero.lstrip("-").isdigit() and len(dec) <= 2:
                    v = entero + "." + dec  # coma decimal (ej. "1234,5")
                else:
                    v = entero + dec
        elif n_puntos >= 2:
            v = v.replace(".", "")  # miles chilenos "1.234.567"
        elif n_puntos == 1:
            entero, dec = v.split(".")
            if entero.lstrip("-").isdigit() and len(dec) == 3 and len(entero.lstrip("-")) <= 3:
                v = entero + dec  # miles "1.234" (no decimal, montos en pesos)
        try:
            return float(v)
        except ValueError:
            return 0.0
    return 0.0


def leer_tabla_st(file_bytes):
    """Lee una tabla ST_x detectando columnas por nombre de encabezado.
    Devuelve lista de (cod_recurso, monto_activado, total_gasto_no_activado).
    Lanza ValueError si no logra identificar las columnas necesarias."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]

    idx_recurso = _encontrar_columna(header_row, [["CODIGO", "RECURSO"], ["COD", "RECURSO"]])
    idx_activado = _encontrar_columna(header_row, [["MONTO", "ANUAL", "ACTIVADO"], ["MONTO", "ACTIVADO"]])
    idx_gasto = _encontrar_columna(
        header_row,
        [["TOTAL", "GASTO", "ANUAL"], ["GASTO", "ANUAL", "NO", "ACTIVADO"], ["TOTAL", "GASTO"], ["GASTO"]],
        evitar=["ACTIVADO", "%"],
    )

    faltantes = []
    if idx_recurso is None:
        faltantes.append("CÓDIGO RECURSO")
    if idx_activado is None:
        faltantes.append("MONTO ANUAL ACTIVADO")
    if idx_gasto is None:
        faltantes.append("TOTAL GASTO ANUAL")
    if faltantes:
        raise ValueError(f"No se identificaron las columnas: {', '.join(faltantes)}")

    filas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        cod_recurso = row[idx_recurso]
        if cod_recurso is None:
            continue
        cod_recurso = _a_numero(cod_recurso)
        cod_recurso = int(cod_recurso) if cod_recurso == int(cod_recurso) else cod_recurso
        filas.append((cod_recurso, _a_numero(row[idx_activado]), _a_numero(row[idx_gasto])))
    return filas


def leer_tabla_mei(file_bytes, necesita_obra_nbi=False, necesita_actividad=False):
    """Lee una tabla MEI_x (MEI_1..MEI_4) detectando columnas por nombre de
    encabezado, TOLERANTE a que se inserten columnas nuevas en el medio del
    archivo (ej. una segunda columna 'ID Respaldo'), que es justo lo que
    rompía la lectura por posición fija anterior. Devuelve una lista de
    tuplas uniformes (empresa, periodo, anio, sector, cod_recurso,
    cod_obra_nbi, monto_activado, total_gasto, cod_actividad) — las primeras
    4 posiciones son siempre fijas (estándar SISS), cod_obra_nbi es None si
    necesita_obra_nbi=False, y cod_actividad es None si
    necesita_actividad=False (cada tabla MEI usa solo una de las dos)."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]

    idx_recurso = _encontrar_columna(header_row, [["CODIGO", "RECURSO"], ["COD", "RECURSO"]])
    idx_activado = _encontrar_columna(header_row, [["MONTO", "ANUAL", "ACTIVADO"], ["MONTO", "ACTIVADO"]])
    idx_gasto = _encontrar_columna(
        header_row,
        [["TOTAL", "GASTO", "ANUAL"], ["GASTO", "ANUAL", "NO", "ACTIVADO"], ["TOTAL", "GASTO"], ["GASTO"]],
        evitar=["ACTIVADO", "%"],
    )

    faltantes = []
    if idx_recurso is None:
        faltantes.append("CÓDIGO RECURSO")
    if idx_activado is None:
        faltantes.append("MONTO ANUAL ACTIVADO")
    if idx_gasto is None:
        faltantes.append("TOTAL GASTO ANUAL")

    idx_obra_nbi = None
    if necesita_obra_nbi:
        idx_obra_nbi = _encontrar_columna(header_row, [["OBRA", "TIPO", "NBI"], ["OBRA", "NBI"], ["ID", "OBRA"]])
        if idx_obra_nbi is None:
            faltantes.append("ID OBRA TIPO NBI")

    idx_actividad = None
    if necesita_actividad:
        idx_actividad = _encontrar_columna(header_row, [["CODIGO", "ACTIVIDAD"], ["COD", "ACTIVIDAD"]])
        if idx_actividad is None:
            faltantes.append("CÓDIGO ACTIVIDAD")

    if faltantes:
        raise ValueError(f"No se identificaron las columnas: {', '.join(faltantes)}")

    filas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        cod_recurso = row[idx_recurso]
        if cod_recurso is None:
            continue
        cod_recurso = _a_numero(cod_recurso)
        cod_recurso = int(cod_recurso) if cod_recurso == int(cod_recurso) else cod_recurso
        cod_obra_nbi = row[idx_obra_nbi] if idx_obra_nbi is not None else None
        cod_actividad = row[idx_actividad] if idx_actividad is not None else None
        filas.append((
            row[0], row[1], row[2], row[3],
            cod_recurso, cod_obra_nbi, _a_numero(row[idx_activado]), _a_numero(row[idx_gasto]), cod_actividad,
        ))
    return filas


def safe_leer_tabla_mei(file_bytes, etiqueta, avisos, necesita_obra_nbi=False, necesita_actividad=False):
    """Como leer_tabla_mei, pero tolera archivo ausente (None) o con error de
    lectura: registra un aviso y devuelve lista vacía en vez de fallar."""
    if file_bytes is None:
        avisos.append(f"⚠️ Falta la tabla **{etiqueta}** — se excluye del cálculo (esa familia/recurso queda en 0 o incompleto).")
        return []
    try:
        filas = leer_tabla_mei(file_bytes, necesita_obra_nbi=necesita_obra_nbi, necesita_actividad=necesita_actividad)
    except Exception as e:
        avisos.append(f"⚠️ No se pudo leer **{etiqueta}** ({e}) — se excluye del cálculo.")
        return []
    if len(filas) == 0:
        avisos.append(f"ℹ️ La tabla **{etiqueta}** está vacía (sin filas de datos).")
    return filas


def leer_tabla_plana(file_bytes):
    """Lee una tabla 'plana' (GGM_x, OGG_5: sin apertura por servicio/actividad
    propia) detectando columnas por nombre de encabezado, TOLERANTE a que se
    inserten o agreguen columnas (incluso columnas vacías al final, que es
    justo lo que rompía la lectura por índice negativo r[-1]/r[-3] anterior).
    Devuelve tuplas uniformes (empresa, periodo, anio, sector, cod_recurso,
    monto_activado, total_gasto)."""
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]

    idx_recurso = _encontrar_columna(header_row, [["CODIGO", "RECURSO"], ["COD", "RECURSO"]])
    idx_activado = _encontrar_columna(header_row, [["MONTO", "ANUAL", "ACTIVADO"], ["MONTO", "ACTIVADO"]])
    idx_gasto = _encontrar_columna(
        header_row,
        [["TOTAL", "GASTO", "ANUAL"], ["GASTO", "ANUAL", "NO", "ACTIVADO"], ["TOTAL", "GASTO"], ["GASTO"]],
        evitar=["ACTIVADO", "%"],
    )
    faltantes = []
    if idx_recurso is None:
        faltantes.append("CÓDIGO RECURSO")
    if idx_activado is None:
        faltantes.append("MONTO ANUAL ACTIVADO")
    if idx_gasto is None:
        faltantes.append("TOTAL GASTO ANUAL")
    if faltantes:
        raise ValueError(f"No se identificaron las columnas: {', '.join(faltantes)}")

    filas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        cod_recurso = row[idx_recurso]
        if cod_recurso is None:
            continue
        cod_recurso = _a_numero(cod_recurso)
        cod_recurso = int(cod_recurso) if cod_recurso == int(cod_recurso) else cod_recurso
        filas.append((
            row[0], row[1], row[2], row[3],
            cod_recurso, _a_numero(row[idx_activado]), _a_numero(row[idx_gasto]),
        ))
    return filas


def safe_leer_tabla_plana(file_bytes, etiqueta, avisos):
    """Como leer_tabla_plana, pero tolera archivo ausente (None) o con error
    de lectura: registra un aviso y devuelve lista vacía en vez de fallar."""
    if file_bytes is None:
        avisos.append(f"⚠️ Falta la tabla **{etiqueta}** — se excluye del cálculo (esa familia/recurso queda en 0 o incompleto).")
        return []
    try:
        filas = leer_tabla_plana(file_bytes)
    except Exception as e:
        avisos.append(f"⚠️ No se pudo leer **{etiqueta}** ({e}) — se excluye del cálculo.")
        return []
    if len(filas) == 0:
        avisos.append(f"ℹ️ La tabla **{etiqueta}** está vacía (sin filas de datos).")
    return filas


# ============================================================================
# LISTADO DE ID RESPALDO -> MCO (tablas de sustento)
# ============================================================================
TABLAS_CON_RESPALDO = [
    "MEI_1", "MEI_2", "MEI_4",
    "GPA_1", "GPA_2", "GPA_3", "GPA_4", "GPA_5", "GPA_6",
    "GGV_3", "OGG_5",
]
TABLAS_MEI_RESPALDO = {"MEI_1", "MEI_2", "MEI_4"}

MCO_ESTANDAR = {
    1: ("MCO_4", "Contratos servicios recibidos"),
    2: ("MCO_9", "Orden de compra servicio"),
    3: ("MCO_12", "Facturas servicios"),
    9999: ("MCO_12", "Facturas servicios"),  # OGG_5: fondos rotatorios, reembolsos, rendiciones, etc.
}
MCO_MEI = {
    1: ("MCO_6", "Contrato Suministro insumos"),
    2: ("MCO_9", "Orden de compra servicio"),
    3: ("MCO_13", "Facturas de suministros"),
}


def inferir_tipo_respaldo_ogg(id_respaldo):
    """OGG_5 no trae columna TIPO RESPALDO en el estándar SISS. Se infiere
    a partir del texto del propio ID RESPALDO, reconociendo el código FA/CO/OC
    tanto al inicio del texto ('FA-123') como en medio ('S-FA-123'):
    - 'CO-' al inicio o '-CO-' en medio -> 1 (Contrato)
    - 'OC-' al inicio o '-OC-' en medio -> 2 (Orden de Compra)
    - 'FA-' al inicio o '-FA-' en medio -> 3 (Factura)
    - si no calza con ninguno -> 9999 (fondos rotatorios, reembolsos,
      rendiciones, dietas, patentes, etc. -- no tienen contrato/OC/factura
      de proveedor asociada, se declaran igual como Factura por defecto)."""
    texto = str(id_respaldo).upper()
    if texto.startswith("CO-") or "-CO-" in texto:
        return 1
    if texto.startswith("OC-") or "-OC-" in texto:
        return 2
    if texto.startswith("FA-") or "-FA-" in texto:
        return 3
    return 9999


def leer_id_tipo_respaldo(file_bytes, nombre_tabla, avisos):
    """Lee de una tabla las columnas ID RESPALDO y TIPO RESPALDO por nombre
    de encabezado (robusto a mayúsculas/minúsculas y variantes). Devuelve
    lista de (tabla, id_respaldo, tipo_respaldo). Para OGG_5 (que no trae
    TIPO RESPALDO en el estándar SISS), el tipo se infiere del propio texto
    del ID RESPALDO vía inferir_tipo_respaldo_ogg()."""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws = wb.active
        header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
    except Exception as e:
        avisos.append(f"⚠️ Respaldos: no se pudo leer **{nombre_tabla}** ({e}).")
        return []

    idx_id = _encontrar_columna(header_row, [["ID", "RESPALDO"]])
    idx_tipo = _encontrar_columna(header_row, [["TIPO", "RESPALDO"]])
    if idx_id is None:
        avisos.append(f"⚠️ Respaldos: **{nombre_tabla}** no tiene columna ID RESPALDO reconocible.")
        return []
    if idx_tipo is None and nombre_tabla != "OGG_5":
        avisos.append(f"ℹ️ Respaldos: **{nombre_tabla}** no tiene columna TIPO RESPALDO — sus ID Respaldo se listan sin MCO asignado.")
    if idx_tipo is None and nombre_tabla == "OGG_5":
        avisos.append(f"ℹ️ Respaldos: **{nombre_tabla}** no trae TIPO RESPALDO (estándar SISS) — se infiere desde el texto del ID RESPALDO (-CO-/-OC-/-FA-; el resto queda como 9999).")

    filas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        id_resp = row[idx_id]
        if id_resp is None or str(id_resp).strip() == "":
            continue
        if idx_tipo is not None:
            tipo_resp = row[idx_tipo]
        elif nombre_tabla == "OGG_5":
            tipo_resp = inferir_tipo_respaldo_ogg(id_resp)
        else:
            tipo_resp = None
        if tipo_resp is not None:
            try:
                tipo_resp = int(_a_numero(tipo_resp))
            except Exception:
                pass
        filas.append((nombre_tabla, id_resp, tipo_resp))
    return filas


def armar_listado_respaldos(archivos_disponibles):
    """archivos_disponibles: dict {nombre_tabla: file_bytes}. Devuelve
    (filas, avisos) con filas = [(tabla, id_respaldo, tipo_respaldo,
    codigo_mco, descripcion_mco), ...], únicas por (tabla, id_respaldo,
    tipo_respaldo), ordenadas por tabla e ID Respaldo."""
    avisos = []
    todas_filas = []
    for tabla in TABLAS_CON_RESPALDO:
        if tabla not in archivos_disponibles:
            avisos.append(f"ℹ️ Respaldos: **{tabla}** no está cargada — se omite del listado.")
            continue
        todas_filas.extend(leer_id_tipo_respaldo(archivos_disponibles[tabla], tabla, avisos))

    vistos = set()
    resultado = []
    tipos_no_reconocidos = set()
    for tabla, id_resp, tipo_resp in todas_filas:
        clave = (tabla, id_resp, tipo_resp)
        if clave in vistos:
            continue
        vistos.add(clave)
        mapa = MCO_MEI if tabla in TABLAS_MEI_RESPALDO else MCO_ESTANDAR
        if tipo_resp in mapa:
            mco_codigo, mco_desc = mapa[tipo_resp]
        else:
            mco_codigo, mco_desc = None, None
            if tipo_resp is not None:
                tipos_no_reconocidos.add((tabla, tipo_resp))
        resultado.append((tabla, id_resp, tipo_resp, mco_codigo, mco_desc))

    for tabla, tipo_resp in sorted(tipos_no_reconocidos):
        avisos.append(f"⚠️ Respaldos: **{tabla}** tiene Tipo Respaldo={tipo_resp} no reconocido (se esperaba 1, 2 o 3) — MCO sin asignar en esas filas.")

    resultado.sort(key=lambda r: (r[0], str(r[1])))
    return resultado, avisos


def build_excel_respaldos(filas, avisos):
    """Arma el libro Excel del listado ID RESPALDO -> MCO."""
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Respaldos_MCO"
    headers = ["TABLA", "ID RESPALDO", "TIPO RESPALDO", "MCO", "DESCRIPCIÓN MCO"]
    ws.append(headers)
    for cell in ws[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    for tabla, id_resp, tipo_resp, mco_codigo, mco_desc in filas:
        ws.append([tabla, id_resp, tipo_resp, mco_codigo or "(sin asignar)", mco_desc or ""])
    widths = [12, 30, 14, 12, 32]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    if avisos:
        ws2 = wb.create_sheet("Avisos")
        ws2.append(["Aviso"])
        ws2["A1"].font = Font(bold=True)
        for a in avisos:
            ws2.append([a])
        ws2.column_dimensions["A"].width = 110

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


# ============================================================================
# CONSOLIDADO ID SERVICIO / SUBSERVICIO TERCERIZADO
# ============================================================================
_ST_SOLO_SUBSERVICIO = [f"ST_{i}" for i in range(3, 11)]  # ST_3..ST_10
_ST_AMBOS = [t for t in ST_TABLES if t not in _ST_SOLO_SUBSERVICIO]  # ST_11..ST_34
_TABLAS_SOLO_SERVICIO = ["GGI_1", "GGI_2", "GGM_1", "GGM_2", "GGM_3"]

TABLAS_SERVICIO_TERCERIZADO = _ST_SOLO_SUBSERVICIO + _ST_AMBOS + _TABLAS_SOLO_SERVICIO

MCO_SERVICIO_TERCERIZADO = {
    1: ("MCO_4", "Contrato Servicios"),
    2: ("MCO_9", "Orden de Compra servicios"),
    3: ("MCO_12", "Factura Servicios"),
}


def inferir_tipo_respaldo_servicio(id_servicio):
    """Igual criterio que inferir_tipo_respaldo_ogg, aplicado sobre el ID
    SERVICIO TERCERIZADO: reconoce CO-/OC-/FA- al inicio o rodeado de
    guiones. Si no calza con ninguno -> 9999 (sin MCO fijo; se debe asignar
    manualmente)."""
    texto = str(id_servicio).upper()
    if texto.startswith("CO-") or "-CO-" in texto:
        return 1
    if texto.startswith("OC-") or "-OC-" in texto:
        return 2
    if texto.startswith("FA-") or "-FA-" in texto:
        return 3
    return 9999


def leer_tabla_servicio_tercerizado(file_bytes, nombre_tabla, tiene_servicio, tiene_subservicio, avisos):
    """Lee CÓDIGO RECURSO, ID SERVICIO TERCERIZADO / ID SUBSERVICIO
    TERCERIZADO (el que exista en la tabla), MONTO ANUAL ACTIVADO y TOTAL
    GASTO ANUAL, por nombre de columna. Completa el campo faltante:
    - Si solo hay subservicio (ej. 'SS-CO-123'): servicio = subservicio sin
      la primera 'S' (-> 'S-CO-123').
    - Si solo hay servicio (ej. 'S-CO-123'): subservicio = 'S' + servicio
      (-> 'SS-CO-123').
    Devuelve lista de (tabla, cod_recurso, id_servicio, id_subservicio,
    gasto_anual, monto_activado)."""
    try:
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws = wb.active
        header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]
    except Exception as e:
        avisos.append(f"⚠️ Servicio Tercerizado: no se pudo leer **{nombre_tabla}** ({e}).")
        return []

    idx_recurso = _encontrar_columna(header_row, [["CODIGO", "RECURSO"], ["COD", "RECURSO"]])
    idx_activado = _encontrar_columna(header_row, [["MONTO", "ANUAL", "ACTIVADO"], ["MONTO", "ACTIVADO"]])
    idx_gasto = _encontrar_columna(
        header_row,
        [["TOTAL", "GASTO", "ANUAL"], ["GASTO", "ANUAL", "NO", "ACTIVADO"], ["TOTAL", "GASTO"], ["GASTO"]],
        evitar=["ACTIVADO", "%"],
    )
    idx_servicio = _encontrar_columna(header_row, [["ID", "SERVICIO", "TERCERIZADO"]], evitar=["SUB"]) if tiene_servicio else None
    idx_subservicio = _encontrar_columna(header_row, [["ID", "SUBSERVICIO", "TERCERIZADO"]]) if tiene_subservicio else None

    faltantes = []
    if idx_recurso is None:
        faltantes.append("CÓDIGO RECURSO")
    if idx_activado is None:
        faltantes.append("MONTO ANUAL ACTIVADO")
    if idx_gasto is None:
        faltantes.append("TOTAL GASTO ANUAL")
    if tiene_servicio and idx_servicio is None:
        faltantes.append("ID SERVICIO TERCERIZADO")
    if tiene_subservicio and idx_subservicio is None:
        faltantes.append("ID SUBSERVICIO TERCERIZADO")
    if faltantes:
        avisos.append(f"⚠️ Servicio Tercerizado: **{nombre_tabla}** no tiene columna(s): {', '.join(faltantes)} — se omite.")
        return []

    filas = []
    sin_s_inicial = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        cod_recurso = row[idx_recurso]
        if cod_recurso is None:
            continue
        cod_recurso = _a_numero(cod_recurso)
        cod_recurso = int(cod_recurso) if cod_recurso == int(cod_recurso) else cod_recurso
        gasto = _a_numero(row[idx_gasto])
        activado = _a_numero(row[idx_activado])

        id_serv = row[idx_servicio] if idx_servicio is not None else None
        id_subserv = row[idx_subservicio] if idx_subservicio is not None else None

        if id_serv is None and id_subserv is not None:
            texto_sub = str(id_subserv).strip()
            if texto_sub.upper().startswith("S"):
                id_serv = texto_sub[1:]
            else:
                id_serv = texto_sub
                sin_s_inicial += 1
        elif id_subserv is None and id_serv is not None:
            id_subserv = "S" + str(id_serv).strip()

        if id_serv is None and id_subserv is None:
            continue

        filas.append((nombre_tabla, cod_recurso, id_serv, id_subserv, gasto, activado))

    if sin_s_inicial:
        avisos.append(
            f"⚠️ Servicio Tercerizado: en **{nombre_tabla}**, {sin_s_inicial} ID Subservicio no empezaban con "
            f"'S' — se usaron tal cual para derivar el ID Servicio (revisar formato)."
        )
    return filas


def armar_trazabilidad_servicio_tercerizado(archivos_disponibles):
    """archivos_disponibles: dict {nombre_tabla: file_bytes}. Devuelve
    (filas_trazabilidad, filas_agregado, avisos).
    filas_trazabilidad: una fila por cada fila fuente, sin consolidar:
      (tabla, cod_recurso, id_servicio, id_subservicio, gasto_anual, monto_activado)
    filas_agregado: consolidado por ID SERVICIO TERCERIZADO (sumando gasto y
    activado de todas sus ocurrencias), con Tipo Respaldo y MCO:
      (id_servicio, gasto_total, monto_activado_total, tipo_respaldo, mco, descripcion_mco)
    """
    avisos = []
    todas_filas = []

    for tabla in _ST_SOLO_SUBSERVICIO:
        if tabla not in archivos_disponibles:
            avisos.append(f"ℹ️ Servicio Tercerizado: **{tabla}** no está cargada — se omite.")
            continue
        todas_filas.extend(leer_tabla_servicio_tercerizado(
            archivos_disponibles[tabla], tabla, tiene_servicio=False, tiene_subservicio=True, avisos=avisos))

    for tabla in _ST_AMBOS:
        if tabla not in archivos_disponibles:
            avisos.append(f"ℹ️ Servicio Tercerizado: **{tabla}** no está cargada — se omite.")
            continue
        todas_filas.extend(leer_tabla_servicio_tercerizado(
            archivos_disponibles[tabla], tabla, tiene_servicio=True, tiene_subservicio=True, avisos=avisos))

    for tabla in _TABLAS_SOLO_SERVICIO:
        if tabla not in archivos_disponibles:
            avisos.append(f"ℹ️ Servicio Tercerizado: **{tabla}** no está cargada — se omite.")
            continue
        todas_filas.extend(leer_tabla_servicio_tercerizado(
            archivos_disponibles[tabla], tabla, tiene_servicio=True, tiene_subservicio=False, avisos=avisos))

    # --- Hoja 1: trazabilidad (sin consolidar) ---
    filas_trazabilidad = todas_filas

    # --- Hoja 2: agregado por ID SERVICIO TERCERIZADO ---
    agregado = defaultdict(lambda: [0.0, 0.0])  # id_servicio -> [gasto, activado]
    for _, _, id_serv, _, gasto, activado in todas_filas:
        agregado[id_serv][0] += gasto
        agregado[id_serv][1] += activado

    filas_agregado = []
    for id_serv, (gasto_tot, activado_tot) in agregado.items():
        tipo = inferir_tipo_respaldo_servicio(id_serv)
        if tipo in MCO_SERVICIO_TERCERIZADO:
            mco_codigo, mco_desc = MCO_SERVICIO_TERCERIZADO[tipo]
        else:
            mco_codigo, mco_desc = "Asignar MCO", ""
        filas_agregado.append((id_serv, gasto_tot, activado_tot, tipo, mco_codigo, mco_desc))

    filas_trazabilidad.sort(key=lambda r: (r[0], str(r[2])))
    filas_agregado.sort(key=lambda r: str(r[0]))

    return filas_trazabilidad, filas_agregado, avisos


def build_excel_servicio_tercerizado(filas_trazabilidad, filas_agregado, avisos):
    wb = openpyxl.Workbook()
    ws1 = wb.active
    ws1.title = "Trazabilidad"
    headers1 = ["TABLA", "CÓDIGO RECURSO", "ID SERVICIO TERCERIZADO", "ID SUBSERVICIO TERCERIZADO", "GASTO ANUAL", "MONTO ACTIVADO"]
    ws1.append(headers1)
    for cell in ws1[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    for tabla, cod_recurso, id_serv, id_subserv, gasto, activado in filas_trazabilidad:
        ws1.append([tabla, cod_recurso, id_serv, id_subserv, round(gasto, 2), round(activado, 2)])
    widths1 = [12, 16, 26, 30, 16, 16]
    for i, w in enumerate(widths1, start=1):
        ws1.column_dimensions[get_column_letter(i)].width = w
    ws1.freeze_panes = "A2"
    for row in ws1.iter_rows(min_row=2, min_col=5, max_col=6):
        for cell in row:
            cell.number_format = "#,##0"

    ws2 = wb.create_sheet("Agregado")
    headers2 = ["ID SERVICIO TERCERIZADO", "GASTO ANUAL TOTAL", "MONTO ACTIVADO TOTAL", "TIPO RESPALDO", "MCO", "DESCRIPCIÓN MCO"]
    ws2.append(headers2)
    for cell in ws2[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    for id_serv, gasto_tot, activado_tot, tipo, mco_codigo, mco_desc in filas_agregado:
        ws2.append([id_serv, round(gasto_tot, 2), round(activado_tot, 2), tipo, mco_codigo, mco_desc])
    widths2 = [26, 18, 18, 14, 14, 26]
    for i, w in enumerate(widths2, start=1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"
    for row in ws2.iter_rows(min_row=2, min_col=2, max_col=3):
        for cell in row:
            cell.number_format = "#,##0"

    if avisos:
        ws3 = wb.create_sheet("Avisos")
        ws3.append(["Aviso"])
        ws3["A1"].font = Font(bold=True)
        for a in avisos:
            ws3.append([a])
        ws3.column_dimensions["A"].width = 110

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def identificar_tabla(nombre_archivo, tablas_conocidas):
    """Extrae el nombre de tabla (ej. 'ST_12', 'GPA_3') de un archivo subido,
    sin importar mayúsculas, extensión o sufijos. Devuelve None si no
    matchea ninguna tabla conocida de la lista dada."""
    base = nombre_archivo.upper().replace(".XLSX", "").replace(".XLS", "")
    for tabla in tablas_conocidas:
        if base == tabla or base.startswith(tabla + "_") or base.startswith(tabla + "-"):
            return tabla
    return None


def identificar_tabla_st(nombre_archivo):
    return identificar_tabla(nombre_archivo, ST_TABLES)


def identificar_tabla_gpa(nombre_archivo):
    return identificar_tabla(nombre_archivo, GPA_TABLES)


def read_rows(file_bytes):
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    return [r for r in ws.iter_rows(min_row=2, values_only=True) if r[0] is not None]


def familia(cod_servicio):
    return cod_servicio // 100


def redondear_para_sumar_100(valores, decimales=4):
    """Redondea una lista de fracciones que en teoría suman ~1.0 usando el
    método de mayor residuo (largest remainder / Hare-Niemeyer), de forma
    que la suma de los valores redondeados sea EXACTAMENTE 1.0 al número de
    decimales indicado. Evita el arrastre de error que ocurre al redondear
    cada valor de forma independiente (ej. una suma final de 99.99% o
    100.01% en vez de 100.00%)."""
    if not valores:
        return []
    escala = 10 ** decimales
    total_objetivo = round(sum(valores) * escala)
    escalados = [v * escala for v in valores]
    pisos = [int(x) for x in escalados]
    residuos = [x - p for x, p in zip(escalados, pisos)]
    suma_pisos = sum(pisos)
    faltante = total_objetivo - suma_pisos
    orden = sorted(range(len(valores)), key=lambda i: residuos[i], reverse=True)
    ajustados = list(pisos)
    for i in range(faltante):
        ajustados[orden[i % len(valores)]] += 1
    return [a / escala for a in ajustados]


def fusionar_cola_larga(items, piso=0.01, id_fusion="-1"):
    """Si hay demasiadas sub-claves para que todas alcancen el piso mínimo
    individualmente (n * piso >= 1.0), fusiona las de MENOR monto bajo
    id_fusion ('-1', el código que el maestro SISS reserva para 'sin cliente
    específico'), dejando como entradas individuales solo las que sí pueden
    alcanzar el piso con margen para variación proporcional.

    items: [(sub, monto), ...] donde 'sub' puede ser un valor simple (ej. el
    ID cliente) o una tupla (ej. (cod_servicio, id_cliente)). En este último
    caso, la fusión agrupa la 'cola' POR EL RESTO DE LA TUPLA (ej. por
    servicio) y reemplaza solo el último componente por id_fusion, de forma
    que cada fila fusionada siga siendo válida (ej. conserva su servicio).

    Devuelve (nuevos_items, se_fusiono)."""
    n = len(items)
    max_individuales = int(1.0 / piso) - 1
    if n <= max_individuales:
        return items, False

    items_ordenados = sorted(items, key=lambda x: x[1], reverse=True)
    individuales = items_ordenados[:max_individuales - 1]
    cola = items_ordenados[max_individuales - 1:]

    nuevos = list(individuales)
    if cola and isinstance(cola[0][0], tuple):
        # Agrupar la cola por el resto de la tupla (ej. por servicio),
        # reemplazando solo el último componente (ej. el cliente) por id_fusion.
        cola_por_grupo = defaultdict(float)
        for sub, monto in cola:
            clave_grupo = sub[:-1]
            cola_por_grupo[clave_grupo] += monto
        existentes = {s[0][:-1]: i for i, s in enumerate(nuevos) if s[0][-1] == id_fusion}
        for clave_grupo, monto in cola_por_grupo.items():
            nuevo_sub = clave_grupo + (id_fusion,)
            if clave_grupo in existentes:
                idx = existentes[clave_grupo]
                nuevos[idx] = (nuevos[idx][0], nuevos[idx][1] + monto)
            else:
                nuevos.append((nuevo_sub, monto))
    else:
        monto_cola = sum(m for _, m in cola)
        fusion_existente = False
        for i, (cid, m) in enumerate(nuevos):
            if cid == id_fusion:
                nuevos[i] = (cid, m + monto_cola)
                fusion_existente = True
                break
        if not fusion_existente:
            nuevos.append((id_fusion, monto_cola))
    return nuevos, True


def aplicar_piso_minimo(fracciones, piso=0.01):
    """Ajusta una lista de fracciones (que suman ~1.0) para que NINGUNA
    quede por debajo de 'piso' (ej. 0.01 = 1%), redistribuyendo
    proporcionalmente el resto entre las demás ("water-filling"). Se usa en
    CYG_9 para cumplir la regla del maestro SISS de no declarar registros
    con % de dedicación igual a 0 (ni prácticamente 0), sin dejar de listar
    a todos los clientes que efectivamente reciben una porción del gasto.
    Si hay demasiadas entradas para que todas alcancen el piso (n*piso >= 1),
    se reparte equitativamente entre todas como mejor esfuerzo."""
    n = len(fracciones)
    if n == 0:
        return []
    if n * piso >= 1.0:
        return [1.0 / n] * n

    resultado = [0.0] * n
    restantes = list(range(n))
    masa_restante = 1.0

    while True:
        suma_raw = sum(fracciones[i] for i in restantes)
        candidatos_bajo_piso = []
        for i in restantes:
            val = (fracciones[i] / suma_raw * masa_restante) if suma_raw > 0 else (masa_restante / len(restantes))
            if val < piso:
                candidatos_bajo_piso.append(i)

        if not candidatos_bajo_piso:
            for i in restantes:
                resultado[i] = (fracciones[i] / suma_raw * masa_restante) if suma_raw > 0 else (masa_restante / len(restantes))
            break

        for i in candidatos_bajo_piso:
            resultado[i] = piso
            restantes.remove(i)
            masa_restante -= piso

        if not restantes:
            break
        if len(restantes) * piso >= masa_restante:
            for i in restantes:
                resultado[i] = masa_restante / len(restantes)
            break

    return resultado


def normalizar_con_piso_minimo(items_por_clave, piso=0.01, decimales=4):
    """Como _normalizar_por_clave, pero garantiza que ningún % quede por
    debajo de 'piso' (ej. 0.01 = 1%). Si hay demasiadas sub-claves (ej.
    clientes) para que todas alcancen el piso, primero fusiona la cola de
    menor monto bajo el identificador '-1' (fusionar_cola_larga), y luego
    aplica el piso mínimo (aplicar_piso_minimo) sobre el conjunto resultante.
    Devuelve también, por clave, si hubo fusión (para poder avisar)."""
    resultado = {}
    hubo_fusion = {}
    for clave, items in items_por_clave.items():
        total = sum(g for _, g in items)
        if total == 0:
            continue
        items_fusionados, se_fusiono = fusionar_cola_larga(items, piso=piso, id_fusion="-1")
        hubo_fusion[clave] = se_fusiono
        fracciones_raw = [g / total for _, g in items_fusionados]
        fracciones_piso = aplicar_piso_minimo(fracciones_raw, piso)
        pcts = redondear_para_sumar_100(fracciones_piso, decimales)
        resultado[clave] = [(sub, pct, pct * total) for (sub, _), pct in zip(items_fusionados, pcts)]
    return resultado, hubo_fusion


def procesar_familia_plana(agg, EMPRESA, PERIODO, ANIO, SECTOR, tablas_specs, params):
    by_recurso = defaultdict(lambda: [0.0, 0.0])
    for filas, idx_r, idx_a, idx_g in tablas_specs:
        for r in filas:
            cod_recurso = r[idx_r]
            monto_act = _a_numero(r[idx_a])
            total_gasto = _a_numero(r[idx_g])
            by_recurso[cod_recurso][0] += total_gasto
            by_recurso[cod_recurso][1] += monto_act

    for cod_recurso, (gasto_no_act, monto_act) in by_recurso.items():
        overrides = params.get(cod_recurso, [])
        pct_reg = 1.0 - sum(p for _, p in overrides)
        fam = familia(1101)
        k = (EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam)
        agg[k][0] += gasto_no_act * pct_reg
        agg[k][1] += monto_act  # monto activado SIEMPRE 100% al servicio regulado por defecto
        for cod_serv_noreg, pct in overrides:
            fam2 = familia(cod_serv_noreg)
            k2 = (EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam2)
            agg[k2][0] += gasto_no_act * pct
            # el monto activado no se traspasa a servicios no regulados
    return by_recurso


def safe_read_rows(file_bytes, etiqueta, avisos):
    """Como read_rows, pero tolera archivo ausente (None) o con error de
    lectura: registra un aviso y devuelve lista vacía en vez de fallar."""
    if file_bytes is None:
        avisos.append(f"⚠️ Falta la tabla **{etiqueta}** — se excluye del cálculo (esa familia/recurso queda en 0 o incompleto).")
        return []
    try:
        filas = read_rows(file_bytes)
    except Exception as e:
        avisos.append(f"⚠️ No se pudo leer **{etiqueta}** ({e}) — se excluye del cálculo.")
        return []
    if len(filas) == 0:
        avisos.append(f"ℹ️ La tabla **{etiqueta}** está vacía (sin filas de datos).")
    return filas


def build_rep2(fb, ggm_params, ogg_params, mei_params, st_files, st_params, gpa_files):
    avisos = []

    grh8 = safe_read_rows(fb.get("grh8"), "GRH_8", avisos)
    grh11 = safe_read_rows(fb.get("grh11"), "GRH_11", avisos)
    gcp4 = safe_read_rows(fb.get("gcp4"), "GCP_4", avisos)
    gcp5 = safe_read_rows(fb.get("gcp5"), "GCP_5", avisos)
    ggv4 = safe_read_rows(fb.get("ggv4"), "GGV_4", avisos)
    ggv5 = safe_read_rows(fb.get("ggv5"), "GGV_5", avisos)
    ggi5 = safe_read_rows(fb.get("ggi5"), "GGI_5", avisos)
    ggm_tablas = [safe_leer_tabla_plana(fb.get(f"ggm{i}"), f"GGM_{i}", avisos) for i in range(1, 6)]
    ogg5 = safe_leer_tabla_plana(fb.get("ogg5"), "OGG_5", avisos)
    mei1 = safe_leer_tabla_mei(fb.get("mei1"), "MEI_1", avisos, necesita_obra_nbi=True)
    mei2 = safe_leer_tabla_mei(fb.get("mei2"), "MEI_2", avisos, necesita_actividad=True)
    mei3 = safe_leer_tabla_mei(fb.get("mei3"), "MEI_3", avisos, necesita_actividad=True)
    mei4 = safe_leer_tabla_mei(fb.get("mei4"), "MEI_4", avisos, necesita_actividad=True)

    agg = defaultdict(lambda: [0.0, 0.0])

    # EMPRESA/PERIODO/AÑO/SECTOR: se toman del primer archivo disponible,
    # ya que si falta GRH_8 no podemos asumir esos valores desde ahí.
    EMPRESA = PERIODO = ANIO = SECTOR = None
    for candidato in [grh8, gcp4, ggv4, ggi5, ogg5] + ggm_tablas + [mei1, mei2, mei3, mei4]:
        if candidato:
            EMPRESA, PERIODO, ANIO, SECTOR = candidato[0][0], candidato[0][1], candidato[0][2], candidato[0][3]
            break
    if EMPRESA is None:
        for filas_st in st_files.values():
            if filas_st:
                EMPRESA, PERIODO, ANIO, SECTOR = filas_st[0][0], filas_st[0][1], filas_st[0][2], filas_st[0][3]
                break

    # --- GRH ---
    shares_persona = defaultdict(list)
    for r in grh11:
        cod_reg, cod_noreg, pct = r[7], r[8], r[9]
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        shares_persona[(r[4], r[5])].append((cod_serv, pct))

    for r in grh8:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r
        for cod_serv, pct in shares_persona.get((persona, cargo), []):
            fam = familia(cod_serv)
            k = (empresa, periodo, anio, sector, cod_recurso, fam)
            agg[k][0] += total_gasto * pct
            agg[k][1] += monto_act * pct

    # --- GCP ---
    for r in gcp5:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        fam = familia(cod_serv)
        k = (empresa, periodo, anio, sector, cod_recurso, fam)
        agg[k][0] += gasto_no_act

    shares_gcp_recurso = defaultdict(list)
    for r in gcp5:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        shares_gcp_recurso[(persona, cargo, cod_recurso)].append((cod_serv, pct))

    for r in gcp4:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r
        if monto_act == 0:
            continue
        for cod_serv, pct in shares_gcp_recurso.get((persona, cargo, cod_recurso), []):
            fam = familia(cod_serv)
            k = (empresa, periodo, anio, sector, cod_recurso, fam)
            agg[k][1] += monto_act * pct

    # --- GGV ---
    shares_activo = defaultdict(list)
    for r in ggv5:
        empresa, periodo, anio, sector, id_activo, total_no_act, cod_reg, cod_noreg, pct = r[:9]
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        shares_activo[id_activo].append((cod_serv, pct))

    for r in ggv4:
        empresa, periodo, anio, sector, id_activo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r[:10]
        for cod_serv, pct in shares_activo.get(id_activo, []):
            fam = familia(cod_serv)
            k = (empresa, periodo, anio, sector, cod_recurso, fam)
            agg[k][0] += total_gasto * pct
            agg[k][1] += monto_act * pct

    # --- GGI ---
    for r in ggi5:
        empresa, periodo, anio, sector, id_inmueble, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r[:10]
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        fam = familia(cod_serv)
        k = (empresa, periodo, anio, sector, cod_recurso, fam)
        agg[k][0] += gasto_no_act

    # --- GGM ---
    ggm_by_recurso = procesar_familia_plana(
        agg, EMPRESA, PERIODO, ANIO, SECTOR,
        [(t, 4, 5, 6) for t in ggm_tablas],  # lectura robusta: recurso=4, activado=5, gasto=6
        ggm_params
    )

    # --- OGG ---
    ogg_by_recurso = procesar_familia_plana(
        agg, EMPRESA, PERIODO, ANIO, SECTOR,
        [(ogg5, 4, 5, 6)],  # lectura robusta: recurso=4, activado=5, gasto=6
        ogg_params
    )

    # --- MEI (lectura robusta por nombre de columna: recurso=4, activado=6, gasto=7) ---
    mei_by_recurso = procesar_familia_plana(
        agg, EMPRESA, PERIODO, ANIO, SECTOR,
        [(mei1, 4, 6, 7), (mei2, 4, 6, 7), (mei3, 4, 6, 7), (mei4, 4, 6, 7)],
        mei_params
    )

    # --- ST: Servicios Tercerizados (29 tablas, lectura por nombre de columna) ---
    st_by_recurso = defaultdict(lambda: [0.0, 0.0])
    tablas_st_faltantes = [t for t in ST_TABLES if t not in st_files]
    if tablas_st_faltantes:
        avisos.append(
            f"⚠️ Faltan {len(tablas_st_faltantes)} tabla(s) ST: {', '.join(tablas_st_faltantes)} "
            "— se excluyen del cálculo (sus códigos de recurso quedan en 0)."
        )
    for nombre_tabla, filas_st in st_files.items():
        if len(filas_st) == 0:
            avisos.append(f"ℹ️ La tabla **{nombre_tabla}** está vacía (sin filas de datos).")
            continue
        for cod_recurso, monto_act, total_gasto in filas_st:
            overrides = st_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            fam = familia(1101)
            k = (EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam)
            agg[k][0] += total_gasto * pct_reg
            agg[k][1] += monto_act  # monto activado SIEMPRE 100% al servicio regulado por defecto
            st_by_recurso[cod_recurso][0] += total_gasto * pct_reg
            st_by_recurso[cod_recurso][1] += monto_act
            for cod_serv_noreg, pct in overrides:
                fam2 = familia(cod_serv_noreg)
                k2 = (EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam2)
                agg[k2][0] += total_gasto * pct
                # el monto activado no se traspasa a servicios no regulados
                st_by_recurso[cod_recurso][0] += total_gasto * pct

    # --- GPA: Gasto Prestaciones Asociadas (6 tablas, 100% a servicio fijo por tabla) ---
    gpa_by_recurso = defaultdict(lambda: [0.0, 0.0])
    gpa_detalle = []  # (tabla, cod_recurso, servicio_asignado, gasto, activado) para reporte
    tablas_gpa_faltantes = [t for t in GPA_TABLES if t not in gpa_files]
    if tablas_gpa_faltantes:
        avisos.append(
            f"⚠️ Faltan {len(tablas_gpa_faltantes)} tabla(s) GPA: {', '.join(tablas_gpa_faltantes)} "
            "— se excluyen del cálculo (sus códigos de recurso quedan en 0)."
        )
    for nombre_tabla, filas_gpa in gpa_files.items():
        if len(filas_gpa) == 0:
            avisos.append(f"ℹ️ La tabla **{nombre_tabla}** está vacía (sin filas de datos).")
            continue
        servicio_fijo = GPA_TABLE_TO_SERVICIO[nombre_tabla]
        fam = familia(servicio_fijo)
        tabla_g, tabla_a = 0.0, 0.0
        recursos_en_tabla = set()
        for cod_recurso, monto_act, total_gasto in filas_gpa:
            k = (EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam)
            agg[k][0] += total_gasto
            agg[k][1] += monto_act
            gpa_by_recurso[cod_recurso][0] += total_gasto
            gpa_by_recurso[cod_recurso][1] += monto_act
            tabla_g += total_gasto
            tabla_a += monto_act
            recursos_en_tabla.add(cod_recurso)
        recurso_label = ", ".join(str(c) for c in sorted(recursos_en_tabla))
        gpa_detalle.append((nombre_tabla, recurso_label, servicio_fijo, tabla_g, tabla_a))

    # --- Normalizar % por recurso (ajustado para sumar EXACTAMENTE 100%) ---
    grupos_por_recurso = defaultdict(list)  # cod_recurso -> [(key_completa, g, a), ...]
    for key, (g, a) in agg.items():
        cod_recurso = key[4]
        grupos_por_recurso[cod_recurso].append((key, g, a))

    final_rows = []
    for cod_recurso, items in grupos_por_recurso.items():
        items = [(k, g, a) for k, g, a in items if not (abs(g) < 1e-9 and abs(a) < 1e-9)]
        if not items:
            continue
        total_g = sum(g for _, g, _ in items)
        total_a = sum(a for _, _, a in items)

        idx_g_pos = [i for i, (_, g, _) in enumerate(items) if abs(g) > 1e-9]
        fracciones_g = [items[i][1] / total_g for i in idx_g_pos] if total_g else []
        pcts_g_corr = redondear_para_sumar_100(fracciones_g, 4) if fracciones_g else []
        pct_g_final = {i: 0.0 for i in range(len(items))}
        for idx, pct in zip(idx_g_pos, pcts_g_corr):
            pct_g_final[idx] = pct

        idx_a_pos = [i for i, (_, _, a) in enumerate(items) if abs(a) > 1e-9]
        fracciones_a = [items[i][2] / total_a for i in idx_a_pos] if total_a else []
        pcts_a_corr = redondear_para_sumar_100(fracciones_a, 4) if fracciones_a else []
        pct_a_final = {i: 0.0 for i in range(len(items))}
        for idx, pct in zip(idx_a_pos, pcts_a_corr):
            pct_a_final[idx] = pct

        for i, (key, g, a) in enumerate(items):
            empresa, periodo, anio, sector, _, fam = key
            final_rows.append([
                empresa, periodo, anio, sector, cod_recurso, fam,
                round(pct_g_final[i] * 100, 2), round(g, 2),
                round(pct_a_final[i] * 100, 2), round(a, 2),
            ])

    final_rows.sort(key=lambda r: (r[4], r[5]))

    # --- Validaciones (solo entre familias con datos disponibles) ---
    recursos_grh = set(r[6] for r in grh8)
    recursos_gcp = set(r[6] for r in gcp4)
    recursos_ggv = set(r[5] for r in ggv4)
    recursos_ggi = set(r[5] for r in ggi5)
    recursos_ggm = set(ggm_by_recurso.keys())
    recursos_ogg = set(ogg_by_recurso.keys())
    recursos_mei = set(mei_by_recurso.keys())
    recursos_st = set(st_by_recurso.keys())
    recursos_gpa = set(gpa_by_recurso.keys())

    def validar(recursos, sum_gasto_fuente, sum_act_fuente):
        sum_gasto_rep2 = sum(r[7] for r in final_rows if r[4] in recursos)
        sum_act_rep2 = sum(r[9] for r in final_rows if r[4] in recursos)
        return {
            "diff_gasto": sum_gasto_rep2 - sum_gasto_fuente,
            "diff_act": sum_act_rep2 - sum_act_fuente,
        }

    checks = {}
    if grh8:
        checks["GRH"] = validar(recursos_grh, sum(r[10] for r in grh8), sum(r[8] for r in grh8))
    if gcp4:
        checks["GCP"] = validar(recursos_gcp, sum(r[10] for r in gcp4), sum(r[8] for r in gcp4))
    if ggv4:
        checks["GGV"] = validar(recursos_ggv, sum(r[9] for r in ggv4), sum(r[7] for r in ggv4))
    if ggi5:
        checks["GGI"] = validar(recursos_ggi, sum(r[6] for r in ggi5), 0)
    if ggm_by_recurso:
        checks["GGM"] = validar(recursos_ggm, sum(v[0] for v in ggm_by_recurso.values()), sum(v[1] for v in ggm_by_recurso.values()))
    if ogg_by_recurso:
        checks["OGG"] = validar(recursos_ogg, sum(v[0] for v in ogg_by_recurso.values()), sum(v[1] for v in ogg_by_recurso.values()))
    if mei_by_recurso:
        checks["MEI"] = validar(recursos_mei, sum(v[0] for v in mei_by_recurso.values()), sum(v[1] for v in mei_by_recurso.values()))
    if st_by_recurso:
        checks["ST"] = validar(recursos_st, sum(v[0] for v in st_by_recurso.values()), sum(v[1] for v in st_by_recurso.values()))
    if gpa_by_recurso:
        checks["GPA"] = validar(recursos_gpa, sum(v[0] for v in gpa_by_recurso.values()), sum(v[1] for v in gpa_by_recurso.values()))

    familia_map = {}
    for c in recursos_grh: familia_map[c] = "GRH - Gastos Recursos Humanos"
    for c in recursos_gcp: familia_map[c] = "GCP - Gastos Generales de Personal"
    for c in recursos_ggv: familia_map[c] = "GGV - Gastos Generales Vehículos y Equipos"
    for c in recursos_ggi: familia_map[c] = "GGI - Gastos Generales Bienes Inmuebles"
    for c in recursos_ggm: familia_map[c] = "GGM - Gastos Generales Bienes Muebles"
    for c in recursos_ogg: familia_map[c] = "OGG - Otros Gastos Generales"
    for c in recursos_mei: familia_map[c] = "MEI - Materiales e Insumos"
    for c in recursos_st: familia_map[c] = "ST - Servicios Tercerizados"
    for c in recursos_gpa: familia_map[c] = "GPA - Gasto Prestaciones Asociadas"

    by_recurso_planas = {"GGM": ggm_by_recurso, "OGG": ogg_by_recurso, "MEI": mei_by_recurso, "ST": st_by_recurso}
    return final_rows, checks, familia_map, by_recurso_planas, gpa_detalle, avisos


def build_excel(final_rows, familia_map, by_recurso_planas, params_by_familia, gpa_detalle=None, avisos=None, template_bytes=None):

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "REP_2"

    header_fill = PatternFill("solid", start_color="1F4E78", end_color="1F4E78")
    header_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
    data_font = Font(name="Arial", size=10)
    thin = Side(style="thin", color="D9D9D9")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    ws.append(HEADERS)
    for c in range(1, len(HEADERS) + 1):
        cell = ws.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws.row_dimensions[1].height = 30

    for row in final_rows:
        ws.append(row)

    for r in range(2, ws.max_row + 1):
        for c in range(1, len(HEADERS) + 1):
            cell = ws.cell(row=r, column=c)
            cell.font = data_font
            cell.border = border
            if c in (7, 9):
                cell.number_format = "0.00"
            elif c in (8, 10):
                cell.number_format = '#,##0;(#,##0);"-"'
            else:
                cell.alignment = Alignment(horizontal="center")

    widths = [14, 16, 14, 20, 14, 22, 18, 18, 18, 18]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"

    # Resumen por familia
    ws2 = wb.create_sheet("Resumen_por_familia_gasto")
    ws2.append(["Familia de Gasto", "GASTO ANUAL (no activado)", "MONTO ACTIVADO", "TOTAL"])
    for c in range(1, 5):
        cell = ws2.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")
        cell.border = border

    resumen = defaultdict(lambda: [0.0, 0.0])
    for row in final_rows:
        fg = familia_map.get(row[4], "Otro")
        resumen[fg][0] += row[7]
        resumen[fg][1] += row[9]

    orden_familias = [
        "GRH - Gastos Recursos Humanos", "GCP - Gastos Generales de Personal",
        "GGV - Gastos Generales Vehículos y Equipos", "GGI - Gastos Generales Bienes Inmuebles",
        "GGM - Gastos Generales Bienes Muebles", "OGG - Otros Gastos Generales",
        "MEI - Materiales e Insumos", "ST - Servicios Tercerizados",
        "GPA - Gasto Prestaciones Asociadas",
    ]
    for fg in orden_familias:
        if fg in resumen:
            g, a = resumen[fg]
            ws2.append([fg, round(g, 2), round(a, 2), round(g + a, 2)])
    total_g = sum(v[0] for v in resumen.values())
    total_a = sum(v[1] for v in resumen.values())
    ws2.append(["TOTAL REP_2", round(total_g, 2), round(total_a, 2), round(total_g + total_a, 2)])

    for r in range(2, ws2.max_row + 1):
        for c in range(1, 5):
            cell = ws2.cell(row=r, column=c)
            cell.font = Font(name="Arial", size=10, bold=(r == ws2.max_row))
            cell.border = border
            if c in (2, 3, 4):
                cell.number_format = '#,##0;(#,##0);"-"'
    ws2.column_dimensions["A"].width = 42
    for col in ["B", "C", "D"]:
        ws2.column_dimensions[col].width = 20

    # Parametrización unificada GGM+OGG+MEI
    ws3 = wb.create_sheet("Parametrizacion_familias_planas")
    ws3.append(["Familia", "CÓDIGO RECURSO", "GASTO ANUAL TOTAL (100%)", "CÓDIGO SERVICIO NO REGULADO", "SERVICIO NO REGULADO", "% ASIGNADO"])
    for c in range(1, 7):
        cell = ws3.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border
    ws3.row_dimensions[1].height = 30

    for nombre_familia, by_recurso in by_recurso_planas.items():
        params = params_by_familia.get(nombre_familia, {})
        for cod_recurso in sorted(by_recurso.keys()):
            gasto_total = by_recurso[cod_recurso][0] + by_recurso[cod_recurso][1]
            overrides = params.get(cod_recurso, [])
            if not overrides:
                ws3.append([nombre_familia, cod_recurso, round(gasto_total, 2), "(sin parametrizar -> 100% a servicio 1101)", "", ""])
            else:
                for cod_serv, pct in overrides:
                    ws3.append([nombre_familia, cod_recurso, round(gasto_total, 2), cod_serv, SERVICIOS_NO_REGULADOS.get(cod_serv, ""), pct])

    for r in range(2, ws3.max_row + 1):
        for c in range(1, 7):
            cell = ws3.cell(row=r, column=c)
            cell.font = data_font
            cell.border = border
            if c == 3:
                cell.number_format = '#,##0;(#,##0);"-"'
            if c == 6 and isinstance(ws3.cell(row=r, column=6).value, float):
                cell.number_format = "0.00"
    ws3.column_dimensions["A"].width = 10
    ws3.column_dimensions["B"].width = 16
    ws3.column_dimensions["C"].width = 22
    ws3.column_dimensions["D"].width = 42
    ws3.column_dimensions["E"].width = 55
    ws3.column_dimensions["F"].width = 12

    # Asignación GPA (fija, sin parametrización de servicios no regulados)
    if gpa_detalle:
        ws3b = wb.create_sheet("Asignacion_GPA")
        ws3b.append(["Tabla", "CÓDIGO(S) RECURSO", "SERVICIO REGULADO ASIGNADO (100%)", "GASTO ANUAL", "MONTO ACTIVADO"])
        for c in range(1, 6):
            cell = ws3b.cell(row=1, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        ws3b.row_dimensions[1].height = 30
        for tabla, cod_recurso, servicio, g, a in sorted(gpa_detalle):
            ws3b.append([tabla, cod_recurso, servicio, round(g, 2), round(a, 2)])
        for r in range(2, ws3b.max_row + 1):
            for c in range(1, 6):
                cell = ws3b.cell(row=r, column=c)
                cell.font = data_font
                cell.border = border
                if c in (4, 5):
                    cell.number_format = '#,##0;(#,##0);"-"'
        ws3b.column_dimensions["A"].width = 12
        ws3b.column_dimensions["B"].width = 16
        ws3b.column_dimensions["C"].width = 32
        ws3b.column_dimensions["D"].width = 18
        ws3b.column_dimensions["E"].width = 18

    # Catálogo
    ws4 = wb.create_sheet("Catalogo_Servicios_No_Regulados")
    ws4.append(["CÓDIGO SERVICIO NO REGULADO", "SERVICIO NO REGULADO"])
    for c in range(1, 3):
        cell = ws4.cell(row=1, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
    for cod, desc in SERVICIOS_NO_REGULADOS.items():
        ws4.append([cod, desc])
    for r in range(2, ws4.max_row + 1):
        for c in range(1, 3):
            cell = ws4.cell(row=r, column=c)
            cell.font = data_font
            cell.border = border
    ws4.column_dimensions["A"].width = 22
    ws4.column_dimensions["B"].width = 70

    # Diccionario SISS
    if template_bytes is not None:
        try:
            tmpl = openpyxl.load_workbook(io.BytesIO(template_bytes), data_only=True)
            src_name = "Diccionario_REP_2" if "Diccionario_REP_2" in tmpl.sheetnames else tmpl.sheetnames[0]
            src_ws = tmpl[src_name]
            dst_ws = wb.create_sheet("Diccionario_REP_2")
            for row in src_ws.iter_rows():
                for cell in row:
                    dst_ws.cell(row=cell.row, column=cell.column, value=cell.value)
        except Exception:
            pass

    # Avisos de carga (tablas faltantes o vacías, generación parcial)
    if avisos:
        ws5 = wb.create_sheet("Avisos_Carga")
        ws5.append(["Aviso"])
        cell = ws5.cell(row=1, column=1)
        cell.font = header_font
        cell.fill = header_fill
        cell.border = border
        for aviso in avisos:
            ws5.append([aviso.replace("**", "")])
        for r in range(2, ws5.max_row + 1):
            cell = ws5.cell(row=r, column=1)
            cell.font = data_font
            cell.border = border
            cell.alignment = Alignment(wrap_text=True, vertical='top')
        ws5.column_dimensions["A"].width = 110

    out = io.BytesIO()
    wb.save(out)
    out.seek(0)
    return out


def panel_parametrizacion(nombre, recursos_disponibles, session_key):
    """Renderiza un panel de parametrización reutilizable para GGM/OGG/MEI."""
    if session_key not in st.session_state:
        st.session_state[session_key] = []

    st.markdown(f"**{nombre}**")
    col1, col2, col3, col4 = st.columns([1.2, 2.5, 1, 0.8])
    with col1:
        sel_recurso = st.selectbox("Código Recurso", recursos_disponibles, key=f"sel_recurso_{session_key}")
    with col2:
        sel_servicio = st.selectbox(
            "Servicio no regulado destino",
            options=list(SERVICIOS_NO_REGULADOS.keys()),
            format_func=lambda c: f"{c} - {SERVICIOS_NO_REGULADOS[c]}",
            key=f"sel_servicio_{session_key}",
        )
    with col3:
        sel_pct = st.number_input(
            "% asignado", min_value=0.0, max_value=100.0, value=10.0,
            step=0.0000000001, format="%.10f", key=f"sel_pct_{session_key}",
            help="Puedes ingresar hasta 10 decimales para minimizar la diferencia de redondeo en el monto resultante (con montos de cientos de millones, 6 decimales aún puede dejar una diferencia de $10-30; con 9-10 decimales queda en centavos). La columna '% ASIGNADO' del reporte final siempre se mostrará con 2 decimales (formato regulatorio), pero el MONTO se calcula con la precisión completa que ingreses aquí.",
        )
    with col4:
        st.write("")
        st.write("")
        if st.button("Agregar", key=f"add_{session_key}"):
            st.session_state[session_key].append((sel_recurso, sel_servicio, sel_pct / 100.0))

    if st.session_state[session_key]:
        for i, (cod_r, cod_s, pct) in enumerate(st.session_state[session_key]):
            c1, c2, c3, c4 = st.columns([1.2, 2.5, 1, 0.8])
            c1.write(cod_r)
            c2.write(f"{cod_s} - {SERVICIOS_NO_REGULADOS.get(cod_s, '')}")
            c3.write(f"{pct:.9%}")
            if c4.button("Quitar", key=f"del_{session_key}_{i}"):
                st.session_state[session_key].pop(i)
                st.rerun()
    else:
        st.caption(f"Sin parametrización: 100% de todos los recursos {nombre} va al servicio 1101.")

    params = defaultdict(list)
    for cod_r, cod_s, pct in st.session_state[session_key]:
        params[cod_r].append((cod_s, pct))
    return dict(params)



# ============================================================================
# REP_3 - Costos y Gastos No Activados de Servicios Regulados - Proceso
# ============================================================================
# Códigos de actividad válidos según MAE_1 del Maestro SISS (617 códigos).
# Se usa para validar que los códigos de actividad de las tablas ST (y
# alertar si alguno no corresponde a una actividad SISS reconocida).
CODIGOS_ACTIVIDAD_VALIDOS = frozenset([
    1010101, 1010201, 1010301, 1010302, 1010303, 1010304, 1010305, 1010306, 1010307, 1020101, 1020201, 1020301, 1020302, 1020303, 1020304,
    1030101, 1030201, 1030202, 1030301, 1030302, 1030303, 1030304, 1030401, 1030402, 1030403, 1030404, 1030501, 1030502, 1030503, 1030504,
    1030505, 1030601, 1030602, 1030603, 1030604, 1030605, 1030701, 1030702, 1030703, 1030704, 1030705, 1030801, 1030802, 1030803, 1030901,
    1030902, 1030903, 1031001, 1031002, 1031101, 1031201, 1031202, 1040101, 1040201, 1040301, 1040302, 1040303, 1040304, 1040305, 1040306,
    1040307, 1040308, 1040309, 1040401, 1050101, 1050102, 1060101, 1060102, 1060201, 1060202, 1060203, 1060204, 1060205, 1060301, 1060302,
    2010101, 2010201, 2010301, 2010302, 2010303, 2010304, 2010305, 2010306, 2010307, 2020101, 2020102, 2020103, 2020201, 2020301, 2030101,
    2030201, 2030301, 2040101, 2040102, 2040103, 2040104, 2040201, 2040301, 2040401, 2040402, 2040403, 2040404, 2040405, 2040501, 2040601,
    2040701, 2040702, 2040703, 2040704, 2040705, 2040801, 2040802, 2040901, 2040902, 2040903, 2040904, 2041001, 3010101, 3010201, 3010301,
    3010401, 3010501, 3010601, 3010701, 3010801, 3010901, 3011001, 3011101, 3011201, 3011301, 3020101, 3020201, 3020301, 3020401, 3020501,
    3030101, 3030201, 3030301, 3030302, 3030303, 3030304, 3030401, 3030402, 3030403, 3030404, 3030501, 3030502, 3030503, 3030504, 3030505,
    3030601, 3030602, 3030603, 3030701, 3030702, 3030703, 3030801, 3030802, 3030901, 3031001, 3031002, 3040101, 3040201, 3040301, 3040401,
    3040501, 3050101, 3050201, 3050301, 3050302, 3050303, 3050401, 3050402, 3050501, 3050502, 3060101, 3060201, 3060301, 3060302, 3060303,
    3060304, 3060401, 3060402, 3060403, 3060404, 3060501, 3060502, 3060503, 3060504, 3070101, 3070201, 3070301, 3070302, 3080101, 3080102,
    3080103, 3090101, 3090201, 3090301, 3090302, 3090303, 3090401, 3090402, 3090403, 3090501, 3090502, 3090503, 3090504, 3100101, 3100201,
    3100301, 3100302, 3100303, 3100304, 3100305, 3100401, 3100402, 3100403, 3100404, 3100405, 3100501, 3100502, 3100503, 3100504, 3110101,
    3110201, 3110301, 3110302, 3110303, 3110304, 3120101, 3120201, 3120301, 3120401, 3120402, 3120403, 3120404, 3120405, 3120501, 3120601,
    3130101, 3130201, 3130301, 3130401, 3130501, 4010101, 4010102, 4010103, 4010104, 4010105, 4010201, 4010202, 4010203, 4010204, 4010205,
    4010206, 4010207, 4010301, 4010302, 4010303, 4010304, 4010401, 4010402, 4010403, 4010404, 4010405, 4010406, 4010407, 4010408, 4020101,
    4020102, 4020103, 4020201, 4020202, 4020203, 4020301, 4020302, 4020303, 4020304, 4020305, 4020306, 4020307, 4020308, 4020309, 4020310,
    4020311, 4020312, 4020313, 4020314, 4020315, 4020316, 4020317, 4020318, 4020319, 4020320, 4020321, 4020322, 4020323, 4020324, 4020325,
    4020326, 4020327, 4020401, 4020402, 4020403, 4020501, 4020502, 4020503, 4020601, 4020602, 4030101, 4030102, 4030103, 4030104, 4030105,
    4030106, 4030107, 4030108, 4030201, 4030202, 4030203, 4030204, 4030205, 4030206, 4030207, 4030208, 4030209, 4030210, 4030211, 4030301,
    4030302, 4030303, 4030304, 4030305, 4030306, 4030307, 4030308, 4030309, 4030401, 4030402, 4030403, 4030404, 5010101, 5010102, 5010103,
    5010104, 5010105, 5010106, 5010107, 5010108, 5010201, 5010202, 5010203, 5010204, 5010205, 5010206, 5010301, 5010302, 5010303, 5010304,
    5010305, 5010306, 5010307, 5010308, 5010401, 5010402, 5010403, 5010404, 5010405, 5010406, 5010407, 5010501, 5010502, 5010503, 5010504,
    5010505, 5010506, 5010601, 5010602, 5010603, 5010604, 5010605, 5010606, 5010607, 5010608, 5020101, 5020102, 5020103, 5020104, 5020105,
    5020106, 5020107, 5020108, 5020109, 5020110, 5020111, 5020112, 5020113, 5020114, 5020115, 5020201, 5020202, 5020203, 5020204, 5020205,
    5020206, 5020207, 5020208, 5020209, 5020210, 5020211, 5020212, 5020301, 5020302, 5020303, 5020304, 5020305, 5020306, 5030101, 5030102,
    5030103, 5030104, 5030105, 5030106, 5030107, 5030108, 5030109, 5030110, 5030111, 5030112, 5030113, 5030114, 5030115, 5030116, 5030117,
    5030118, 5030119, 5030120, 5030121, 5030122, 5030123, 5030124, 5030125, 5030201, 5030202, 5030203, 5030204, 5030205, 5030206, 5030207,
    5030208, 5030209, 5030210, 5030211, 5030212, 5030213, 5040101, 5040102, 5040103, 5040104, 5040105, 5040106, 5040107, 5040201, 5040202,
    5040203, 5040204, 5040205, 5040206, 5040207, 5040208, 5040209, 5040210, 5040301, 5040302, 5040303, 5040304, 5050101, 5050102, 5050103,
    5050104, 5050105, 5050106, 5050107, 5050108, 5050109, 5050110, 5050111, 5050112, 5050113, 5050114, 5050115, 5050116, 5060101, 5060102,
    5060103, 5060104, 5060105, 5060106, 5060107, 5060201, 5060202, 5060203, 5060204, 5060301, 5060302, 5060303, 5060304, 5060401, 5060402,
    5060403, 5060404, 5060405, 5060406, 5060501, 5060502, 5060503, 5060504, 5060505, 5060506, 5060507, 5060508, 5060509, 6010101, 6010102,
    6010103, 6010104, 6010105, 6010106, 6010107, 6010108, 6010109, 6010110, 6010111, 6010112, 6010113, 6010114, 6010201, 6010202, 6010203,
    6010204, 6010301, 6010302, 6010303, 6010304, 6010305, 6010306, 6010401, 6010402, 6010403, 6010404, 6010501, 6010502, 6010503, 6010504,
    6010505, 6010506, 6010507, 6010508, 6010509, 6010601, 6010602, 7010101, 7010102, 7010103, 7010104, 7010105, 7010201, 7010202, 7010203,
    7010204, 7010205, 7010206, 7010207, 7010301, 7010302, 7010303, 7010304, 7010305, 7020101, 7020102, 7020103, 7020104, 7020105, 7020106,
    7030101, 7030102, 7030103, 7030104, 7030105, 7030106, 7030201, 7030202, 7030203, 7030204, 7030301, 7030302, 7030303, 7030304, 7040101,
    7040102, 7040103, 7040104, 7040105, 7040201, 7040202, 7040203, 7040204, 7040205, 7050101, 7050102, 7050103, 7050201, 7050202, 7050301,
    7050302, 7060101,
])

# Catálogo de procesos (para nombrar los códigos de proceso en reportes)
PROCESO_NOMBRE = {
    101: "Operación Infraestructura de Apoyo AP", 102: "Gestión de Recursos Hídricos",
    103: "Operación Infraestructura de Capacidad AP", 104: "Operación Redes y Conducciones AP",
    105: "Operación Infraestructura de Capacidad AP", 106: "Operación Redes y Conducciones AP",
    201: "Operación Infraestructura de Apoyo AS", 202: "Operación Infraestructura de Capacidad AS",
    203: "Operación Redes y Conducciones AS", 204: "Operación Infraestructura de Capacidad AS",
    301: "Mantención Preventiva de Infraestructura de Apoyo", 302: "Mantención Correctiva de Infraestructura de Apoyo",
    303: "Mantención Preventiva de Infraestructura de Capacidad AP", 304: "Mantención Correctiva de Infraestructura de Capacidad AP",
    305: "Mantención Preventiva Redes y Conducciones AP", 306: "Mantención Correctiva Redes y Conducciones AP",
    307: "Mantención Preventiva Conexiones AP", 308: "Mantención Correctiva Conexiones AP",
    309: "Mantención Preventiva Redes y Conducciones AS", 310: "Mantención Correctiva Redes y Conducciones AS",
    311: "Mantención Correctiva Conexiones AS", 312: "Mantención Preventiva de Infraestructura de Capacidad AS",
    313: "Mantención Correctiva de Infraestructura de Capacidad AS",
    401: "Ciclo de Recaudación", 402: "Gestión Comercial", 403: "Gestión de Incorporación de Clientes",
    501: "Dirección Superior", 502: "Administración y Finanzas", 503: "Recursos Humanos",
    504: "Abastecimiento y Servicios Generales", 505: "Informática", 506: "Planificación y Desarrollo",
    601: "Prestaciones Asociadas",
    701: "Planificación y Diseño", 702: "Construcción de Obras", 703: "Inspección Técnica de Obras",
    704: "Telemetría, Comunicación y Sistemas de Información", 705: "Adquisición Bienes Muebles e Inmuebles",
    706: "Adquisición de Derechos de Agua",
}


def cod_proceso_de_actividad(cod_actividad):
    """Código de Proceso = primeros 3 dígitos del Código Actividad (7 dígitos),
    según MAE_1 del Maestro SISS."""
    return int(str(int(cod_actividad)).zfill(7)[:3])


def nombre_proceso(cod_proceso):
    return PROCESO_NOMBRE.get(cod_proceso, f"Proceso {cod_proceso}")


# --- Defaults de proceso para familias SIN tabla de apertura por actividad ---
# GGM: costos de informática/telecomunicaciones -> Informática (505);
#      materiales de oficina/laboratorio -> Abastecimiento y Servicios Generales (504)
DEFAULT_PROCESO_GGM = {
    2401: 505, 2402: 505, 2403: 505, 2404: 505, 2405: 505,
    2406: 505, 2407: 505, 2408: 505, 2409: 505,
    2410: 504, 2411: 504,
}
# OGG: gastos de Directorio -> Dirección Superior (501); resto -> Administración y Finanzas (502)
DEFAULT_PROCESO_OGG = {2501: 501, 2502: 501}
# Todo recurso OGG no listado explícitamente usa 502 por defecto (se resuelve en el código)
DEFAULT_PROCESO_OGG_FALLBACK = 502

# MEI_1 (4101 Productos químicos): sin columna de actividad. Se usan para
# tratamiento de agua potable (cloración/fluoración, proceso 103) y de aguas
# servidas (pretratamiento/tratamiento, proceso 204). Referencia empírica:
# la distribución real de MEI_2 (Energía Eléctrica, que SÍ trae actividad)
# entre estos mismos procesos de tratamiento es 103≈51% / 204≈23% (proporción
# ≈70/30 entre ambos) — se usa esa proporción como default razonable.
DEFAULT_PROCESO_MEI1 = {4101: [(103, 0.70), (204, 0.30)]}

# GPA: cada tabla corresponde 1 a 1 a una "Prestación Asociada" específica
# (Control de Riles, Mantención de Grifos, Corte y Reposición, Revisión de
# Proyectos, Verificación de Medidores, Otras Prestaciones), TODAS bajo el
# macroproceso 6 "Prestaciones Asociadas" según MAE_1. Default: 100% -> 601.
DEFAULT_PROCESO_GPA = 601
RECURSOS_GPA_REP3 = [6101, 6201, 6301, 6401, 6501, 6601]


def leer_tabla_actividad_st(file_bytes):
    """Lee una tabla ST_x para REP_3: además de CÓDIGO RECURSO / MONTO ANUAL
    ACTIVADO / TOTAL GASTO ANUAL (igual que para REP_2), detecta la columna
    CÓDIGO ACTIVIDAD (tolerante a variantes: 'CODIGO ACTIVIDAD',
    'CODIGO DE ACTIVIDAD', 'CÓDIGO DE ACTIVIDAD', etc. y errores de tipeo).
    Devuelve lista de (cod_recurso, monto_activado, gasto_no_activado, cod_actividad)
    y una lista de avisos (columna faltante, código de actividad o recurso
    no reconocido según MAE_1/MAE_2)."""
    avisos_tabla = []
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    ws = wb.active
    header_row = [c.value for c in next(ws.iter_rows(min_row=1, max_row=1))]

    idx_recurso = _encontrar_columna(header_row, [["CODIGO", "RECURSO"], ["COD", "RECURSO"]])
    idx_activado = _encontrar_columna(header_row, [["MONTO", "ANUAL", "ACTIVADO"], ["MONTO", "ACTIVADO"]])
    idx_gasto = _encontrar_columna(
        header_row,
        [["TOTAL", "GASTO", "ANUAL"], ["GASTO", "ANUAL", "NO", "ACTIVADO"], ["TOTAL", "GASTO"], ["GASTO"]],
        evitar=["ACTIVADO", "%"],
    )
    idx_actividad = _encontrar_columna(
        header_row,
        [["CODIGO", "ACTIVIDAD"], ["COD", "ACTIVIDAD"], ["CODIGO", "DE", "ACTIVIDAD"]],
    )

    faltantes = []
    if idx_recurso is None:
        faltantes.append("CÓDIGO RECURSO")
    if idx_activado is None:
        faltantes.append("MONTO ANUAL ACTIVADO")
    if idx_gasto is None:
        faltantes.append("TOTAL GASTO ANUAL")
    if idx_actividad is None:
        faltantes.append("CÓDIGO ACTIVIDAD")
    if faltantes:
        raise ValueError(f"No se identificaron las columnas: {', '.join(faltantes)}")

    filas = []
    recursos_invalidos = set()
    actividades_invalidas = set()
    for row in ws.iter_rows(min_row=2, values_only=True):
        if row[0] is None:
            continue
        cod_recurso = row[idx_recurso]
        cod_actividad = row[idx_actividad]
        if cod_recurso is None or cod_actividad is None:
            continue
        cod_recurso_num = _a_numero(cod_recurso)
        cod_recurso_num = int(cod_recurso_num) if cod_recurso_num == int(cod_recurso_num) else cod_recurso_num
        cod_actividad_num = int(_a_numero(cod_actividad))

        if cod_recurso_num not in RECURSO_NOMBRE:
            recursos_invalidos.add(cod_recurso_num)
        if cod_actividad_num not in CODIGOS_ACTIVIDAD_VALIDOS:
            actividades_invalidas.add(cod_actividad_num)

        filas.append((cod_recurso_num, _a_numero(row[idx_activado]), _a_numero(row[idx_gasto]), cod_actividad_num))

    if recursos_invalidos:
        avisos_tabla.append(
            f"CÓDIGO(S) RECURSO no reconocido(s) según MAE_2: {sorted(recursos_invalidos)}"
        )
    if actividades_invalidas:
        avisos_tabla.append(
            f"CÓDIGO(S) ACTIVIDAD no reconocido(s) según MAE_1: {sorted(actividades_invalidas)}"
        )
    return filas, avisos_tabla

def build_rep3(fb3, ggm_proceso_params, ogg_proceso_params, ggm_params, ogg_params, st_params, st_files_raw,
               mei_proceso_params=None, gpa_proceso_params=None, mei_params=None, gpa_files_raw=None):
    """
    Construye la tabla REP_3 (Costos y Gastos No Activados de Servicios
    Regulados - Proceso), abriendo el gasto regulado (familias 11 y 12)
    de REP_2 en Código de Proceso, para las familias GRH, GCP, GGV, GGI
    (con tablas de apertura por actividad dedicadas), GGM y OGG (con
    asignación por defecto/parametrizable, ya que no tienen tabla de
    apertura por actividad) y ST (que trae CÓDIGO ACTIVIDAD en la misma
    tabla, detectado de forma flexible/tolerante a variantes de nombre).

    fb3: dict con bytes de GRH_8, GRH_11, GRH_12, GCP_5, GCP_6, GGV_4, GGV_5,
         GGV_6, GGI_5, GGI_6, GGM_1..5, OGG_5 (todos opcionales -> None si
         no se cargó)
    st_files_raw: dict {nombre_tabla: file_bytes} de las tablas ST subidas
                  (se leen aquí con leer_tabla_actividad_st, no con
                  leer_tabla_st, ya que se necesita también CÓDIGO ACTIVIDAD)
    """
    avisos3 = []
    mei_proceso_params = mei_proceso_params or {}
    gpa_proceso_params = gpa_proceso_params or {}
    mei_params = mei_params or {}
    gpa_files_raw = gpa_files_raw or {}
    agg3 = defaultdict(float)  # (cod_recurso, familia_servicio, cod_proceso) -> gasto

    def safe_rows(clave, etiqueta):
        b = fb3.get(clave)
        if b is None:
            avisos3.append(f"⚠️ REP_3: falta **{etiqueta}** — esa familia queda excluida de REP_3.")
            return []
        try:
            return read_rows(b)
        except Exception as e:
            avisos3.append(f"⚠️ REP_3: no se pudo leer **{etiqueta}** ({e}).")
            return []

    REGULADOS = {1101, 1201, 1202, 1203, 1204, 1205}

    # ================= GRH =================
    grh8 = safe_rows("grh8", "GRH_8")
    grh11 = safe_rows("grh11", "GRH_11")
    grh12 = safe_rows("grh12", "GRH_12")

    if grh8 and grh11:
        shares_servicio_reg = defaultdict(list)
        for r in grh11:
            cod_reg, cod_noreg, pct = r[7], r[8], r[9]
            cod = cod_reg if cod_reg != -1 else cod_noreg
            if cod in REGULADOS:
                shares_servicio_reg[(r[4], r[5])].append((cod, pct))

        shares_proceso_grh = defaultdict(lambda: defaultdict(float))
        for r in grh12:
            persona, cargo, cod_actividad, pct_act = r[4], r[5], r[7], r[8]
            shares_proceso_grh[(persona, cargo)][cod_proceso_de_actividad(cod_actividad)] += pct_act

        personas_sin_grh12 = set()
        for r in grh8:
            empresa, periodo, anio, sector, persona, cargo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r
            servicios = shares_servicio_reg.get((persona, cargo), [])
            if not servicios:
                continue
            procesos = shares_proceso_grh.get((persona, cargo))
            if not procesos:
                personas_sin_grh12.add((persona, cargo))
                continue
            for cod_serv, pct_serv in servicios:
                fam = familia(cod_serv)
                gasto_rf = total_gasto * pct_serv
                for proceso, pct_proc in procesos.items():
                    agg3[(cod_recurso, fam, proceso)] += gasto_rf * pct_proc
        if personas_sin_grh12:
            avisos3.append(f"⚠️ GRH: {len(personas_sin_grh12)} persona(s) con dedicación regulada sin apertura en GRH_12; se excluyen de REP_3.")

    # ================= GCP =================
    gcp5 = safe_rows("gcp5", "GCP_5")
    gcp6 = safe_rows("gcp6", "GCP_6")

    if gcp5:
        shares_proceso_gcp = defaultdict(lambda: defaultdict(float))
        for r in gcp6:
            persona, cargo, cod_recurso_g6, cod_actividad, pct = r[4], r[5], r[6], r[8], r[9]
            shares_proceso_gcp[(persona, cargo, cod_recurso_g6)][cod_proceso_de_actividad(cod_actividad)] += pct

        faltan_gcp6 = 0
        for r in gcp5:
            empresa, periodo, anio, sector, persona, cargo, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r
            if cod_reg == -1 or cod_reg not in REGULADOS:
                continue
            fam = familia(cod_reg)
            procesos = shares_proceso_gcp.get((persona, cargo, cod_recurso))
            if not procesos:
                faltan_gcp6 += 1
                continue
            for proceso, pct_proc in procesos.items():
                agg3[(cod_recurso, fam, proceso)] += gasto_no_act * pct_proc
        if faltan_gcp6:
            avisos3.append(f"⚠️ GCP: {faltan_gcp6} combinación(es) persona-cargo-recurso reguladas sin apertura en GCP_6; se excluyen de REP_3.")

    # ================= GGV =================
    ggv4 = safe_rows("ggv4", "GGV_4")
    ggv5 = safe_rows("ggv5", "GGV_5")
    ggv6 = safe_rows("ggv6", "GGV_6")

    if ggv4 and ggv5:
        shares_servicio_ggv = defaultdict(list)
        for r in ggv5:
            empresa, periodo, anio, sector, id_activo, total_no_act, cod_reg, cod_noreg, pct = r[:9]
            if cod_reg in REGULADOS:
                shares_servicio_ggv[id_activo].append((cod_reg, pct))

        shares_proceso_ggv = defaultdict(lambda: defaultdict(float))
        for r in ggv6:
            id_activo, cod_actividad, pct = r[4], r[6], r[7]
            shares_proceso_ggv[id_activo][cod_proceso_de_actividad(cod_actividad)] += pct

        activos_sin_ggv6 = set()
        for r in ggv4:
            empresa, periodo, anio, sector, id_activo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r[:10]
            for cod_serv, pct_serv in shares_servicio_ggv.get(id_activo, []):
                fam = familia(cod_serv)
                gasto_rf = total_gasto * pct_serv
                procesos = shares_proceso_ggv.get(id_activo)
                if not procesos:
                    activos_sin_ggv6.add(id_activo)
                    continue
                for proceso, pct_proc in procesos.items():
                    agg3[(cod_recurso, fam, proceso)] += gasto_rf * pct_proc
        if activos_sin_ggv6:
            avisos3.append(f"⚠️ GGV: {len(activos_sin_ggv6)} activo(s) con dedicación regulada sin apertura en GGV_6; se excluyen de REP_3.")

    # ================= GGI =================
    ggi5 = safe_rows("ggi5", "GGI_5")
    ggi6 = safe_rows("ggi6", "GGI_6")

    if ggi5:
        shares_proceso_ggi = defaultdict(lambda: defaultdict(float))
        for r in ggi6:
            id_inmueble, cod_proceso_directo, pct = r[4], r[6], r[7]
            shares_proceso_ggi[id_inmueble][cod_proceso_directo] += pct

        inmuebles_sin_ggi6 = set()
        for r in ggi5:
            empresa, periodo, anio, sector, id_inmueble, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r[:10]
            if cod_reg == -1 or cod_reg not in REGULADOS:
                continue
            fam = familia(cod_reg)
            procesos = shares_proceso_ggi.get(id_inmueble)
            if not procesos:
                inmuebles_sin_ggi6.add(id_inmueble)
                continue
            for proceso, pct_proc in procesos.items():
                agg3[(cod_recurso, fam, proceso)] += gasto_no_act * pct_proc
        if inmuebles_sin_ggi6:
            avisos3.append(f"⚠️ GGI: {len(inmuebles_sin_ggi6)} inmueble(s) con dedicación regulada sin apertura en GGI_6; se excluyen de REP_3.")

    # ================= GGM (sin tabla de actividad: default/parametrizable) =================
    ggm_tablas = [safe_leer_tabla_plana(fb3.get(f"ggm{i}"), f"GGM_{i}", avisos3) for i in range(1, 6)]
    if any(ggm_tablas):
        ggm_regulado_por_recurso = defaultdict(float)  # solo la porción a familia 11 (servicio 1101 u overrides de GGM_PARAMS)
        for tabla in ggm_tablas:
            for r in tabla:
                cod_recurso, total_gasto = r[4], r[6]
                overrides = ggm_params.get(cod_recurso, [])
                pct_reg = 1.0 - sum(p for _, p in overrides)  # % que va a servicio 1101 (familia 11)
                ggm_regulado_por_recurso[cod_recurso] += total_gasto * pct_reg

        for cod_recurso, gasto in ggm_regulado_por_recurso.items():
            if gasto == 0:
                continue
            overrides_proceso = ggm_proceso_params.get(cod_recurso)
            if overrides_proceso:
                pct_default = 1.0 - sum(p for _, p in overrides_proceso)
                proceso_default = DEFAULT_PROCESO_GGM.get(cod_recurso, 504)
                agg3[(cod_recurso, 11, proceso_default)] += gasto * pct_default
                for proceso, pct in overrides_proceso:
                    agg3[(cod_recurso, 11, proceso)] += gasto * pct
            else:
                proceso_default = DEFAULT_PROCESO_GGM.get(cod_recurso, 504)
                agg3[(cod_recurso, 11, proceso_default)] += gasto

    # ================= OGG (sin tabla de actividad: default/parametrizable) =================
    ogg5 = safe_leer_tabla_plana(fb3.get("ogg5"), "OGG_5", avisos3)
    if ogg5:
        ogg_regulado_por_recurso = defaultdict(float)
        for r in ogg5:
            cod_recurso, monto_act, total_gasto = r[4], r[5], r[6]
            overrides = ogg_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            ogg_regulado_por_recurso[cod_recurso] += total_gasto * pct_reg

        for cod_recurso, gasto in ogg_regulado_por_recurso.items():
            if gasto == 0:
                continue
            overrides_proceso = ogg_proceso_params.get(cod_recurso)
            proceso_default = DEFAULT_PROCESO_OGG.get(cod_recurso, DEFAULT_PROCESO_OGG_FALLBACK)
            if overrides_proceso:
                pct_default = 1.0 - sum(p for _, p in overrides_proceso)
                agg3[(cod_recurso, 11, proceso_default)] += gasto * pct_default
                for proceso, pct in overrides_proceso:
                    agg3[(cod_recurso, 11, proceso)] += gasto * pct
            else:
                agg3[(cod_recurso, 11, proceso_default)] += gasto

    # ================= ST (trae CÓDIGO ACTIVIDAD en la misma tabla) =================
    avisos_st_detalle = []
    for nombre_tabla, file_bytes in st_files_raw.items():
        try:
            filas_st, avisos_tabla = leer_tabla_actividad_st(file_bytes)
        except Exception as e:
            avisos3.append(f"⚠️ ST (REP_3): no se pudo leer **{nombre_tabla}** ({e}); se excluye.")
            continue
        for aviso in avisos_tabla:
            avisos_st_detalle.append(f"{nombre_tabla}: {aviso}")
        for cod_recurso, monto_act, total_gasto, cod_actividad in filas_st:
            overrides = st_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)  # % regulado (servicio 1101, familia 11)
            gasto_regulado_fila = total_gasto * pct_reg
            if gasto_regulado_fila == 0:
                continue
            proceso = cod_proceso_de_actividad(cod_actividad)
            agg3[(cod_recurso, 11, proceso)] += gasto_regulado_fila
    if avisos_st_detalle:
        avisos3.append(f"⚠️ ST (REP_3): validación de códigos — {'; '.join(avisos_st_detalle[:10])}" + (" ..." if len(avisos_st_detalle) > 10 else ""))

    # ================= MEI =================
    # MEI_1 (4101, Productos químicos): SIN columna de actividad -> default/parametrizable.
    # MEI_2 (4102), MEI_3 (4103), MEI_4 (4104-4106): SÍ traen CÓDIGO ACTIVIDAD
    # por fila (como ST) -> se deriva el proceso directamente.
    mei1 = safe_leer_tabla_mei(fb3.get("mei1"), "MEI_1", avisos3, necesita_obra_nbi=True)
    mei2 = safe_leer_tabla_mei(fb3.get("mei2"), "MEI_2", avisos3, necesita_actividad=True)
    mei3 = safe_leer_tabla_mei(fb3.get("mei3"), "MEI_3", avisos3, necesita_actividad=True)
    mei4 = safe_leer_tabla_mei(fb3.get("mei4"), "MEI_4", avisos3, necesita_actividad=True)

    if mei1:
        mei1_regulado_por_recurso = defaultdict(float)
        for r in mei1:
            cod_recurso, total_gasto = r[4], r[7]
            overrides = mei_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            mei1_regulado_por_recurso[cod_recurso] += total_gasto * pct_reg
        for cod_recurso, gasto in mei1_regulado_por_recurso.items():
            if gasto == 0:
                continue
            overrides_proceso = mei_proceso_params.get(cod_recurso)
            if overrides_proceso:
                pct_default = 1.0 - sum(p for _, p in overrides_proceso)
                default_split = DEFAULT_PROCESO_MEI1.get(cod_recurso, [(504, 1.0)])
                for proceso, pct_d in default_split:
                    agg3[(cod_recurso, 11, proceso)] += gasto * pct_default * pct_d
                for proceso, pct in overrides_proceso:
                    agg3[(cod_recurso, 11, proceso)] += gasto * pct
            else:
                for proceso, pct_d in DEFAULT_PROCESO_MEI1.get(cod_recurso, [(504, 1.0)]):
                    agg3[(cod_recurso, 11, proceso)] += gasto * pct_d

    def procesar_mei_con_actividad(tabla, idx_recurso, idx_gasto, idx_actividad, etiqueta):
        if not tabla:
            return
        for r in tabla:
            cod_recurso, total_gasto, cod_actividad = r[idx_recurso], r[idx_gasto], r[idx_actividad]
            if cod_actividad is None:
                continue
            overrides = mei_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            gasto_regulado = total_gasto * pct_reg
            if gasto_regulado == 0:
                continue
            proceso = cod_proceso_de_actividad(cod_actividad)
            agg3[(cod_recurso, 11, proceso)] += gasto_regulado

    procesar_mei_con_actividad(mei2, 4, 7, 8, "MEI_2")
    procesar_mei_con_actividad(mei3, 4, 7, 8, "MEI_3")
    procesar_mei_con_actividad(mei4, 4, 7, 8, "MEI_4")

    # ================= GPA =================
    # Cada tabla GPA_x tiene 100% de su gasto en un servicio regulado FIJO
    # (1201-1205) y, por defecto, se asigna 100% al proceso 601 "Prestaciones
    # Asociadas" (parametrizable por código de recurso).
    tablas_gpa_faltantes = [t for t in GPA_TABLES if t not in gpa_files_raw]
    if tablas_gpa_faltantes and gpa_files_raw:
        avisos3.append(f"⚠️ REP_3-GPA: faltan {len(tablas_gpa_faltantes)} tabla(s): {', '.join(tablas_gpa_faltantes)}.")
    for nombre_tabla, file_bytes in gpa_files_raw.items():
        servicio_fijo = GPA_TABLE_TO_SERVICIO.get(nombre_tabla)
        if servicio_fijo is None:
            continue
        fam_gpa = familia(servicio_fijo)
        try:
            filas_gpa = leer_tabla_st(file_bytes)
        except Exception as e:
            avisos3.append(f"⚠️ REP_3-GPA: no se pudo leer {nombre_tabla} ({e}).")
            continue
        gpa_gasto_por_recurso = defaultdict(float)
        for cod_recurso, monto_act, total_gasto in filas_gpa:
            gpa_gasto_por_recurso[cod_recurso] += total_gasto
        for cod_recurso, gasto in gpa_gasto_por_recurso.items():
            if gasto == 0:
                continue
            overrides_proceso = gpa_proceso_params.get(cod_recurso)
            if overrides_proceso:
                pct_default = 1.0 - sum(p for _, p in overrides_proceso)
                agg3[(cod_recurso, fam_gpa, DEFAULT_PROCESO_GPA)] += gasto * pct_default
                for proceso, pct in overrides_proceso:
                    agg3[(cod_recurso, fam_gpa, proceso)] += gasto * pct
            else:
                agg3[(cod_recurso, fam_gpa, DEFAULT_PROCESO_GPA)] += gasto

    # ================= Normalizar (sumar EXACTAMENTE 100% por recurso+familia) =================
    grupos = defaultdict(list)
    for (cod_recurso, fam, proceso), gasto in agg3.items():
        if abs(gasto) < 1e-6:
            continue
        grupos[(cod_recurso, fam)].append((proceso, gasto))

    EMPRESA = PERIODO = ANIO = SECTOR = None
    for clave in ["grh8", "gcp5", "ggv4", "ggi5", "ggm1", "ogg5"]:
        b = fb3.get(clave)
        if b:
            filas_tmp = read_rows(b)
            if filas_tmp:
                EMPRESA, PERIODO, ANIO, SECTOR = filas_tmp[0][0], filas_tmp[0][1], filas_tmp[0][2], filas_tmp[0][3]
                break

    final_rows3 = []
    for (cod_recurso, fam), items in grupos.items():
        total_grupo = sum(g for _, g in items)
        fracciones = [g / total_grupo for _, g in items]
        pcts = redondear_para_sumar_100(fracciones, 4)
        for (proceso, gasto), pct in zip(items, pcts):
            final_rows3.append([EMPRESA, PERIODO, ANIO, SECTOR, cod_recurso, fam, proceso, round(pct * 100, 2), round(gasto, 2)])

    final_rows3.sort(key=lambda r: (r[4], r[5], r[6]))
    return final_rows3, avisos3



# ============================================================================
# CYG - Costos y Gastos por Recurso (CYG_1 a CYG_4, CYG_8, CYG_9)
# ============================================================================
# CYG_1: Recursos - Servicios Continuos AP/AS (servicios 1101, 1102)
# CYG_2: Recursos - Prestaciones Asociadas Reguladas (servicios 1201-1205)
# CYG_3: Recursos - Prestaciones Asociadas No Reguladas (servicios 2101-2117)
# CYG_4: Recursos - Servicios No Regulados (servicios 2201-2214)
# CYG_8: Recursos de Servicios Continuos - Actividades (solo gasto regulado)
# CYG_9: Recursos de Servicios No Regulados - Servicios Prestados a Terceros
#        (abre por ID CLIENTE usando MCO_42 + ING_4)
#
# IMPORTANTE: a diferencia de REP_2 (que colapsa los servicios en su FAMILIA
# 11/12/22) y de REP_3 (que colapsa las actividades en su PROCESO), las
# tablas CYG necesitan el código de SERVICIO ESPECÍFICO (ej. 1201, no solo
# "12") y el código de ACTIVIDAD ESPECÍFICO (ej. 4010101, no solo "401").
# Por eso CYG se construye SIEMPRE desde las tablas fuente, nunca desde
# REP_2/REP_3 ya agregados. REP_2 y REP_3 solo se usan como checkpoint de
# validación de cuadratura.

# --- Mapeo CÓDIGO OBRA TIPO NBI -> CÓDIGO ACTIVIDAD (exclusivo de MEI_1) ---
# Aportado por el usuario: reemplaza cualquier supuesto/estimación previa.
# El código de proceso se deriva de esta actividad (primeros 3 dígitos).
OBRA_NBI_A_ACTIVIDAD = {
    101: 1030401,  # Captación en Río (proceso 103)
    102: 1030402,  # Captación en Canal (proceso 103)
    103: 1030403,  # Captación en Lago o Embalse (proceso 103)
    104: 1030404,  # Captación en Mar (proceso 103)
    201: 1030302,  # Captación mediante Drenes y Galerías (proceso 103)
    202: 1030304,  # Captación mediante Punteras (proceso 103)
    203: 1030303,  # Captación mediante Sondajes (proceso 103)
    204: 1030301,  # Captación mediante Norias (proceso 103)
    301: 1031101,  # Plantas Elevadoras de Agua Potable Tipo A (proceso 103)
    302: 1031101,  # Plantas Elevadoras de Agua Potable Tipo B (proceso 103)
    303: 1031101,  # Plantas Elevadoras de Agua Potable Tipo C (proceso 103)
    304: 1031101,  # Plantas Elevadoras de Agua Potable Tipo D (proceso 103)
    305: 1031101,  # Plantas Elevadoras de Agua Potable Tipo E (proceso 103)
    351: 2020301,  # Plantas Elevadoras de Aguas Servidas (proceso 202)
    401: 1050101,  # Estanques Semienterrados y Enterrados (proceso 105)
    402: 1050102,  # Estanques Elevados (proceso 105)
    501: 1030601,  # Plantas de Tratamiento de Agua Potable excepto Osmosis Inversa (proceso 103)
    502: 1030604,  # Plantas de Tratamiento de Agua Potable de Osmosis Inversa (proceso 103)
    601: 1030801,  # Sistemas de Desinfección de Agua Potable (proceso 103)
    701: 1030902,  # Sistemas de Fluoración (proceso 103)
    801: 1060101,  # Red de Distribución (proceso 106)
    901: 2030301,  # Red de Recolección (proceso 203)
    1001: 3080101,  # Arranques (proceso 308)
    1002: 3070302,  # Medidores (proceso 307)
    1003: 3110301,  # Uniones Domiciliarias (proceso 311)
    1101: 1040101,  # Conducciones de AP (proceso 104)
    1151: 2030101,  # Conducciones de AS (proceso 203)
    1201: 2040404,  # Tabla General de Sistemas de Tratamiento de Aguas Servidas (proceso 204)
    1202: 2040501,  # Pretratamiento de Aguas Servidas (proceso 204)
    1203: 2040401,  # Tratamiento Primario de Aguas Servidas (proceso 204)
    1204: 2040401,  # Tratamiento Secundario de Aguas Servidas (proceso 204)
    1205: 2040404,  # Desinfección y Decloración en Plantas de Tratamiento de Aguas Servidas (proceso 204)
    1206: 2040404,  # Línea de Lodos de Plantas de Tratamiento de Aguas Servidas (proceso 204)
    1207: 2040601,  # Emisario Submarino - Información General (proceso 204)
    1208: 2040601,  # Emisario Submarino - Información por Tramos (proceso 204)
    1209: 2041001,  # Control Olores, Generador y Aforos en Plantas de Tratamiento de Aguas Servidas (proceso 204)
    1402: 3010901,  # Macromedidores (proceso 301)
    1403: 3011101,  # Reductoras de Presión (proceso 301)
    1404: 3011001,  # Anti Golpe de Ariete (proceso 301)
}

# --- Catálogo completo Proceso -> [Actividades] (para defaults GGI/GGM/OGG/GPA) ---
PROCESO_A_ACTIVIDADES = {
    101: [1010101, 1010201, 1010301, 1010302, 1010303, 1010304, 1010305, 1010306, 1010307],
    102: [1020101, 1020201, 1020301, 1020302, 1020303, 1020304],
    103: [1030101, 1030201, 1030202, 1030301, 1030302, 1030303, 1030304, 1030401, 1030402, 1030403, 1030404, 1030501, 1030502, 1030503, 1030504, 1030505, 1030601, 1030602, 1030603, 1030604, 1030605, 1030701, 1030702, 1030703, 1030704, 1030705, 1030801, 1030802, 1030803, 1030901, 1030902, 1030903, 1031001, 1031002, 1031101, 1031201, 1031202],
    104: [1040101, 1040201, 1040301, 1040302, 1040303, 1040304, 1040305, 1040306, 1040307, 1040308, 1040309, 1040401],
    105: [1050101, 1050102],
    106: [1060101, 1060102, 1060201, 1060202, 1060203, 1060204, 1060205, 1060301, 1060302],
    201: [2010101, 2010201, 2010301, 2010302, 2010303, 2010304, 2010305, 2010306, 2010307],
    202: [2020101, 2020102, 2020103, 2020201, 2020301],
    203: [2030101, 2030201, 2030301],
    204: [2040101, 2040102, 2040103, 2040104, 2040201, 2040301, 2040401, 2040402, 2040403, 2040404, 2040405, 2040501, 2040601, 2040701, 2040702, 2040703, 2040704, 2040705, 2040801, 2040802, 2040901, 2040902, 2040903, 2040904, 2041001],
    301: [3010101, 3010201, 3010301, 3010401, 3010501, 3010601, 3010701, 3010801, 3010901, 3011001, 3011101, 3011201, 3011301],
    302: [3020101, 3020201, 3020301, 3020401, 3020501],
    303: [3030101, 3030201, 3030301, 3030302, 3030303, 3030304, 3030401, 3030402, 3030403, 3030404, 3030501, 3030502, 3030503, 3030504, 3030505, 3030601, 3030602, 3030603, 3030701, 3030702, 3030703, 3030801, 3030802, 3030901, 3031001, 3031002],
    304: [3040101, 3040201, 3040301, 3040401, 3040501],
    305: [3050101, 3050201, 3050301, 3050302, 3050303, 3050401, 3050402, 3050501, 3050502],
    306: [3060101, 3060201, 3060301, 3060302, 3060303, 3060304, 3060401, 3060402, 3060403, 3060404, 3060501, 3060502, 3060503, 3060504],
    307: [3070101, 3070201, 3070301, 3070302],
    308: [3080101, 3080102, 3080103],
    309: [3090101, 3090201, 3090301, 3090302, 3090303, 3090401, 3090402, 3090403, 3090501, 3090502, 3090503, 3090504],
    310: [3100101, 3100201, 3100301, 3100302, 3100303, 3100304, 3100305, 3100401, 3100402, 3100403, 3100404, 3100405, 3100501, 3100502, 3100503, 3100504],
    311: [3110101, 3110201, 3110301, 3110302, 3110303, 3110304],
    312: [3120101, 3120201, 3120301, 3120401, 3120402, 3120403, 3120404, 3120405, 3120501, 3120601],
    313: [3130101, 3130201, 3130301, 3130401, 3130501],
    401: [4010101, 4010102, 4010103, 4010104, 4010105, 4010201, 4010202, 4010203, 4010204, 4010205, 4010206, 4010207, 4010301, 4010302, 4010303, 4010304, 4010401, 4010402, 4010403, 4010404, 4010405, 4010406, 4010407, 4010408],
    402: [4020101, 4020102, 4020103, 4020201, 4020202, 4020203, 4020301, 4020302, 4020303, 4020304, 4020305, 4020306, 4020307, 4020308, 4020309, 4020310, 4020311, 4020312, 4020313, 4020314, 4020315, 4020316, 4020317, 4020318, 4020319, 4020320, 4020321, 4020322, 4020323, 4020324, 4020325, 4020326, 4020327, 4020401, 4020402, 4020403, 4020501, 4020502, 4020503, 4020601, 4020602],
    403: [4030101, 4030102, 4030103, 4030104, 4030105, 4030106, 4030107, 4030108, 4030201, 4030202, 4030203, 4030204, 4030205, 4030206, 4030207, 4030208, 4030209, 4030210, 4030211, 4030301, 4030302, 4030303, 4030304, 4030305, 4030306, 4030307, 4030308, 4030309, 4030401, 4030402, 4030403, 4030404],
    501: [5010101, 5010102, 5010103, 5010104, 5010105, 5010106, 5010107, 5010108, 5010201, 5010202, 5010203, 5010204, 5010205, 5010206, 5010301, 5010302, 5010303, 5010304, 5010305, 5010306, 5010307, 5010308, 5010401, 5010402, 5010403, 5010404, 5010405, 5010406, 5010407, 5010501, 5010502, 5010503, 5010504, 5010505, 5010506, 5010601, 5010602, 5010603, 5010604, 5010605, 5010606, 5010607, 5010608],
    502: [5020101, 5020102, 5020103, 5020104, 5020105, 5020106, 5020107, 5020108, 5020109, 5020110, 5020111, 5020112, 5020113, 5020114, 5020115, 5020201, 5020202, 5020203, 5020204, 5020205, 5020206, 5020207, 5020208, 5020209, 5020210, 5020211, 5020212, 5020301, 5020302, 5020303, 5020304, 5020305, 5020306],
    503: [5030101, 5030102, 5030103, 5030104, 5030105, 5030106, 5030107, 5030108, 5030109, 5030110, 5030111, 5030112, 5030113, 5030114, 5030115, 5030116, 5030117, 5030118, 5030119, 5030120, 5030121, 5030122, 5030123, 5030124, 5030125, 5030201, 5030202, 5030203, 5030204, 5030205, 5030206, 5030207, 5030208, 5030209, 5030210, 5030211, 5030212, 5030213],
    504: [5040101, 5040102, 5040103, 5040104, 5040105, 5040106, 5040107, 5040201, 5040202, 5040203, 5040204, 5040205, 5040206, 5040207, 5040208, 5040209, 5040210, 5040301, 5040302, 5040303, 5040304],
    505: [5050101, 5050102, 5050103, 5050104, 5050105, 5050106, 5050107, 5050108, 5050109, 5050110, 5050111, 5050112, 5050113, 5050114, 5050115, 5050116],
    506: [5060101, 5060102, 5060103, 5060104, 5060105, 5060106, 5060107, 5060201, 5060202, 5060203, 5060204, 5060301, 5060302, 5060303, 5060304, 5060401, 5060402, 5060403, 5060404, 5060405, 5060406, 5060501, 5060502, 5060503, 5060504, 5060505, 5060506, 5060507, 5060508, 5060509],
    601: [6010101, 6010102, 6010103, 6010104, 6010105, 6010106, 6010107, 6010108, 6010109, 6010110, 6010111, 6010112, 6010113, 6010114, 6010201, 6010202, 6010203, 6010204, 6010301, 6010302, 6010303, 6010304, 6010305, 6010306, 6010401, 6010402, 6010403, 6010404, 6010501, 6010502, 6010503, 6010504, 6010505, 6010506, 6010507, 6010508, 6010509, 6010601, 6010602],
    701: [7010101, 7010102, 7010103, 7010104, 7010105, 7010201, 7010202, 7010203, 7010204, 7010205, 7010206, 7010207, 7010301, 7010302, 7010303, 7010304, 7010305],
    702: [7020101, 7020102, 7020103, 7020104, 7020105, 7020106],
    703: [7030101, 7030102, 7030103, 7030104, 7030105, 7030106, 7030201, 7030202, 7030203, 7030204, 7030301, 7030302, 7030303, 7030304],
    704: [7040101, 7040102, 7040103, 7040104, 7040105, 7040201, 7040202, 7040203, 7040204, 7040205],
    705: [7050101, 7050102, 7050103, 7050201, 7050202, 7050301, 7050302],
    706: [7060101],
}

def actividades_de_proceso(cod_proceso):
    return PROCESO_A_ACTIVIDADES.get(cod_proceso, [])

# --- Defaults de actividad para familias sin apertura (equitativo dentro del proceso) ---
# GGI (proceso 504, 100% de sus 21 actividades, sin excluir ninguna)
DEFAULT_ACTIVIDADES_GGI = {504: actividades_de_proceso(504)}

# GGM: 2401-2409 -> proceso 505 (16 act.); 2410-2411 -> proceso 504 (21 act., completas)
DEFAULT_ACTIVIDADES_GGM = {505: actividades_de_proceso(505), 504: actividades_de_proceso(504)}

# OGG: 2501-2502 -> proceso 501 (43 act.); resto -> proceso 502 (33 act.)
DEFAULT_ACTIVIDADES_OGG = {501: actividades_de_proceso(501), 502: actividades_de_proceso(502)}

# GPA: cada tabla reparte equitativamente SOLO entre las actividades de su
# propio subproceso (ej. GPA_1/6101 -> solo las 14 actividades de "Control
# Directo de Riles", no las 39 de todo el proceso 601).
def _acts_subproceso(prefijo5, catalogo_601):
    return [a for a in catalogo_601 if str(a)[:5] == prefijo5]

_acts_601 = actividades_de_proceso(601)
DEFAULT_ACTIVIDADES_GPA = {
    6101: _acts_subproceso("60101", _acts_601),
    6201: _acts_subproceso("60102", _acts_601),
    6301: _acts_subproceso("60103", _acts_601),
    6401: _acts_subproceso("60104", _acts_601),
    6501: _acts_subproceso("60105", _acts_601),
    6601: _acts_subproceso("60106", _acts_601),
}

def build_cyg_core(fb, ggm_params, ogg_params, mei_params, st_params, st_files_raw,
                    ggm_proceso_params, ogg_proceso_params, mei_proceso_params, gpa_proceso_params,
                    gpa_files_raw):
    """
    Construye, desde las tablas fuente (NO desde REP_2/REP_3), dos
    agregaciones con granularidad completa:

    agg_serv[(cod_servicio_especifico, cod_recurso)] = [gasto_no_act, monto_act]
        -> insumo para CYG_1, CYG_2, CYG_3, CYG_4 (agrupado por servicio
           específico, ej. 1201 separado de 1202).

    agg_act[(cod_actividad_especifica, cod_recurso)] = [gasto_no_act, monto_act]
        -> insumo para CYG_8 (SOLO la porción de gasto regulada, familias 11/12).

    Devuelve también 'avisos' con lo que falta o no se pudo determinar.
    """
    avisos = []
    agg_serv = defaultdict(lambda: [0.0, 0.0])
    agg_act = defaultdict(lambda: [0.0, 0.0])
    REGULADOS = {1101, 1201, 1202, 1203, 1204, 1205}

    def safe_rows(clave, etiqueta):
        b = fb.get(clave)
        if b is None:
            avisos.append(f"CYG: falta **{etiqueta}** — esa familia queda excluida.")
            return []
        try:
            return read_rows(b)
        except Exception as e:
            avisos.append(f"CYG: no se pudo leer **{etiqueta}** ({e}).")
            return []

    # ================= GRH =================
    grh8 = safe_rows("grh8", "GRH_8")
    grh11 = safe_rows("grh11", "GRH_11")
    grh12 = safe_rows("grh12", "GRH_12")

    shares_servicio_grh = defaultdict(list)  # (persona,cargo) -> [(servicio, pct), ...] TODOS los servicios
    for r in grh11:
        cod_reg, cod_noreg, pct = r[7], r[8], r[9]
        cod = cod_reg if cod_reg != -1 else cod_noreg
        shares_servicio_grh[(r[4], r[5])].append((cod, pct))

    shares_proceso_grh = defaultdict(lambda: defaultdict(float))  # (persona,cargo) -> {actividad: pct}
    # OJO: GRH_12 da % por ACTIVIDAD (no colapsado a proceso) -- se usa tal cual para CYG_8.
    shares_actividad_grh = defaultdict(lambda: defaultdict(float))
    for r in grh12:
        persona, cargo, cod_actividad, pct_act = r[4], r[5], r[7], r[8]
        shares_actividad_grh[(persona, cargo)][int(cod_actividad)] += pct_act

    personas_sin_grh12 = set()
    for r in grh8:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r
        for cod_serv, pct_serv in shares_servicio_grh.get((persona, cargo), []):
            gasto_serv = total_gasto * pct_serv
            monto_act_serv = monto_act * pct_serv
            agg_serv[(cod_serv, cod_recurso)][0] += gasto_serv
            agg_serv[(cod_serv, cod_recurso)][1] += monto_act_serv
            if cod_serv in REGULADOS:
                actividades = shares_actividad_grh.get((persona, cargo))
                if not actividades:
                    personas_sin_grh12.add((persona, cargo))
                    continue
                for cod_actividad, pct_act_i in actividades.items():
                    agg_act[(cod_actividad, cod_recurso)][0] += gasto_serv * pct_act_i
                    agg_act[(cod_actividad, cod_recurso)][1] += monto_act_serv * pct_act_i
    if personas_sin_grh12:
        avisos.append(f"CYG_8-GRH: {len(personas_sin_grh12)} persona(s) reguladas sin apertura en GRH_12.")

    # ================= GCP =================
    gcp5 = safe_rows("gcp5", "GCP_5")
    gcp6 = safe_rows("gcp6", "GCP_6")

    shares_actividad_gcp = defaultdict(lambda: defaultdict(float))  # (persona,cargo,recurso) -> {actividad: pct}
    for r in gcp6:
        persona, cargo, cod_recurso_g6, cod_actividad, pct = r[4], r[5], r[6], r[8], r[9]
        shares_actividad_gcp[(persona, cargo, cod_recurso_g6)][int(cod_actividad)] += pct

    faltan_gcp6 = 0
    for r in gcp5:
        empresa, periodo, anio, sector, persona, cargo, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        agg_serv[(cod_serv, cod_recurso)][0] += gasto_no_act
        # GCP_5 no trae desglose de monto activado por servicio -> se deja en 0 aquí
        if cod_reg in REGULADOS:
            actividades = shares_actividad_gcp.get((persona, cargo, cod_recurso))
            if not actividades:
                faltan_gcp6 += 1
                continue
            for cod_actividad, pct_a in actividades.items():
                agg_act[(cod_actividad, cod_recurso)][0] += gasto_no_act * pct_a
    if faltan_gcp6:
        avisos.append(f"CYG_8-GCP: {faltan_gcp6} combinación(es) reguladas sin apertura en GCP_6.")

    # ================= GGV =================
    ggv4 = safe_rows("ggv4", "GGV_4")
    ggv5 = safe_rows("ggv5", "GGV_5")
    ggv6 = safe_rows("ggv6", "GGV_6")

    shares_servicio_ggv = defaultdict(list)
    for r in ggv5:
        empresa, periodo, anio, sector, id_activo, total_no_act, cod_reg, cod_noreg, pct = r[:9]
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        shares_servicio_ggv[id_activo].append((cod_serv, pct))

    shares_actividad_ggv = defaultdict(lambda: defaultdict(float))
    for r in ggv6:
        id_activo, cod_actividad, pct = r[4], r[6], r[7]
        shares_actividad_ggv[id_activo][int(cod_actividad)] += pct

    activos_sin_ggv6 = set()
    for r in ggv4:
        empresa, periodo, anio, sector, id_activo, cod_recurso, monto_anual, monto_act, pct_act, total_gasto = r[:10]
        for cod_serv, pct_serv in shares_servicio_ggv.get(id_activo, []):
            gasto_serv = total_gasto * pct_serv
            monto_act_serv = monto_act * pct_serv
            agg_serv[(cod_serv, cod_recurso)][0] += gasto_serv
            agg_serv[(cod_serv, cod_recurso)][1] += monto_act_serv
            if cod_serv in REGULADOS:
                actividades = shares_actividad_ggv.get(id_activo)
                if not actividades:
                    activos_sin_ggv6.add(id_activo)
                    continue
                for cod_actividad, pct_a in actividades.items():
                    agg_act[(cod_actividad, cod_recurso)][0] += gasto_serv * pct_a
                    agg_act[(cod_actividad, cod_recurso)][1] += monto_act_serv * pct_a
    if activos_sin_ggv6:
        avisos.append(f"CYG_8-GGV: {len(activos_sin_ggv6)} activo(s) regulados sin apertura en GGV_6.")

    # ================= GGI =================
    ggi5 = safe_rows("ggi5", "GGI_5")
    ggi6 = safe_rows("ggi6", "GGI_6")

    shares_proceso_ggi = defaultdict(lambda: defaultdict(float))
    for r in ggi6:
        id_inmueble, cod_proceso_directo, pct = r[4], r[6], r[7]
        shares_proceso_ggi[id_inmueble][cod_proceso_directo] += pct

    for r in ggi5:
        empresa, periodo, anio, sector, id_inmueble, cod_recurso, gasto_no_act, cod_reg, cod_noreg, pct = r[:10]
        cod_serv = cod_reg if cod_reg != -1 else cod_noreg
        agg_serv[(cod_serv, cod_recurso)][0] += gasto_no_act
        if cod_reg in REGULADOS:
            # GGI no tiene tabla de actividad -> reparto equitativo por defecto
            # entre las actividades del proceso 504 (parametrizable)
            overrides = DEFAULT_ACTIVIDADES_GGI  # {proceso: [actividades]}; sin parametrización propia aún
            actividades_504 = overrides.get(504, [])
            if actividades_504:
                pct_igual = 1.0 / len(actividades_504)
                for cod_actividad in actividades_504:
                    agg_act[(cod_actividad, cod_recurso)][0] += gasto_no_act * pct_igual

    # ================= GGM =================
    ggm_tablas = [safe_leer_tabla_plana(fb.get(f"ggm{i}"), f"GGM_{i}", avisos) for i in range(1, 6)]
    ggm_regulado_por_recurso = defaultdict(lambda: [0.0, 0.0])  # [gasto_no_act, monto_act]
    for tabla in ggm_tablas:
        for r in tabla:
            cod_recurso, monto_act, total_gasto = r[4], r[5], r[6]
            overrides = ggm_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            agg_serv[(1101, cod_recurso)][0] += total_gasto * pct_reg
            agg_serv[(1101, cod_recurso)][1] += monto_act  # monto activado SIEMPRE 100% regulado
            for cod_serv_over, pct_o in overrides:
                agg_serv[(cod_serv_over, cod_recurso)][0] += total_gasto * pct_o
                # el monto activado no se traspasa a servicios no regulados
            ggm_regulado_por_recurso[cod_recurso][0] += total_gasto * pct_reg
            ggm_regulado_por_recurso[cod_recurso][1] += monto_act

    def _repartir_actividades_igual(cod_recurso, gasto_g, gasto_a, proceso_default, overrides_proceso, defaults_map):
        if gasto_g == 0 and gasto_a == 0:
            return
        if overrides_proceso:
            pct_default = 1.0 - sum(p for _, p in overrides_proceso)
            acts_default = defaults_map.get(proceso_default) or actividades_de_proceso(proceso_default)
            if acts_default:
                pct_igual = pct_default / len(acts_default)
                for a in acts_default:
                    agg_act[(a, cod_recurso)][0] += gasto_g * pct_igual
                    agg_act[(a, cod_recurso)][1] += gasto_a * pct_igual
            for proc_ov, pct_ov in overrides_proceso:
                acts_ov = actividades_de_proceso(proc_ov)
                if acts_ov:
                    pct_igual_ov = pct_ov / len(acts_ov)
                    for a in acts_ov:
                        agg_act[(a, cod_recurso)][0] += gasto_g * pct_igual_ov
                        agg_act[(a, cod_recurso)][1] += gasto_a * pct_igual_ov
        else:
            acts_default = defaults_map.get(proceso_default) or actividades_de_proceso(proceso_default)
            if acts_default:
                pct_igual = 1.0 / len(acts_default)
                for a in acts_default:
                    agg_act[(a, cod_recurso)][0] += gasto_g * pct_igual
                    agg_act[(a, cod_recurso)][1] += gasto_a * pct_igual

    for cod_recurso, (gasto_g, gasto_a) in ggm_regulado_por_recurso.items():
        proceso_default = DEFAULT_PROCESO_GGM.get(cod_recurso, 504)
        overrides_proceso = ggm_proceso_params.get(cod_recurso)
        _repartir_actividades_igual(cod_recurso, gasto_g, gasto_a, proceso_default, overrides_proceso, DEFAULT_ACTIVIDADES_GGM)

    # ================= OGG =================
    ogg5 = safe_leer_tabla_plana(fb.get("ogg5"), "OGG_5", avisos)
    ogg_regulado_por_recurso = defaultdict(lambda: [0.0, 0.0])
    for r in ogg5:
        cod_recurso, monto_act, total_gasto = r[4], r[5], r[6]
        overrides = ogg_params.get(cod_recurso, [])
        pct_reg = 1.0 - sum(p for _, p in overrides)
        agg_serv[(1101, cod_recurso)][0] += total_gasto * pct_reg
        agg_serv[(1101, cod_recurso)][1] += monto_act  # monto activado SIEMPRE 100% regulado
        for cod_serv_over, pct_o in overrides:
            agg_serv[(cod_serv_over, cod_recurso)][0] += total_gasto * pct_o
            # el monto activado no se traspasa a servicios no regulados
        ogg_regulado_por_recurso[cod_recurso][0] += total_gasto * pct_reg
        ogg_regulado_por_recurso[cod_recurso][1] += monto_act

    for cod_recurso, (gasto_g, gasto_a) in ogg_regulado_por_recurso.items():
        proceso_default = DEFAULT_PROCESO_OGG.get(cod_recurso, DEFAULT_PROCESO_OGG_FALLBACK)
        overrides_proceso = ogg_proceso_params.get(cod_recurso)
        _repartir_actividades_igual(cod_recurso, gasto_g, gasto_a, proceso_default, overrides_proceso, DEFAULT_ACTIVIDADES_OGG)

    # ================= MEI =================
    mei1 = safe_leer_tabla_mei(fb.get("mei1"), "MEI_1", avisos, necesita_obra_nbi=True)
    mei2 = safe_leer_tabla_mei(fb.get("mei2"), "MEI_2", avisos, necesita_actividad=True)
    mei3 = safe_leer_tabla_mei(fb.get("mei3"), "MEI_3", avisos, necesita_actividad=True)
    mei4 = safe_leer_tabla_mei(fb.get("mei4"), "MEI_4", avisos, necesita_actividad=True)

    # --- MEI_1: mapeo EXACTO vía CÓDIGO OBRA TIPO NBI (no estimación) ---
    obras_no_mapeadas = set()
    for r in mei1:
        cod_recurso, cod_obra_nbi, monto_act, total_gasto = r[4], r[5], r[6], r[7]
        overrides = mei_params.get(cod_recurso, [])
        pct_reg = 1.0 - sum(p for _, p in overrides)
        gasto_reg = total_gasto * pct_reg
        act_reg = monto_act  # monto activado SIEMPRE 100% regulado
        agg_serv[(1101, cod_recurso)][0] += gasto_reg
        agg_serv[(1101, cod_recurso)][1] += act_reg
        for cod_serv_over, pct_o in overrides:
            agg_serv[(cod_serv_over, cod_recurso)][0] += total_gasto * pct_o
            # el monto activado no se traspasa a servicios no regulados
        if gasto_reg == 0 and act_reg == 0:
            continue
        cod_actividad = OBRA_NBI_A_ACTIVIDAD.get(cod_obra_nbi)
        if cod_actividad is None:
            obras_no_mapeadas.add(cod_obra_nbi)
            continue
        agg_act[(cod_actividad, cod_recurso)][0] += gasto_reg
        agg_act[(cod_actividad, cod_recurso)][1] += act_reg
    if obras_no_mapeadas:
        avisos.append(f"CYG_8-MEI_1: código(s) OBRA TIPO NBI sin mapeo conocido: {sorted(obras_no_mapeadas)}.")

    def procesar_mei_con_actividad(tabla, idx_recurso, idx_gasto, idx_activado, idx_actividad):
        for r in tabla:
            cod_recurso, total_gasto, monto_act, cod_actividad = r[idx_recurso], r[idx_gasto], r[idx_activado], r[idx_actividad]
            if cod_actividad is None:
                continue
            overrides = mei_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            gasto_reg = total_gasto * pct_reg
            act_reg = monto_act  # monto activado SIEMPRE 100% regulado
            agg_serv[(1101, cod_recurso)][0] += gasto_reg
            agg_serv[(1101, cod_recurso)][1] += act_reg
            for cod_serv_over, pct_o in overrides:
                agg_serv[(cod_serv_over, cod_recurso)][0] += total_gasto * pct_o
                # el monto activado no se traspasa a servicios no regulados
            if gasto_reg == 0 and act_reg == 0:
                continue
            agg_act[(int(cod_actividad), cod_recurso)][0] += gasto_reg
            agg_act[(int(cod_actividad), cod_recurso)][1] += act_reg

    procesar_mei_con_actividad(mei2, 4, 7, 6, 8)
    procesar_mei_con_actividad(mei3, 4, 7, 6, 8)
    procesar_mei_con_actividad(mei4, 4, 7, 6, 8)

    # ================= ST =================
    tablas_st_faltantes = [t for t in ST_TABLES if t not in st_files_raw]
    if tablas_st_faltantes and st_files_raw:
        avisos.append(f"CYG: faltan {len(tablas_st_faltantes)} tabla(s) ST: {', '.join(tablas_st_faltantes)}.")
    for nombre_tabla, file_bytes in st_files_raw.items():
        try:
            filas_st, avisos_tabla = leer_tabla_actividad_st(file_bytes)
        except Exception as e:
            avisos.append(f"CYG-ST: no se pudo leer {nombre_tabla} ({e}).")
            continue
        for a in avisos_tabla:
            avisos.append(f"CYG-ST {nombre_tabla}: {a}")
        for cod_recurso, monto_act, total_gasto, cod_actividad in filas_st:
            overrides = st_params.get(cod_recurso, [])
            pct_reg = 1.0 - sum(p for _, p in overrides)
            gasto_reg = total_gasto * pct_reg
            act_reg = monto_act  # monto activado SIEMPRE 100% regulado
            agg_serv[(1101, cod_recurso)][0] += gasto_reg
            agg_serv[(1101, cod_recurso)][1] += act_reg
            for cod_serv_over, pct_o in overrides:
                agg_serv[(cod_serv_over, cod_recurso)][0] += total_gasto * pct_o
                # el monto activado no se traspasa a servicios no regulados
            if gasto_reg == 0 and act_reg == 0:
                continue
            agg_act[(cod_actividad, cod_recurso)][0] += gasto_reg
            agg_act[(cod_actividad, cod_recurso)][1] += act_reg

    # ================= GPA =================
    tablas_gpa_faltantes = [t for t in GPA_TABLES if t not in gpa_files_raw]
    if tablas_gpa_faltantes and gpa_files_raw:
        avisos.append(f"CYG-GPA: faltan {len(tablas_gpa_faltantes)} tabla(s): {', '.join(tablas_gpa_faltantes)}.")
    for nombre_tabla, file_bytes in gpa_files_raw.items():
        servicio_fijo = GPA_TABLE_TO_SERVICIO.get(nombre_tabla)
        if servicio_fijo is None:
            continue
        try:
            filas_gpa = leer_tabla_st(file_bytes)
        except Exception as e:
            avisos.append(f"CYG-GPA: no se pudo leer {nombre_tabla} ({e}).")
            continue
        gpa_por_recurso = defaultdict(lambda: [0.0, 0.0])
        for cod_recurso, monto_act, total_gasto in filas_gpa:
            agg_serv[(servicio_fijo, cod_recurso)][0] += total_gasto
            agg_serv[(servicio_fijo, cod_recurso)][1] += monto_act
            gpa_por_recurso[cod_recurso][0] += total_gasto
            gpa_por_recurso[cod_recurso][1] += monto_act
        for cod_recurso, (gasto_g, gasto_a) in gpa_por_recurso.items():
            overrides_proceso = gpa_proceso_params.get(cod_recurso)
            acts_default_map = {DEFAULT_PROCESO_GPA: DEFAULT_ACTIVIDADES_GPA.get(cod_recurso, [])}
            _repartir_actividades_igual(cod_recurso, gasto_g, gasto_a, DEFAULT_PROCESO_GPA, overrides_proceso, acts_default_map)

    EMPRESA = PERIODO = ANIO = SECTOR = None
    for clave in ["grh8", "gcp5", "ggv4", "ggi5", "ggm1", "ogg5"]:
        b = fb.get(clave)
        if b:
            filas_tmp = read_rows(b)
            if filas_tmp:
                EMPRESA, PERIODO, ANIO, SECTOR = filas_tmp[0][0], filas_tmp[0][1], filas_tmp[0][2], filas_tmp[0][3]
                break

    return agg_serv, agg_act, avisos, (EMPRESA, PERIODO, ANIO, SECTOR)

def _normalizar_por_clave(items_por_clave, decimales=4):
    """items_por_clave: {clave_agrupadora: [(sub_clave, gasto), ...]}
    Devuelve {clave_agrupadora: [(sub_clave, pct, gasto), ...]} con pct
    sumando EXACTAMENTE 1.0 dentro de cada clave_agrupadora (método de mayor
    residuo)."""
    resultado = {}
    for clave, items in items_por_clave.items():
        total = sum(g for _, g in items)
        if total == 0:
            continue
        fracciones = [g / total for _, g in items]
        pcts = redondear_para_sumar_100(fracciones, decimales)
        resultado[clave] = [(sub, pct, g) for (sub, g), pct in zip(items, pcts)]
    return resultado


def build_cyg_1_a_4(agg_serv, empresa_periodo_anio_sector):
    """Arma CYG_1 (1101,1102), CYG_2 (1201-1205), CYG_3 (2101-2117),
    CYG_4 (2201-2214). Normaliza % por SERVICIO ESPECÍFICO (no por recurso)."""
    EMPRESA, PERIODO, ANIO, SECTOR = empresa_periodo_anio_sector
    rangos = {
        "CYG_1": set([1101, 1102]),
        "CYG_2": set([1201, 1202, 1203, 1204, 1205]),
        "CYG_3": set(range(2101, 2118)),
        "CYG_4": set(range(2201, 2215)),
    }
    resultados = {}
    for nombre_tabla, servicios_validos in rangos.items():
        # separar no-activado y activado (normalización independiente, igual que REP_2)
        items_no_act = defaultdict(list)  # servicio -> [(recurso, gasto), ...]
        items_act = defaultdict(list)
        for (cod_serv, cod_recurso), (g, a) in agg_serv.items():
            if cod_serv not in servicios_validos:
                continue
            if abs(g) > 1e-9:
                items_no_act[cod_serv].append((cod_recurso, g))
            if abs(a) > 1e-9:
                items_act[cod_serv].append((cod_recurso, a))

        norm_no_act = _normalizar_por_clave(items_no_act)
        norm_act = _normalizar_por_clave(items_act)

        filas = defaultdict(lambda: [0.0, 0.0, 0.0, 0.0])  # (servicio,recurso) -> [pct_no_act,monto_no_act,pct_act,monto_act]
        for servicio, lista in norm_no_act.items():
            for recurso, pct, g in lista:
                filas[(servicio, recurso)][0] = pct
                filas[(servicio, recurso)][1] = round(g, 2)
        for servicio, lista in norm_act.items():
            for recurso, pct, a in lista:
                filas[(servicio, recurso)][2] = pct
                filas[(servicio, recurso)][3] = round(a, 2)

        final_rows = []
        for (servicio, recurso), (pct_na, monto_na, pct_a, monto_a) in sorted(filas.items()):
            if monto_na == 0 and monto_a == 0:
                continue
            final_rows.append([EMPRESA, PERIODO, ANIO, SECTOR, servicio, recurso,
                                round(pct_na * 100, 2), monto_na, round(pct_a * 100, 2), monto_a])
        resultados[nombre_tabla] = final_rows
    return resultados


def build_cyg_8(agg_act, empresa_periodo_anio_sector):
    """CYG_8: (actividad, recurso, %no_act, monto_no_act, %act, monto_act).
    Normaliza % por RECURSO (sin distinguir familia de servicio)."""
    EMPRESA, PERIODO, ANIO, SECTOR = empresa_periodo_anio_sector

    items_no_act = defaultdict(list)  # recurso -> [(actividad, gasto), ...]
    items_act = defaultdict(list)
    for (cod_actividad, cod_recurso), (g, a) in agg_act.items():
        if abs(g) > 1e-9:
            items_no_act[cod_recurso].append((cod_actividad, g))
        if abs(a) > 1e-9:
            items_act[cod_recurso].append((cod_actividad, a))

    norm_no_act = _normalizar_por_clave(items_no_act)
    norm_act = _normalizar_por_clave(items_act)

    filas = defaultdict(lambda: [0.0, 0.0, 0.0, 0.0])  # (actividad,recurso) -> [pct_no_act,monto_no_act,pct_act,monto_act]
    for recurso, lista in norm_no_act.items():
        for actividad, pct, g in lista:
            filas[(actividad, recurso)][0] = pct
            filas[(actividad, recurso)][1] = round(g, 2)
    for recurso, lista in norm_act.items():
        for actividad, pct, a in lista:
            filas[(actividad, recurso)][2] = pct
            filas[(actividad, recurso)][3] = round(a, 2)

    final_rows = []
    for (actividad, recurso), (pct_na, monto_na, pct_a, monto_a) in sorted(filas.items()):
        if monto_na == 0 and monto_a == 0:
            continue
        final_rows.append([EMPRESA, PERIODO, ANIO, SECTOR, actividad, recurso,
                            round(pct_na * 100, 2), monto_na, round(pct_a * 100, 2), monto_a])
    return final_rows


PATRON_RUT_CHILENO = re.compile(r"^\d{6,9}-[0-9kK]$")


def es_rut_valido(id_cliente):
    """True si id_cliente tiene formato de RUT chileno válido (ej.
    '12345678-9' o '12345678-K'). Se usa para filtrar, en ING_4, las filas
    que corresponden a clientes reales (excluyendo provisiones, notas
    contables, referencias de facturación u otros textos que no son RUT)."""
    if id_cliente is None:
        return False
    return bool(PATRON_RUT_CHILENO.match(str(id_cliente).strip()))


def build_cyg_9(agg_serv, mco42_bytes, ing4_bytes, empresa_periodo_anio_sector):
    """CYG_9: abre servicios NO regulados (21xx, 22xx) por ID CLIENTE, usando
    MCO_42 (servicio -> id_ingreso) e ING_4 (id_ingreso -> id_cliente + monto
    anual de ingreso) para repartir proporcionalmente al ingreso de cada
    cliente. Solo se consideran clientes con ID CLIENTE en formato RUT
    chileno válido (ej. '12345678-9' o '12345678-K'); las demás filas de
    ING_4 (provisiones, notas contables, referencias de facturación, etc.)
    se EXCLUYEN del cálculo de proporción y se reportan en 'avisos' para que
    se corrijan en el origen. Si un servicio no está cubierto por MCO_42/
    ING_4 (o solo tiene entradas no-RUT), se declara con ID CLIENTE = '-1'
    (100%), según lo dispuesto en el diccionario SISS."""
    EMPRESA, PERIODO, ANIO, SECTOR = empresa_periodo_anio_sector
    avisos = []

    if mco42_bytes is None or ing4_bytes is None:
        avisos.append("CYG_9: falta MCO_42 y/o ING_4 — todos los clientes se declaran como '-1'.")
        mco = []
        ing = []
    else:
        mco = read_rows(mco42_bytes)
        ing = read_rows(ing4_bytes)

    servicio_a_ingreso = {r[4]: r[5] for r in mco}
    clientes_por_ingreso = defaultdict(list)
    no_rut_encontrados = []  # (id_ingreso, id_cliente, monto) para el aviso
    for r in ing:
        id_ingreso, id_cliente, monto = r[5], r[6], r[9]
        if es_rut_valido(id_cliente):
            clientes_por_ingreso[id_ingreso].append((id_cliente, monto))
        else:
            no_rut_encontrados.append((id_ingreso, id_cliente, monto))

    if no_rut_encontrados:
        monto_excluido = sum(m for _, _, m in no_rut_encontrados)
        ejemplos = ", ".join(f"'{c}'" for _, c, _ in no_rut_encontrados[:5])
        avisos.append(
            f"CYG_9: se excluyeron {len(no_rut_encontrados)} fila(s) de ING_4 con ID CLIENTE "
            f"sin formato RUT válido (monto total excluido de la base de reparto: "
            f"{monto_excluido:,.0f}). Revisar y corregir en el origen. Ejemplos: {ejemplos}"
            + (", ..." if len(no_rut_encontrados) > 5 else "") + "."
        )

    servicios_no_reg = set(range(2101, 2118)) | set(range(2201, 2215))

    filas = []
    for (cod_serv, cod_recurso), (g, a) in agg_serv.items():
        if cod_serv not in servicios_no_reg:
            continue
        if g == 0 and a == 0:
            continue

        id_ingreso = servicio_a_ingreso.get(cod_serv)
        clientes = clientes_por_ingreso.get(id_ingreso) if id_ingreso is not None else None

        if not clientes:
            # sin cobertura MCO_42/ING_4 -> 100% a ID CLIENTE "-1"
            filas.append([cod_serv, cod_recurso, "-1", g, a])
            continue

        total_ingreso = sum(c[1] for c in clientes)
        if total_ingreso == 0:
            filas.append([cod_serv, cod_recurso, "-1", g, a])
            continue

        # Reparto proporcional SIN redondear aún (el redondeo final, con piso
        # mínimo, se aplica una sola vez más abajo al consolidar por recurso;
        # redondear aquí prematuramente a 4 decimales podía dejar en 0.0000 —
        # y por lo tanto eliminar— a clientes muy pequeños relativos al total
        # de ingreso, antes de que el piso mínimo alcanzara a rescatarlos).
        for id_cliente, monto in clientes:
            pct_cliente = monto / total_ingreso
            filas.append([cod_serv, cod_recurso, id_cliente, g * pct_cliente, a * pct_cliente])

    # Consolidar filas duplicadas (mismo servicio,recurso,cliente pudiera repetirse si
    # el mismo cliente aparece en más de un ID INGRESO -- no debería, pero por robustez)
    consolidado = defaultdict(lambda: [0.0, 0.0])
    for cod_serv, cod_recurso, id_cliente, g, a in filas:
        consolidado[(cod_serv, cod_recurso, id_cliente)][0] += g
        consolidado[(cod_serv, cod_recurso, id_cliente)][1] += a

    # Normalizar % por RECURSO (no por servicio+recurso): el % representa la
    # porción del RECURSO asignada a cada cliente, combinando todos los
    # servicios no regulados que usan ese recurso.
    items_no_act = defaultdict(list)
    items_act = defaultdict(list)
    for (cod_serv, cod_recurso, id_cliente), (g, a) in consolidado.items():
        if abs(g) > 1e-9:
            items_no_act[cod_recurso].append(((cod_serv, id_cliente), g))
        if abs(a) > 1e-9:
            items_act[cod_recurso].append(((cod_serv, id_cliente), a))

    norm_no_act, fusion_no_act = normalizar_con_piso_minimo(items_no_act, piso=0.0001)
    norm_act, fusion_act = normalizar_con_piso_minimo(items_act, piso=0.0001)

    recursos_fusionados = {rec for rec, f in fusion_no_act.items() if f} | {rec for rec, f in fusion_act.items() if f}
    if recursos_fusionados:
        avisos.append(
            f"CYG_9: {len(recursos_fusionados)} recurso(s) con demasiados clientes para "
            f"que todos alcancen el piso mínimo de 0,01% individualmente; los clientes de "
            f"menor monto se agruparon bajo ID CLIENTE '-1'. Recursos afectados: "
            f"{sorted(recursos_fusionados)}."
        )

    filas_finales = defaultdict(lambda: [0.0, 0.0, 0.0, 0.0])
    for cod_recurso, lista in norm_no_act.items():
        for (cod_serv, id_cliente), pct, g in lista:
            filas_finales[(cod_serv, cod_recurso, id_cliente)][0] = pct
            filas_finales[(cod_serv, cod_recurso, id_cliente)][1] = round(g, 2)
    for cod_recurso, lista in norm_act.items():
        for (cod_serv, id_cliente), pct, a in lista:
            filas_finales[(cod_serv, cod_recurso, id_cliente)][2] = pct
            filas_finales[(cod_serv, cod_recurso, id_cliente)][3] = round(a, 2)

    final_rows = []
    for (cod_serv, cod_recurso, id_cliente), (pct_na, monto_na, pct_a, monto_a) in sorted(
            filas_finales.items(), key=lambda x: (x[0][0], x[0][1], str(x[0][2]))):
        if monto_na == 0 and monto_a == 0:
            continue
        final_rows.append([EMPRESA, PERIODO, ANIO, SECTOR, id_cliente, cod_serv, cod_recurso,
                            round(pct_na * 100, 2), monto_na, round(pct_a * 100, 2), monto_a])

    return final_rows, avisos



# ============================================================================
# CARGA UNIFICADA: una sola carpeta comprimida (.zip) con todas las tablas
# ============================================================================
# El usuario sube UN solo archivo .zip (su carpeta completa, con subcarpetas
# como "ST/", "GPA/", etc. si quiere organizarlas así -- no importa la
# estructura de carpetas, solo el NOMBRE de cada archivo). El sistema busca,
# dentro del zip, cada tabla que necesita, tolerando texto adicional en el
# nombre (ej. "MEI_1_2025.xlsx" o "MEI_1 (v2) 2025-07.xlsx" calzan con "MEI_1").

# Tablas "simples" (nombre único, sin variantes ST_x/GPA_x)
TABLAS_SIMPLES = [
    "GRH_1", "GRH_8", "GRH_11", "GRH_12",
    "GCP_4", "GCP_5", "GCP_6",
    "GGV_3", "GGV_4", "GGV_5", "GGV_6",
    "GGI_1", "GGI_2", "GGI_5", "GGI_6",
    "GGM_1", "GGM_2", "GGM_3", "GGM_4", "GGM_5",
    "OGG_5",
    "MEI_1", "MEI_2", "MEI_3", "MEI_4",
    "MCO_42", "ING_4",
]

# Prefijos de familia reconocidos del SCR (Sistema Contable Regulatorio). Se
# usan como respaldo genérico en identificar_tabla_generico(): cualquier
# archivo "{FAMILIA}_{NUMERO}..." con una de estas familias se reconoce
# automáticamente como esa tabla, AUNQUE NO esté en TABLAS_SIMPLES/ST_TABLES/
# GPA_TABLES explícitamente. Así, tablas nuevas del SCR que se incorporen a
# futuras funciones (ej. una GGI_3 o GRH_2 que hoy no existen) se reconocen
# solas al cargarlas, sin tener que acordarse de agregarlas aquí a mano.
FAMILIAS_SCR_CONOCIDAS = ["GRH", "GCP", "GGV", "GGI", "GGM", "OGG", "MEI", "ST", "GPA", "MCO", "ING"]


# Catálogo completo para el emparejamiento (simples + ST_x + GPA_x)
def _catalogo_completo():
    return TABLAS_SIMPLES + ST_TABLES + GPA_TABLES


# =====================================================================
# CATÁLOGO OFICIAL DE TABLAS DEL SCR (Sistema Contable Regulatorio)
# Fuente: articles-19850_SistContReg235.xlsx, hoja "Tablas" (219 tablas).
# Es la fuente de verdad para saber qué tablas EXISTEN de verdad — se usa
# para restringir el reconocimiento automático de archivos (ver
# identificar_tabla_generico) a solo códigos que realmente existen en el
# SCR, evitando "reconocer" tablas inventadas como GGI_8 (que no existe).
# =====================================================================
TABLAS_SCR_OFICIAL = {
    'ACT_1': 'Activos Totales',
    'ACT_2': 'Activos Sanitarios',
    'ACT_3': 'Activos No Sanitarios',
    'ACT_4': 'Proyectos en Desarrollo',
    'ACT_5': 'Activos en Comodato (Propiedad de Terceros)',
    'ACT_6': 'Activos Dados en Explotación a Terceros ',
    'ACT_7': 'Activos Sanitarios - Servicios',
    'ACT_8': 'Activos No Sanitarios - Servicios',
    'AUX_1': 'Control de calidad AP - Muestreo',
    'AUX_2': 'Control de calidad AP - Análisis de laboratorio',
    'AUX_3': 'Control de calidad AS - Muestreo',
    'AUX_4': 'Control de calidad AS - Análisis de laboratorio',
    'AUX_5': 'Servicios de operación de redes y conexiones',
    'AUX_6': 'Servicios de operación de infraestructura',
    'AUX_7': 'Servicios de gestión de lodos',
    'AUX_8': 'Servicios de control y monitoreo ambiental',
    'AUX_9': 'Gestión de residuos sólidos',
    'AUX_10': 'Mantención de infraestructura capacidad y conducciones',
    'AUX_11': 'Mantención de recintos',
    'AUX_12': 'Mantención de servidumbres',
    'AUX_13': 'Mantención de Estanques de Regulación',
    'AUX_14': 'Mantención Red Distribución - Limpieza de Cámaras',
    'AUX_15': 'Mantención Red Distribución - Renovación de Cámaras',
    'AUX_16': 'Mantención Red Distribución - Detección de Fugas',
    'AUX_17': 'Operación Red de Distribución - Control de presiones',
    'AUX_18': 'Mantención Red Distribución y conducciones AP - Reparación de Roturas',
    'AUX_19': 'Mantención Red Distribución - Reparación de Arranques',
    'AUX_20': 'Mantención Red Distribución - Renovación de Arranques',
    'AUX_21': 'Mantención Red Distribución - Verificación y Control de Medidores',
    'AUX_22': 'Mantención Red Distribución - Recambio de Medidores',
    'AUX_23': 'Mantención Red Recolección - Inspección Televisiva de Colectores',
    'AUX_24': 'Mantención Red Recolección - Desobstrucción de Colectores',
    'AUX_25': 'Mantención Red Recolección - Limpieza y Reparación de Colectores',
    'AUX_26': 'Mantención Red Recolección - Limpieza y Reparación de Cámaras',
    'AUX_27': 'Mantención Red Recolección - Renovación de Cámaras',
    'AUX_28': 'Mantención de UD - Desobstrucción, reparación y renovación de Uniones Domiciliaras',
    'AUX_29': 'Ciclo Recaudación - Lectura de Medidores',
    'AUX_30': 'Ciclo Recaudación - Reparto Boletas y Otros Documentos',
    'AUX_31': 'Ciclo Recaudación - Cobranzas',
    'AUX_32': 'Ciclo Recaudación - Castigo de Deuda Clientes',
    'AUX_33': 'Gestión Comercial - Atención Presencial',
    'AUX_34': 'Gestión Comercial - Atención Telefónica',
    'AUX_35': 'Gestión Comercial - Atención No Presencial',
    'AUX_36': 'Gestión Comercial - Recepción y Estudio de Factibilidades',
    'AUX_37': 'Gestión Comercial - Aprobaciones de Proyectos de Conexión',
    'AUX_38': 'Gestión Comercial - Inspecciones Comerciales',
    'AUX_39': 'Prestaciones Asociadas - Control Directo de RILES',
    'AUX_40': 'Prestaciones Asociadas - Mantención de Grifos',
    'AUX_41': 'Prestaciones Asociadas - Corte y Reposición',
    'AUX_42': 'Prestaciones Asociadas - Revisión de Proyectos de Construcción',
    'AUX_43': 'Prestaciones Asociadas - Verificación de Medidores a Solicitud del Cliente',
    'AUX_44': 'Bitácora Mantención',
    'AUX_45': 'Bitácora Mantención-Materiales y Repuestos',
    'AUX_46': 'Bitácora Operación',
    'AUX_47': 'Bitácora Operación-Materiales y Repuestos',
    'CTB_1': 'EERR - CMF',
    'CTB_2': 'EEFF - CMF',
    'CTB_3': 'Cuentas Libro Mayor - Ingresos',
    'CTB_4': 'Cuentas Libro Mayor - Costos y Gastos',
    'CTB_5': 'Cuentas Libro Mayor - Activos',
    'CTB_6': 'Cuentas Libro Mayor - Pasivos',
    'CTB_7': 'Cuentas Libro Mayor - Patrimonio',
    'CTB_8': 'Cuentas Libro Remuneraciones',
    'CTB_9': 'Centros de Costos',
    'CYG_1': 'Recursos - Servicios Continuos de Agua Potable y Alcantarillado',
    'CYG_2': 'Recursos - Prestaciones Asociadas Reguladas',
    'CYG_3': 'Recursos - Prestaciones Asociadas No Reguladas',
    'CYG_4': 'Recursos - Servicios No Regulados',
    'CYG_5': 'Recursos de Servicios Continuos - Emergencias',
    'CYG_6': 'Recursos de Servicios Continuos - Sequía',
    'CYG_7': 'Recursos de Servicios Continuos - Turbiedad Extrema',
    'CYG_8': 'Recursos de Servicios Continuos - Actividades',
    'CYG_9': 'Recursos de Servicios No Regulados - Servicios Prestados a Terceros',
    'GCP_1': 'Capacitación',
    'GCP_2': 'Viajes y viáticos',
    'GCP_3': 'Accesorios Personal',
    'GCP_4': 'Resumen Gastos Generales Personal',
    'GCP_5': 'Gastos Generales Personal: Servicios',
    'GCP_6': 'Gastos Generales Personal: Actividades',
    'GGI_1': 'Vigilancia Presencial',
    'GGI_2': 'Vigilancia a Distancia',
    'GGI_3': 'Consumos básicos',
    'GGI_4': 'Resumen Gastos Generales Bienes Inmuebles ',
    'GGI_5': 'Gastos Generales Bienes Inmuebles: Servicio',
    'GGI_6': 'Gastos Generales Bienes Inmuebles: Proceso',
    'GGM_1': 'Arriendo de Equipos Informáticos',
    'GGM_2': 'Servicios Informáticos',
    'GGM_3': 'Servicios de Telecomunicaciones',
    'GGM_4': 'Materiales e insumos de oficina, computacionales y bodega',
    'GGM_5': 'Materiales e insumos de laboratorio',
    'GGV_1': 'Gastos Generales Vehículos: Arriendo de Vehículos',
    'GGV_2': 'Gastos Generales Vehículos: Combustibles',
    'GGV_3': 'Gastos Generales Equipos de Generación',
    'GGV_4': 'Resumen Gastos Generales Vehículos y Equipos',
    'GGV_5': 'Gastos Generales Vehículos y Equipos: Servicios',
    'GGV_6': 'Gastos Generales Vehículos y Equipos: Actividades',
    'GPA_1': 'Control Directo de Riles',
    'GPA_2': 'Mantención de Grifos',
    'GPA_3': 'Corte y Reposición',
    'GPA_4': 'Revisión de Proyectos de Construcción',
    'GPA_5': 'Verificación de Medidores',
    'GPA_6': 'Otras Prestaciones Asociadas',
    'GRH_1': 'RRHH-Remuneraciones',
    'GRH_2': 'RRHH-Desagregación Remuneraciones',
    'GRH_3': 'RRHH-Remuneraciones Dotación Honorarios',
    'GRH_4': 'RRHH-Horas Extras',
    'GRH_5': 'RRHH-Indemnizaciones',
    'GRH_6': 'RRHH-Protección social y salud',
    'GRH_7': 'RRHH-Otros Beneficios',
    'GRH_8': 'RRHH-Resumen Gastos Recursos Humanos',
    'GRH_9': 'RRHH-Emergencias y Turbiedad Extrema',
    'GRH_10': 'RRHH-Extras por Sequía',
    'GRH_11': 'RRHH-Asignación Servicios',
    'GRH_12': 'RRHH-Asignación Actividades',
    'GRH_13': 'RRHH-Obras-Actividades',
    'ING_1': 'Ingresos de Servicios Continuos de Agua Potable y Alcantarillado',
    'ING_2': 'Ingresos de Prestaciones Asociadas Reguladas',
    'ING_3': 'Ingresos de Prestaciones Asociadas No Reguladas',
    'ING_4': 'Ingresos de Servicios No Regulados',
    'MEI_1': 'Productos Químicos',
    'MEI_2': 'Energía Eléctrica',
    'MEI_3': 'Materiales y Repuestos',
    'MEI_4': 'Compras de Agua',
    'NCR_1': 'Facturación Servicios de NCR',
    'NCR_2': 'RRHH Servicios de NCR',
    'NCR_3': 'Recursos Generales Servicios de NCR',
    'NCR_4': 'Activos Servicios de NCR',
    'OGG_1': 'Directorio-Dietas',
    'OGG_2': 'Directorio-Otros Gastos',
    'OGG_3': 'Seguros de Infraestructura',
    'OGG_4': 'Seguros de Inmuebles',
    'OGG_5': 'Otros Gastos Generales',
    'OYM_1': 'Resumen Costos de Servicios Operacionales',
    'OYM_2': 'Costos Directos de Servicios Operacionales: Obras-Actividad',
    'PIN_1': 'Proyectos de Inversión: Activos Sanitarios',
    'PIN_2': 'Proyectos de Inversión: Activos No Sanitarios',
    'PIN_3': 'Proyectos de Inversión: Componentes de Obras Sanitarias',
    'PIN_4': 'Proyectos de Inversión: Estudios de Diseño e Ingeniería de Activos Sanitarios',
    'PIN_5': 'Proyectos de Inversión: Inspección Técnica de Obras (ITO)',
    'PIN_6': 'Proyectos de Inversión: TIC - SW Licencias',
    'PIN_7': 'Proyectos de Inversión: TIC - HW Macroinformática',
    'PIN_8': 'Proyectos de Inversión: TIC - HW Microinformática, Comunicaciones y Telecontrol',
    'PIN_9': 'Proyectos de Inversión: TIC - Desarrollos o Implementación',
    'PIN_10': 'Proyectos de Inversión: TIC - Contraparte Interna',
    'PIN_11': 'Proyectos de Inversión: TIC - Consultorías y Asesorías',
    'PIN_12': 'Proyectos de Inversión: Vehículos',
    'PIN_13': 'Proyectos de Inversión: Herramientas y Equipos',
    'PIN_14': 'Proyectos de Inversión: Cuentas de Contabilidad Regulatoria de Activos',
    'PIN_15': 'Proyectos de Inversión: Recursos Activados',
    'PIN_16': 'Proyectos de Inversión: Asignación Estudios',
    'PIN_17': 'Proyectos de Inversión: Asignación ITO',
    'REP_1': 'Resultados por Familia de Servicios',
    'REP_2': 'Costos y Gastos por Familia de Servicios',
    'REP_3': 'Costos y Gastos No Activados de Servicios Regulados - Proceso',
    'REP_4': 'Costos y Gastos No Activados de Servicios Regulados - Puesta en Marcha',
    'REP_5': 'Costos y Gastos - Empresas Relacionadas',
    'REP_6': 'Activos por Familia de Servicios',
    'RUS_1': 'Reutilización de Subproductos - Recursos, inversiones y gastos',
    'RUS_2': 'Reutilización de Subproductos - Actividades',
    'RUS_3': 'RRHH - Reutilización de Subproductos',
    'RUS_4': 'Activos Sanitarios - Reutilización de Subproductos',
    'RUS_5': 'Otros Activos - Reutilización de Subproductos',
    'SPC_1': 'Gastos de Servicios Prestados: Recursos y Actividades',
    'SPC_2': 'RRHH Servicios Prestados',
    'SPC_3': 'Activos NBI Servicios Prestados',
    'SPC_4': 'Activos Sanitarios No Informados en la NBI - Servicios Prestados',
    'SPC_5': 'Otros Activos Servicios Prestados',
    'SRC_1': 'Gastos de Servicios Recibidos: Recursos y Actividad',
    'SRC_2': 'RRHH Servicios Recibidos de Concesionarias y Empresas Relacionadas',
    'SRC_3': 'Activos NBI Servicios Recibidos de Concesionarias y Empresas Relacionadas',
    'SRC_4': 'Activos Sanitarios No Informados en la NBI - Servicios Recibidos de Concesionarias y Empresas Relacionadas',
    'SRC_5': 'Otros Activos Servicios Recibidos de Concesionarias y Empresas Relacionadas',
    'ST_1': 'Servicios Tercerizados',
    'ST_2': 'Servicios Tercerizados Homologados a Recursos y actividades',
    'ST_3': 'Lectura de medidores',
    'ST_4': 'Reparto de boletas y otros documentos',
    'ST_5': 'Suministro e impresión de boletas y otros documentos',
    'ST_6': 'Servicios de recaudación en cajas externas',
    'ST_7': 'Servicios de recaudación en cajas propias',
    'ST_8': 'Servicios de atención telefónica o distante',
    'ST_9': 'Servicios de inspección comercial',
    'ST_10': 'Servicios de cobranza prejudicial',
    'ST_11': 'Asesorías y Estudios',
    'ST_12': 'Servicios de control de calidad de agua potable',
    'ST_13': 'Servicios de control de calidad de aguas servidas',
    'ST_14': 'Servicios de operación de redes y conexiones',
    'ST_15': 'Servicios de operación de infraestructura',
    'ST_16': 'Servicios de operación y mantención de PTAP',
    'ST_17': 'Servicios de operación y mantención de PTAS',
    'ST_18': 'Servicios de transporte y disposición de lodos',
    'ST_19': 'Servicios de transporte y disposición de residuos',
    'ST_20': 'Servicios de control y monitoreo ambiental',
    'ST_21': 'Servicios de Mantención de Infraestructura de Capacidad\xa0y\xa0Conducciones',
    'ST_22': 'Servicios de mantención de redes y conexiones',
    'ST_23': 'Servicios de mantención de redes y conexiones-Reparación de roturas en redes y conducciones',
    'ST_24': 'Servicios de mantención de redes y conexiones-Reparación de roturas-Componentes',
    'ST_25': 'Servicios de mantención de redes y conexiones-Reparación y renovación de arranques',
    'ST_26': 'Servicios de mantención de redes y conexiones-Reparación y renovación de arranques-Componentes',
    'ST_27': 'Servicios de mantención de redes y conexiones-Desobstrucción de colectores',
    'ST_28': 'Servicios de mantención de redes y conexiones-Inspección televisiva de colectores',
    'ST_29': 'Servicios de mantención de redes y conexiones-Limpieza de colectores',
    'ST_30': 'Servicios de mantención de redes y conexiones-Desobstrucción, reparación y renovación de uniones',
    'ST_31': 'Servicios de mantención de recintos',
    'ST_32': 'Servicios de mantención de servidumbres',
    'ST_33': 'Otros servicios operacionales',
    'ST_34': 'Otros servicios no operacionales',
    'ST_35': 'Servicios Tercerizados-Emergencias y Turbiedad Extrema',
    'ST_36': 'Servicios Tercerizados- Sequía Extrema',
    'ST_37': 'Servicios Tercerizados-Asignación de Horas equivalentes a actividades',
    'TZC_1': 'CR Recursos - Cuentas Libro Mayor de Gastos',
    'TZC_2': 'CR Recursos - Cuentas Libro Mayor de Activos (Gastos Activados)',
    'TZC_3': 'CR Ingresos Regulados - Cuentas Libro Mayor de Ingresos',
    'TZC_4': 'CR Ingresos No Regulados - Cuentas Libro Mayor de Ingresos',
    'TZC_5': 'CR Activos - Cuentas Libro Mayor de Activos',
    'TZC_6': 'EERR - Libro Mayor',
    'TZC_7': 'EEFF - Libro Mayor',
    'TZC_8': 'Centros de Costos - Libro Mayor de Costos y Gastos',
    'TZC_9': 'Libro de Remuneraciones - Libro Mayor de Costos y Gastos',
    'TZC_10': 'Provisiones Contables de Costos y Gastos-Recursos',
}

# =====================================================================
# CATÁLOGO DE MCO (Maestros de Contratos y Órdenes / tablas de sustento)
# Fuente: articles-19850_MaestrosCon2305.xlsx, hoja "IndiceMaestrosEmpresa"
# (56 tablas MCO_1 a MCO_56). Referencia para las funciones de asignación
# de respaldos (ID Respaldo / ID Servicio Tercerizado -> MCO).
# =====================================================================
MCO_CATALOGO_COMPLETO = {
    'MCO_1': 'CONTRATOS DE SERVICIOS NO REGULADOS PRESTADOS',
    'MCO_2': 'LICITACIONES',
    'MCO_3': 'COMPRA DIRECTA',
    'MCO_4': 'CONTRATOS DE SERVICIOS RECIBIDOS VIGENTES',
    'MCO_5': 'CONTRATOS DE OBRAS',
    'MCO_6': 'CONTRATOS DE SUMINISTROS DE OBRAS E INSUMOS',
    'MCO_7': 'PRESUPUESTOS DE OBRAS',
    'MCO_8': 'ÓRDENES DE COMPRA DE OBRAS',
    'MCO_9': 'ÓRDENES DE TRABAJO O COMPRA DE SERVICIOS',
    'MCO_10': 'ÓRDENES DE COMPRA DE SUMINISTROS PARA OBRAS',
    'MCO_11': 'FACTURAS DE OBRAS',
    'MCO_12': 'FACTURAS DE SERVICIOS',
    'MCO_13': 'FACTURAS DE SUMINISTROS',
    'MCO_14': 'PROVEEDORES',
    'MCO_15': 'UNIDADES ORGANIZACIONALES',
    'MCO_16': 'DOTACIÓN PERSONAL INTERNO',
    'MCO_17': 'DOTACIÓN PERSONAL EXTERNO: SUMINISTRO DE PERSONAL',
    'MCO_18': 'DOTACIÓN PERSONAL EXTERNO: SERVICIOS TERCERIZADOS (PERMANENTES)',
    'MCO_19': 'DOTACIÓN PERSONAL DE REEMPLAZO',
    'MCO_20': 'RRHH-DOTACIÓN ZONALES',
    'MCO_21': 'CARGOS EMPRESA',
    'MCO_22': 'ZONALES',
    'MCO_23': 'INMUEBLES',
    'MCO_24': 'EDIFICACIONES',
    'MCO_25': 'TERRENOS',
    'MCO_26': 'URBANIZACIONES',
    'MCO_27': 'VEHÍCULOS',
    'MCO_28': 'GENERADORES',
    'MCO_29': 'MOBILIARIO',
    'MCO_30': 'HERRAMIENTAS Y EQUIPOS',
    'MCO_31': 'PLATAFORMAS TIC MACROINFORMÁTICA',
    'MCO_32': 'PLATAFORMAS TIC SEGÚN SUBPROCESO',
    'MCO_33': 'HW MACROINFORMÁTICA',
    'MCO_34': 'HW MICROINFORMÁTICA',
    'MCO_35': 'HW COMUNICACIONES Y TELECONTROL',
    'MCO_36': 'SW-LICENCIAS',
    'MCO_37': 'SW-DESARROLLOS A MEDIDA',
    'MCO_38': 'HERRAMIENTAS TIC ASIGNADAS A PERSONAS',
    'MCO_39': 'SERVIDUMBRES',
    'MCO_40': 'DERECHOS DE AGUA',
    'MCO_41': 'OTROS INTANGIBLES',
    'MCO_42': 'SERVICIOS E INGRESOS NO REGULADOS',
    'MCO_43': 'SECTORES Y PUNTOS DE CONTROL DE PRESIÓN',
    'MCO_44': 'EMERGENCIAS',
    'MCO_45': 'TURBIEDAD EXTREMA',
    'MCO_46': 'CLIENTES PERSONAS JURÍDICAS',
    'MCO_47': 'EMPRESAS RELACIONADAS',
    'MCO_48': 'OFICINAS MÓVILES',
    'MCO_49': 'CLIENTES REGULADOS INCORPORADOS',
    'MCO_50': 'PROCESOS DE REUTILIZACIÓN DE SUBPRODUCTOS',
    'MCO_51': 'GESTIÓN DOCUMENTAL: CONTRATOS',
    'MCO_52': 'GESTIÓN DOCUMENTAL: LICITACIONES',
    'MCO_53': 'GESTIÓN DOCUMENTAL: PRESUPUESTOS',
    'MCO_54': 'GESTIÓN DOCUMENTAL: ÓRDENES DE COMPRA',
    'MCO_55': 'GESTIÓN DOCUMENTAL: FACTURAS',
    'MCO_56': 'GESTIÓN DOCUMENTAL: INFORMACIÓN COMPLEMENTARIA',
}

# =====================================================================
# ÍNDICE DE MAESTROS SISS (tablas MAE_1 a MAE_15)
# Fuente: articles-19850_MaestrosSISS2312.xlsx, hoja "IndiceMaestrosSISS".
# Solo el índice (nombre de cada tabla) -- el contenido específico de
# cada MAE (ej. MAE_1 Cuentas de Actividades, MAE_2 Cuentas de Recursos)
# ya está embebido parcialmente en otras partes del código (ej.
# PROCESO_A_ACTIVIDADES, RECURSO_NOMBRE) y se puede seguir ampliando en
# futuras funciones según se necesite.
# =====================================================================
MAE_INDICE = {'MAE_1': 'Cuentas de Actividades', 'MAE_2': 'Cuentas de Recursos', 'MAE_3': 'Cuentas de Servicios e Ingresos Regulados', 'MAE_4': 'Cuentas de Servicios No Regulados', 'MAE_5': 'Cuentas de Activos', 'MAE_6': 'Empresas', 'MAE_7': 'Cargos SISS', 'MAE_8': 'Localidades', 'MAE_9': 'Comunas', 'MAE_10': 'Grupos Tarifarios', 'MAE_11': 'Sistemas Tarifarios', 'MAE_12': 'Código Sector Decreto Tarifario', 'MAE_13': 'Clasificación de Activos NBI', 'MAE_14': 'Clasificación de Activos No Sanitarios', 'MAE_15': 'Código Partidas usadas en la SVI'}

# =====================================================================
# ÍNDICE DE TIPIFICACIONES / PARÁMETROS SISS (tablas PARAM_1 a PARAM_114)
# Fuente: articles-19850_MaestrosSISS2312.xlsx, hoja "ÍndiceTipificaciones".
# Cada PARAM_N es un campo tipificado (ej. PARAM_12 = "Tipo Cargo
# Eléctrico", el mismo campo usado en MEI_2). Índice de referencia para
# futuras funciones que necesiten decodificar estos campos.
# =====================================================================
TIPIFICACIONES_INDICE = {
    'PARAM_1': 'TIPO ACCESORIO',
    'PARAM_2': 'TIPO ACTIVIDAD DE CONTROL DIRECTO DE RILES',
    'PARAM_3': 'TIPO ACTIVIDAD DE MANTENCIÓN DE INFRAESTRUCTURA',
    'PARAM_4': 'TIPO ACTIVIDAD DE OPERACIÓN',
    'PARAM_5': 'TIPO ADQUISICIÓN',
    'PARAM_6': 'TIPO ARRIENDO EQUIPO INFORMÁTICO',
    'PARAM_7': 'TIPO ASESORIA Y ESTUDIO',
    'PARAM_8': 'TIPO ASIGNACIÓN',
    'PARAM_9': 'TIPO ATENCIÓN',
    'PARAM_10': 'TIPO AUTONOMÍA',
    'PARAM_11': 'TIPO CARGO DIRECTORIO',
    'PARAM_12': 'TIPO CARGO ELÉCTRICO',
    'PARAM_13': 'TIPO CATEGORÍA PROYECTO',
    'PARAM_14': 'TIPO CATEGORÍA PROYECTO NO SANITARIO',
    'PARAM_15': 'TIPO CAUSA INTERVENCIÓN ARRANQUE',
    'PARAM_16': 'TIPO CAUSA INTERVENCIÓN DE UD',
    'PARAM_17': 'TIPO CAUSA ROTURA',
    'PARAM_18': 'TIPO CIFRAS AFECTADAS POR EL CARGO',
    'PARAM_19': 'TIPO CLIENTE',
    'PARAM_20': 'TIPO COMBUSTIBLE O ENERGÍA',
    'PARAM_21': 'TIPO COMPLEJIDAD DE GESTIÓN',
    'PARAM_22': 'TIPO COMPONENTE DE COSTO REPARACIÓN',
    'PARAM_23': 'TIPO COMPONENTE INVERSIÓN TELECONTROL',
    'PARAM_24': 'TIPO COMPONENTE REMUNERACIONES',
    'PARAM_25': 'TIPO CONCEPTO COMPENSACIONES RRHH',
    'PARAM_26': 'TIPO CONOCIMIENTOS',
    'PARAM_27': 'TIPO CONSULTORÍA INVERSIÓN TIC',
    'PARAM_28': 'TIPO CONSUMO BÁSICO',
    'PARAM_29': 'TIPO CONTACTOS',
    'PARAM_30': 'TIPO CONTRATO',
    'PARAM_31': 'TIPO CONTRATO PERSONAS',
    'PARAM_32': 'TIPO CONTROL DE CALIDAD AP',
    'PARAM_33': 'TIPO CONTROL DE CALIDAD AS',
    'PARAM_34': 'TIPO CUENTA LIBRO MAYOR',
    'PARAM_35': 'TIPO DERECHOS SOBRE EL ACTIVO O BIEN',
    'PARAM_36': 'TIPO DOCUMENTO',
    'PARAM_37': 'TIPO DOCUMENTO LICITACIÓN',
    'PARAM_38': 'TIPO DOMINIO DE IDIOMA',
    'PARAM_39': 'TIPO ELEMENTO OBSTRUCTIVO',
    'PARAM_40': 'TIPO EQUIPO',
    'PARAM_41': 'TIPO ESPECIALIDAD',
    'PARAM_42': 'TIPO ESTADO PROYECTO',
    'PARAM_43': 'TIPO ESTUDIOS',
    'PARAM_44': 'TIPO EXPERIENCIA',
    'PARAM_45': 'TIPO FACTURA',
    'PARAM_46': 'TIPO FINANCIAMIENTO PROYECTO',
    'PARAM_47': 'TIPO FUNCIÓN',
    'PARAM_48': 'TIPO GRUPO COMPONENTE',
    'PARAM_49': 'TIPO HABILIDAD LIDERAZGO',
    'PARAM_50': 'TIPO HERRAMIENTA O EQUIPO',
    'PARAM_51': 'TIPO IMPRESORAS Y PLOTTERS',
    'PARAM_52': 'TIPO INMUEBLE',
    'PARAM_53': 'TIPO INSPECCIÓN TELEVISIVA',
    'PARAM_54': 'TIPO INSTANCIA CORTE',
    'PARAM_55': 'TIPO LLAMADA O CONTACTO',
    'PARAM_56': 'TIPO MANTENCIÓN GRIFOS',
    'PARAM_57': 'TIPO MANTENCIÓN RECINTOS',
    'PARAM_58': 'TIPO MATERIAL E INSUMO DE LABORATORIO',
    'PARAM_59': 'TIPO MATERIAL E INSUMOS DE OFICINA Y BODEGA',
    'PARAM_60': 'TIPO MATERIAL TUBERÍA',
    'PARAM_61': 'TIPO MOBILIARIO',
    'PARAM_62': 'TIPO MODALIDAD PAGO',
    'PARAM_63': 'TIPO MODALIDAD SERVICIO MANTENCIÓN',
    'PARAM_64': 'TIPO MONEDA',
    'PARAM_65': 'TIPO MONITOREO',
    'PARAM_66': 'TIPO MOTIVO EGRESO',
    'PARAM_67': 'TIPO MUESTREO AS',
    'PARAM_68': 'TIPO NIVEL DE TERCERIZACIÓN',
    'PARAM_69': 'TIPO OPERACIÓN DE EQUIPO',
    'PARAM_70': 'TIPO OPERACIÓN DE INFRAESTRUCTURA',
    'PARAM_71': 'TIPO ORDEN DE COMPRA',
    'PARAM_72': 'TIPO ORIGEN HORAS EXTRA',
    'PARAM_73': 'TIPO OTROS BENEFICIOS ADICIONALES',
    'PARAM_74': 'TIPO PERIODICIDAD CONTRATO SERVICIOS',
    'PARAM_75': 'TIPO PERSONAL',
    'PARAM_76': 'TIPO PERSONAL CONTACTADO',
    'PARAM_77': 'TIPO PÓLIZA',
    'PARAM_78': 'TIPO PONDERADOR SOLUCIÓN DE PROBLEMAS',
    'PARAM_79': 'TIPO PRODUCTO QUÍMICO',
    'PARAM_80': 'TIPO PROVEEDOR O CLIENTE',
    'PARAM_81': 'TIPO PROYECTO',
    'PARAM_82': 'TIPO PLANTA',
    'PARAM_83': 'TIPO RECAUDADOR',
    'PARAM_84': 'TIPO RED',
    'PARAM_85': 'TIPO REMUNERACIÓN DIRECTOR',
    'PARAM_86': 'TIPO REPARACIÓN ARRANQUE',
    'PARAM_87': 'TIPO RESIDUO',
    'PARAM_88': 'TIPO RESPALDO',
    'PARAM_89': 'TIPO RESPONSABILIDAD EMERGENCIA',
    'PARAM_90': 'TIPO RESPONSABILIDAD POR RESULTADOS',
    'PARAM_91': 'TIPO SERVICIO',
    'PARAM_92': 'TIPO SERVICIO DE DOCUMENTO',
    'PARAM_93': 'TIPO SERVICIO DE MANTENCIÓN DE REDES',
    'PARAM_94': 'TIPO SERVICIO INFORMÁTICO',
    'PARAM_95': 'TIPO SERVIDUMBRE',
    'PARAM_96': 'TIPO SISTEMA OPERATIVO',
    'PARAM_97': 'TIPO SOFTWARE',
    'PARAM_98': 'TIPO SOLUCIÓN DE PROBLEMAS',
    'PARAM_99': 'TIPO SUPERFICIE ROTURA',
    'PARAM_100': 'TIPO SUPERFICIE TERRENO',
    'PARAM_101': 'TIPO SW MACROINFORMÁTICA',
    'PARAM_102': 'TIPO TAMAÑO DE SUPERVISIÓN',
    'PARAM_103': 'TIPO TARIFA ELÉCTRICA',
    'PARAM_104': 'TIPO TECNOLOGÍA DE OPERACIÓN',
    'PARAM_105': 'TIPO TECNOLOGÍA DETECCIÓN DE FUGAS',
    'PARAM_106': 'TIPO UNIDAD ADQUISICIÓN TIC',
    'PARAM_107': 'TIPO UNIDAD CONCENTRACIÓN',
    'PARAM_108': 'TIPO UNIDADES DE CAUDAL',
    'PARAM_109': 'TIPO URBANIZACIÓN INMUEBLE',
    'PARAM_110': 'TIPO USO DE HERRAMIENTA O EQUIPO',
    'PARAM_111': 'TIPO USO EDIFICACIÓN',
    'PARAM_112': 'TIPO VEHÍCULO',
    'PARAM_113': 'TIPO VIAJE',
    'PARAM_114': 'TIPO ZONAL',
}


# =====================================================================
# CONTENIDO COMPLETO DE LAS TABLAS MAE (Maestro SISS) -- 13 de las 15 MAE.
# Fuente: articles-19850_MaestrosSISS2312.xlsx, hojas MAE_1 a MAE_14.
# MAE_7 (Cargos SISS, ~7.787 filas) y MAE_15 (Partidas SVI, ~14.134 filas)
# NO se embebieron por su tamaño -- si se necesitan, subir el archivo
# maestro original para consultarlas puntualmente.
# =====================================================================
MAE_DATOS = {
    'MAE_1': {
        "nombre": 'Cuentas de Actividades',
        "headers": ['CÓDIGO MACROPROCESO', 'MACROPROCESO', 'CÓDIGO PROCESO', 'PROCESO', 'CÓDIGO SUBPROCESO', 'SUBPROCESO', 'CÓDIGO ACTIVIDAD', 'ACTIVIDAD', 'PUESTA EN MARCHA'],
        "filas": [
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', '10101', 'Planificación y Control-Operación Infraestructura de Apoyo AP', '1010101', 'Planificación y Control-Operación Infraestructura de Apoyo AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10102, 'Supervisión de la Operación-Infraestructura de Apoyo AP', '1010201', 'Supervisión de la Operación-Infraestructura de Apoyo AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', '1010301', 'Operación Grupos Electrógenos AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010302, 'Operación Macromedidor AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010303, 'Operación Sistema Antigolpe de Ariete AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010304, 'Operación Estaciones Reductoras de Presión AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010305, 'Operación Centro de Control General - Sistema de Telemetria, Telecontrol y Automatización AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010306, 'Operación Centro de Control Local - Sistema de Telemetria, Telecontrol y Automatización AP', 0),
            (1, 'Suministro de Agua Potable', '101', 'Operación Infraestructura de Apoyo AP', 10103, 'Operación Infraestructura de Apoyo AP', 1010307, 'Operación Otra Infraestructura de Apoyo AP', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', '10201', 'Planificación y Control-Gestión de Recursos Hídricos', '1020101', 'Planificación y Control-Gestión de Recursos Hídricos', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', 10202, 'Supervisión-Gestión de Recursos Hídricos', '1020201', 'Supervisión-Gestión de Recursos Hídricos', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', 10203, 'Gestión de Recursos Hídricos', '1020301', 'Seguimiento de fuentes (condiciones hídricas, evolución de calidad de agua de la fuente, plan de macro medición de fuentes)', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', 10203, 'Gestión de Recursos Hídricos', 1020302, 'Definición y gestión técnica de recursos hídricos en base a pautas de ingeniería operacional', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', 10203, 'Gestión de Recursos Hídricos', 1020303, 'Realizar estudios hidrológicos e hidrogeológicos internos y externos', 0),
            (1, 'Suministro de Agua Potable', 102, 'Gestión de Recursos Hídricos', 10203, 'Gestión de Recursos Hídricos', 1020304, 'Compra de agua potable y cruda', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', '10301', 'Planificación y Control de la Operación Infraestructura de Capacidad AP', '1030101', 'Planificación y Control-Operación Captación de Agua', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10302, 'Supervisión de la Operación-Infraestructura de Capacidad AP', '1030201', 'Supervisión de la Operación-Captación de Agua', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10302, 'Supervisión de la Operación-Infraestructura de Capacidad AP', 1030202, 'Control y monitoreo ambiental', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10303, 'Operación Captación Subterránea', '1030301', 'Operación Norias', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10303, 'Operación Captación Subterránea', 1030302, 'Operación Drenes', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10303, 'Operación Captación Subterránea', 1030303, 'Operación Sondajes', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10303, 'Operación Captación Subterránea', 1030304, 'Operación Punteras', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10304, 'Operación Captación Superficial', '1030401', 'Operación Captación Superficial en río', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10304, 'Operación Captación Superficial', 1030402, 'Operación Captación Superficial en Canal', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10304, 'Operación Captación Superficial', 1030403, 'Operación Captación Embalses y tranques', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10304, 'Operación Captación Superficial', 1030404, 'Operación Captación en mar', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10305, 'Control de calidad AP Captación', '1030501', 'Supervisión de contrato de laboratorio externo captación', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10305, 'Control de calidad AP Captación', 1030502, 'Muestreo AP captación normativo', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10305, 'Control de calidad AP Captación', 1030503, 'Muestreo AP captación control de procesos', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10305, 'Control de calidad AP Captación', 1030504, 'Análisis de laboratorio AP captación normativo', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10305, 'Control de calidad AP Captación', 1030505, 'Análisis de laboratorio AP captación control de procesos', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10306, 'Operación plantas de tratamiento AP', '1030601', 'Operación plantas de tratamiento AP Compacta', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10306, 'Operación plantas de tratamiento AP', 1030602, 'Operación plantas de tratamiento AP Convencional', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10306, 'Operación plantas de tratamiento AP', 1030603, 'Operación plantas de tratamiento AP Osmosis Inversa (agua de mar)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10306, 'Operación plantas de tratamiento AP', 1030604, 'Operación plantas de tratamiento AP Osmosis Inversa (aguas salobres)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10306, 'Operación plantas de tratamiento AP', 1030605, 'Operación plantas de tratamiento AP Abatimiento de Arsénico', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10307, 'Control de calidad AP producción', '1030701', 'Supervisión de contrato de laboratorio externo producción', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10307, 'Control de calidad AP producción', 1030702, 'Muestreo AP producción normativo', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10307, 'Control de calidad AP producción', 1030703, 'Muestreo AP producción control de procesos', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10307, 'Control de calidad AP producción', 1030704, 'Análisis de laboratorio AP producción normativo', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10307, 'Control de calidad AP producción', 1030705, 'Análisis de laboratorio AP producción control de procesos', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10308, 'Operación Sistema Cloración ', '1030801', 'Operación Sistema Cloración (gas cloro en cilindro)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10308, 'Operación Sistema Cloración ', 1030802, 'Operación Sistema Cloración (gas cloro en contenedores)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10308, 'Operación Sistema Cloración ', 1030803, 'Operación Sistema Cloración (hipoclorito)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10309, 'Operación Sistema Fluoración', '1030901', 'Operación Sistema Fluoración (fluoruro de sodio)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10309, 'Operación Sistema Fluoración', 1030902, 'Operación Sistema Fluoración (silicofluoruro de sodio)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10309, 'Operación Sistema Fluoración', 1030903, 'Operación Sistema Fluoración (ácido fluorsilícico)', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10310, 'Operación Estanques de producción', '1031001', 'Operación Estanques de producción Semienterrados', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10310, 'Operación Estanques de producción', 1031002, 'Operación Estanques de producción Elevados', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10311, 'Operación Plantas Elevadoras AP', '1031101', 'Operación Plantas Elevadoras AP', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10312, 'Gestión de residuos sólidos', '1031201', 'Transporte de residuos sólidos', 0),
            (1, 'Suministro de Agua Potable', 103, 'Operación Infraestructura de Capacidad AP', 10312, 'Gestión de residuos sólidos', 1031202, 'Disposición de residuos sólidos', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', '10401', 'Operación Conducciones en Presión AP', '1040101', 'Operación Conducciones en Presión AP', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10402, 'Operación Acueductos AP', '1040201', 'Operación Acueductos AP', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', '1040301', 'Planificación y Control-Operación Distribución de AP', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040302, 'Control operativo de oferta y demanda (proyección de demanda y control de caudales)', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040303, 'Monitoreo en tiempo real (transmisión y seguimiento de datos con sensores de presión y sectorización). Incluye seguimiento a nivel de KPI.', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040304, 'Propuesta de mejoramiento de red (control de presiones, caudales y modificación y reposición de redes y equipos para su operación)', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040305, 'Investigar metodologías para mejorar la gestión de redes', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040306, 'Entrega información para planes preventivos y de mejoramiento AP (de gestión de presión, detección de fugas)', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040307, 'Revisión de proyectos de modificación de redes AP (incluye impacto de  obras durante y despues de su desarrollo)', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040308, 'Acuartelamiento (para disminuir los clientes afectados se instalan válvulas para "acuartelar" los afectados al mínimo) ', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10403, 'Planificación y Control-Operación Redes y Conducciones AP', 1040309, 'Programación, coordinación y ejecución de intervenciones', 0),
            (1, 'Suministro de Agua Potable', 104, 'Operación Redes y Conducciones AP', 10404, 'Supervisión de la Operación-Redes y Conducciones AP', '1040401', 'Supervisión de la Operación-Redes y Conducciones AP', 0),
            (1, 'Suministro de Agua Potable', 105, 'Operación Infraestructura de Capacidad AP', '10501', 'Operación Estanques de distribución', '1050101', 'Operación Estanques de distribución Semienterrados', 0),
            (1, 'Suministro de Agua Potable', 105, 'Operación Infraestructura de Capacidad AP', '10501', 'Operación Estanques de distribución', 1050102, 'Operación Estanques de distribución Elevados', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', '10601', 'Operación Red de Distribución', '1060101', 'Operación Red de Distribución', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', '10601', 'Operación Red de Distribución', 1060102, 'Control de presiones en la Red Distribución', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10602, 'Control de calidad AP distribución', '1060201', 'Supervisión de contrato de laboratorio externo distribución', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10602, 'Control de calidad AP distribución', 1060202, 'Muestreo AP distribución normativo', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10602, 'Control de calidad AP distribución', 1060203, 'Muestreo AP distribución control de procesos', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10602, 'Control de calidad AP distribución', 1060204, 'Análisis de laboratorio AP distribución normativo', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10602, 'Control de calidad AP distribución', 1060205, 'Análisis de laboratorio AP distribución control de procesos', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10603, 'Gestión de residuos sólidos', '1060301', 'Transporte de residuos sólidos', 0),
            (1, 'Suministro de Agua Potable', 106, 'Operación Redes y Conducciones AP', 10603, 'Gestión de residuos sólidos', 1060302, 'Disposición de residuos sólidos', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', '20101', 'Planificación y Control-Operación Infraestructura de Apoyo AS', '2010101', 'Planificación y Control-Operación Infraestructura de Apoyo AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20102, 'Supervisión de la Operación-Infraestructura de Apoyo AS', '2010201', 'Supervisión de la Operación-Infraestructura de Apoyo AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', '2010301', 'Operación Grupos Electrógenos AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010302, 'Operación Macromedidor AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010303, 'Operación Sistema Antigolpe de Ariete AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010304, 'Operación Estaciones Reductoras de Presión AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010305, 'Operación Centro de Control General - Sistema de Telemetria, Telecontrol y Automatización AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010306, 'Operación Centro de Control Local - Sistema de Telemetria, Telecontrol y Automatización AS', 0),
            (2, 'Saneamiento de Aguas Servidas', '201', 'Operación Infraestructura de Apoyo AS', 20103, 'Operación Infraestructura de Apoyo AS', 2010307, 'Operación Otra Infraestructura de Apoyo AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 202, 'Operación Infraestructura de Capacidad AS', '20201', 'Planificación y Control-Operación Infraestructura de Capacidad AS', '2020101', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 202, 'Operación Infraestructura de Capacidad AS', '20201', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 2020102, 'Entrega información para planes preventivos y de mejoramiento AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 202, 'Operación Infraestructura de Capacidad AS', '20201', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 2020103, 'Programación, coordinación y ejecución de intervenciones', 0),
            (2, 'Saneamiento de Aguas Servidas', 202, 'Operación Infraestructura de Capacidad AS', 20202, 'Supervisión de la Operación-Infraestructura de Capacidad AS', '2020201', 'Supervisión de la Operación-Infraestructura de Capacidad AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 202, 'Operación Infraestructura de Capacidad AS', 20203, 'Operación Plantas Elevadoras AS', '2020301', 'Operación Plantas Elevadoras AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 203, 'Operación Redes y Conducciones AS', '20301', 'Operación Conducciones en Presión AS', '2030101', 'Operación Conducciones en Presión AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 203, 'Operación Redes y Conducciones AS', 20302, 'Operación Acueductos AS', '2030201', 'Operación Acueductos AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 203, 'Operación Redes y Conducciones AS', 20303, 'Operación Red Recolección', '2030301', 'Operación Red Recolección', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', '20401', 'Planificación y Control-Operación Infraestructura de Capacidad AS', '2040101', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', '20401', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 2040102, 'Definición y seguimiento técnico (visitas técnicas a plantas, definición de parámetros de control de procesos y seguimiento de parámetros para mejorar procesos, cumplir normas, ISO, procedimientos)', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', '20401', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 2040103, 'Revisión de Proyectos de RILES (analizar la capacidad de planta para ver factibilidad, disponibilidad para recepcionar en las plantas de camiones limpiafosas)', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', '20401', 'Planificación y Control-Operación Infraestructura de Capacidad AS', 2040104, 'Entrega información para planes preventivos y de mejoramiento AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20402, 'Planificación y Control-Operación Sistemas de Tratamiento AS', '2040201', 'Planificación y Control-Operación Sistemas de Tratamiento AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20403, 'Supervisión de la Operación-Sistemas de Tratamiento AS', '2040301', 'Supervisión de la Operación-Sistemas de Tratamiento AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20404, 'Operación plantas de tratamiento AS', '2040401', 'Operación Sistema Tratamiento Primario AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20404, 'Operación plantas de tratamiento AS', 2040402, 'Operación Plantas Tratamiento AS Lagunas Facultativas', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20404, 'Operación plantas de tratamiento AS', 2040403, 'Operación Plantas Tratamiento AS Lagunas Aireadas', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20404, 'Operación plantas de tratamiento AS', 2040404, 'Operación Plantas Tratamiento AS Lodos Activados', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20404, 'Operación plantas de tratamiento AS', 2040405, 'Operación Plantas Tratamiento AS Lombrifiltros', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20405, 'Operación Sistema Pretratamiento AS', '2040501', 'Operación Sistema Pretratamiento AS', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20406, 'Operación emisarios submarinos', '2040601', 'Operación emisarios submarinos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20407, 'Control de calidad AS', '2040701', 'Supervisión de contrato de laboratorio externo disposición', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20407, 'Control de calidad AS', 2040702, 'Muestreo AS normativo', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20407, 'Control de calidad AS', 2040703, 'Muestreo AS control de procesos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20407, 'Control de calidad AS', 2040704, 'Análisis de laboratorio AS normativo', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20407, 'Control de calidad AS', 2040705, 'Análisis de laboratorio AS control de procesos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20408, 'Gestión de residuos sólidos', '2040801', 'Transporte de residuos sólidos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20408, 'Gestión de residuos sólidos', 2040802, 'Disposición de residuos sólidos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20409, 'Gestión de lodos', '2040901', 'Transporte de lodos a vertederos propios', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20409, 'Gestión de lodos', 2040902, 'Transporte de lodos a vertederos de terceros', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20409, 'Gestión de lodos', 2040903, 'Transporte de lodos a predios agrícolas', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20409, 'Gestión de lodos', 2040904, 'Disposición de lodos', 0),
            (2, 'Saneamiento de Aguas Servidas', 204, 'Operación Infraestructura de Capacidad AS', 20410, 'Control y monitoreo ambiental', '2041001', 'Control y monitoreo ambiental', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', '30101', 'Planificación y Control del Mantenimiento Preventivo-Infraestructura de Apoyo', '3010101', 'Planificación y Control del Mantenimiento Preventivo-Otra Infraestructura de Apoyo', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30102, 'Supervisión del Mantenimiento Preventivo-Infraestructura de Apoyo', '3010201', 'Supervisión del Mantenimiento Preventivo-Otra Infraestructura de Apoyo', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30103, 'Mantención Preventiva de vehículos', '3010301', 'Mantención Preventiva de vehículos', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30104, 'Mantención Preventiva de equipos especiales', '3010401', 'Mantención Preventiva de equipos especiales', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30105, 'Mantención Preventiva de recintos', '3010501', 'Mantención Preventiva de recintos', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30106, 'Mantención Preventiva de servidumbres', '3010601', 'Mantención Preventiva de servidumbres', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30107, 'Mantención Preventiva de Grupos Electrógenos', '3010701', 'Mantención Preventiva de Grupos Electrógenos', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30108, 'Mantención Preventiva de líneas eléctricas y subestaciones', '3010801', 'Mantención Preventiva de líneas eléctricas y subestaciones', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30109, 'Mantención Preventiva de Macromedidores', '3010901', 'Mantención Preventiva de Macromedidores', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30110, 'Mantención Preventiva de Sistema Antigolpe de Ariete', '3011001', 'Mantención Preventiva de Sistema Antigolpe de Ariete', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30111, 'Mantención Preventiva de Estaciones Reductoras de Presión', '3011101', 'Mantención Preventiva de Estaciones Reductoras de Presión', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30112, 'Mantención Preventiva de Sistema de Telemetria', '3011201', 'Mantención Preventiva de Sistema de Telemetria (periféricos, sensores, etc.)', 0),
            (3, 'Mantenimiento', '301', 'Mantención Preventiva de Infraestructura de Apoyo', 30113, 'Mantención Preventiva de Otra Infraestructura de Apoyo', '3011301', 'Mantención Preventiva de Otra Infraestructura de Apoyo', 0),
            (3, 'Mantenimiento', 302, 'Mantención Correctiva de Infraestructura de Apoyo', '30201', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Apoyo', '3020101', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Apoyo', 0),
            (3, 'Mantenimiento', 302, 'Mantención Correctiva de Infraestructura de Apoyo', 30202, 'Supervisión del Mantenimiento Correctivo-Infraestructura de Apoyo', '3020201', 'Supervisión del Mantenimiento Correctivo-Infraestructura de Apoyo', 0),
            (3, 'Mantenimiento', 302, 'Mantención Correctiva de Infraestructura de Apoyo', 30203, 'Mantención Correctiva de Obras Civiles', '3020301', 'Mantención Correctiva de Obras Civiles', 0),
            (3, 'Mantenimiento', 302, 'Mantención Correctiva de Infraestructura de Apoyo', 30204, 'Mantención Correctiva de Equipos', '3020401', 'Mantención Correctiva de Equipos', 0),
            (3, 'Mantenimiento', 302, 'Mantención Correctiva de Infraestructura de Apoyo', 30205, 'Mantención Correctiva de Instalaciones Eléctricas', '3020501', 'Mantención Correctiva de Instalaciones Eléctricas', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', '30301', 'Planificación y Control del Mantenimiento Preventivo-Infraestructura de Capacidad AP', '3030101', 'Planificación y Control del Mantenimiento Preventivo-Infraestructura de Capacidad AP', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30302, 'Supervisión del Mantenimiento Preventivo-Infraestructura de Capacidad AP', '3030201', 'Supervisión del Mantenimiento Preventivo-Infraestructura de Capacidad AP', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30303, 'Mantención Preventiva Captaciones Subterráneas', '3030301', 'Mantención Preventiva Norias', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30303, 'Mantención Preventiva Captaciones Subterráneas', 3030302, 'Mantención Preventiva Drenes', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30303, 'Mantención Preventiva Captaciones Subterráneas', 3030303, 'Mantención Preventiva Sondajes', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30303, 'Mantención Preventiva Captaciones Subterráneas', 3030304, 'Mantención Preventiva Punteras', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30304, 'Mantención Preventiva Captaciones Superficiales', '3030401', 'Mantención Preventiva Captación Superficial en río', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30304, 'Mantención Preventiva Captaciones Superficiales', 3030402, 'Mantención Preventiva Captación Superficial en Canal', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30304, 'Mantención Preventiva Captaciones Superficiales', 3030403, 'Mantención Preventiva Captación Embalses y tranques', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30304, 'Mantención Preventiva Captaciones Superficiales', 3030404, 'Mantención Preventiva Captación en mar', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30305, 'Mantención Preventiva plantas de tratamiento AP', '3030501', 'Mantención preventiva plantas de tratamiento AP Compacta', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30305, 'Mantención Preventiva plantas de tratamiento AP', 3030502, 'Mantención preventiva plantas de tratamiento AP Convencional', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30305, 'Mantención Preventiva plantas de tratamiento AP', 3030503, 'Mantención preventiva plantas de tratamiento AP Osmosis Inversa (agua de mar)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30305, 'Mantención Preventiva plantas de tratamiento AP', 3030504, 'Mantención preventiva plantas de tratamiento AP Osmosis Inversa (aguas salobres)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30305, 'Mantención Preventiva plantas de tratamiento AP', 3030505, 'Mantención preventiva plantas de tratamiento AP Abatimiento de Arsénico', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30306, 'Mantención Preventiva Sistema Cloración ', '3030601', 'Mantención Preventiva Sistema Cloración (gas cloro en cilindro)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30306, 'Mantención Preventiva Sistema Cloración ', 3030602, 'Mantención Preventiva Sistema Cloración (gas cloro en contenedores)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30306, 'Mantención Preventiva Sistema Cloración ', 3030603, 'Mantención Preventiva Sistema Cloración (hipoclorito)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30307, 'Mantención Preventiva Sistema Fluoración', '3030701', 'Mantención Preventiva Sistema Fluoración (fluoruro de sodio)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30307, 'Mantención Preventiva Sistema Fluoración', 3030702, 'Mantención Preventiva Sistema Fluoración (silicofluoruro de sodio)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30307, 'Mantención Preventiva Sistema Fluoración', 3030703, 'Mantención Preventiva Sistema Fluoración (ácido fluorsilícico)', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30308, 'Mantención Preventiva Estanques de producción', '3030801', 'Mantención Preventiva Estanques de producción Semienterrados', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30308, 'Mantención Preventiva Estanques de producción', 3030802, 'Mantención Preventiva Estanques de producción Elevados', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30309, 'Mantención Preventiva Plantas Elevadoras AP', '3030901', 'Mantención Preventiva Plantas Elevadoras AP', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30310, 'Mantención Preventiva Estanques de distribución', '3031001', 'Mantención Preventiva Estanques de distribución Semienterrados', 0),
            (3, 'Mantenimiento', 303, 'Mantención Preventiva de Infraestructura de Capacidad AP', 30310, 'Mantención Preventiva Estanques de distribución', 3031002, 'Mantención Preventiva Estanques de distribución Elevados', 0),
            (3, 'Mantenimiento', 304, 'Mantención Correctiva de Infraestructura de Capacidad AP', '30401', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Capacidad AP', '3040101', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Capacidad AP', 0),
            (3, 'Mantenimiento', 304, 'Mantención Correctiva de Infraestructura de Capacidad AP', 30402, 'Supervisión del Mantenimiento Correctivo-Infraestructura de Capacidad AP', '3040201', 'Supervisión del Mantenimiento Correctivo-Infraestructura de Capacidad AP', 0),
            (3, 'Mantenimiento', 304, 'Mantención Correctiva de Infraestructura de Capacidad AP', 30403, 'Mantención Correctiva de Obras Civiles', '3040301', 'Mantención Correctiva de Obras Civiles', 0),
            (3, 'Mantenimiento', 304, 'Mantención Correctiva de Infraestructura de Capacidad AP', 30404, 'Mantención Correctiva de Equipos', '3040401', 'Mantención Correctiva de Equipos', 0),
            (3, 'Mantenimiento', 304, 'Mantención Correctiva de Infraestructura de Capacidad AP', 30405, 'Mantención Correctiva de Instalaciones Eléctricas', '3040501', 'Mantención Correctiva de Instalaciones Eléctricas', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', '30501', 'Planificación y Control del Mantenimiento Preventivo-Redes y Conducciones AP', '3050101', 'Planificación y Control del Mantenimiento Preventivo-Redes y Conducciones AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30502, 'Supervisión del Mantenimiento Preventivo-Redes y Conducciones AP', '3050201', 'Supervisión del Mantenimiento Preventivo-Redes y Conducciones AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30503, 'Mantención Preventiva Conducciones en Presión AP', '3050301', 'Limpieza preventiva de cámaras Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30503, 'Mantención Preventiva Conducciones en Presión AP', 3050302, 'Detección de fugas en Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30503, 'Mantención Preventiva Conducciones en Presión AP', 3050303, 'Control de presiones en Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30504, 'Mantención Preventiva Acueductos AP', '3050401', 'Limpieza de cámaras Acueductos AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30504, 'Mantención Preventiva Acueductos AP', 3050402, 'Detección de fugas Acueductos AP', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30505, 'Mantención Preventiva Red Distribución', '3050501', 'Limpieza de cámaras Red Distribución', 0),
            (3, 'Mantenimiento', 305, 'Mantención Preventiva Redes y Conducciones AP', 30505, 'Mantención Preventiva Red Distribución', 3050502, 'Detección de fugas en Red Distribución', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', '30601', 'Planificación y Control del Mantenimiento Correctivo-Redes y Conducciones AP', '3060101', 'Planificación y Control del Mantenimiento Correctivo-Redes y Conducciones AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30602, 'Supervisión del Mantenimiento Correctivo-Redes y Conducciones AP', '3060201', 'Supervisión del Mantenimiento Correctivo-Redes y Conducciones AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30603, 'Mantención Correctiva Conducciones en Presión AP', '3060301', 'Reparación de cámaras Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30603, 'Mantención Correctiva Conducciones en Presión AP', 3060302, 'Renovación de cámaras Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30603, 'Mantención Correctiva Conducciones en Presión AP', 3060303, 'Diagnóstico de fallas en Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30603, 'Mantención Correctiva Conducciones en Presión AP', 3060304, 'Reparación de roturas en Conducciones en Presión AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30604, 'Mantención Correctiva Acueductos AP', '3060401', 'Reparación de cámaras Acueductos AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30604, 'Mantención Correctiva Acueductos AP', 3060402, 'Renovación de cámaras Acueductos AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30604, 'Mantención Correctiva Acueductos AP', 3060403, 'Diagnóstico de fallas Acueductos AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30604, 'Mantención Correctiva Acueductos AP', 3060404, 'Reparación de roturas Acueductos AP', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30605, 'Mantención Correctiva Red Distribución', '3060501', 'Reparación de cámaras Red Distribución', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30605, 'Mantención Correctiva Red Distribución', 3060502, 'Renovación de cámaras Red Distribución', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30605, 'Mantención Correctiva Red Distribución', 3060503, 'Diagnóstico de fallas en Red distribución', 0),
            (3, 'Mantenimiento', 306, 'Mantención Correctiva Redes y Conducciones AP', 30605, 'Mantención Correctiva Red Distribución', 3060504, 'Reparación de roturas en Red Distribución', 0),
            (3, 'Mantenimiento', 307, 'Mantención Preventiva Conexiones AP', '30701', 'Planificación y Control del Mantenimiento Preventivo-Conexiones AP', '3070101', 'Planificación y Control del Mantenimiento Preventivo-Conexiones AP', 0),
            (3, 'Mantenimiento', 307, 'Mantención Preventiva Conexiones AP', 30702, 'Supervisión del Mantenimiento Preventivo-Conexiones AP', '3070201', 'Supervisión del Mantenimiento Preventivo-Conexiones AP', 0),
            (3, 'Mantenimiento', 307, 'Mantención Preventiva Conexiones AP', 30703, 'Control de medidores (Ord. SISS N°688/2004)', '3070301', 'Verificación y control de medidores (Ord. SISS N°688/2004)', 0),
            (3, 'Mantenimiento', 307, 'Mantención Preventiva Conexiones AP', 30703, 'Control de medidores (Ord. SISS N°688/2004)', 3070302, 'Recambio de medidores (Ord. SISS N°688/2004)', 0),
            (3, 'Mantenimiento', 308, 'Mantención Correctiva Conexiones AP', '30801', 'Mantención Correctiva de arranques', '3080101', 'Reparación de arranques', 0),
            (3, 'Mantenimiento', 308, 'Mantención Correctiva Conexiones AP', '30801', 'Mantención Correctiva de arranques', 3080102, 'Renovación de arranques', 0),
            (3, 'Mantenimiento', 308, 'Mantención Correctiva Conexiones AP', '30801', 'Mantención Correctiva de arranques', 3080103, 'Diagnóstico de fallas en arranques', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', '30901', 'Planificación y Control del Mantenimiento Preventivo-Redes y Conducciones AS', '3090101', 'Planificación y Control del Mantenimiento Preventivo-Redes y Conducciones AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30902, 'Supervisión del Mantenimiento Preventivo-Redes y Conducciones AS', '3090201', 'Supervisión del Mantenimiento Preventivo-Redes y Conducciones AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30903, 'Mantención Preventiva Conducciones en Presión AS', '3090301', 'Inspección televisiva de Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30903, 'Mantención Preventiva Conducciones en Presión AS', 3090302, 'Limpieza de Conducciones en Presión AS (preventivo con camiones jet)', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30903, 'Mantención Preventiva Conducciones en Presión AS', 3090303, 'Limpieza de cámaras Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30904, 'Mantención Preventiva Acueductos AS', '3090401', 'Inspección televisiva Acueductos AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30904, 'Mantención Preventiva Acueductos AS', 3090402, 'Limpieza Acueductos AS (preventivo con camiones jet)', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30904, 'Mantención Preventiva Acueductos AS', 3090403, 'Limpieza de cámaras Acueductos AS', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30905, 'Mantención Preventiva Red Recolección', '3090501', 'Inspección televisiva de Red Recolección', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30905, 'Mantención Preventiva Red Recolección', 3090502, 'Limpieza de Red Recolección (preventivo con camiones jet)', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30905, 'Mantención Preventiva Red Recolección', 3090503, 'Diagnóstico de obstrucciones y fallas en Red Recolección', 0),
            (3, 'Mantenimiento', 309, 'Mantención Preventiva Redes y Conducciones AS', 30905, 'Mantención Preventiva Red Recolección', 3090504, 'Limpieza de cámaras Red Recolección', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', '31001', 'Planificación y Control del Mantenimiento Correctivo-Redes y Conducciones AS', '3100101', 'Planificación y Control del Mantenimiento Correctivo-Redes y Conducciones AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31002, 'Supervisión del Mantenimiento Correctivo-Redes y Conducciones AS', '3100201', 'Supervisión del Mantenimiento Correctivo-Redes y Conducciones AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31003, 'Mantención Correctiva Conducciones en Presión AS', '3100301', 'Desobstrucción de Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31003, 'Mantención Correctiva Conducciones en Presión AS', 3100302, 'Reparación de Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31003, 'Mantención Correctiva Conducciones en Presión AS', 3100303, 'Diagnóstico de obstrucciones y fallas en Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31003, 'Mantención Correctiva Conducciones en Presión AS', 3100304, 'Reparación de cámaras Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31003, 'Mantención Correctiva Conducciones en Presión AS', 3100305, 'Renovación de cámaras Conducciones en Presión AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31004, 'Mantención Correctiva Acueductos AS', '3100401', 'Desobstrucción Acueductos AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31004, 'Mantención Correctiva Acueductos AS', 3100402, 'Reparación Acueductos AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31004, 'Mantención Correctiva Acueductos AS', 3100403, 'Diagnóstico de obstrucciones y fallas Acueductos AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31004, 'Mantención Correctiva Acueductos AS', 3100404, 'Reparación de cámaras Acueductos AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31004, 'Mantención Correctiva Acueductos AS', 3100405, 'Renovación de cámaras Acueductos AS', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31005, 'Mantención Correctiva Red Recolección', '3100501', 'Desobstrucción de Red Recolección', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31005, 'Mantención Correctiva Red Recolección', 3100502, 'Reparación de Red Recolección', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31005, 'Mantención Correctiva Red Recolección', 3100503, 'Reparación de cámaras Red Recolección', 0),
            (3, 'Mantenimiento', 310, 'Mantención Correctiva Redes y Conducciones AS', 31005, 'Mantención Correctiva Red Recolección', 3100504, 'Renovación de cámaras Red Recolección', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', '31101', 'Planificación y Control del Mantenimiento Correctivo-Conexiones AS', '3110101', 'Planificación y Control del Mantenimiento Correctivo-Conexiones AS', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', 31102, 'Supervisión del Mantenimiento Correctivo-Conexiones AS', '3110201', 'Supervisión del Mantenimiento Correctivo-Conexiones AS', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', 31103, 'Mantención Correctiva de uniones domiciliarias', '3110301', 'Desobstrucción de uniones', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', 31103, 'Mantención Correctiva de uniones domiciliarias', 3110302, 'Reparación de uniones', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', 31103, 'Mantención Correctiva de uniones domiciliarias', 3110303, 'Renovación de uniones', 0),
            (3, 'Mantenimiento', 311, 'Mantención Correctiva Conexiones AS', 31103, 'Mantención Correctiva de uniones domiciliarias', 3110304, 'Diagnóstico de obstrucciones y fallas en uniones', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', '31201', 'Planificación y Control del Mantenimiento Preventivo-Infraestructura de Capacidad AS', '3120101', 'Planificación y Control del Mantenimiento Preventivo-Infraestructura de Capacidad AS', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31202, 'Supervisión del Mantenimiento Preventivo-Infraestructura de Capacidad AS', '3120201', 'Supervisión del Mantenimiento Preventivo-Infraestructura de Capacidad AS', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31203, 'Mantención Preventiva Plantas Elevadoras AS', '3120301', 'Mantención Plantas Elevadoras AS', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31204, 'Mantención Preventiva plantas de tratamiento AS', '3120401', 'Mantención Sistema Tratamiento Primario AS', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31204, 'Mantención Preventiva plantas de tratamiento AS', 3120402, 'Mantención Plantas Tratamiento AS Lagunas Facultativas', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31204, 'Mantención Preventiva plantas de tratamiento AS', 3120403, 'Mantención Plantas Tratamiento AS Lagunas Aireadas', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31204, 'Mantención Preventiva plantas de tratamiento AS', 3120404, 'Mantención Plantas Tratamiento AS Lodos Activados', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31204, 'Mantención Preventiva plantas de tratamiento AS', 3120405, 'Mantención Plantas Tratamiento AS Lombrifiltros', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31205, 'Mantención Preventiva Sistema Pretratamiento AS', '3120501', 'Mantención Sistema Pretratamiento AS', 0),
            (3, 'Mantenimiento', 312, 'Mantención Preventiva de Infraestructura de Capacidad AS', 31206, 'Mantención Preventiva de emisarios submarinos', '3120601', 'Mantención de emisarios submarinos', 0),
            (3, 'Mantenimiento', 313, 'Mantención Correctiva de Infraestructura de Capacidad AS', '31301', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Capacidad AS', '3130101', 'Planificación y Control del Mantenimiento Correctivo-Infraestructura de Capacidad AS', 0),
            (3, 'Mantenimiento', 313, 'Mantención Correctiva de Infraestructura de Capacidad AS', 31302, 'Supervisión del Mantenimiento Correctivo-Infraestructura de Capacidad AS', '3130201', 'Supervisión del Mantenimiento Correctivo-Infraestructura de Capacidad AS', 0),
            (3, 'Mantenimiento', 313, 'Mantención Correctiva de Infraestructura de Capacidad AS', 31303, 'Mantención Correctiva de Obras Civiles', '3130301', 'Mantención Correctiva de Obras Civiles', 0),
            (3, 'Mantenimiento', 313, 'Mantención Correctiva de Infraestructura de Capacidad AS', 31304, 'Mantención Correctiva de Equipos', '3130401', 'Mantención Correctiva de Equipos', 0),
            (3, 'Mantenimiento', 313, 'Mantención Correctiva de Infraestructura de Capacidad AS', 31305, 'Mantención Correctiva de Instalaciones Eléctricas', '3130501', 'Mantención Correctiva de Instalaciones Eléctricas', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', '40101', 'Medición', '4010101', 'Administración y Control-Medición', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', '40101', 'Medición', 4010102, 'Lectura manual', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', '40101', 'Medición', 4010103, 'Lectura remota', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', '40101', 'Medición', 4010104, 'Verificación de lecturas anormales', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', '40101', 'Medición', 4010105, 'Inspección de medidores', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', '4010201', 'Administración y Control-Facturación', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010202, 'Facturación', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010203, 'Cambios de Tarifas', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010204, 'Refacturación', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010205, 'Reparto de Boletas y Otros Documentos', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010206, 'Envío de Boletas por courier', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40102, 'Facturación', 4010207, 'Facturación Corporativa', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40103, 'Recaudación', '4010301', 'Administración y Control-Recaudación', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40103, 'Recaudación', 4010302, 'Recaudación propia presencial', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40103, 'Recaudación', 4010303, 'Recaudación externa presencial', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40103, 'Recaudación', 4010304, 'Recaudación externa digital', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', '4010401', 'Administración y Control-Cobranza', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010402, 'Administración Subsidios', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010403, 'Notificación Suspensión Suministro', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010404, 'Cobranza interna ', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010405, 'Cobranza prejudicial', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010406, 'Cobranza judicial', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010407, 'Castigo de deuda clientes', 0),
            (4, 'Atención Integral de Clientes', '401', 'Ciclo de Recaudación', 40104, 'Cobranza', 4010408, 'Informar internamente reposición de suministro', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', '40201', 'Estudios Comerciales', '4020101', 'Administración y Control-Estudios Comerciales', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', '40201', 'Estudios Comerciales', 4020102, 'Estudios Comerciales', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', '40201', 'Estudios Comerciales', 4020103, 'Definir propuesta de valor por segmento', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40202, 'Difusión Comercial', '4020201', 'Administración y Control-Difusión Comercial', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40202, 'Difusión Comercial', 4020202, 'Difusión Comercial (masivas el uno esta por canal)', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40202, 'Difusión Comercial', 4020203, 'Publicación de Tarifas', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', '4020301', 'Definición de estrategias y políticas de atención al cliente', 1),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020302, 'Administración y Control-Atención de clientes', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020303, 'Supervisión de la atención presencial en oficinas comerciales', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020304, 'Derivación de atenciones de caracter operacional', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020305, 'Bajada y Control de Políticas de Atención a clientes por cualquier canal', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020306, 'Atención presencial ', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020307, 'Sucursal web', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020308, 'Atención telefónica inbound ', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020309, 'Atención virtual por Internet - mail', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020310, 'Atención por correspondencia', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020311, 'Atención telefónica outbound', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020312, 'Atención vía chat bot (inteligencia artificial)', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020313, 'Atenciones de autoconsulta', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020314, 'App mobile ', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020315, 'Notificaciones push', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020316, 'Gestión de aviso a aplicaciones de tráfico', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020317, 'Generación y actualización de BD única de clientes', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020318, 'Campañas de contactibilidad', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020319, 'Atención SMS numero aleatorio', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020320, 'Atención SMS número registrado', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020321, 'Comunicaciones segmentadas en redes sociales', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020322, 'Atención a clientes en redes sociales ', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020323, 'Atención SERNAC y SISS', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020324, 'Atención a consultas de clientes y potenciales clientes en el proceso de incorporación', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020325, 'Atención de Clientes - consulta planos (con Archivo Técnico)', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020326, 'Atención de Clientes - venta planos (con Archivo Técnico)', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40203, 'Atención de Clientes', 4020327, 'Mantención del archivo técnico', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40204, 'Inspección Comercial Terreno', '4020401', 'Administración y Control-Inspección Comercial Terreno', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40204, 'Inspección Comercial Terreno', 4020402, 'Inspección Comercial Terreno', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40204, 'Inspección Comercial Terreno', 4020403, 'Detección  de Ilícitos y Conexiones Fraudulentas', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40205, 'Atención de Requerimientos', '4020501', 'Administración y Control-Atención de Requerimientos', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40205, 'Atención de Requerimientos', 4020502, 'Requerimientos Técnicos Complejos', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40205, 'Atención de Requerimientos', 4020503, 'Requerimientos Técnicos Simples', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40206, 'Gestión de Archivos Comerciales', '4020601', 'Administración y Control-Gestión de Archivos Comerciales', 0),
            (4, 'Atención Integral de Clientes', 402, 'Gestión Comercial', 40206, 'Gestión de Archivos Comerciales', 4020602, 'Mantención de información comercial de clientes', 0),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', '4030101', 'Administración y Control-Solicitud y obtención del certificado de factibilidad', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030102, 'Evaluación de casos de nuevos negocios regulados', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030103, 'Verificar antecedentes del propietario, proyectista y del o los inmuebles', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030104, 'Verificar que la localización del terreno de la urbanización o proyecto domiciliario se encuentre dentro del territorio operacional de la empresa', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030105, 'Verificar y comparar las capacidades hidráulicas de la zona de emplazamiento para determinar si se requieren obras adicionales', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030106, 'Definir y avisar a responsable las necesidades de nueva infraestructura por factibilidad de alto impacto', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030107, 'Confeccionar y emitir el certificado de factibilidad', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', '40301', 'Solicitud y obtención del certificado de factibilidad', 4030108, 'Respuesta en casos de no factibilidad a posibles clientes regulados (dentro del área de concesión)', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', '4030201', 'Administración y Control-Revisión de Proyectos de urbanización o domiciliario', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030202, 'Aprobaciones de proyectos de conexión', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030203, 'Solicitud y obtención del servicio provisorio', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030204, 'Revisión y verificación diseño de proyectos sanitarios de acuerdo a legislación vigente', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030205, 'Revisión de permisos complementarios', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030206, 'Elaborar informe con observaciones y recomendaciones técnicas de modificación', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030207, 'Verificar las modificaciones planteadas al proyecto', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030208, 'Confección de carta de aprobación y/o archivo del proyecto', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030209, 'Gestión AFR', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030210, 'Elaboración y Gestión de Contratos, Convenios y Acuerdos (ATOs, 52BIS, acuerdos especiales)', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40302, 'Revisión y aprobación del proyecto de urbanización o domiciliario', 4030211, 'Revisión y aprobación o rechazo de modificaciones de proyectos de clientes existentes', 0),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', '4030301', 'Administración y Control-Inspección de obras y emisión de certificado de recepción', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030302, 'Coordinación de Inspección Técnica de obras (aprobación de ITO y revisión de informes)', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030303, 'Inspección técnica en terreno', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030304, 'Recepción de antecedentes técnicos y elaboración de acta de recepción de la obras', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030305, 'Recepción de antecedentes y emisión de autorización de conexión', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030306, 'Recepción al término del proyecto', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030307, 'Confección del certificado de recepción e instalación de obras', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030308, 'Confección de documento de servicio provisorio', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40303, 'Inspección de obras y emisión de certificado de recepción', 4030309, 'Gestión de corte para conexión', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40304, 'Enrolamiento de nuevos clientes y actualización de catastro de obras', '4030401', 'Administración y Control-Enrolamiento de nuevos clientes y actualización de catastro de obras', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40304, 'Enrolamiento de nuevos clientes y actualización de catastro de obras', 4030402, 'Ingreso de los clientes al sistema comercial como cliente provisorio', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40304, 'Enrolamiento de nuevos clientes y actualización de catastro de obras', 4030403, 'Actualización de catastro de obras (incluye planilla de aporte de terceros y GIS)', 1),
            (4, 'Atención Integral de Clientes', 403, 'Gestión de Incorporación de Clientes', 40304, 'Enrolamiento de nuevos clientes y actualización de catastro de obras', 4030404, 'Incorporación de las obras nuevas a la georeferenciazación', 1),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', '5010101', 'Administración y Control-Gestión de Comunicaciones', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010102, 'Definición de plan de comunicaciones', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010103, 'Ejecución de plan de comunicaciones en prensa y medios ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010104, 'Gestion de comunicaciones digitales masivas', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010105, 'Gestion de comunicaciones digitales individuales', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010106, 'Relaciones Publicas ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010107, 'Desarrollo de la Imagen Corporativa ', 1),
            (5, 'Soporte', '501', 'Dirección Superior', '50101', 'Gestión de Comunicaciones', 5010108, 'Comunicaciones Internas', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', '5010201', 'Administración y Control-Gestión de la Responsabilidad Social', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', 5010202, 'Responsabilidad Social o Valor Compartido, medición de sostenibilidad', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', 5010203, 'Definición de estrategia  y políticas de RSE, valor compartido, sostenibilidad ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', 5010204, 'Ejecución de estrategia y políticas de RSE, valor compartido, sustentabilidad', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', 5010205, 'Relación con la comunidad y grupos de interés  ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50102, 'Gestión de la Resposabilidad Social', 5010206, 'Gestión de inclusión y diversidad', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', '5010301', 'Administración y Control-Asesoría Legal', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010302, 'Asesorias Jurídicas ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010303, 'Elaboración de contratos laborales', 1),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010304, 'Asesorias Laborales', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010305, 'Defensa Judicial ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010306, 'Asesoría legal para asuntos inmobiliarios ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010307, 'Asesoría en bases de licitación respecto a ley de subcontratación ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50103, 'Asesoría Legal', 5010308, 'Asesoría en bases de licitación ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', '5010401', 'Administración y Control-Compliance', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010402, 'Implementar un Modelo de Prevención de Delito ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010403, 'Reportar a UAF operaciones sospechosas, ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010404, 'Identificar y coordinar, las necesidades de diseño, actualización o regulación de instrumentos, políticas y procedimientos, que aseguren el cumplimiento normativo.', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010405, 'Revisar y gestionar las instancias/canales, que permitan efectuar denuncias respecto a eventuales irregularidades o incumplimientos normativos.', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010406, 'Asegurar la correcta comunicación, despliegue y aplicación de normativas definidas para los funcionarios.', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50104, 'Compliance (colaboración con Auditoría en la elaboración del plan)', 5010407, 'Verificar el cumplimiento del Código de Ética, ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', '5010501', 'Diseño procedimiento de auditoría interna y externa', 1),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', 5010502, 'Administración y Control-Auditoría', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', 5010503, 'Planificación de Auditoría', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', 5010504, 'Control de Normas y Procedimientos', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', 5010505, 'Verificación de Estados Financieros (externalizado)', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50105, 'Auditoría', 5010506, 'Reportes de Auditoría', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', '5010601', 'Administración y Control-Gestión de Riesgo', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010602, 'Definición del modelo de control de gestión', 1),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010603, 'Establecimiento de políticas y directrices de gestión de riesgos', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010604, 'Generación de la metodología de gestión de riesgos', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010605, 'Aplicación y supervisión de aplicación de metodología en todos procesos ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010606, 'Capacitación en gestión de riesgos', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010607, 'Reportabilidad a distintos niveles ', 0),
            (5, 'Soporte', '501', 'Dirección Superior', 50106, 'Gestión de Riesgo', 5010608, 'Auditorías de Control de Riesgos, sobre la base del modelo de riesgos corporativo', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', '5020101', 'Planificación y organización operacional y financiera de la empresa', 1),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020102, 'Administración y control-Gestión Financiera', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020103, 'Elaboración de presupuesto', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020104, 'Control de presupuesto ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020105, 'Definición de opción de seguros o autoseguros ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020106, 'Atención de clientes externos ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020107, 'Análisis y control financiero ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020108, 'Evaluación económica y financiera de proyectos', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020109, 'Evaluación Expost de Proyectos', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020110, 'Memorias anuales ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020111, 'Apoyo a estudios tarifarios y a la revisión de normativa regulatoria ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020112, 'Gestión de Accionistas e Inversionistas ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020113, 'Gestión de instituciones financieras', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020114, 'Gestión financiera de los AFR ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', '50201', 'Gestión Financiera', 5020115, 'Preparación de informe corporativos', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', '5020201', 'Administración y Control-Tesorería', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020202, 'Administración del flujo de caja anual, mensual y diario ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020203, 'Pago con cheques ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020204, 'Pago vales vista manuales ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020205, 'Pago con transferencia de fondos en línea ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020206, 'Conciliaciones bancarias', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020207, 'Control de la recaudación ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020208, 'Control y registros antecedentes legales ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020209, 'Gestiones bancarias', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020210, 'Gestión de valores', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020211, 'Gestión y control de boletas de garantías', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50202, 'Tesorería', 5020212, 'Gestión de excedentes caja  ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', '5020301', 'Administración y Control-Contabilidad e Impuestos', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', 5020302, 'Contabilidad mensual ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', 5020303, 'Control del activo fijo ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', 5020304, 'Atención clientes externos ', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', 5020305, 'Emitir estados financieros, información contable', 0),
            (5, 'Soporte', 502, 'Administración y Finanzas', 50203, 'Contabilidad e Impuestos', 5020306, 'Impuestos/patentes, declaraciones mensuales y declaraciones juradas', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', '5030101', 'Administración y Control-Gestión de Recursos Humanos ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030102, 'Gestión de relaciones con Sindicatos', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030103, 'Gestión de cargos críticos, talentos, sucesores', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030104, 'Diseño y definición de procesos, estructura organizacional y manual de cargos', 1),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030105, 'Reclutamiento y Selección de Personal', 1),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030106, 'Gestión de Contratos del Personal', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030107, 'Contratación de servicios transitorios ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030108, 'Gestión de becas laborales de estudio ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030109, 'Control horario y control de jornada excepcional', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030110, 'Atención a fiscalizaciones de la Seremi de Salud, la Inspección del Trabajo y Mutualidad, juicios laborales', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030111, 'Servicios al personal', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030112, 'Gestión de licencias médicas ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030113, 'Formación y capacitación ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030114, 'Gestión del conocimiento', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030115, 'Gestión cultural estratégica', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030116, 'Gestión de clima', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030117, 'Pago de remuneraciones e incentivos', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030118, 'Diseño y definición de un sistema de compensaciones', 1),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030119, 'Gestión del Sistema de Compensaciones', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030120, 'Administración de Beneficios ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030121, 'Diseño y definición de un sistema de evaluación de desempeño', 1),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030122, 'Gestión del Desempeño ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030123, 'Desvinculación del personal', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030124, 'Apoyo de análisis de negociaciones colectivas', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', '50301', 'Gestión de Recursos Humanos ', 5030125, 'Asesorías a la línea en la gestión de su personal', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', '5030201', 'Administración y Control-Prevención de Riesgos', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030202, 'Análisis y evaluación de cumplimiento legal de normativa', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030203, 'Identificación de peligros y evaluación de riesgos', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030204, 'Asesoría de unidades operativas (para que tengan procedimientos claros y los respecten, a través de ITO pueden pedir detener una obra)', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030205, 'Asesoría de comites paritarios ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030206, 'Capacitación y desarrollo de competencias en seguridad y salud ocupacional ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030207, 'Control de contratistas ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030208, 'Gestión de higiene y salud ocupacional ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030209, 'Gestión de incidentes o cuasi incidentes ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030210, 'Administración de relaciones con Mutuarias', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030211, 'Reportabilidad de accidentabilidad, ', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030212, 'Atención a fiscalizaciones de la Seremi de Salud , la Inspección del Trabajo y Mutualidad', 0),
            (5, 'Soporte', 503, 'Recursos Humanos', 50302, 'Prevención de Riesgos', 5030213, 'Simulacros de emergencias y accidentes', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', '5040101', 'Administración y Control-Gestión de Abastecimiento', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040102, 'Compras via licitaciones y contrato marco ', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040103, 'Compras sin licitación ', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040104, 'Gestión Proveedores', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040105, 'Gestión de Contratos', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040106, 'Administración de Bodegas', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', '50401', 'Gestión de Abastecimiento', 5040107, 'Despacho y Entregas', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', '5040201', 'Administración y Control-Gestión de Servicios Generales', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040202, 'Aseo y Áreas Verdes', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040203, 'Vigilancia y Seguridad', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040204, 'Valija y Correspondencia', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040205, 'Habilitación de Inmuebles', 1),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040206, 'Mantenimiento de inmuebles, equipos de oficina y otros similares.', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040207, 'Gestión del Transporte', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040208, 'Otros servicios y abastecimientos generales ', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040209, 'Gestión de combustibles', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50402, 'Gestión de Servicios Generales', 5040210, 'Gestión de abastecimiento alternativo. ', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50403, 'Gestión de Activos', '5040301', 'Administración y Control-Gestión de Activos', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50403, 'Gestión de Activos', 5040302, 'Administración de Activo Fijo Físico', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50403, 'Gestión de Activos', 5040303, 'Venta de activos', 0),
            (5, 'Soporte', 504, 'Abastecimiento y Servicios Generales', 50403, 'Gestión de Activos', 5040304, 'Administración de Seguros y gestión de siniestros', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', '5050101', 'Administración y Control-Gestión Informática', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050102, 'Investigación y recomendación de nuevas tecnologías para TI', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050103, 'Análisis de requerimientos tecnológicos y selección de alternativas', 1),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050104, 'Desarrollo de Sistemas', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050105, 'Soporte a usuarios ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050106, 'Custodia de programas y datos ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050107, 'Ingeniería de sistemas ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050108, 'Administración de red de comunicaciones ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050109, 'Administración de bases de datos ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050110, 'Control de Calidad  ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050111, 'Administración de roles y perfiles ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050112, 'Generar información para la gestión ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050113, 'Integración e implementación de soluciones ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050114, 'Administración de seguridad ', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050115, 'Mantenimiento de Sistemas', 0),
            (5, 'Soporte', 505, 'Informática', '50501', 'Gestión Informática', 5050116, 'Administración de procesos ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', '5060101', 'Administración y Control-Ingeniería Operacional ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060102, 'Estimación de recursos hídricos y definición de variables de operación', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060103, 'Optimizar Producción AP ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060104, 'Ingenieria operacional AS ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060105, 'Gestión de redes', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060106, 'Planificar el mantenimiento preventivo y correctivo', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', '50601', 'Ingeniería Operacional', 5060107, 'Desarrollo y mantención de modelos de operaciones ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50602, 'Gestión de Calidad', '5060201', 'Administración y Control-Gestión de Calidad', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50602, 'Gestión de Calidad', 5060202, 'Administración de la Calidad', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50602, 'Gestión de Calidad', 5060203, 'Auditorías normas de calidad', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50602, 'Gestión de Calidad', 5060204, 'Certificación Normas ISO', 1),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50603, 'Gestión Medioambiental', '5060301', 'Administración y Control-Gestión Medioambiental', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50603, 'Gestión Medioambiental', 5060302, 'Gestión Medioambiental ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50603, 'Gestión Medioambiental', 5060303, 'Auditorías normas medio ambientales', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50603, 'Gestión Medioambiental', 5060304, 'Estudios de impacto ambiental', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', '5060401', 'Administración y Control-Planificación', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', 5060402, 'Generación estrategia LP y actualización anual', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', 5060403, 'Alineamiento estratégico  y despliegues', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', 5060404, 'Monitoreo de la estrategia ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', 5060405, 'Gestión general de la innovación', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50604, 'Planificación', 5060406, 'Gestión Documental', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', '5060501', 'Administración y Control-Gestión Regulatoria', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060502, 'Gestión del proceso Tarifario', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060503, 'Actualización de modelos técnicos y tarifarios', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060504, 'Asesoría en marco regulatorio a pedido de clientes internos y externos ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060505, 'Gestión con el regulador y preparación informes y correspondencia para la SISS', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060506, 'Actualización de planes de desarrollo', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060507, 'Documentación de normativa y oficios SISS', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060508, 'Gestión anual del Plan de Desarrollo ', 0),
            (5, 'Soporte', 506, 'Planificación y Desarrollo', 50605, 'Gestión Regulatoria', 5060509, 'Gestión con otras autoridades y preparación informes y correspondencia con otras autoridades', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', '6010101', 'Administración y Control-Control Directo de Riles', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010102, 'Muestreo-Batch', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010103, 'Muestreo-8 horas', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010104, 'Muestreo-12 horas', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010105, 'Muestreo-24 horas', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010106, 'Análisis-G1', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010107, 'Análisis-G2', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010108, 'Análisis-G3', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010109, 'Análisis-G4', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010110, 'Análisis-G5', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010111, 'Análisis-G6', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010112, 'Análisis-G7', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010113, 'Labores Administrativas', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', '60101', 'Control Directo de Riles', 6010114, 'Ejecutar Control Directo de Riles', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60102, 'Mantención de Grifos', '6010201', 'Administración y Control-Mantención de Grifos', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60102, 'Mantención de Grifos', 6010202, 'Inspección visual', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60102, 'Mantención de Grifos', 6010203, 'Mantención menor', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60102, 'Mantención de Grifos', 6010204, 'Mantención mayor', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', '6010301', 'Administración y Control-Corte y Reposición', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', 6010302, 'Visita de Corte', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', 6010303, 'Corte 1a Instancia', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', 6010304, 'Corte 2a Instancia', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', 6010305, 'Reposición 1a Instancia', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60103, 'Corte y Reposición', 6010306, 'Reposición 2a Instancia', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60104, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', '6010401', 'Administración y Control-Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60104, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 6010402, 'Revisión Proyectos: Inversión <= 500 UF', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60104, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 6010403, 'Revisión Proyectos: 500 UF < Inversión < 15.000 UF', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60104, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 6010404, 'Revisión Proyectos: Inversión >= 15.000 UF', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', '6010501', 'Administración y Control-Verificación de Medidores a Solicitud del Cliente', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010502, 'Verificación Medidores Metrológica Medidor Patrón 13 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010503, 'Verificación Medidores Metrológica Medidor Patrón 19 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010504, 'Verificación Medidores Metrológica Medidor Patrón 25 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010505, 'Verificación Medidores Metrológica Medidor Patrón 38 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010506, 'Verificación Medidores Metrológica Banco Prueba fijo 50 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010507, 'Verificación Medidores Metrológica Banco Prueba fijo 80 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010508, 'Verificación Medidores Metrológica Banco Prueba fijo 100 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60105, 'Verificación de Medidores a Solicitud del Cliente', 6010509, 'Verificación Medidores Metrológica Banco Prueba fijo 150 mm', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60106, 'Otras Prestaciones Asociadas', '6010601', 'Administración y Control-Otras Prestaciones Asociadas', 0),
            (6, 'Prestaciones Asociadas', '601', 'Prestaciones Asociadas', 60106, 'Otras Prestaciones Asociadas', 6010602, 'Ejecutar Otras Prestaciones Asociadas', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', '70101', 'Estudios Preliminares,  Hidrológicos e Hidrogeológicos', '7010101', 'Administración y Control-Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', '70101', 'Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 7010102, 'Realización de Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', '70101', 'Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 7010103, 'Licitación-Estudios Preliminares, Hidrológicos e Hidrogeológicos', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', '70101', 'Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 7010104, 'Administración de Contrato-Estudios Preliminares, Hidrológicos e Hidrogeológicos', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', '70101', 'Estudios Preliminares,  Hidrológicos e Hidrogeológicos', 7010105, 'Contraparte Interna-Estudios Preliminares, Hidrológicos e Hidrogeológicos', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', '7010201', 'Administración y Control-Diseños de Obras', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010202, 'Realización de Diseño de Obras', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010203, 'Control de la cartera de proyectos de diseño', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010204, 'Gestión de permisos ', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010205, 'Licitación-Diseños de Obras', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010206, 'Administración de Contrato-Diseños de Obras', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70102, 'Diseños de Obras', 7010207, 'Contraparte Interna-Diseños de Obras', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70103, 'Estudios o Declaraciones de Impacto Ambiental', '7010301', 'Administración y Control-Estudios o Declaraciones de Impacto Ambiental', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70103, 'Estudios o Declaraciones de Impacto Ambiental', 7010302, 'Realización de Estudios o Declaraciones de Impacto Ambiental', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70103, 'Estudios o Declaraciones de Impacto Ambiental', 7010303, 'Licitación-Estudios o Declaraciones de Impacto Ambiental', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70103, 'Estudios o Declaraciones de Impacto Ambiental', 7010304, 'Administración de Contrato-Estudios o Declaraciones de Impacto Ambiental', 0),
            (7, 'Proyectos e Inversiones', '701', 'Planificación y Diseño', 70103, 'Estudios o Declaraciones de Impacto Ambiental', 7010305, 'Contraparte Interna-Estudios o Declaraciones de Impacto Ambiental', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', '7020101', 'Administración y Control-Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 7020102, 'Realización de Construcción de Obras', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 7020103, 'Control de la cartera de proyectos de construcción', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 7020104, 'Licitación-Construcción de Obras', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 7020105, 'Administración de Contrato y Control-Construcción de Obras', 0),
            (7, 'Proyectos e Inversiones', 702, 'Construcción de Obras', '70201', 'Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)', 7020106, 'Contraparte Interna-Construcción de Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', '7030101', 'Administración y Control- Inspección Técnica de Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', 7030102, 'Realización Inspección Técnica de Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', 7030103, 'Documentación técnico económica de la infraestructura', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', 7030104, 'Licitación-Inspección Técnica de Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', 7030105, 'Administración de Contrato y Control-Inspección Técnica de Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', '70301', ' Inspección Técnica de Obras', 7030106, 'Contraparte Interna-Inspección Técnica Obras', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70302, 'Recepción y Puesta en Marcha de Nueva Infraestructura', '7030201', 'Administración y Control-Recepción y Puesta en Marcha de Nueva Infraestructura', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70302, 'Recepción y Puesta en Marcha de Nueva Infraestructura', 7030202, 'Marcha blanca de nueva infraestructura sanitaria', 1),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70302, 'Recepción y Puesta en Marcha de Nueva Infraestructura', 7030203, 'Puesta en marcha de nueva infraestructura  ', 1),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70302, 'Recepción y Puesta en Marcha de Nueva Infraestructura', 7030204, 'Recepción de redes propias o de terceros ', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70303, 'Gestión de la Georeferenciación', '7030301', 'Administración y Control-Gestión de la Georeferenciación', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70303, 'Gestión de la Georeferenciación', 7030302, 'Incorporación de las obras nuevas a la georeferenciazación', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70303, 'Gestión de la Georeferenciación', 7030303, 'Gestión del catastro GIS', 0),
            (7, 'Proyectos e Inversiones', 703, 'Inspección Técnica de Obras', 70303, 'Gestión de la Georeferenciación', 7030304, 'Generación de reportes desde GIS', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', '70401', 'Inversiones en Telemetría y Telecontrol', '7040101', 'Administración y Control-Inversiones en Telemetría y Telecontrol', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', '70401', 'Inversiones en Telemetría y Telecontrol', 7040102, 'Desarrollo, Implementación y Aseguramiento de la Calidad- Inversiones en Telemetría y Telecontrol', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', '70401', 'Inversiones en Telemetría y Telecontrol', 7040103, 'Licitación y adquisición-Inversiones en Telemetría y Telecontrol', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', '70401', 'Inversiones en Telemetría y Telecontrol', 7040104, 'Administración de Contrato y Control-Inversiones en Telemetría y Telecontrol', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', '70401', 'Inversiones en Telemetría y Telecontrol', 7040105, 'Contraparte Interna-Inversiones en Telemetría y Telecontrol', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', 70402, 'Inversiones en Comunicaciones y Sistemas de Información', '7040201', 'Administración y Control-Inversiones en Comunicaciones y Sistemas de Información', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', 70402, 'Inversiones en Comunicaciones y Sistemas de Información', 7040202, 'Desarrollo, Implementación y Aseguramiento de la Calidad-Inversiones en Comunicaciones y Sistemas de Información', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', 70402, 'Inversiones en Comunicaciones y Sistemas de Información', 7040203, 'Licitación y adquisición-Inversiones en Comunicaciones y Sistemas de Información', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', 70402, 'Inversiones en Comunicaciones y Sistemas de Información', 7040204, 'Administración de Contrato y Control-Inversiones en Comunicaciones y Sistemas de Información', 0),
            (7, 'Proyectos e Inversiones', 704, 'Telemetría, Comunicación y Sistemas de Información', 70402, 'Inversiones en Comunicaciones y Sistemas de Información', 7040205, 'Contraparte Interna-Inversiones en Comunicaciones y Sistemas de Información', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', '70501', 'Adquisición Inmuebles', '7050101', 'Administración y Control-Adquisición Inmuebles', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', '70501', 'Adquisición Inmuebles', 7050102, 'Estudios de Factibilidad Interna y de Mercado-Adquisición de Inmuebles', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', '70501', 'Adquisición Inmuebles', 7050103, 'Adquisición Inmuebles', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', 70502, 'Adquisición Vehículos y equipos', '7050201', 'Administración y Control-Adquisición Vehículos y equipos', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', 70502, 'Adquisición Vehículos y equipos', 7050202, 'Licitación y adquisición Vehículos y Equipos', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', 70503, 'Adquisición Bienes Muebles e Insumos', '7050301', 'Administración y Control-Adquisición Bienes Muebles e Insumos', 0),
            (7, 'Proyectos e Inversiones', 705, 'Adquisición Bienes Muebles e Inmuebles', 70503, 'Adquisición Bienes Muebles e Insumos', 7050302, 'Licitación y adquisición Bienes Muebles e Insumos', 0),
            (7, 'Proyectos e Inversiones', 706, 'Adquisición de Derechos de Agua', '70601', 'Adquisición de Derechos de Agua', '7060101', 'Adquisición y gestión de derechos de aguas (planificación, búsqueda y compra de derechos de aguas, regularización de títulos, gestión expedientes DGA, control de publicaciones y presentación de oposiciones en diario oficial, traslado de derechos, estudio títulos) ', 0),
        ],
    },
    'MAE_2': {
        "nombre": 'Cuentas de Recursos',
        "headers": ['CÓDIGO CATEGORÍA', 'CATEGORÍA', 'CÓDIGO SUBCATEGORÍA', 'SUBCATEGORÍA', 'CÓDIGO RECURSO', 'RECURSO'],
        "filas": [
            (1, 'Recursos Humanos', 11, 'Remuneraciones', 1101, 'Remuneraciones'),
            (1, 'Recursos Humanos', 11, 'Remuneraciones', 1102, 'Honorarios'),
            (1, 'Recursos Humanos', 11, 'Remuneraciones', 1103, 'Horas Extras'),
            (1, 'Recursos Humanos', 12, 'Beneficios', 1201, 'Indemnizaciones'),
            (1, 'Recursos Humanos', 12, 'Beneficios', 1202, 'Seguro de Cesantía'),
            (1, 'Recursos Humanos', 12, 'Beneficios', 1203, 'Seguro de Accidentes'),
            (1, 'Recursos Humanos', 12, 'Beneficios', 1204, 'Seguro de Invalidez y Sobrevivencia'),
            (1, 'Recursos Humanos', 12, 'Beneficios', 1205, 'Otros Beneficios Adicionales'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2101, 'Alimentacíón'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2102, 'Capacitación'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2103, 'Pasajes'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2104, 'Alojamientos'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2105, 'Viáticos'),
            (2, 'Gastos Generales', 21, 'Gastos Generales Personal', 2106, 'Accesorios de personal'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2201, 'Arriendo de vehículos y maquinarias'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2202, 'Combustible'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2203, 'Permisos de circulación'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2204, 'Revisión técnica'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2205, 'Seguros'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2206, 'Mantención Preventiva (no incluye combustible)'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2207, 'Recobros y Mantención Correctiva'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2208, 'Peajes'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2209, 'Tags'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2210, 'Implementación tag'),
            (2, 'Gastos Generales', 22, 'Gastos Generales Vehículos y Equipos', 2211, 'Gasto identificación de vehículos'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2301, 'Arriendo de inmuebles'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2302, 'Consumos básicos'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2303, 'Servicio de aseo'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2304, 'Materiales de aseo'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2305, 'Mantención de inmuebles'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2306, 'Mantención de extintores'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2307, 'Mantención de areas verdes'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2308, 'Vigilancia presencial'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2309, 'Vigilancia a distancia'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2310, 'Contribuciones'),
            (2, 'Gastos Generales', 23, 'Gastos Generales Bienes Inmuebles', 2311, 'Comisiones de Corretaje de Inmuebles por compra o arriendo'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2401, 'Arriendo de equipos informáticos'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2402, 'Servicios informáticos'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2403, 'Telefonía Fija'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2404, 'Enlasces de Internet Fijos'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2405, 'Enlaces de datos entre inmuebles (incluye enlaces satelitales)'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2406, 'Telefonía y Banda Ancha Móvil'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2407, 'Tarjetas SIM para aplicaciones M2M'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2408, 'Radio trunking de voz y datos'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2409, 'Telefonía Satelital'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2410, 'Materiales e insumos de oficina, computacionales y bodega'),
            (2, 'Gastos Generales', 24, 'Gastos Generales Bienes Muebles', 2411, 'Materiales e insumos de laboratorio'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2501, 'Dietas del Directorio'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2502, 'Gastos de Representación Directorio'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2503, 'Patentes Comerciales'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2504, 'Servicios de Imprenta, Fotocopiado y Reproducción'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2505, 'Trámites y gastos notariales'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2506, 'Actuaciones Judiciales'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2507, 'Inscripciones'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2508, 'Peritajes'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2509, 'Tasaciones'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2510, 'Enlaces Satelitales'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2511, 'Líneas Transmisión de Datos'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2512, 'Enlaces de Internet'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2513, 'Impuestos por uso de Canales de Radiofrecuencia'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2514, 'Fletes'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2515, 'Transporte de Correspondencia (incluye servicios postales y mensajería)'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2516, 'Transporte de Muestras de Laboratorio'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2517, 'Seguros de infraestructura de capacidad'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2518, 'Seguros de inmuebles'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2519, 'Seguros de redes'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2520, 'Autoseguro'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2521, 'Seguros de Responsabilidad Civil'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2522, 'Seguros Menores'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2523, 'Deducibles pagados'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2524, 'Publicidad y Avisos (Radio, TV, Diarios u otros medios)'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2525, 'Diseño gráfico'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2526, 'Rotulaciones gráficas'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2527, 'Materiales de difusión'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2528, 'Auspicios y aportes'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2529, 'Donaciones'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2530, 'Eventos comunitarios'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2531, 'Eventos corporativos (juntas de accionistas, cenas fin de año, etc.)'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2532, 'Campañas de educación'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2533, 'Materiales de campañas'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2534, 'Derechos de asociaciones y canalistas'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2535, 'Derechos de afiliaciones'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2536, 'Derechos SERVIU o Vialidad'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2537, 'Permisos municipales'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2538, 'Canon anual por activos en comodato'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2539, 'Fondo Fijo Rotativo'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2540, 'Garantías a favor de SISS'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2541, 'Gastos financieros asociados a Garantías SERVIU o Vialidad'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2542, 'Multas'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2543, 'Indemnizaciones a terceros'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2544, 'Suscripciones'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2545, 'Impuestos pagados'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2546, 'Servicios Bancarios'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2547, 'Operaciones Financieras'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2548, 'Castigo incobrables'),
            (2, 'Gastos Generales', 25, 'Otros Gastos Generales', 2549, 'Otros Gastos Generales'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3101, 'Lectura de medidores'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3102, 'Reparto de boletas y otros documentos'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3103, 'Suministro e impresión de boletas y otros documentos'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3104, 'Servicios de recaudación en cajas externas'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3105, 'Servicios de recaudación en cajas propias'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3106, 'Servicios de atención telefónica o distante'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3107, 'Servicios de inspección comercial'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3108, 'Servicios de cobranza prejudicial'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3109, 'Servicios de gestión'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3110, 'Servicios de transporte de personas (buses de acercamiento, radiotaxis, etc.)'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3111, 'Servicios de almacenamiento y bodegaje'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3112, 'Servicios de traslados de mercancías'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3113, 'Servicios de Procesamiento, Archivo y Digitación de Datos'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3114, 'Auditorías a los Estados Financieros'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3115, 'Clasificación de Riesgo'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3116, 'Administración del Registro de Accionistas'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3117, 'Asesorías Tributarias y Contables'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3118, 'Gestión de Recursos Hídricos'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3119, 'Administración del Rol Privado'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3120, 'Selección de Personal'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3121, 'Auditorías Sistemas de Calidad'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3122, 'Asesorías en Servicio al Cliente'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3123, 'Planes de Desarrollo'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3124, 'Estudios Tarifarios'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3125, 'Comisiones de Expertos en Procesos Tarifarios'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3126, 'Defensa por acciones de responsabilidad civil'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3127, 'Defensa de derechos sobre inmuebles'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3128, 'Defensa en juicios laborales'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3129, 'Reclamaciones tributarias'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3130, 'Asesoría y defensa en procesos penales'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3131, 'Laboral Permanente y Negociación Colectiva'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3132, 'Informes Legales o en Derecho'),
            (3, 'Servicios de terceros y asesorías no operacionales', 31, 'Servicios No Operacionales', 3133, 'Otros Servicios No Operacionales'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4101, 'Productos químicos'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4102, 'Energía Eléctrica'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4103, 'Materiales y repuestos'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4104, 'Compra de agua cruda'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4105, 'Compra de agua potable'),
            (4, 'Materiales e Insumos Operacionales', 41, 'Materiales e Insumos ', 4106, 'Arriendo de derechos de agua'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5101, 'Servicios de control de calidad de agua potable'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5102, 'Servicios de control de calidad de agua servidas'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5103, 'Servicios de interconexión AP'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5104, 'Servicios de interconexión AS'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5105, 'Servicios de operación de redes y conexiones'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5106, 'Servicios de operación de infraestructura'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5107, 'Servicios de transporte y disposición de lodos'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5108, 'Servicios de control y monitoreo ambiental'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5109, 'Servicios de mantención de infraestructura'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5110, 'Servicios de mantención de redes y conexiones'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5111, 'Servicios de mantención de recintos'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5112, 'Servicios de mantención de servidumbres'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5113, 'Concesiones marítimas'),
            (5, 'Servicios de Terceros asociados a la Operación y Mantención', 51, 'Servicios Operacionales', 5114, 'Otros Servicios Operacionales'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 61, 'Servicios de Terceros asociados a Control Directo de Riles', 6101, 'Servicios de Terceros asociados a Control Directo de Riles'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 62, 'Servicios de Terceros asociados a Mantención de Grifos', 6201, 'Servicios de Terceros asociados a Mantención de Grifos'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 63, 'Servicios de Terceros asociados a Corte y Reposición', 6301, 'Servicios de Terceros asociados a Corte y Reposición'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 64, 'Servicios de Terceros asociados a Revisión de Proyectos de Construcción', 6401, 'Servicios de Terceros asociados a Revisión de Proyectos de Construcción'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 65, 'Servicios de Terceros asociados a Verificación de Medidores', 6501, 'Servicios de Terceros asociados a Verificación de Medidores'),
            (6, 'Servicios de Terceros de Prestaciones Asociadas', 66, 'Servicios de Terceros asociados a Otras Prestaciones Asociadas', 6601, 'Servicios de Terceros asociados a Otras Prestaciones Asociadas'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 71, 'Servicios de Terceros asociados a Planificación y Diseño', 7101, 'Servicios de Terceros asociados a Estudios preliminares,  Hidrológicos e Hidrogeológicos'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 71, 'Servicios de Terceros asociados a Planificación y Diseño', 7102, 'Servicios de Terceros asociados a Diseños de obras'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 71, 'Servicios de Terceros asociados a Planificación y Diseño', 7103, 'Servicios de Terceros asociados a Impacto Ambiental'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 72, 'Servicios de Terceros asociados a Construcción de Obras', 7201, 'Servicios de Terceros asociados a Construcción de Obras (Obras Civiles, Eléctricas, de Control, Montaje de Equipos, Mitigación Ambiental, Compensación Ambiental, etc.)'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 73, 'Servicios de Terceros asociados a Inspección Técnica de Obras', 7301, 'Servicios de Terceros asociados a Inspección Técnica de Obras'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 74, 'Servicios de Terceros asociados a Telemetría, Comunicación y Sistemas de Información', 7401, 'Servicios de Terceros asociados a Inversiones en Telemetría y Telecontrol'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 74, 'Servicios de Terceros asociados a Telemetría, Comunicación y Sistemas de Información', 7402, 'Servicios de Terceros asociados a Inversiones en Comunicaciones y Sistemas de Información'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 75, 'Servicios de Terceros asociados a Adquisición Bienes Muebles e Inmuebles', 7501, 'Servicios de Terceros asociados a Adquisición Inmuebles'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 75, 'Servicios de Terceros asociados a Adquisición Bienes Muebles e Inmuebles', 7502, 'Servicios de Terceros asociados a Adquisición Vehículos y equipos'),
            (7, 'Servicios de Terceros asociados a Proyectos e Inversiones', 75, 'Servicios de Terceros asociados a Adquisición Bienes Muebles e Inmuebles', 7503, 'Servicios de Terceros asociados a Adquisición Bienes Muebles e Insumos'),
        ],
    },
    'MAE_3': {
        "nombre": 'Cuentas de Servicios e Ingresos Regulados',
        "headers": ['CÓDIGO FAMILIA SERVICIOS REGULADOS', 'FAMILIA SERVICIOS REGULADOS', 'CÓDIGO SERVICIO REGULADO', 'SERVICIO REGULADO', 'CÓDIGO INGRESO REGULADO', 'INGRESO REGULADO'],
        "filas": [
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', '110101', 'Cargo Fijo mensual por cliente'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110102, 'Cargo variable por consumo de agua potable (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110103, 'Cargo variable por consumo de agua potable en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110104, 'Cargo variable por consumo de agua potable en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110105, 'Cargo variable por sobreconsumo de agua potable en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110106, 'Cargo variable por servicio de recolección de aguas servidas (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110107, 'Cargo variable por servicio de recolección de aguas servidas en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110108, 'Cargo variable por servicio de recolección de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110109, 'Cargo variable por sobreconsumo de servicio de recolección de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110110, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110111, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110112, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1101, 'Servicios de Agua Potable y Alcantarillado Usuarios Finales', 110113, 'Cargo variable por sobreconsumo de servicio de tratamiento y disposición de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', '110201', 'Cargo Fijo mensual por cliente'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110202, 'Cargo variable por consumo de agua potable (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110203, 'Cargo variable por consumo de agua potable en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110204, 'Cargo variable por consumo de agua potable en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110205, 'Cargo variable por sobreconsumo de agua potable en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110206, 'Cargo variable por servicio de recolección de aguas servidas (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110207, 'Cargo variable por servicio de recolección de aguas servidas en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110208, 'Cargo variable por servicio de recolección de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110209, 'Cargo variable por sobreconsumo de servicio de recolección de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110210, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas (plano)'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110211, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas en período no punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110212, 'Cargo variable por servicio de tratamiento y disposición de aguas servidas en período punta'),
            (11, 'SERVICIOS CONTINUOS DE AGUA POTABLE Y ALCANTARILLADO', 1102, 'Servicios de Agua Potable y Alcantarillado a Empresas Interconectadas', 110213, 'Cargo variable por sobreconsumo de servicio de tratamiento y disposición de aguas servidas en período punta'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', '120101', 'Muestreo Batch'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120102, 'Muestreo Batch + pH/T°C '),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120103, 'Muestreo Batch + Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120104, 'Muestreo Batch + pH/T°C+Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120105, 'Muestreo Batch + otros'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120106, 'Muestreo 8 Hrs'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120107, 'Muestreo 8 Hrs + pH/T°C'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120108, 'Muestreo 8 Hrs + Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120109, 'Muestreo 8 Hrs + pH/T°C+Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120110, 'Muestreo 8 Hrs + otros'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120111, 'Muestreo 12 Hrs'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120112, 'Muestreo 12 Hrs + pH/T°C'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120113, 'Muestreo 12 Hrs + Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120114, 'Muestreo 12 Hrs + pH/T°C+Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120115, 'Muestreo 12 Hrs + otros'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120116, 'Muestreo 24 Hrs'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120117, 'Muestreo 24 Hrs + pH/T°C'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120118, 'Muestreo 24 Hrs + Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120119, 'Muestreo 24 Hrs + pH/T°C+Q'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120120, 'Muestreo 24 Hrs + otros'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120121, 'Análisis-G1'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120122, 'Análisis-G2'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120123, 'Análisis-G3'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120124, 'Análisis-G4'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120125, 'Análisis-G5'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120126, 'Análisis-G6'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120127, 'Análisis-G7'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1201, 'Control Directo de Riles', 120128, 'Labores Administrativas'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1202, 'Mantención de Grifos', '120201', 'Inspección visual'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1202, 'Mantención de Grifos', 120202, 'Mantención menor'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1202, 'Mantención de Grifos', 120203, 'Mantención mayor'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1203, 'Corte y Reposición', '120301', 'Visita de Corte'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1203, 'Corte y Reposición', 120302, 'Corte 1a Instancia'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1203, 'Corte y Reposición', 120303, 'Corte 2a Instancia'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1203, 'Corte y Reposición', 120304, 'Reposición 1a Instancia'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1203, 'Corte y Reposición', 120305, 'Reposición 2a Instancia'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1204, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', '120401', 'Revisión Proyectos: Inversión <= 500 UF'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1204, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 120402, 'Revisión Proyectos: 500 UF < Inversión < 15.000 UF'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1204, 'Revisión de Proyectos de modificación servicios de agua potable y de alcantarillado.', 120403, 'Revisión Proyectos: Inversión >= 15.000 UF'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', '120501', 'Verificación Medidores Metrológica Medidor Patrón 13 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120502, 'Verificación Medidores Metrológica Medidor Patrón 19 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120503, 'Verificación Medidores Metrológica Medidor Patrón 25 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120504, 'Verificación Medidores Metrológica Medidor Patrón 38 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120505, 'Verificación Medidores Metrológica Banco Prueba fijo 50 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120506, 'Verificación Medidores Metrológica Banco Prueba fijo 80 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120507, 'Verificación Medidores Metrológica Banco Prueba fijo 100 mm'),
            (12, 'PRESTACIONES ASOCIADAS REGULADAS', 1205, 'Verificación de Medidores', 120508, 'Verificación Medidores Metrológica Banco Prueba fijo 150 mm'),
        ],
    },
    'MAE_4': {
        "nombre": 'Cuentas de Servicios No Regulados',
        "headers": ['CÓDIGO FAMILIA SERVICIOS NO REGULADOS', 'FAMILIA SERVICIOS NO REGULADOS', 'CÓDIGO SERVICIO NO REGULADO', 'SERVICIO NO REGULADO'],
        "filas": [
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2101, 'INSPECCIÓN DE INSTALACIÓN DOMICILIARIA'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2102, 'VERIFICACIÓN DE FUNCIONAMIENTO DEL MEDIDOR A TRAVÉS DEL MEDIDOR PATRÓN'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2103, 'REPARACIÓN DE INSTALACIÓN DOMICILIARIA DE AGUA POTABLE'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2104, 'DETECCIÓN DE FUGA INTRADOMICILIARIA DE AGUA POTABLE'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2105, 'CONSTRUCCIÓN DE ARRANQUE'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2106, 'ELIMINACIÓN DE ARRANQUE DE AGUA POTABLE A SOLICITUD DEL CLIENTE'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2107, 'CAMBIO DE MEDIDOR'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2108, 'VENTA DE MEDIDOR'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2109, 'CAMBIO DE UBICACIÓN DE MEDIDOR DE AGUA POTABLE'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2110, 'CONEXIÓN DE REDES A EJECUTAR CON PERSONAL DE LA EMPRESA'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2111, 'DESOBSTRUCCIÓN DE UNIÓN DOMICILIARIA'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2112, 'LIMPIEZA DE FOSA SÉPTICA'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2113, 'CONSTRUCCIÓN DE UD'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2114, 'REPARACIÓN DE CÁMARAS DE INSPECCIÓN DOMICILIARIAS'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2115, 'CONSTRUCCIÓN DE CÁMARAS DE INSPECCIÓN DOMICILIARIAS'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2116, 'EMPALMES DE COLECTORES LOTEOS NUEVOS REALIZADOS POR PERSONAL DE LA EMPRESA'),
            (21, 'PRESTACIONES ASOCIADAS NO REGULADAS', 2117, 'OTRAS PRESTACIONES EXPRESAMENTE VINCULADAS AL INCISO 2° DEL ARTÍCULO 21 DE LA LEY DE TARIFAS'),
            (22, 'SERVICIOS NO REGULADOS', 2201, 'VENTA DE AGUA Y SERVICIOS DE ALCANTARILLADO'),
            (22, 'SERVICIOS NO REGULADOS', 2202, 'VENTA DE SUBPRODUCTOS'),
            (22, 'SERVICIOS NO REGULADOS', 2203, 'TRATAMIENTO DE RILES'),
            (22, 'SERVICIOS NO REGULADOS', 2204, 'SERVICIOS DE ASESORÍA Y GESTIÓN NO APR (PARA TERCEROS, RELACIONADOS O NO)'),
            (22, 'SERVICIOS NO REGULADOS', 2205, 'SERVICIOS DE ASESORÍA Y GESTIÓN DE APR'),
            (22, 'SERVICIOS NO REGULADOS', 2206, 'SERVICIOS DE INGENIERÍA'),
            (22, 'SERVICIOS NO REGULADOS', 2207, 'SERVICIOS DE CONSTRUCCIÓN E INSTALACIÓN'),
            (22, 'SERVICIOS NO REGULADOS', 2208, 'SERVICIOS DE ANÁLISIS DE LABORATORIO'),
            (22, 'SERVICIOS NO REGULADOS', 2209, 'SERVICIOS DE REPARACIONES, OPERACIONES Y MANTENIMIENTOS'),
            (22, 'SERVICIOS NO REGULADOS', 2210, 'SERVICIOS DE INSPECCIONES Y CERTIFICACIONES'),
            (22, 'SERVICIOS NO REGULADOS', 2211, 'VENTA DE EQUIPOS Y PIEZAS'),
            (22, 'SERVICIOS NO REGULADOS', 2212, 'ARRIENDO DE ACTIVOS'),
            (22, 'SERVICIOS NO REGULADOS', 2213, 'MONITOREOS AMBIENTALES'),
            (22, 'SERVICIOS NO REGULADOS', 2214, 'OTROS SERVICIOS NO REGULADOS'),
        ],
    },
    'MAE_5': {
        "nombre": 'Cuentas de Activos',
        "headers": ['CÓDIGO FAMILIA CUENTAS ACTIVOS', 'FAMILIA CUENTAS ACTIVOS', 'CÓDIGO CUENTA CONTABILIDAD REGULATORIA DE ACTIVO', 'CUENTA DE CONTABILIDAD REGULATORIA DE ACTIVO'],
        "filas": [
            (100, 'Captaciones Superficiales', '101', 'Captación en Río '),
            (100, 'Captaciones Superficiales', '102', 'Captación en Canal'),
            (100, 'Captaciones Superficiales', '103', 'Captación en Lago o Embalse'),
            (100, 'Captaciones Superficiales', '104', 'Captación en Mar'),
            (200, 'Captaciones Subterráneas', 201, 'Captación mediante Drenes y Galerías'),
            (200, 'Captaciones Subterráneas', 202, 'Captación mediante Punteras'),
            (200, 'Captaciones Subterráneas', 203, 'Captación mediante Sondajes'),
            (200, 'Captaciones Subterráneas', 204, 'Captación mediante Norias'),
            (300, 'Plantas Elevadoras de Agua Potable', 301, 'Plantas Elevadoras de Agua Potable Tipo A'),
            (300, 'Plantas Elevadoras de Agua Potable', 302, 'Plantas Elevadoras de Agua Potable Tipo B'),
            (300, 'Plantas Elevadoras de Agua Potable', 303, 'Plantas Elevadoras de Agua Potable Tipo C'),
            (300, 'Plantas Elevadoras de Agua Potable', 304, 'Plantas Elevadoras de Agua Potable Tipo D'),
            (300, 'Plantas Elevadoras de Agua Potable', 305, 'Plantas Elevadoras de Agua Potable Tipo E'),
            (400, 'Plantas Elevadoras de Aguas Servidas', '401', 'Plantas Elevadoras de Aguas Servidas'),
            (500, 'Estanques', '501', 'Estanques Semienterrados '),
            (500, 'Estanques', '502', 'Estanques Elevados'),
            (600, 'Plantas de Tratamiento de Agua Potable', '601', 'PTAP Compacta'),
            (600, 'Plantas de Tratamiento de Agua Potable', '602', 'PTAP Convencional'),
            (600, 'Plantas de Tratamiento de Agua Potable', '603', 'PTAP Nanofiltración'),
            (600, 'Plantas de Tratamiento de Agua Potable', '604', 'PTAP Osmosis Inversa'),
            (700, 'Sistemas de Desinfección', '701', 'Sistemas de Desinfección de Agua Potable'),
            (800, 'Sistemas de Fluoración', '801', 'Sistemas de Fluoración'),
            (900, 'Red de Distribución', '901', 'Red de Distribución'),
            (1000, 'Red de Recolección', '1001', 'Red de Recolección'),
            (1100, 'Colectores Unitarios', '1101', 'Colectores Unitarios'),
            (1200, 'Conexiones Domiciliarias', '1201', 'Arranques'),
            (1200, 'Conexiones Domiciliarias', '1202', 'Medidores'),
            (1200, 'Conexiones Domiciliarias', '1203', 'Uniones Domiciliarias'),
            (1300, 'Conducciones de Agua Potable', '1301', 'Conducciones de AP en Acueductos'),
            (1300, 'Conducciones de Agua Potable', '1302', 'Conducciones de AP en Presión'),
            (1400, 'Conducciones de Aguas Servidas', '1401', 'Conducciones de AS en Acueductos'),
            (1400, 'Conducciones de Aguas Servidas', '1402', 'Conducciones de AS en Presión'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1501', 'PTAS Lagunas Facultativas'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1502', 'PTAS Lagunas Aireadas'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1503', 'PTAS Lodos Activados'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1504', 'PTAS Tratamiento Primario'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1505', 'PTAS Lombrifiltros'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1506', 'Pretratamiento AS'),
            (1500, 'Sistemas de Tratamiento de Aguas Servidas', '1507', 'Emisario Submarino'),
            (1600, 'Infraestructura de apoyo', '1601', 'Macromedidores'),
            (1600, 'Infraestructura de apoyo', '1602', 'Reductoras de Presión'),
            (1600, 'Infraestructura de apoyo', '1603', 'Anti Golpe de Ariete'),
            (1600, 'Infraestructura de apoyo', '1605', 'Grupos Electrógenos'),
            (1600, 'Infraestructura de apoyo', '1606', 'Subestaciones'),
            (1600, 'Infraestructura de apoyo', '1607', 'Telemetría'),
            (1700, 'Inmuebles y Terrenos', '1701', 'Terrenos'),
            (1700, 'Inmuebles y Terrenos', '1702', 'Edificaciones'),
            (1700, 'Inmuebles y Terrenos', '1703', 'Urbanizaciones'),
            (1800, 'Maquinarias, Camiones y vehículos', '1801', 'Automóvil'),
            (1800, 'Maquinarias, Camiones y vehículos', '1802', 'Furgón'),
            (1800, 'Maquinarias, Camiones y vehículos', '1803', 'Camionetas'),
            (1800, 'Maquinarias, Camiones y vehículos', '1804', 'Camiones'),
            (1800, 'Maquinarias, Camiones y vehículos', '1805', 'Maquinarias'),
            (1800, 'Maquinarias, Camiones y vehículos', '1806', 'Otras Maquinarias, Camiones y vehículos'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1901', 'Software Macroinformática'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1902', 'Hardware Macroinformática'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1903', 'Software Macroinformática'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1904', 'Hardware Microinformática'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1905', 'Equipos de Comunicaciones'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1906', 'Hardware Telecontrol (Telemetría)'),
            (1900, 'Tecnología de Información y Comunicaciones (TIC)', '1907', 'Software Telecontrol (Telemetría)'),
            (2000, 'Bienes Muebles', '2001', 'Equipos de Oficina, Bodega y Laboratorio'),
            (2000, 'Bienes Muebles', '2002', 'Equipos Operativos'),
            (2000, 'Bienes Muebles', '2003', 'Mobiliario'),
            (2000, 'Bienes Muebles', '2004', 'Herramientas'),
            (2000, 'Bienes Muebles', '2005', 'Otros Bienes Muebles'),
            (2100, 'Activos Intangibles', '2101', 'Derechos de Agua'),
            (2100, 'Activos Intangibles', '2102', 'Derechos de explotación de concesiones'),
            (2100, 'Activos Intangibles', '2103', 'Otros Activos Intangibles'),
            (2100, 'Activos Intangibles', '2104', 'Servidumbres'),
            (2200, 'Otros Activos No Sanitarios', '2201', 'Plantas de Conversión de Gas'),
            (2200, 'Otros Activos No Sanitarios', '2202', 'Centrales de Generación Eléctrica'),
            (2200, 'Otros Activos No Sanitarios', '2203', 'Otros Activos No Sanitarios'),
            (2300, 'Proyectos en Desarrollo', '2301', 'Infraestructura sanitaria'),
            (2300, 'Proyectos en Desarrollo', '2302', 'Infraestructura no sanitaria'),
        ],
    },
    'MAE_6': {
        "nombre": 'Empresas',
        "headers": ['CÓDIGO EMPRESA', 'EMPRESA'],
        "filas": [
            ('001', 'Aguas del Altiplano S.A.'),
            ('002', 'AQUABIO S.A.'),
            ('003', 'Aguas de Antofagasta S.A.'),
            ('004', 'Tratacal S.A.'),
            ('005', 'Empresa Concesionaria de Servicios Sanitarios  S.A. ECONSSA (Ex ESSAN)'),
            ('006', 'Aguas Chañar S.A.'),
            ('007', 'Aguas del Valle S.A.'),
            ('008', 'Aguas La Serena S.A.'),
            ('009', 'Empresa de Servicios Totoralillo ESSETO S.A.'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.'),
            ('011', 'ESVAL S.A.'),
            ('012', 'Cooperativa de AP Santo Domingo Coopagua Ltda.'),
            ('013', 'Asociación de Vecinos Población Mirasol de Algarrobo'),
            ('014', 'Comunidad Balneario Brisas de Mirasol'),
            ('015', 'Corporación Balneario Algarrobo Norte'),
            ('016', 'E.A.P. Los Molles S.A.'),
            ('017', 'Sociedad Agrícola y Servicios Isla de Pascua S.A.'),
            ('018', 'Inmobiliaria Norte Mar S.A.'),
            ('019', 'Lago Peñuelas S.A.'),
            ('020', 'ESSBÍO  S.A.'),
            ('021', 'Nuevosur S.A.'),
            ('022', 'Cooperativa Comuna de Sagrada Familia Ltda.'),
            ('023', 'Aguas del Centro S.A.'),
            ('024', 'Cooperativa de la Comunidad Maule Ltda.'),
            ('025', 'Cooperativa de la Comunidad de Sarmiento Ltda.'),
            ('026', 'Aguas San Pedro S.A.'),
            ('027', 'Aguas Araucanía S.A.'),
            ('028', 'Aguas Quepe S.A.'),
            ('029', 'Empresa de Servicios Sanitarios de Los Lagos ESSAL  S.A.'),
            ('030', 'Aguas Patagonia de Aysén S.A.'),
            ('031', 'Aguas Magallanes S.A.'),
            ('032', 'Aguas Décima S.A.'),
            ('033', 'Aguas Andinas S.A.'),
            ('034', 'Aguas Cordillera S.A.'),
            ('035', 'Aguas Manquehue S.A.'),
            ('036', 'Sembcorp Aguas Chacabuco S.A.'),
            ('037', 'Sembcorp Aguas Lampa S.A.'),
            ('038', 'Servicio Municipal de Ap y AlC. de Maipú  SMAPA'),
            ('039', 'Sembcorp Aguas Santiago S.A.'),
            ('040', 'Aguas Santiago Poniente ASP S.A.'),
            ('041', 'Comunidad de Servicios Remodelación San Borja COSSBO'),
            ('042', 'Empresa de AP Lo Aguirre Emapal S.A.'),
            ('043', 'Explotaciones Sanitarias ESSA S.A.'),
            ('044', 'Empresa particular de AP y ALC La Leonera S.A.'),
            ('045', 'Melipilla Norte S.A.'),
            ('046', 'Alberto Planella Ortiz Servicio de AP Santa Rosa del Peral'),
            ('047', 'Servicios Sanitarios Larapinta Selar S.A. '),
            ('048', 'Empresa de Servicios Sanitarios Lo Prado SEPRA S.A. '),
            ('049', 'Novaguas S.A.'),
            ('050', 'Huertos Familiares S.A.'),
            ('051', 'Aguas de Colina S.A.'),
            ('052', 'BCC S.A.'),
            ('053', 'Servicios Sanitarios de la Estación S.A.'),
            ('054', 'Empres de Agua Potable Izarra de lo Aguirre S.A.'),
            ('055', 'Servicios Sanitarios Llanos del Solar S.A.'),
            ('056', 'Sanitaria Aguas Lampa S.A.'),
            ('057', 'Aguas de Las Lilas S.A.'),
            ('058', 'Aguas Santiago Norte S.A.'),
        ],
    },
    'MAE_8': {
        "nombre": 'Localidades',
        "headers": ['CÓDIGO LOCALIDAD', 'NOMBRE LOCALIDAD'],
        "filas": [
            (1, 'ACHAO'),
            (3, 'AGUAS COLINA'),
            (4, 'AGUAS CORDILLERA'),
            (6, 'AGUAS LA SERENA'),
            (7, 'PARQUE INDUSTRIAL CORONEL'),
            (8, 'SARMIENTO'),
            (9, 'ALERCE'),
            (10, 'ALGARROBO'),
            (11, 'ALGARROBO NORTE'),
            (12, 'ALMENDRAL'),
            (13, 'ALTO HOSPICIO'),
            (14, 'ANCUD'),
            (15, 'ANDACOLLO'),
            (16, 'ANGOL'),
            (17, 'ANTOFAGASTA'),
            (17, 'ANTOFAGASTA'),
            (18, 'AP BARRANCAS'),
            (19, 'ARAUCO'),
            (20, 'ARICA'),
            (21, 'ARTIFICIO'),
            (22, 'BALMACEDA'),
            (23, 'BOCA DE RAPEL'),
            (24, 'BUIN - PAINE - LINDEROS - MAIPO - ALTO JAHUEL'),
            (25, 'BULNES'),
            (26, 'CABILDO'),
            (27, 'CABRERO'),
            (28, 'CACHAGUA'),
            (29, 'CAJON'),
            (30, 'CALAMA'),
            (30, 'CALAMA'),
            (31, 'CALBUCO'),
            (32, 'CALDERA'),
            (33, 'CALERA DE TANGO'),
            (34, 'CALLE LARGA'),
            (35, 'CANELA ALTA'),
            (36, 'CANELA BAJA'),
            (37, 'CAPITAN PASTENE'),
            (38, 'CARAHUE'),
            (39, 'CARAMPANGUE'),
            (40, 'CARTAGENA'),
            (41, 'CASABLANCA'),
            (42, 'CASTRO'),
            (43, 'CATEMU'),
            (44, 'CAUQUENES'),
            (45, 'CAÑETE'),
            (46, 'CHAITEN'),
            (47, 'CHANCO'),
            (48, 'CHAÑARAL'),
            (49, 'CHAÑARAL ALTO'),
            (50, 'CHEPICA'),
            (51, 'CHEPICAL'),
            (52, 'CHERQUENCO'),
            (53, 'CHICUREO'),
            (54, 'CHIGUAYANTE'),
            (55, 'CHILE CHICO'),
            (56, 'CHILLAN'),
            (57, 'CHIMBARONGO'),
            (58, 'CHINCOLCO'),
            (59, 'CHOL CHOL'),
            (60, 'CHONCHI'),
            (61, 'COBQUECURA'),
            (62, 'COCHRANE'),
            (63, 'CODEGUA'),
            (64, 'COELEMU'),
            (65, 'COIHUECO'),
            (66, 'COINCO'),
            (67, 'COLINA - ESMERALDA'),
            (68, 'COLLIPULLI'),
            (69, 'COLTAUCO'),
            (70, 'COMBARBALA'),
            (71, 'CONCEPCION'),
            (72, 'CONCON'),
            (73, 'CONSTITUCION'),
            (74, 'CONTULMO'),
            (75, 'COPIAPO'),
            (76, 'COQUIMBO'),
            (77, 'CORONEL'),
            (78, 'CORRAL'),
            (79, 'CORTE ALTO'),
            (80, 'COSSBO'),
            (81, 'COYA'),
            (82, 'COYHAIQUE'),
            (83, 'CUNCO'),
            (84, 'CURACAUTIN'),
            (85, 'CURACAVI'),
            (86, 'CURANILAHUE'),
            (87, 'CURANIPE'),
            (88, 'CURAUMA'),
            (89, 'CUREPTO'),
            (90, 'CURICO'),
            (91, 'DALCAHUE'),
            (92, 'DICHATO'),
            (93, 'DIEGO DE ALMAGRO'),
            (94, 'DOÑIHUE'),
            (95, 'EL CARMEN'),
            (96, 'EL CHAMISERO'),
            (97, 'EL MONTE - EL PAICO'),
            (98, 'EL PALQUI'),
            (99, 'EL QUISCO'),
            (100, 'EL SALADO'),
            (101, 'EL TABO'),
            (102, 'EMPEDRADO'),
            (103, 'ERCILLA'),
            (104, 'FLORIDA'),
            (105, 'FREIRE'),
            (106, 'FREIRINA'),
            (107, 'FRESIA'),
            (108, 'FRUTILLAR'),
            (109, 'FUTALEUFU'),
            (110, 'FUTRONO'),
            (111, 'GALVARINO'),
            (112, 'GORBEA'),
            (113, 'GRAN SANTIAGO'),
            (114, 'GRANEROS'),
            (115, 'GUALLECO'),
            (116, 'GUANAQUEROS'),
            (117, 'HANGA ROA'),
            (118, 'HIJUELAS'),
            (119, 'HUALAÑE'),
            (120, 'HUALQUI'),
            (121, 'HUAMALATA'),
            (122, 'HUARA'),
            (123, 'HUASCO'),
            (124, 'HUEPIL'),
            (125, 'HUERTOS FAMILIARES'),
            (126, 'ILLAPEL'),
            (127, 'ILOCA'),
            (128, 'INCA DE ORO'),
            (129, 'IQUIQUE'),
            (130, 'ISLA DE MAIPO'),
            (131, 'ISLA NEGRA'),
            (132, 'LA CALERA'),
            (133, 'LA CRUZ'),
            (134, 'LA HUAYCA'),
            (135, 'LA LAGUNA'),
            (136, 'LA LEONERA'),
            (137, 'LA LIGUA'),
            (138, 'LA PUNTA'),
            (139, 'LA SERENA'),
            (140, 'LA TIRANA'),
            (141, 'LA UNION'),
            (142, 'LABRANZA'),
            (143, 'LAGO RANCO'),
            (144, 'LAJA'),
            (145, 'LAMPA'),
            (146, 'LANCO'),
            (147, 'LARAPINTA'),
            (148, 'LAS BRISAS DE MIRASOL'),
            (149, 'LAS CABRAS'),
            (150, 'LAS CRUCES'),
            (151, 'LAS LOMAS DE MAIPU'),
            (152, 'LAS VERTIENTES - EL CANELO - LA OBRA'),
            (153, 'LASTARRIA'),
            (154, 'LAUTARO'),
            (155, 'LEBU'),
            (156, 'LICAN RAY'),
            (157, 'LICANTEN'),
            (158, 'LIMACHE'),
            (159, 'LINARES'),
            (160, 'LLANQUIHUE'),
            (161, 'LLAY LLAY'),
            (162, 'LO AGUIRRE'),
            (163, 'LO BARNECHEA'),
            (164, 'LO MIRANDA'),
            (165, 'LO PRADO'),
            (166, 'LOLOL'),
            (167, 'LOMAS COLORADAS'),
            (169, 'LONCOCHE'),
            (170, 'LONGAVI'),
            (171, 'LONQUIMAY'),
            (172, 'LONTUE'),
            (173, 'LOS ALAMOS'),
            (174, 'LOS ANDES'),
            (175, 'LOS ANGELES'),
            (176, 'LOS BOSQUINOS'),
            (177, 'LOS LAGOS'),
            (178, 'LOS MOLLES'),
            (179, 'LOS MUERMOS'),
            (180, 'LOS QUEÑES'),
            (181, 'LOS SAUCES'),
            (182, 'LOS TRAPENSES'),
            (183, 'LOS VILOS'),
            (184, 'LOTA'),
            (185, 'LOTEO SANTO TOMAS'),
            (186, 'LUMACO'),
            (187, 'MACHALI'),
            (188, 'MAFIL'),
            (189, 'MAIPU'),
            (190, 'MALLOA'),
            (191, 'MALLOCO - PEÑAFLOR'),
            (193, 'MATILLA'),
            (194, 'COOPERATIVA MAULE'),
            (195, 'MAULLIN'),
            (196, 'MEJILLONES'),
            (197, 'MELIPILLA'),
            (198, 'MELIPILLA NORTE'),
            (199, 'MININCO'),
            (200, 'MIRASOL ALGARROBO'),
            (201, 'MOLINA'),
            (202, 'MONTE AGUILA'),
            (203, 'MONTE PATRIA'),
            (204, 'MULCHEN'),
            (205, 'NACIMIENTO'),
            (206, 'NANCAGUA'),
            (207, 'NAVIDAD'),
            (208, 'NEGRETE'),
            (209, 'NINHUE'),
            (210, 'NOGALES'),
            (211, 'NUEVA IMPERIAL'),
            (212, 'NUEVA TOLTEN'),
            (213, 'OLIVAR ALTO'),
            (214, 'OSORNO'),
            (215, 'OVALLE'),
            (216, 'PADRE HURTADO'),
            (217, 'PADRE LAS CASAS'),
            (218, 'PAIHUANO'),
            (219, 'PAILLACO'),
            (220, 'PALMILLA'),
            (221, 'PAN DE AZUCAR'),
            (222, 'PANGUIPULLI'),
            (223, 'PAPUDO'),
            (224, 'PARRAL'),
            (225, 'PELARCO'),
            (226, 'PELEQUEN'),
            (227, 'PELLUHUE'),
            (228, 'PEMUCO'),
            (229, 'PENCO - LIRQUEN'),
            (230, 'PERALILLO'),
            (231, 'PERALILLO DE VICUÑA'),
            (232, 'PETORCA'),
            (233, 'PEUMO'),
            (234, 'PICA'),
            (235, 'PICHIDANGUI'),
            (236, 'PICHIDEGUA'),
            (237, 'PICHILEMU'),
            (238, 'PILLANLELBUN'),
            (239, 'PINTO'),
            (240, 'PIRQUE'),
            (241, 'PISAGUA'),
            (242, 'PITRUFQUEN'),
            (243, 'PLACILLA'),
            (244, 'PLACILLA DE LA LIGUA'),
            (245, 'PLACILLA DE PEÑUELAS'),
            (245, 'PLACILLA DE PEÑUELAS'),
            (247, 'POBLACION'),
            (248, 'POMAIRE'),
            (249, 'PORVENIR'),
            (250, 'POZO ALMONTE'),
            (251, 'PUCHUNCAVI'),
            (252, 'PUCON'),
            (253, 'PUENTE NEGRO'),
            (254, 'PUERTO AYSEN'),
            (255, 'PUERTO CHACABUCO'),
            (256, 'PUERTO CISNES'),
            (257, 'PUERTO IBAÑEZ'),
            (258, 'PUERTO MONTT'),
            (259, 'PUERTO NATALES'),
            (260, 'PUERTO SAAVEDRA'),
            (261, 'PUERTO VARAS'),
            (262, 'PUNITAQUI'),
            (263, 'PUNTA ARENAS'),
            (264, 'PUNTA DE TRALCA'),
            (266, 'PUREN'),
            (267, 'PURRANQUE'),
            (268, 'PUTAENDO'),
            (269, 'PUTU'),
            (271, 'QUELLON'),
            (272, 'QUEPE'),
            (273, 'QUILACO'),
            (274, 'QUILICURA'),
            (275, 'QUILLECO'),
            (276, 'QUILLON'),
            (277, 'QUILLOTA'),
            (278, 'QUILPUE'),
            (279, 'QUINTA DE TILCOCO'),
            (280, 'QUINTERO'),
            (281, 'QUIRIHUE'),
            (282, 'QUITRATUE'),
            (283, 'RAFAEL'),
            (284, 'RAMADILLAS'),
            (285, 'RANCAGUA'),
            (286, 'RAUCO'),
            (287, 'REAL CURIMON'),
            (288, 'RENAICO'),
            (289, 'RENGO'),
            (290, 'REQUINOA'),
            (291, 'RETIRO'),
            (292, 'REÑACA'),
            (293, 'RINCONADA'),
            (294, 'LA RINCONADA'),
            (295, 'RIO BUENO'),
            (296, 'RIO NEGRO'),
            (297, 'ROCAS DE SANTO DOMINGO'),
            (298, 'ROMERAL'),
            (299, 'ROSARIO'),
            (300, 'SAGRADA FAMILIA'),
            (301, 'SALAMANCA'),
            (302, 'SAN ANTONIO'),
            (303, 'SAN CARLOS'),
            (304, 'SAN CLEMENTE'),
            (305, 'SAN ESTEBAN'),
            (306, 'SAN FELIPE'),
            (307, 'SAN FERNANDO'),
            (308, 'SAN FRANCISCO DE MOSTAZAL'),
            (309, 'SAN GABRIEL'),
            (310, 'SAN IGNACIO'),
            (311, 'SAN ISIDRO'),
            (312, 'SAN JAVIER'),
            (313, 'SAN JOSE DE LA MARIQUINA'),
            (314, 'SAN JOSE DE MAIPO'),
            (315, 'SAN PABLO'),
            (316, 'SAN PEDRO'),
            (317, 'SAN PEDRO DE LA PAZ'),
            (318, 'SAN RAFAEL'),
            (319, 'SAN ROSENDO'),
            (320, 'SAN SEBASTIAN'),
            (321, 'SAN VICENTE DE TAGUA TAGUA'),
            (322, 'SANTA BARBARA'),
            (323, 'SANTA CLARA'),
            (324, 'SANTA CRUZ'),
            (325, 'SANTA JUANA'),
            (326, 'SANTA MARIA'),
            (327, 'SANTA MARIA DE MANQUEHUE'),
            (328, 'SANTA ROSA DEL PERAL'),
            (330, 'SOTAQUI'),
            (331, 'TAL TAL'),
            (332, 'TALAGANTE'),
            (333, 'TALCA'),
            (334, 'TALCAHUANO'),
            (335, 'TEMUCO'),
            (336, 'TENO'),
            (337, 'TIERRA AMARILLA'),
            (338, 'TIL TIL'),
            (339, 'TOCOPILLA'),
            (340, 'TOME'),
            (341, 'TONGOY'),
            (342, 'TOTORALILLO'),
            (343, 'TRAIGUEN'),
            (344, 'TUCAPEL'),
            (345, 'VALDIVIA'),
            (346, 'VALDIVIA DE PAINE'),
            (347, 'VALLE ESCONDIDO'),
            (348, 'VALLE GRANDE ETAPA III'),
            (348, 'VALLE GRANDE ETAPA III'),
            (350, 'VALLENAR'),
            (351, 'VALPARAISO'),
            (352, 'VICTORIA'),
            (353, 'VICUÑA'),
            (354, 'VILCUN'),
            (355, 'VILLA ALEGRE'),
            (356, 'VILLA ALEMANA'),
            (357, 'VILLA LOS DOMINICOS'),
            (358, 'VILLARRICA'),
            (359, 'VIÑA DEL MAR'),
            (360, 'YERBAS BUENAS'),
            (361, 'YUMBEL'),
            (362, 'YUNGAY'),
            (363, 'ZAPALLAR'),
            (364, 'ÑIPAS'),
            (367, 'AYRES DE CHICUREO'),
            (368, 'CHINQUIHUE'),
            (369, 'JARDINES DE LA ESTACION'),
            (370, 'VALLE GRANDE'),
            (372, 'ALERCE SUR'),
            (373, 'LOTEO DOÑA CARMEN (SARMIENTO)'),
            (374, 'IZARRA DE LO AGUIRRE'),
            (378, 'PUNTA DE PARRA'),
            (382, 'ESTACION BUIN'),
            (384, 'LOS ALAMOS DE COLINA'),
            (387, 'RODELILLO'),
            (388, 'PUERTAS DE PADRE HURTADO'),
        ],
    },
    'MAE_9': {
        "nombre": 'Comunas',
        "headers": ['CÓDIGO COMUNA', 'NOMBRE COMUNA'],
        "filas": [
            (15101, 'Arica'),
            (15102, 'Camarones'),
            (15201, 'Putre'),
            (15202, 'General Lagos'),
            ('01101', 'Iquique'),
            ('01107', 'Alto Hospicio'),
            ('01401', 'Pozo Almonte'),
            ('01402', 'Camiña'),
            ('01403', 'Colchane'),
            ('01404', 'Huara'),
            ('01405', 'Pica'),
            ('02101', 'Antofagasta'),
            ('02102', 'Mejillones'),
            ('02103', 'Sierra Gorda'),
            ('02104', 'Taltal'),
            ('02201', 'Calama'),
            ('02202', 'Ollagüe'),
            ('02203', 'San Pedro de Atacama'),
            ('02301', 'Tocopilla'),
            ('02302', 'María Elena'),
            ('03101', 'Copiapó'),
            ('03102', 'Caldera'),
            ('03103', 'Tierra Amarilla'),
            ('03201', 'Chañaral'),
            ('03202', 'Diego de Almagro'),
            ('03301', 'Vallenar'),
            ('03302', 'Alto del Carmen'),
            ('03303', 'Freirina'),
            ('03304', 'Huasco'),
            ('04101', 'La Serena'),
            ('04102', 'Coquimbo'),
            ('04103', 'Andacollo'),
            ('04104', 'La Higuera'),
            ('04105', 'Paiguano'),
            ('04106', 'Vicuña'),
            ('04201', 'Illapel'),
            ('04202', 'Canela'),
            ('04203', 'Los Vilos'),
            ('04204', 'Salamanca'),
            ('04301', 'Ovalle'),
            ('04302', 'Combarbalá'),
            ('04303', 'Monte Patria'),
            ('04304', 'Punitaqui'),
            ('04305', 'Río Hurtado'),
            ('05101', 'Valparaíso'),
            ('05102', 'Casablanca'),
            ('05103', 'Concón'),
            ('05104', 'Juan Fernández'),
            ('05105', 'Puchuncaví'),
            ('05801', 'Quilpué'),
            ('05107', 'Quintero'),
            ('05804', 'Villa Alemana'),
            ('05109', 'Viña del Mar'),
            ('05201', 'Isla de Pascua'),
            ('05301', 'Los Andes'),
            ('05302', 'Calle Larga'),
            ('05303', 'Rinconada'),
            ('05304', 'San Esteban'),
            ('05401', 'La Ligua'),
            ('05402', 'Cabildo'),
            ('05403', 'Papudo'),
            ('05404', 'Petorca'),
            ('05405', 'Zapallar'),
            ('05501', 'Quillota'),
            ('05502', 'Calera'),
            ('05503', 'Hijuelas'),
            ('05504', 'La Cruz'),
            ('05802', 'Limache'),
            ('05506', 'Nogales'),
            ('05803', 'Olmué'),
            ('05601', 'San Antonio'),
            ('05602', 'Algarrobo'),
            ('05603', 'Cartagena'),
            ('05604', 'El Quisco'),
            ('05605', 'El Tabo'),
            ('05606', 'Santo Domingo'),
            ('05701', 'San Felipe'),
            ('05702', 'Catemu'),
            ('05703', 'Llaillay'),
            ('05704', 'Panquehue'),
            ('05705', 'Putaendo'),
            ('05706', 'Santa María'),
            ('06101', 'Rancagua'),
            ('06102', 'Codegua'),
            ('06103', 'Coinco'),
            ('06104', 'Coltauco'),
            ('06105', 'Doñihue'),
            ('06106', 'Graneros'),
            ('06107', 'Las Cabras'),
            ('06108', 'Machalí'),
            ('06109', 'Malloa'),
            ('06110', 'Mostazal'),
            ('06111', 'Olivar'),
            ('06112', 'Peumo'),
            ('06113', 'Pichidegua'),
            ('06114', 'Quinta de Tilcoco'),
            ('06115', 'Rengo'),
            ('06116', 'Requínoa'),
            ('06117', 'San Vicente'),
            ('06201', 'Pichilemu'),
            ('06202', 'La Estrella'),
            ('06203', 'Litueche'),
            ('06204', 'Marchihue'),
            ('06205', 'Navidad'),
            ('06206', 'Paredones'),
            ('06301', 'San Fernando'),
            ('06302', 'Chépica'),
            ('06303', 'Chimbarongo'),
            ('06304', 'Lolol'),
            ('06305', 'Nancagua'),
            ('06306', 'Palmilla'),
            ('06307', 'Peralillo'),
            ('06308', 'Placilla'),
            ('06309', 'Pumanque'),
            ('06310', 'Santa Cruz'),
            ('07101', 'Talca'),
            ('07102', 'Constitución'),
            ('07103', 'Curepto'),
            ('07104', 'Empedrado'),
            ('07105', 'Maule'),
            ('07106', 'Pelarco'),
            ('07107', 'Pencahue'),
            ('07108', 'Río Claro'),
            ('07109', 'San Clemente'),
            ('07110', 'San Rafael'),
            ('07201', 'Cauquenes'),
            ('07202', 'Chanco'),
            ('07203', 'Pelluhue'),
            ('07301', 'Curicó'),
            ('07302', 'Hualañé'),
            ('07303', 'Licantén'),
            ('07304', 'Molina'),
            ('07305', 'Rauco'),
            ('07306', 'Romeral'),
            ('07307', 'Sagrada Familia'),
            ('07308', 'Teno'),
            ('07309', 'Vichuquén'),
            ('07401', 'Linares'),
            ('07402', 'Colbún'),
            ('07403', 'Longaví'),
            ('07404', 'Parral'),
            ('07405', 'Retiro'),
            ('07406', 'San Javier'),
            ('07407', 'Villa Alegre'),
            ('07408', 'Yerbas Buenas'),
            ('08101', 'Concepción'),
            ('08102', 'Coronel'),
            ('08103', 'Chiguayante'),
            ('08104', 'Florida'),
            ('08105', 'Hualqui'),
            ('08106', 'Lota'),
            ('08107', 'Penco'),
            ('08108', 'San Pedro de la Paz'),
            ('08109', 'Santa Juana'),
            ('08110', 'Talcahuano'),
            ('08111', 'Tomé'),
            ('08112', 'Hualpén'),
            ('08201', 'Lebu'),
            ('08202', 'Arauco'),
            ('08203', 'Cañete'),
            ('08204', 'Contulmo'),
            ('08205', 'Curanilahue'),
            ('08206', 'Los Álamos'),
            ('08207', 'Tirúa'),
            ('08301', 'Los Ángeles'),
            ('08302', 'Antuco'),
            ('08303', 'Cabrero'),
            ('08304', 'Laja'),
            ('08305', 'Mulchén'),
            ('08306', 'Nacimiento'),
            ('08307', 'Negrete'),
            ('08308', 'Quilaco'),
            ('08309', 'Quilleco'),
            ('08310', 'San Rosendo'),
            ('08311', 'Santa Bárbara'),
            ('08312', 'Tucapel'),
            ('08313', 'Yumbel'),
            ('08314', 'Alto Biobío'),
            ('08401', 'Chillán'),
            ('08402', 'Bulnes'),
            ('08403', 'Cobquecura'),
            ('08404', 'Coelemu'),
            ('08405', 'Coihueco'),
            ('08406', 'Chillán Viejo'),
            ('08407', 'El Carmen'),
            ('08408', 'Ninhue'),
            ('08409', 'Ñiquén'),
            ('08410', 'Pemuco'),
            ('08411', 'Pinto'),
            ('08412', 'Portezuelo'),
            ('08413', 'Quillón'),
            ('08414', 'Quirihue'),
            ('08415', 'Ránquil'),
            ('08416', 'San Carlos'),
            ('08417', 'San Fabián'),
            ('08418', 'San Ignacio'),
            ('08419', 'San Nicolás'),
            ('08420', 'Treguaco'),
            ('08421', 'Yungay'),
            ('09101', 'Temuco'),
            ('09102', 'Carahue'),
            ('09103', 'Cunco'),
            ('09104', 'Curarrehue'),
            ('09105', 'Freire'),
            ('09106', 'Galvarino'),
            ('09107', 'Gorbea'),
            ('09108', 'Lautaro'),
            ('09109', 'Loncoche'),
            ('09110', 'Melipeuco'),
            ('09111', 'Nueva Imperial'),
            ('09112', 'Padre Las Casas'),
            ('09113', 'Perquenco'),
            ('09114', 'Pitrufquén'),
            ('09115', 'Pucón'),
            ('09116', 'Saavedra'),
            ('09117', 'Teodoro Schmidt'),
            ('09118', 'Toltén'),
            ('09119', 'Vilcún'),
            ('09120', 'Villarrica'),
            ('09121', 'Cholchol'),
            ('09201', 'Angol'),
            ('09202', 'Collipulli'),
            ('09203', 'Curacautín'),
            ('09204', 'Ercilla'),
            ('09205', 'Lonquimay'),
            ('09206', 'Los Sauces'),
            ('09207', 'Lumaco'),
            ('09208', 'Purén'),
            ('09209', 'Renaico'),
            ('09210', 'Traiguén'),
            ('09211', 'Victoria'),
            (14101, 'Valdivia'),
            (14102, 'Corral'),
            (14103, 'Lanco'),
            (14104, 'Los Lagos'),
            (14105, 'Máfil'),
            (14106, 'Mariquina'),
            (14107, 'Paillaco'),
            (14108, 'Panguipulli'),
            (14201, 'La Unión'),
            (14202, 'Futrono'),
            (14203, 'Lago Ranco'),
            (14204, 'Río Bueno'),
            (10101, 'Puerto Montt'),
            (10102, 'Calbuco'),
            (10103, 'Cochamó'),
            (10104, 'Fresia'),
            (10105, 'Frutillar'),
            (10106, 'Los Muermos'),
            (10107, 'Llanquihue'),
            (10108, 'Maullín'),
            (10109, 'Puerto Varas'),
            (10201, 'Castro'),
            (10202, 'Ancud'),
            (10203, 'Chonchi'),
            (10204, 'Curaco de Vélez'),
            (10205, 'Dalcahue'),
            (10206, 'Puqueldón'),
            (10207, 'Queilén'),
            (10208, 'Quellón'),
            (10209, 'Quemchi'),
            (10210, 'Quinchao'),
            (10301, 'Osorno'),
            (10302, 'Puerto Octay'),
            (10303, 'Purranque'),
            (10304, 'Puyehue'),
            (10305, 'Río Negro'),
            (10306, 'San Juan de la Costa'),
            (10307, 'San Pablo'),
            (10401, 'Chaitén'),
            (10402, 'Futaleufú'),
            (10403, 'Hualaihué'),
            (10404, 'Palena'),
            (11101, 'Coyhaique'),
            (11102, 'Lago Verde'),
            (11201, 'Aysén'),
            (11202, 'Cisnes'),
            (11203, 'Guaitecas'),
            (11301, 'Cochrane'),
            (11302, 'O’Higgins'),
            (11303, 'Tortel'),
            (11401, 'Chile Chico'),
            (11402, 'Río Ibáñez'),
            (12101, 'Punta Arenas'),
            (12102, 'Laguna Blanca'),
            (12103, 'Río Verde'),
            (12104, 'San Gregorio'),
            (12201, 'Cabo de Hornos (Ex - Navarino)'),
            (12202, 'Antártica'),
            (12301, 'Porvenir'),
            (12302, 'Primavera'),
            (12303, 'Timaukel'),
            (12401, 'Natales'),
            (12402, 'Torres del Paine'),
            (13101, 'Santiago'),
            (13102, 'Cerrillos'),
            (13103, 'Cerro Navia'),
            (13104, 'Conchalí'),
            (13105, 'El Bosque'),
            (13106, 'Estación Central'),
            (13107, 'Huechuraba'),
            (13108, 'Independencia'),
            (13109, 'La Cisterna'),
            (13110, 'La Florida'),
            (13111, 'La Granja'),
            (13112, 'La Pintana'),
            (13113, 'La Reina'),
            (13114, 'Las Condes'),
            (13115, 'Lo Barnechea'),
            (13116, 'Lo Espejo'),
            (13117, 'Lo Prado'),
            (13118, 'Macul'),
            (13119, 'Maipú'),
            (13120, 'Ñuñoa'),
            (13121, 'Pedro Aguirre Cerda'),
            (13122, 'Peñalolén'),
            (13123, 'Providencia'),
            (13124, 'Pudahuel'),
            (13125, 'Quilicura'),
            (13126, 'Quinta Normal'),
            (13127, 'Recoleta'),
            (13128, 'Renca'),
            (13129, 'San Joaquín'),
            (13130, 'San Miguel'),
            (13131, 'San Ramón'),
            (13132, 'Vitacura'),
            (13201, 'Puente Alto'),
            (13202, 'Pirque'),
            (13203, 'San José de Maipo'),
            (13301, 'Colina'),
            (13302, 'Lampa '),
            (13303, 'Tiltil'),
            (13401, 'San Bernardo'),
            (13402, 'Buin'),
            (13403, 'Calera de Tango'),
            (13404, 'Paine'),
            (13501, 'Melipilla'),
            (13502, 'Alhué'),
            (13503, 'Curacaví'),
            (13504, 'María Pinto'),
            (13505, 'San Pedro'),
            (13601, 'Talagante'),
            (13602, 'El Monte'),
            (13603, 'Isla de Maipo'),
            (13604, 'Padre Hurtado'),
            (13605, 'Peñaflor'),
        ],
    },
    'MAE_10': {
        "nombre": 'Grupos Tarifarios',
        "headers": ['CÓDIGO GRUPO TARIFARIO', 'NOMBRE GRUPO TARIFARIO'],
        "filas": [
            (1, 'A1 SECTOR CHINQUIHUE '),
            (2, 'AGRICOLA Y SERVICIOS ISLA DE PASCUA '),
            (3, 'AGUAS CORDILLERA INCLUYE SECTOR LOS DOMINICOS '),
            (4, 'AGUAS DE COLINA '),
            (5, 'AGUAS DECIMA '),
            (6, 'AGUAS LA SERENA '),
            (7, 'AGUAS SANTIAGO PONIENTE '),
            (8, 'ANTOFAGASTA '),
            (9, 'ASOCIACION DE VECINOS POBLACION MIRASOL DE ALGARROBO '),
            (10, 'CALAMA '),
            (11, 'COMUNIDAD BALNEARIO BRISAS DE MIRASOL '),
            (12, 'COMUNIDAD SERVICIOS REMODELACION SAN BORJA '),
            (13, 'CONCESIONES PARQUE INDUSTRIAL CORONEL Y SAN PEDRO CORONEL '),
            (14, 'COOP. DE SERVICIO DE ABASTECIMIENTO DE AP Y ALCANTARILLADO DE LA COMUNIDAD DE SAGRADA FAMILIA '),
            (15, 'COOP. DE SERVICIO DE ABASTECIMIENTO DE AP Y SANEAMIENTO AMBIENTAL DE LA COMUNIDAD DE MAULE '),
            (16, 'COOPAGUA '),
            (17, 'COOPERATIVA SARMIENTO '),
            (18, 'CORPORACION BALNEARIO ALGARROBO NORTE '),
            (19, 'EMPRESA DE AGUA POTABLE LO AGUIRRE '),
            (20, 'EMPRESA DE AGUA POTABLE LOS MOLLES '),
            (21, 'EMPRESA DE AGUA POTABLE MELIPILLA NORTE '),
            (22, 'EMPRESA DE SERVICIOS SANITARIOS LO PRADO '),
            (23, 'EMPRESA DE SERVICIOS TOTORALILLO '),
            (24, 'EMPRESA PARTICULAR DE AGUA POTABLE Y ALCANTARILLADO LA LEONERA '),
            (25, 'EMPRESA SANITARIA LAGO PEÑUELAS '),
            (26, 'EXPLOTACIONES SANITARIAS '),
            (27, 'GRUPO 1  ANTOFAGASTA '),
            (28, 'GRUPO 1  ARAUCANIA '),
            (29, 'GRUPO 1  CHAÑAR '),
            (30, 'GRUPO 1  DEL ALTIPLANO '),
            (31, 'GRUPO 1  DEL VALLE '),
            (32, 'GRUPO 1  NUEVO SUR  '),
            (33, 'GRUPO 1  PATAGONIA'),
            (34, 'GRUPO 1 ESSAL '),
            (35, 'GRUPO 1 ESVAL '),
            (36, 'GRUPO 1 VI REGION ESSBIO '),
            (37, 'GRUPO 1 VIII REGION ESSBIO '),
            (38, 'GRUPO 1: GRAN SANTIAGO  ANDINAS  '),
            (39, 'GRUPO 1: SANTA MARIA Y LOS TRAPENSES  MANQUEHUE '),
            (40, 'GRUPO 2 ANTOFAGASTA '),
            (41, 'GRUPO 2 ARAUCANIA '),
            (42, 'GRUPO 2 CHAÑAR '),
            (43, 'GRUPO 2 DEL ALTIPLANO '),
            (44, 'GRUPO 2 DEL VALLE '),
            (45, 'GRUPO 2 NUEVO SUR  '),
            (46, 'GRUPO 2 PATAGONIA'),
            (47, 'GRUPO 2 ESSAL '),
            (48, 'GRUPO 2 ESVAL '),
            (49, 'GRUPO 2 VI REGION ESSBIO '),
            (50, 'GRUPO 2 VIII REGION ESSBIO '),
            (51, 'GRUPO 2: CIUDAD DE CHICUREO  MANQUEHUE '),
            (52, 'GRUPO 2: LOCALIDADES  ANDINAS  '),
            (53, 'GRUPO 3  ARAUCANIA '),
            (54, 'GRUPO 3  DEL ALTIPLANO '),
            (55, 'GRUPO 3  DEL VALLE '),
            (56, 'GRUPO 3  NUEVO SUR  '),
            (57, 'GRUPO 3 ESSAL '),
            (58, 'GRUPO 3 ESVAL '),
            (59, 'GRUPO 4 ESSAL '),
            (60, 'GRUPO 4 ESVAL '),
            (61, 'GRUPO 5 ESVAL '),
            (62, 'GRUPO 6 ESVAL '),
            (63, 'IZARRA DE LO AGUIRRE '),
            (64, 'LABRANZA '),
            (65, 'LAS LOMAS DE MAIPU '),
            (66, 'LOMAS COLORADAS '),
            (67, 'LOS ALAMOS DE COLINA '),
            (68, 'LOS BOSQUINOS '),
            (69, 'LOTEO EL CHAMISERO '),
            (70, 'LOTEO SANTA ROSA DEL PERAL '),
            (71, 'LOTEO SANTO TOMAS '),
            (72, 'MAIPU '),
            (73, 'PAN DE AZUCAR SECTORES B1, C Y D '),
            (74, 'PICHIDANGUI '),
            (75, 'PILLANLELBUN '),
            (76, 'PORVENIR '),
            (77, 'PUERTAS DE PADRE HURTADO '),
            (78, 'PUERTO NATALES '),
            (79, 'PUNTA ARENAS '),
            (80, 'QUEPE '),
            (81, 'SECTOR ALERCE X REGION '),
            (82, 'SECTOR AYRES DE COLINA '),
            (83, 'SECTOR ESTACION '),
            (84, 'SECTOR HUERTOS FAMILIARES '),
            (85, 'SECTOR JARDINES DE LA ESTACION '),
            (86, 'SECTOR LA RINCONADA '),
            (87, 'SECTOR LOTEO DOÑA CARMEN '),
            (88, 'SECTOR LOTEO RODELILLO '),
            (89, 'SECTOR VALLE GRANDE COMUNA DE LAMPA '),
            (90, 'SERVICIOS SANITARIOS LARAPINTA '),
            (91, 'SERVICOMUNAL '),
            (92, 'SERVILAMPA '),
            (93, 'SISTEMA LO BARNECHEA '),
            (94, 'SISTEMA VALLE ESCONDIDO '),
            (95, 'VALLE GRANDE ETAPA III '),
        ],
    },
    'MAE_11': {
        "nombre": 'Sistemas Tarifarios',
        "headers": ['CÓDIGO SISTEMA TARIFARIO', 'NOMBRE SISTEMA TARIFARIO'],
        "filas": [
            (1, 'BUIN - PAINE - LINDEROS - MAIPO - ALTO JAHUEL'),
            (2, 'CURACAVI'),
            (3, 'EL CANELO - LAS VERTIENTES - LA OBRA'),
            (5, 'GRAN SANTIAGO'),
            (6, 'ISLA DE MAIPO'),
            (7, 'MELIPILLA'),
            (8, 'POMAIRE'),
            (9, 'SAN GABRIEL'),
            (10, 'SAN JOSE DE MAIPO - GUAYACAN - EL CAMPITO'),
            (11, 'TALAGANTE - PADRE HURTADO - PEÑAFLOR - MALLOCO - CALERA DE TANGO - EL MONTE - EL PAICO'),
            (12, 'TIL TIL'),
            (13, 'VALDIVIA DE PAINE'),
            (14, 'GRAN SISTEMA NORTE'),
            (15, 'TAL TAL'),
            (16, 'ANGOL'),
            (17, 'CAJON'),
            (18, 'CAPITAN PASTENE'),
            (19, 'CARAHUE'),
            (20, 'CHERQUENCO'),
            (21, 'CHOL CHOL'),
            (22, 'COLLIPULLI'),
            (23, 'CUNCO'),
            (24, 'CURACAUTIN'),
            (25, 'ERCILLA'),
            (26, 'FREIRE'),
            (27, 'GALVARINO'),
            (28, 'GORBEA'),
            (29, 'LASTARRIA'),
            (30, 'LAUTARO'),
            (31, 'LICAN RAY'),
            (32, 'LONCOCHE'),
            (33, 'LONQUIMAY'),
            (34, 'LOS SAUCES'),
            (35, 'LUMACO'),
            (36, 'MININCO'),
            (37, 'NUEVA IMPERIAL'),
            (38, 'NUEVA TOLTEN'),
            (40, 'PITRUFQUEN'),
            (41, 'PUCON'),
            (42, 'PUERTO SAAVEDRA'),
            (43, 'PUREN'),
            (44, 'QUITRATUE'),
            (45, 'RENAICO'),
            (46, 'TEMUCO'),
            (47, 'TRAIGUEN'),
            (48, 'VICTORIA'),
            (49, 'VILCUN'),
            (50, 'VILLARRICA'),
            (51, 'COPIAPO - TIERRA AMARILLA'),
            (52, 'CALDERA - CHAÑARAL'),
            (53, 'VALLENAR'),
            (54, 'HUASCO - FREIRINA'),
            (55, 'DIEGO DE ALMAGRO - EL SALADO'),
            (56, 'INCA DE ORO'),
            (57, 'AGUAS CORDILLERA'),
            (58, 'AGUAS DECIMA'),
            (59, 'IQUIQUE - ALTO HOSPICIO - POZO ALMONTE'),
            (60, 'HUARA'),
            (62, 'LA HUAYCA'),
            (63, 'LA TIRANA'),
            (64, 'PICA - MATILLA'),
            (66, 'PISAGUA'),
            (68, 'ARICA'),
            (69, 'LA SERENA - COQUIMBO'),
            (70, 'ANDACOLLO'),
            (71, 'CANELA ALTA'),
            (72, 'CANELA BAJA'),
            (73, 'CHAÑARAL ALTO'),
            (74, 'COMBARBALA'),
            (75, 'EL PALQUI'),
            (76, 'OVALLE - HUAMALATA'),
            (77, 'MONTE PATRIA'),
            (78, 'PAIHUANO'),
            (79, 'PERALILLO DE VICUÑA'),
            (80, 'PUNITAQUI'),
            (81, 'LOS VILOS'),
            (82, 'SALAMANCA'),
            (83, 'SOTAQUI'),
            (84, 'VICUÑA'),
            (85, 'ILLAPEL'),
            (86, 'PUNTA ARENAS'),
            (87, 'PUERTO NATALES'),
            (88, 'PORVENIR'),
            (89, 'CIUDAD DE CHICUREO'),
            (90, 'LOS TRAPENSES'),
            (92, 'LOTEO EL CHAMISERO'),
            (93, 'VALLE GRANDE ETAPA III'),
            (94, 'BALMACEDA'),
            (95, 'CHILE CHICO'),
            (96, 'PUERTO CHACABUCO'),
            (97, 'COCHRANE'),
            (98, 'COYHAIQUE'),
            (99, 'PUERTO AYSEN'),
            (100, 'PUERTO CISNES'),
            (101, 'PUERTO IBAÑEZ'),
            (102, 'COOPAGUA'),
            (104, 'ACHAO'),
            (105, 'ALERCE'),
            (106, 'ANCUD'),
            (107, 'CALBUCO'),
            (108, 'CASTRO'),
            (109, 'CHONCHI'),
            (110, 'CORRAL'),
            (112, 'DALCAHUE'),
            (113, 'FRESIA'),
            (114, 'FRUTILLAR'),
            (115, 'FUTALEUFU'),
            (116, 'FUTRONO'),
            (118, 'LAGO RANCO'),
            (119, 'LANCO'),
            (121, 'LOS LAGOS'),
            (122, 'LOS MUERMOS'),
            (123, 'MAFIL'),
            (124, 'MAULLIN'),
            (125, 'OSORNO'),
            (126, 'PAILLACO'),
            (127, 'PANGUIPULLI'),
            (128, 'PUERTO MONTT'),
            (128, 'PUERTO MONTT'),
            (129, 'PUERTO VARAS - LLANQUIHUE'),
            (130, 'PURRANQUE - CORTE ALTO'),
            (131, 'QUELLON'),
            (132, 'LA UNION - RIO BUENO'),
            (133, 'RIO NEGRO'),
            (134, 'SAN JOSE DE LA MARIQUINA'),
            (135, 'SAN PABLO'),
            (136, 'BOCA DE RAPEL - NAVIDAD'),
            (137, 'CHEPICA'),
            (138, 'CHIMBARONGO'),
            (139, 'CODEGUA - LA PUNTA'),
            (140, 'COINCO'),
            (141, 'COLTAUCO'),
            (142, 'COYA'),
            (143, 'DOÑIHUE'),
            (144, 'GRANEROS'),
            (145, 'LAS CABRAS'),
            (146, 'LO MIRANDA'),
            (147, 'LOLOL'),
            (148, 'MALLOA'),
            (149, 'NANCAGUA'),
            (150, 'OLIVAR ALTO'),
            (151, 'PELEQUEN'),
            (152, 'PERALILLO'),
            (153, 'PEUMO'),
            (154, 'PICHIDEGUA'),
            (155, 'PICHILEMU'),
            (156, 'PLACILLA'),
            (157, 'POBLACION'),
            (158, 'PUENTE NEGRO'),
            (159, 'QUINTA DE TILCOCO'),
            (160, 'RANCAGUA - MACHALI'),
            (161, 'RENGO'),
            (162, 'REQUINOA'),
            (163, 'ROSARIO'),
            (164, 'SAN FERNANDO'),
            (165, 'SAN FRANCISCO DE MOSTAZAL'),
            (166, 'SAN VICENTE DE TAGUA TAGUA'),
            (167, 'SANTA CRUZ - PALMILLA'),
            (168, 'ARAUCO'),
            (169, 'BULNES'),
            (170, 'CABRERO'),
            (171, 'CAÑETE'),
            (173, 'CHILLAN'),
            (174, 'COBQUECURA'),
            (175, 'COELEMU'),
            (176, 'COIHUECO'),
            (177, 'CONCEPCION'),
            (178, 'CONTULMO'),
            (179, 'CORONEL'),
            (180, 'CURANILAHUE'),
            (181, 'DICHATO'),
            (182, 'EL CARMEN'),
            (183, 'FLORIDA'),
            (184, 'HUALQUI'),
            (185, 'HUEPIL'),
            (186, 'LAJA'),
            (187, 'SMAPA'),
            (188, 'LEBU'),
            (189, 'LOMAS COLORADAS'),
            (190, 'LOS ALAMOS - CERRO ALTO - TRES PINOS'),
            (191, 'LOS ANGELES'),
            (192, 'LOTA'),
            (193, 'MONTE AGUILA'),
            (194, 'MULCHEN'),
            (195, 'NACIMIENTO'),
            (196, 'NEGRETE'),
            (197, 'NINHUE'),
            (198, 'ÑIPAS'),
            (199, 'PEMUCO'),
            (201, 'PINTO'),
            (202, 'QUILACO'),
            (203, 'QUILLECO'),
            (204, 'QUILLON'),
            (205, 'QUIRIHUE'),
            (206, 'RAFAEL'),
            (207, 'SAN CARLOS'),
            (208, 'SAN IGNACIO'),
            (209, 'SAN PEDRO DE LA PAZ'),
            (210, 'SANTA BARBARA'),
            (211, 'SANTA CLARA'),
            (212, 'SANTA JUANA'),
            (213, 'TOME'),
            (214, 'TUCAPEL'),
            (215, 'YUMBEL'),
            (216, 'YUNGAY'),
            (217, 'GRAN VALPARAISO - LIMACHE - LA CALERA - QUILLOTA - LA CRUZ'),
            (218, 'CASABLANCA'),
            (219, 'LLAY LLAY'),
            (220, 'CATEMU'),
            (221, 'SAN ESTEBAN'),
            (222, 'PUTAENDO'),
            (223, 'SANTA MARIA DE MANQUEHUE - HUECHURABA - VITACURA'),
            (224, 'LOS ANDES - SAN FELIPE - SANTA MARIA'),
            (225, 'CABILDO'),
            (226, 'PETORCA'),
            (227, 'CHINCOLCO'),
            (229, 'LA LIGUA - PUCHUNCAVI - LITORAL NORTE'),
            (230, 'LITORAL SUR'),
            (231, 'TENO'),
            (232, 'RAUCO'),
            (233, 'HUALAÑE'),
            (234, 'ROMERAL'),
            (235, 'LICANTEN'),
            (236, 'CURICO'),
            (237, 'LOS QUEÑES'),
            (238, 'ILOCA'),
            (239, 'MOLINA'),
            (240, 'CUREPTO'),
            (241, 'PUTU'),
            (242, 'SAN RAFAEL'),
            (243, 'GUALLECO'),
            (244, 'CONSTITUCION'),
            (245, 'PELARCO'),
            (246, 'TALCA'),
            (247, 'SAN CLEMENTE'),
            (248, 'SAN JAVIER'),
            (249, 'EMPEDRADO'),
            (250, 'VILLA ALEGRE'),
            (251, 'CHANCO'),
            (252, 'YERBAS BUENAS'),
            (253, 'PELLUHUE'),
            (254, 'LINARES'),
            (255, 'CURANIPE'),
            (256, 'LONGAVI'),
            (257, 'CAUQUENES'),
            (258, 'RETIRO'),
            (259, 'PARRAL'),
            (260, 'COLINA - ESMERALDA'),
            (261, 'LAMPA'),
            (262, 'AGUAS COLINA'),
            (264, 'AGUAS LA SERENA'),
            (265, 'ALERCE SUR'),
            (266, 'ALGARROBO NORTE'),
            (269, 'AYRES DE COLINA'),
            (270, 'BARRANCAS'),
            (271, 'CHAITEN'),
            (272, 'CIUDAD SATELITE LARAPINTA'),
            (273, 'COOPERATIVA MAULE'),
            (274, 'COOPERATIVA SAGRADA FAMILIA'),
            (275, 'COOPERATIVA SARMIENTO '),
            (276, 'COSSBO'),
            (277, 'ESTACION BUIN'),
            (280, 'HUERTOS FAMILIARES'),
            (281, 'IZARRA DE LO AGUIRRE'),
            (282, 'JARDINES DE LA ESTACION'),
            (283, 'LA LEONERA'),
            (284, 'LABRANZA'),
            (285, 'LAS BRISAS DE MIRASOL'),
            (287, 'LO AGUIRRE'),
            (288, 'LO BARNECHEA'),
            (289, 'LO PRADO'),
            (291, 'LONTUE'),
            (292, 'LOS ALAMOS DE COLINA'),
            (293, 'LOS MOLLES'),
            (294, 'LOTEO DOÑA CARMEN (SARMIENTO)'),
            (295, 'LOTEO INDUSTRIAL'),
            (296, 'LOTEO SANTO TOMAS'),
            (297, 'MELIPILLA NORTE'),
            (298, 'MIRASOL ALGARROBO'),
            (299, 'NOGALES'),
            (300, 'NOVAGUAS'),
            (301, 'PAN DE AZUCAR SECTORES B1-C-D'),
            (302, 'PARQUE INDUSTRIAL CORONEL'),
            (303, 'PICHIDANGUI'),
            (304, 'PILLANLELBUN'),
            (305, 'QUEPE'),
            (306, 'RODELILLO'),
            (308, 'SAN PEDRO'),
            (309, 'SANTA ROSA DEL PERAL'),
            (310, 'SASIPA'),
            (312, 'TOTORALILLO'),
            (313, 'VALLE ESCONDIDO'),
            (313, 'PUERTAS DE PADRE HURTADO'),
        ],
    },
    'MAE_12': {
        "nombre": 'Código Sector Decreto Tarifario',
        "headers": ['CÓDIGO EMPRESA', 'EMPRESA', 'CÓDIGO SECTOR DECRETO TARIFARIO', 'NOMBRE SECTOR DECRETO TARIFARIO'],
        "filas": [
            ('001', 'Aguas del Altiplano S.A.', '01', 'No Aplica'),
            ('002', 'AQUABIO S.A.', '01', 'No Aplica'),
            ('003', 'Aguas de Antofagasta S.A.', '01', 'No Aplica'),
            ('004', 'Tratacal S.A.', '01', 'No Aplica'),
            ('005', 'Empresa Concesionaria de Servicios Sanitarios  S.A. ECONSSA (Ex ESSAN)', '01', 'No Aplica'),
            ('006', 'Aguas Chañar S.A.', '01', 'No Aplica'),
            ('007', 'Aguas del Valle S.A.', '01', 'No Aplica'),
            ('008', 'Aguas La Serena S.A.', '01', 'No Aplica'),
            ('009', 'Empresa de Servicios Totoralillo ESSETO S.A.', '01', 'No Aplica'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '01', 'Los Molles'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '02', 'Pichidangui'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '03', 'Labranza'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '04', 'Pillanlelbun'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '05', 'Puertas de Padre Hurtado'),
            ('010', 'Empresa de Servicios Sanitarios San Isidro ESSSI  S.A.', '06', 'Huertos Familiares'),
            ('011', 'ESVAL S.A.', '01', 'No Aplica'),
            ('012', 'Cooperativa de AP Santo Domingo Coopagua Ltda.', '01', 'No Aplica'),
            ('013', 'Asociación de Vecinos Población Mirasol de Algarrobo', '01', 'No Aplica'),
            ('014', 'Comunidad Balneario Brisas de Mirasol', '01', 'No Aplica'),
            ('015', 'Corporación Balneario Algarrobo Norte', '01', 'No Aplica'),
            ('016', 'E.A.P. Los Molles S.A.', '01', 'No Aplica'),
            ('017', 'Sociedad Agrícola y Servicios Isla de Pascua S.A.', '01', 'No Aplica'),
            ('018', 'Inmobiliaria Norte Mar S.A.', '01', 'No Aplica'),
            ('019', 'Lago Peñuelas S.A.', '01', 'No Aplica'),
            ('020', 'ESSBÍO  S.A.', '01', 'No Aplica'),
            ('021', 'Nuevosur S.A.', '01', 'No Aplica'),
            ('022', 'Cooperativa Comuna de Sagrada Familia Ltda.', '01', 'No Aplica'),
            ('023', 'Aguas del Centro S.A.', '01', 'No Aplica'),
            ('024', 'Cooperativa de la Comunidad Maule Ltda.', '01', 'No Aplica'),
            ('025', 'Cooperativa de la Comunidad de Sarmiento Ltda.', '01', 'No Aplica'),
            ('026', 'Aguas San Pedro S.A.', '01', 'Parque Industrial Coronel y San Pedro'),
            ('026', 'Aguas San Pedro S.A.', '02', 'Las Mariposas'),
            ('026', 'Aguas San Pedro S.A.', '03', 'Alerce, Pto. Montt'),
            ('026', 'Aguas San Pedro S.A.', '04', 'Estación de la comuna de Buin'),
            ('026', 'Aguas San Pedro S.A.', '05', 'San Luis y Brisas Norte'),
            ('027', 'Aguas Araucanía S.A.', '01', 'No Aplica'),
            ('028', 'Aguas Quepe S.A.', '01', 'No Aplica'),
            ('029', 'Empresa de Servicios Sanitarios de Los Lagos ESSAL  S.A.', '01', 'No Aplica'),
            ('030', 'Aguas Patagonia de Aysén S.A.', '01', 'No Aplica'),
            ('031', 'Aguas Magallanes S.A.', '01', 'No Aplica'),
            ('032', 'Aguas Décima S.A.', '01', 'No Aplica'),
            ('033', 'Aguas Andinas S.A.', '01', 'No Aplica'),
            ('034', 'Aguas Cordillera S.A.', '01', 'No Aplica'),
            ('035', 'Aguas Manquehue S.A.', '01', 'No Aplica'),
            ('036', 'Sembcorp Aguas Chacabuco S.A.', '01', 'No Aplica'),
            ('037', 'Sembcorp Aguas Lampa S.A.', '01', 'No Aplica'),
            ('038', 'Servicio Municipal de Ap y AlC. de Maipú  SMAPA', '01', 'No Aplica'),
            ('039', 'Sembcorp Aguas Santiago S.A.', '01', 'No Aplica'),
            ('040', 'Aguas Santiago Poniente ASP S.A.', '01', 'No Aplica'),
            ('041', 'Comunidad de Servicios Remodelación San Borja COSSBO', '01', 'No Aplica'),
            ('042', 'Empresa de AP Lo Aguirre Emapal S.A.', '01', 'No Aplica'),
            ('043', 'Explotaciones Sanitarias ESSA S.A.', '01', 'No Aplica'),
            ('044', 'Empresa particular de AP y ALC La Leonera S.A.', '01', 'No Aplica'),
            ('045', 'Melipilla Norte S.A.', '01', 'No Aplica'),
            ('046', 'Alberto Planella Ortiz Servicio de AP Santa Rosa del Peral', '01', 'No Aplica'),
            ('047', 'Servicios Sanitarios Larapinta Selar S.A. ', '01', 'No Aplica'),
            ('048', 'Empresa de Servicios Sanitarios Lo Prado SEPRA S.A. ', '01', 'No Aplica'),
            ('049', 'Novaguas S.A.', '01', 'No Aplica'),
            ('050', 'Huertos Familiares S.A.', '01', 'No Aplica'),
            ('051', 'Aguas de Colina S.A.', '01', 'No Aplica'),
            ('052', 'BCC S.A.', '01', 'No Aplica'),
            ('053', 'Servicios Sanitarios de la Estación S.A.', '01', 'No Aplica'),
            ('054', 'Empres de Agua Potable Izarra de lo Aguirre S.A.', '01', 'No Aplica'),
            ('055', 'Servicios Sanitarios Llanos del Solar S.A.', '01', 'No Aplica'),
            ('056', 'Sanitaria Aguas Lampa S.A.', '01', 'No Aplica'),
            ('057', 'Aguas de Las Lilas S.A.', '01', 'No Aplica'),
            ('058', 'Aguas Santiago Norte S.A.', '01', 'No Aplica'),
        ],
    },
    'MAE_13': {
        "nombre": 'Clasificación de Activos NBI',
        "headers": ['CÓDIGO FAMILIA NBI', 'FAMILIA NBI', 'CÓDIGO OBRA TIPO NBI', 'OBRA TIPO NBI', 'CÓDIGO TABLA NBI', 'TABLA NBI'],
        "filas": [
            (100, 'Captaciones ', '101', 'Captación en Río ', '101', 'Captación en Río '),
            (100, 'Captaciones ', '102', 'Captación en Canal', '102', 'Captación en Canal'),
            (100, 'Captaciones ', '103', 'Captación en Lago o Embalse', '103', 'Captación en Lago o Embalse'),
            (100, 'Captaciones ', '104', 'Captación en Mar', '104', 'Captación en Mar'),
            (200, 'Captaciones Subterráneas', 201, 'Captación mediante Drenes y Galerías', 201, 'Captación mediante Drenes y Galerías'),
            (200, 'Captaciones Subterráneas', 202, 'Captación mediante Punteras', 202, 'Captación mediante Punteras'),
            (200, 'Captaciones Subterráneas', 203, 'Captación mediante Sondajes', 203, 'Captación mediante Sondajes'),
            (200, 'Captaciones Subterráneas', 204, 'Captación mediante Norias', 204, 'Captación mediante Norias'),
            (300, 'Plantas Elevadoras de Agua Potable', 301, 'Plantas Elevadoras de Agua Potable Tipo A', 301, 'Plantas Elevadoras de Agua Potable Tipo A'),
            (300, 'Plantas Elevadoras de Agua Potable', 302, 'Plantas Elevadoras de Agua Potable Tipo B', 302, 'Plantas Elevadoras de Agua Potable Tipo B'),
            (300, 'Plantas Elevadoras de Agua Potable', 303, 'Plantas Elevadoras de Agua Potable Tipo C', 303, 'Plantas Elevadoras de Agua Potable Tipo C'),
            (300, 'Plantas Elevadoras de Agua Potable', 304, 'Plantas Elevadoras de Agua Potable Tipo D', 304, 'Plantas Elevadoras de Agua Potable Tipo D'),
            (300, 'Plantas Elevadoras de Agua Potable', 305, 'Plantas Elevadoras de Agua Potable Tipo E', 305, 'Plantas Elevadoras de Agua Potable Tipo E'),
            (300, 'Plantas Elevadoras de Aguas Servidas', 351, 'Plantas Elevadoras de Aguas Servidas', 351, 'Plantas Elevadoras de Aguas Servidas'),
            (400, 'Estanques', 401, 'Estanques Semienterrados y Enterrados', 401, 'Estanques Semienterrados y Enterrados'),
            (400, 'Estanques', 402, 'Estanques Elevados', 402, 'Estanques Elevados'),
            (500, 'Plantas de Tratamiento de Agua Potable', 501, 'Plantas de Tratamiento de Agua Potable excepto Osmosis Inversa', 501, 'Plantas de Tratamiento de Agua Potable excepto Osmosis Inversa'),
            (500, 'Plantas de Tratamiento de Agua Potable', '502', 'Plantas de Tratamiento de Agua Potable de Osmosis Inversa', '502', 'Plantas de Tratamiento de Agua Potable de Osmosis Inversa'),
            (600, 'Sistemas de Desinfección', 601, 'Sistemas de Desinfección de Agua Potable', 601, 'Sistemas de Desinfección de Agua Potable'),
            (700, 'Sistemas de Fluoración', 701, 'Sistemas de Fluoración', 701, 'Sistemas de Fluoración'),
            (800, 'Red de Distribución', 801, 'Red de Distribución', 801, 'Red de Distribución (Localidad)'),
            (800, 'Red de Distribución', 801, 'Red de Distribución', 802, 'Red de Distribución (Sectores de Red AP por Localidad)'),
            (800, 'Red de Distribución', 801, 'Red de Distribución', 803, 'Red de Distribución (Tuberías)'),
            (900, 'Red de Recolección', 901, 'Red de Recolección', 901, 'Red de Recolección (Localidad)'),
            (900, 'Red de Recolección', 901, 'Red de Recolección', 902, 'Red de Recolección (Sectores de Red AS por Localidad)'),
            (900, 'Red de Recolección', 901, 'Red de Recolección', 903, 'Red de Recolección (Tuberías)'),
            (900, 'Red de Recolección', 901, 'Red de Recolección', 904, 'Red de Recolección (Unitaria 1)'),
            (1000, 'Conexiones Domiciliarias', 1001, 'Arranques', 1001, 'Arranques'),
            (1000, 'Conexiones Domiciliarias', 1002, 'Medidores', 1002, 'Medidores'),
            (1000, 'Conexiones Domiciliarias', 1003, 'Uniones Domiciliarias', 1003, 'Uniones Domiciliarias'),
            (1100, 'Conducciones de Agua Potable', 1101, 'Conducciones de AP', 1101, 'Conducciones de AP (Identificación)'),
            (1100, 'Conducciones de Agua Potable', 1101, 'Conducciones de AP', 1102, 'Conducciones de AP (Tramos)'),
            (1100, 'Conducciones de Aguas Servidas', 1151, 'Conducciones de AS', 1151, 'Conducciones de AS (Identificación)'),
            (1100, 'Conducciones de Aguas Servidas', 1151, 'Conducciones de AS', 1152, 'Conducciones de AS (Tramos)'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1201, 'Tabla General de Sistemas de Tratamiento de Aguas Servidas', 1201, 'Tabla General de Sistemas de Tratamiento de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1202, 'Pretratamiento de Aguas Servidas', 1202, 'Pretratamiento de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1203, 'Tratamiento Primario de Aguas Servidas', 1203, 'Tratamiento Primario de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1204, 'Tratamiento Secundario de Aguas Servidas', 1204, 'Tratamiento Secundario de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1205, 'Desinfección y Decloración en Plantas de Tratamiento de Aguas Servidas', 1205, 'Desinfección y Decloración en Plantas de Tratamiento de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1206, 'Línea de Lodos de Plantas de Tratamiento de Aguas Servidas', 1206, 'Línea de Lodos de Plantas de Tratamiento de Aguas Servidas'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1207, 'Emisario Submarino - Información General', 1207, 'Emisario Submarino - Información General'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1208, 'Emisario Submarino - Información por Tramos', 1208, 'Emisario Submarino - Información por Tramos'),
            (1200, 'Sistemas de Tratamiento de Aguas Servidas', 1209, 'Control Olores, Generador y Aforos en Plantas de Tratamiento de Aguas Servidas', 1209, 'Control Olores, Generador y Aforos en Plantas de Tratamiento de Aguas Servidas'),
            (1400, 'Obras de Apoyo Sanitario', 1402, 'Macromedidores', 1402, 'Macromedidores'),
            (1400, 'Obras de Apoyo Sanitario', 1403, 'Reductoras de Presión', 1403, 'Reductoras de Presión'),
            (1400, 'Obras de Apoyo Sanitario', 1404, 'Anti Golpe de Ariete', 1404, 'Anti Golpe de Ariete'),
            (1400, 'Obras de Apoyo Sanitario', 1405, 'Atraviesos', 1405, 'Atraviesos'),
            (1500, 'Terrenos', 1501, 'Recintos', 1501, 'Recintos'),
            (1500, 'Terrenos', 1501, 'Recintos', 1502, 'Recintos-Obras'),
            (1500, 'Terrenos', 1503, 'Servidumbres', 1503, 'Servidumbres'),
            (1600, 'Obras de Apoyo General', 1601, 'Grupos Electrógenos', 1601, 'Grupos Electrógenos'),
            (1600, 'Obras de Apoyo General', 1602, 'Subestaciones', 1602, 'Subestaciones'),
            (1600, 'Obras de Apoyo General', 1603, 'Telemetría', 1603, 'Telemetría'),
        ],
    },
    'MAE_14': {
        "nombre": 'Clasificación de Activos No Sanitarios',
        "headers": ['CÓDIGO FAMILIA CUENTAS ACTIVOS', 'FAMILIA CUENTAS ACTIVOS', 'CÓDIGO CUENTA CONTABILIDAD REGULATORIA DE ACTIVO', 'CUENTA CONTABILIDAD REGULATORIA DE ACTIVO', 'CÓDIGO DE ACTIVOS NO SANITARIOS', 'ACTIVOS  NO SANITARIOS'],
        "filas": [
            (1700, 'Inmuebles y Terrenos', 1701, 'Terrenos', 170101, 'Terrenos'),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170201, 'Edificación Oficina Corporativa'),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170202, 'Edificación Administrativa - Oficinas Zonales'),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170203, 'Edificación Operativa - Recintos de Obras '),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170204, 'Edificación - Oficinas de Atención de Público'),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170205, 'Bodegas'),
            (1700, 'Inmuebles y Terrenos', 1702, 'Edificaciones', 170206, 'Estacionamientos'),
            (1700, 'Inmuebles y Terrenos', 1703, 'Urbanizaciones', 170207, 'Urbanizaciones'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1801, 'Automóvil', 180101, 'Automóvil'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1802, 'Furgón', 180201, 'Furgón'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1803, 'Camionetas', 180301, 'Camioneta Cabina Simpla (CS)'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1803, 'Camionetas', 180302, 'Camioneta Doble Cabina (DC)'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1804, 'Camiones', 180401, 'Camión'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1804, 'Camiones', 180402, 'Camión Grúa'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1804, 'Camiones', 180403, 'Camión Desobstructor'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1805, 'Maquinarias', 180501, 'Bomba Cambión Desobstructor'),
            (1800, 'Maquinarias, Camiones y Vehículos', 1805, 'Maquinarias', 180502, 'Otro Equipo o Vehículo'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190101, 'Base de Datos'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190102, 'Modelos Matemáticos, Hidrológicos e Hidráulicos'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190103, 'Sistema de Laboratorio'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190104, 'GIS'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190105, 'Sistema Comercial o de Clientes'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190106, 'Call o Contact Center'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190107, 'ERP o Sistema Administrativo'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190108, 'Business Intelligence'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190109, 'E-learning'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190110, 'Sistema Documental'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190111, 'Sistema de Calidad'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190112, 'Servicio de Plataforma Web'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190113, 'Oficina Comercial Virtual'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190114, 'Sistema de Facturación Electrónica'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190115, 'Sistema de Correo Corporativo'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190116, 'Seguridad Informática'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190117, 'Helpdesk Informático'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1901, 'Software Macroinformática', 190118, 'Otros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1902, 'Hardware Macroinformática', 190201, 'Servidor Físico de Procesamiento'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1902, 'Hardware Macroinformática', 190202, 'Servidor Virtual de Procesamiento (Cloud)'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1902, 'Hardware Macroinformática', 190203, 'Almacenamiento Físico (Servidores, Arreglos de Discos, Datacenter)'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1902, 'Hardware Macroinformática', 190204, 'Almacenamiento Virtual (Cloud)'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1902, 'Hardware Macroinformática', 190205, 'Otros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190301, 'Autocad'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190302, 'Sistema Operativo'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190303, 'MS Officce'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190304, 'Project'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190305, 'Aplicación de PDA'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190306, 'SW Autoconsulta'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190307, 'SW Helpdesk Informático'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190308, 'Monitoreo de red'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190309, 'Monitoreo de Bases de Datos'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190310, 'Extractor de Información'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190311, 'SW de análisis, diseño y desarrollo informático.'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190312, 'SW de Gestión de Filas'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1903, 'Software Microinformática', 190313, 'Otros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190401, 'PC - Desktop'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190402, 'Notebooks'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190403, 'Tablets'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190404, 'Scanners'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190405, 'Plotters'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190406, 'Impresoras compartidas'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190407, 'Impresoras individuales'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190408, 'Equipamiento de Teleconferencia'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190409, 'Video cámaras'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190410, 'Multifuncionales'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190411, 'Memocolectores'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190412, 'Lector Código de Barras - Cajeros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190413, 'Verificador de Cheques - Cajeros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190414, 'Dispensador Boletas - Cajeros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190415, 'Autoconsulta - Cajeros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190416, 'Gestores de Filas'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190417, 'Proyectores'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190418, 'Smart TV'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1904, 'Hardware Microinformática', 190419, 'Otros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190501, 'Telefonos Satelitales'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190502, 'Red WiFi'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190503, 'Firewall plantas'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190504, 'Switch cabecera'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190505, 'Switch distribución'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190506, 'Equipos de Red Comunicaciones Lan-Wan'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190507, 'Equipamiento de radio trunking voz y datos'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190508, 'Equpamiento Radio Comunicaciones'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190509, 'Rack'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190510, 'Unidad de respaldo de energía'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1905, 'Equipos de comunicaciones', 190511, 'Otros'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190601, 'Estación Remota Tipo 1'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190602, 'Estación Remota Tipo 2'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190603, 'Servidor SCADA'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190604, 'Cliente SCADA'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190605, 'Servidor BD y Aplicaciones'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1906, 'Hardware Telecontrol (telemetría)', 190606, 'Sensores y Actuadores'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1907, 'Software Telecontrol (telemtría)', 190701, 'Base de Datos'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1907, 'Software Telecontrol (telemtría)', 190702, 'Software Servidor'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1907, 'Software Telecontrol (telemtría)', 190703, 'Software Cliente'),
            (1900, 'Tecnologías de Infdormación y Comunicaciones (TIC)', 1907, 'Software Telecontrol (telemtría)', 190704, 'Ofimática'),
            (2000, 'Bienes Muebles', 2001, 'Equipos de Oficina, Bodega y Laboratorio', 200101, 'Equipos de Oficina'),
            (2000, 'Bienes Muebles', 2001, 'Equipos de Oficina, Bodega y Laboratorio', 200102, 'Equipos de Bodega'),
            (2000, 'Bienes Muebles', 2001, 'Equipos de Oficina, Bodega y Laboratorio', 200103, 'Equipos de Laboratorio'),
            (2000, 'Bienes Muebles', 2002, 'Equipos Operativos', 200201, 'Equipos Operativos'),
            (2000, 'Bienes Muebles', 2003, 'Mobiliario', 200301, 'Mobiliario'),
            (2000, 'Bienes Muebles', 2004, 'Herramientas', 200401, 'Herramientas'),
            (2000, 'Bienes Muebles', 2005, 'Otros bienes muebles', 200501, 'Otros bienes muebles'),
            (2100, 'Activos Intangibles', '2101', 'Derechos de Agua', 210101, 'Derechos de Agua'),
            (2100, 'Activos Intangibles', '2102', 'Derechos de explotación de concesiones', 210201, 'Derechos de explotación de concesiones'),
            (2100, 'Activos Intangibles', '2103', 'Otros Activos Intangibles', 210301, 'Otros Activos Intangibles'),
            (2100, 'Activos Intangibles', '2104', 'Servidumbres', 210401, 'Servidumbres'),
            (2200, 'Otros Activos No Sanitarios', '2201', 'Plantas de Conversión de Gas', 220101, 'Plantas de Conversión de Gas'),
            (2200, 'Otros Activos No Sanitarios', '2202', 'Centrales de Generación Eléctrica', 220201, 'Centrales de Generación Eléctrica'),
            (2200, 'Otros Activos No Sanitarios', '2203', 'Otros Activos No Sanitarios', 220301, 'Otros Activos No Sanitarios'),
        ],
    },
}

# =====================================================================
# CONTENIDO COMPLETO DE LAS TABLAS DE TIPIFICACIONES/PARÁMETROS SISS
# (PARAM_1 a PARAM_114). Fuente: articles-19850_MaestrosSISS2312.xlsx.
# Cada una es un catálogo valor->descripción de un campo tipificado
# (ej. PARAM_12 = Tipo Cargo Eléctrico, el mismo campo de MEI_2).
# =====================================================================
TIPIFICACIONES_DATOS = {
    'PARAM_1': {
        "headers": ['VALOR TIPO ACCESORIO', 'DESCRIPCIÓN TIPO ACCESORIO'],
        "filas": [(1, 'Abrigo'), (2, 'Abrigo impermeable'), (3, 'Antiparras con Protección UV'), (4, 'Arnes de seguridad, tres puntas'), (5, 'Audífonos de seguridad'), (6, 'Botas de Seguridad tipo pescador'), (7, 'Botas de seguridad, tipo minera'), (8, 'Barbiquejo (para afirmar el casco)'), (9, 'Botin Antideslizante'), (10, 'Botin de seguridad'), (11, 'Camisa'), (12, 'Casco de seguridad V - Gard.'), (13, 'Chaleco reflectante.'), (14, 'Chaqueta'), (15, 'Coipa (protección solar trasera)'), (16, 'Delantal Blanco'), (17, 'Filtros para polvo y particulas'), (18, 'Gorro de legionario (protección solar delantera y trasera)'), (19, 'Guante antiácido'), (20, 'Guantes anticorte extralargo'), (21, 'Guante anticorte'), (22, 'Guantes de cabritilla'), (23, 'Guantes descarne con puño'), (24, 'Guantes quirurgicos 100 pares'), (25, 'Lentes con Protección UV'), (26, 'Línea de vida para arnés de seguridad'), (27, 'Mascarilla medio rostro'), (28, 'Máscara Respirador + Filtro'), (29, 'Mascarilla desechable'), (30, 'Overol, tipo piloto con reflectante'), (31, 'Pantalón'), (32, 'Protector solar (loción)'), (33, 'Slack (casaca y pantalón)'), (34, 'Tapones para oidos desechables'), (35, 'Tapones para oidos reutilizables'), (36, 'Toca de Malla'), (37, 'Traje 2 piezas'), (38, 'Traje Impermeable Antiácido'), (39, 'Zapatos')],
    },
    'PARAM_2': {
        "headers": ['VALOR TIPO ACTIVIDAD DE CONTROL DIRECTO DE RILES', 'DESCRIPCIÓN TIPO ACTIVIDAD DE CONTROL DIRECTO DE RILES'],
        "filas": [(1, 'Muestreo'), (2, 'Análisis'), (3, 'Labores Administrativas')],
    },
    'PARAM_3': {
        "headers": ['VALOR TIPO ACTIVIDAD DE MANTENCIÓN DE INFRAESTRUCTURA', 'DESCRIPCIÓN TIPO ACTIVIDAD DE MANTENCIÓN DE INFRAESTRUCTURA', 'INFRAESTRUCTURA A MANTENER', 'MANTENIMIENTO'],
        "filas": [(1, 'Ajuste', 'EQUIPOS', 'PREVENTIVO'), (2, 'Alineamiento', None, None), (3, 'Análisis (incluye análisis de aceite y de vibraciones)', None, None), (4, 'Apriete', None, None), (5, 'Calibración', None, None), (6, 'Cambio o reemplazo (incluye kits)', None, None), (7, 'Contrastación', None, None), (8, 'Desenergización o desconexión', None, None), (9, 'Desinfección', None, None), (10, 'Eliminación de aire', None, None), (11, 'Inspección, revisión, realización de pruebas o comprobación del estado o funcionamiento', None, None), (12, 'Lavado', None, None), (13, 'Limpieza', None, None), (14, 'Lubricación, cambio de aceite o engrase', None, None), (15, 'Medición de aislación', None, None), (16, 'Pintura', None, None), (17, 'Registro de parámetros operacionales', None, None), (18, 'Relleno', None, None), (19, 'Reemplazo o reposición', None, None), (20, 'Reparación', None, None), (21, 'Retiro o desarme para revisión', None, None), (22, 'Reemplazo o reposición', None, 'EMERGENCIA O CORRECTIVO'), (23, 'Reparación', None, None), (24, 'Detección de fallas', 'OBRAS CIVILES', 'PREVENTIVO'), (25, 'Lavado', None, None), (26, 'Limpieza', None, None), (27, 'Pintura', None, None), (28, 'Refuerzo', None, None), (29, 'Reemplazo o reposición', None, None), (30, 'Reparación', None, None), (31, 'Reemplazo o reposición', None, 'EMERGENCIA O CORRECTIVO'), (32, 'Reparación', None, None), (33, 'Desmalezado', 'RECINTOS Y SERVIDUMBRES', 'PREVENTIVO'), (34, 'Control de Plagas', None, None), (35, 'Pintura', None, None), (36, 'Mantención de Jardines', None, None), (37, 'Limpieza', None, None), (38, 'Poda', None, None)],
    },
    'PARAM_4': {
        "headers": ['VALOR TIPO ACTIVIDAD DE OPERACIÓN', 'DESCRIPCIÓN TIPO ACTIVIDAD DE OPERACIÓN', 'CATEGORÍA OPERACIÓN'],
        "filas": [(1, 'Inspección de aliviaderos', 'Inspección y Vigilancia'), (2, 'Inspección de acueducto (interior)', None), (3, 'Inspección de acueducto (exterior)', None), (4, 'Inspección de cámara de captación', None), (5, 'Inspección de colectores', None), (6, 'Inspección de contrato', None), (7, 'Inspección de dren (interior)', None), (8, 'Inspección de equipo e instalaciones anexas', None), (9, 'Inspección de infraestructura operativa de planta', None), (10, 'Inspección de instalaciones en cámara de inyección', None), (11, 'Inspección de instalaciones en sala de bodega', None), (12, 'Inspección de instalaciones en sala de dosificación', None), (13, 'Inspección de instalaciones en sala de preparación', None), (14, 'Inspección de obra de disposición', None), (15, 'Inspección de pozo', None), (16, 'Inspección de puntos de control de presión', None), (17, 'Inspección de sondaje', None), (18, 'Inspección de tuberia de aducción', None), (19, 'Inspección de tuberia de impulsión', None), (20, 'Inspección de zona aguas arriba de captación', None), (21, 'Inspección sectorización', None), (22, 'Revisión de cámara de inspección', None), (23, 'Análisis', 'Lectura o Medición de Variables'), (24, 'Control de consumo eléctrico', None), (25, 'Control de niveles', None), (26, 'Control de stock', None), (27, 'Lectura de medidor de caudal', None), (28, 'Lectura de parámetros', None), (29, 'Lectura de parámetros de equipos en línea', None), (30, 'Medición de calidad del producto', None), (31, 'Medición de caudal instantáneo', None), (32, 'Medición de equipos en línea', None), (33, 'Medición de la condición del AS', None), (34, 'Medición de niveles dinámico y estático', None), (35, 'Medición de parámetros de agua conducida', None), (36, 'Medición de parámetros de control operacional', None), (37, 'Medición de presiones en puntos de control definidos por la SISS', None), (38, 'Medición de variables', None), (39, 'Muestreo', None), (40, 'Abastecimiento de Petróleo', 'Operación de Elementos'), (41, 'Cambio de cilindros', None), (42, 'Cambio de contenedor', None), (43, 'Control de unidades nebulizadoras de sistema para control olores', None), (44, 'Deshidratación de lodos', None), (45, 'Desinfección de pozo', None), (46, 'Desinfección de sondaje', None), (47, 'Limpieza de rejas', None), (48, 'Operación de cambio de consignas', None), (49, 'Operación de limpieza de filtros greensand', None), (50, 'Operación de limpieza de unidades', None), (51, 'Operación de planta', None), (52, 'Operación de regulación de dosificación de productos químicos', None), (53, 'Operación de sifones', None), (54, 'Operación de válvulas', None), (55, 'Operación de válvulas de desagüe', None), (56, 'Operaciones rutinarias para el correcto funcionamiento de captación', None), (57, 'Prueba alarma de fuga gas cloro', None), (58, 'Regulación de caudal', None), (59, 'Simulacro de derrame', None), (60, 'Recepción de producto, operaciones de limpieza y regulación', None), (61, 'Reposición o recarga de unidades nebulizadoras', None), (62, 'Retiro de Basuras', None), (63, 'Retiro de Lodos', None), (64, 'Rotación de equipos', None), (65, 'Transporte y Disposición de Lodos', None), (66, 'Registro de datos, parámetros, lecturas o mediciones', 'Registro e Ingreso de Datos'), (67, 'Registro de estado de infraestructura', None), (68, 'Registro de horómetros', None), (69, 'Registro de mediciones de equipos en línea', None), (70, 'Reporte de novedades', None)],
    },
    'PARAM_5': {
        "headers": ['VALOR TIPO ADQUISICIÓN', 'DESCRIPCIÓN TIPO ADQUISICIÓN'],
        "filas": [(1, 'Servicios'), (2, 'Compra de insumos y suministros'), (3, 'Construcción de obras'), (4, 'Compra de activos')],
    },
    'PARAM_6': {
        "headers": ['VALOR TIPO ARRIENDO EQUIPO INFORMÁTICO', 'DESCRIPCIÓN TIPO ARRIENDO EQUIPO INFORMÁTICO'],
        "filas": [(1, 'Fotocopiadoras'), (2, 'Servidor de Correo'), (3, 'Servidor Sistema Comercial o de Clientes'), (4, 'Servidor GIS'), (5, 'Servidor Sistema Modelos Matemáticos'), (6, 'Servidor Sistema de Laboratorio'), (7, 'Servidor ERP o Sistema Administrativo'), (8, 'Servidor Sistema Documental'), (9, 'Servidor Web'), (10, 'Servidor Sistema de Calidad'), (11, 'Servidor de Red'), (12, 'PC'), (13, 'Notebooks'), (14, 'Scanners'), (15, 'Plotters'), (16, 'Impresoras'), (17, 'Multifuncionales'), (18, 'Servidores de Fax'), (19, 'Memocolectores'), (20, 'Lector Código de Barras - Cajeros'), (21, 'Verificador de Cheques - Cajeros'), (22, 'Dispensador Boletas - Cajeros'), (23, 'Autoconsulta - Cajeros'), (24, 'Equipos de Red Comunicaciones Lan-Wan'), (25, 'Equipamiento de radio trunking voz y datos')],
    },
    'PARAM_7': {
        "headers": ['VALOR TIPO ASESORIA Y ESTUDIO', 'DESCRIPCIÓN TIPO ASESORIA Y ESTUDIO'],
        "filas": [(1, 'Auditoría a los Estados Financieros'), (2, 'Clasificación de Riesgo'), (3, 'Administración del Registro de Accionistas'), (4, 'Asesorías Tributarias y Contables'), (5, 'Gestión de Recursos Hídricos'), (6, 'Administración del Rol Privado'), (7, 'Selección de Personal'), (8, 'Otros ítems de RRHH'), (9, 'Comunicación con Autoridades'), (10, 'Comunicación enfocada en la Comunidad'), (11, 'Auditorías Sistemas de Calidad'), (12, 'Asesorías en Servicios al Cliente'), (13, 'Planes de Desarrollo'), (14, 'Estudios Tarifarios'), (15, 'Comisiones Periciales'), (16, 'Defensa por acciones de responsabilidad civil'), (17, 'Defensa de derechos sobre inmuebles'), (18, 'Defensa en juicios laborales'), (19, 'Defensa en procesos sancionatorios de la SISS'), (20, 'Reclamaciones tributarias'), (21, 'Asesoría y defensa en procesos penales'), (22, 'Laboral Permanente y Negociación Colectiva'), (23, 'Informes Legales o en Derecho')],
    },
    'PARAM_8': {
        "headers": ['VALOR TIPO ASIGNACIÓN', 'DESCRIPCIÓN TIPO ASIGNACIÓN'],
        "filas": [(1, 'Identificador único de la persona o personas asignadas de acuerdo a lo informado en la dotación de RRHH'), (2, 'Identificador único de la o las unidadades organizacionales informadas en el archivo Maestros Concesionaria'), (3, 'Identificador único del o de los procesos a los cuales está asignado el gasto o una proporción de él, según la tabla Procesos')],
    },
    'PARAM_9': {
        "headers": ['VALOR TIPO ATENCIÓN', 'DESCRIPCIÓN TIPO ATENCIÓN'],
        "filas": [(1, 'Consultas'), (2, 'Solicitud de servicios'), (3, 'Reclamos'), (4, 'Atención de emergencias'), (5, 'Cobranza')],
    },
    'PARAM_10': {
        "headers": ['VALOR TIPO AUTONOMÍA', 'DESCRIPCIÓN TIPO AUTONOMÍA'],
        "filas": [(1, 'El titular del Cargo no puede tomar ninguna decisión. Todo está sujeto a las órdenes y aprobación del jefe.'), (2, 'El titular del Cargo puede tomar decisiones menores ajustándose a instrucciones y procedimientos de trabajo muy específicos. Recibe alta supervisión.'), (3, 'El titular del Cargo toma decisiones de acuerdo con instrucciones generales y guías de acción. Se controla periódicamente el desarrollo del trabajo.'), (4, 'El titular del Cargo toma decisiones complejas respaldadas en políticas muy específicas. Se supervisa el cumplimiento de metas.'), (5, 'El titular del Cargo toma decisiones complejas basado en políticas y objetivos funcionales. Se le controlan los resultados finales del área a su cargo.'), (6, 'El titular del Cargo toma decisiones de envergadura para el cumplimiento de los objetivos empresariales. Sólo está sujeto a dirección general.'), (7, 'El titular del Cargo sólo está sujeto a muy amplias políticas y guía general. Tiene una participación muy significativa en la fijación de los objetivos y estrategias globales en función de la misión del negocio.')],
    },
    'PARAM_11': {
        "headers": ['VALOR TIPO CARGO DIRECTORIO', 'DESCRIPCIÓN TIPO CARGO DIRECTORIO'],
        "filas": [(1, 'Presidente'), (2, 'Vicepresidente'), (3, 'Director Titular'), (4, 'Director Suplente')],
    },
    'PARAM_12': {
        "headers": ['VALOR TIPO CARGO ELÉCTRICO', 'DESCRIPCIÓN TIPO CARGO ELÉCTRICO'],
        "filas": [(1, 'Administración del servicio'), (2, 'Arriendo de empalme'), (3, 'Arriendo de medidor'), (4, 'Arriendo de transformador'), (5, 'Cargo Fijo Mensual'), (6, 'Cargo por compras de potencia'), (7, 'Cargo por demanda máxima de potencia contratada'), (8, 'Cargo por demanda máxima de potencia contratada en horas de punta'), (9, 'Cargo por demanda máxima de potencia leída'), (10, 'Cargo por demanda máxima de potencia leída en horas de punta'), (11, 'Cargo por demanda máxima de potencia leída, en su componente de distribución'), (12, 'Cargo por demanda máxima de potencia suministrada'), (13, 'Cargo por demanda máxima de potencia suministrada, en su componente de distribución'), (14, 'Cargo por demanda máxima leída en horas de punta, en su componente de distribución'), (15, 'Cargo por Energía'), (16, 'Cargo por mínimo técnico'), (17, 'Cargo por potencia adicional de invierno en su componente de compras de potencia'), (18, 'Cargo por potencia adicional de invierno en su componente de distribución'), (19, 'Cargo por potencia base en su componente de distribución'), (20, 'Cargo por potencia base en su componente de transmisión'), (21, 'Cargo por potencia contratada'), (22, 'Cargo por potencia de invierno'), (23, 'Cargo por potencia suministrada, en su componente de distribución'), (24, 'Cargo por Servicio Público'), (25, 'Cargo por servicios complementarios'), (26, 'Cargo por Uso del Sistema de Transmisión'), (27, 'Facturación precio estabilizado'), (28, 'Multa o cargo por consumos reactivos'), (29, 'Otro Cargo')],
    },
    'PARAM_13': {
        "headers": ['VALOR TIPO CATEGORÍA PROYECTO', 'DESCRIPCIÓN TIPO CATEGORÍA PROYECTO'],
        "filas": [(1, 'Ampliación de infraestructura (nueva) o adquisición de activos'), (2, 'Mejoramiento de infraestructura o de activos existentes'), (3, 'Reemplazo o reposición de infraestructura o de activos')],
    },
    'PARAM_14': {
        "headers": ['VALOR TIPO CATEGORÍA PROYECTO NO SANITARIO', 'DESCRIPCIÓN TIPO CATEGORÍA PROYECTO NO SANITARIO'],
        "filas": [(1, 'Adquisición de bienes muebles o inmuebles'), (2, 'Adquisición, Desarrollo e/o Implantación de sistemas de TIC'), (3, 'Adquisición de Derechos de Aguas'), (4, 'Adquisición de Vehículos y Maquinarias'), (5, 'Adquisición de Materiales, Repuestos y Equipos'), (6, 'Otros')],
    },
    'PARAM_15': {
        "headers": ['VALOR TIPO CAUSA INTERVENCIÓN ARRANQUE', 'DESCRIPCIÓN TIPO CAUSA INTERVENCIÓN ARRANQUE'],
        "filas": [(1, 'Responsabilidad del cliente'), (2, 'Mala instalación'), (3, 'Término de la vida útil'), (4, 'Asentamiento del terreno'), (5, 'Intervención de terceros identificables'), (6, 'Intervención de terceros no identificables'), (7, 'Factores externos'), (8, 'Otras causas')],
    },
    'PARAM_16': {
        "headers": ['VALOR TIPO CAUSA INTERVENCIÓN DE UD', 'DESCRIPCIÓN TIPO CAUSA INTERVENCIÓN DE UD'],
        "filas": [(1, 'Responsabilidad del cliente'), (2, 'Mala instalación'), (3, 'Término de la vida útil'), (4, 'Asentamiento del terreno'), (5, 'Intervención de terceros identificables'), (6, 'Intervención de terceros no identificables'), (7, 'Factores externos'), (8, 'Otras causas')],
    },
    'PARAM_17': {
        "headers": ['VALOR TIPO CAUSA ROTURA', 'DESCRIPCIÓN TIPO CAUSA ROTURA'],
        "filas": [(1, 'Responsabilidad del cliente'), (2, 'Mala instalación'), (3, 'Término de la vida útil'), (4, 'Asentamiento del terreno'), (5, 'Intervención de terceros identificables'), (6, 'Intervención de terceros no identificables'), (7, 'Factores externos'), (8, 'Otras causas')],
    },
    'PARAM_18': {
        "headers": ['VALOR TIPO CIFRAS AFECTADAS POR EL CARGO', 'DESCRIPCIÓN TIPO CIFRAS AFECTADAS POR EL CARGO'],
        "filas": [(1, 'Irrelevantes'), (2, 'Hasta US$ 2,1 millones al año'), (3, 'De US$ 2,2 a 4,1 millones al año'), (4, 'De US$ 4,2 a 8,3 millones al año'), (5, 'De US$ 8,4 a 16,5 millones al año'), (6, 'De US$ 16,6 a 33,1 millones al año'), (7, 'De US$ 33,2 a 66,1 millones al año'), (8, 'De US$ 66,2 a 132,2 millones al año'), (9, 'De US$ 132,3 a 264,4 millones al año'), (10, 'De US$ 264,5 a 528,9 millones al año'), (11, 'De US$ 529,0 a 1.057,8 millones al año'), (12, 'De US$ 1.057,9 a 2.115,6 millones al año'), (13, 'De US$ 2.115,7 a 4.231,2 millones al año'), (14, 'De US$ 4.231,3 a 8.462,3 millones al año')],
    },
    'PARAM_19': {
        "headers": ['VALOR TIPO CLIENTE', 'DESCRIPCIÓN TIPO CLIENTE'],
        "filas": [(1, 'Residencial'), (2, 'Comercial'), (3, 'Industrial'), (4, 'Institucional'), (5, 'No asociado inmueble')],
    },
    'PARAM_20': {
        "headers": ['VALOR TIPO COMBUSTIBLE O ENERGÍA', 'DESCRIPCIÓN TIPO COMBUSTIBLE O ENERGÍA'],
        "filas": [(1, 'Gasolina'), (2, 'Diesel'), (3, 'GNC'), (4, 'GLP'), (5, 'Kerosene'), (6, 'Electricidad')],
    },
    'PARAM_21': {
        "headers": ['VALOR TIPO COMPLEJIDAD DE GESTIÓN', 'DESCRIPCIÓN TIPO COMPLEJIDAD DE GESTIÓN'],
        "filas": [(1, 'Ejecución de una tarea específica y repetitiva.'), (2, 'Ejecución o supervisión de pocas tareas similares en objetivos y naturaleza.'), (3, 'Ejecución o supervisión de muchas tareas de naturaleza y objetivos similares, generalmente “secciones”'), (4, 'Dirección de un “departamento” que cubre secciones diferentes con objetivos homogéneos.'), (5, 'Dirección de una “gerencia funcional” o dirección de un área que tiene implicaciones significativas en el orden de la planeación general o resultados de la empresa.'), (6, 'Dirección de varias “gerencias funcionales” heterogéneas, lo cual requiere una visión integral de la empresa.'), (7, 'Dirección de todas las funciones que componen una labor empresarial, con miras al cumplimiento de sus objetivos.'), (8, 'Dirección de las funciones tendientes al logro de todos los objetivos estratégicos de varias empresas pertenecientes a un grupo corporativo, con actividades complementarias.'), (9, 'Dirección de las funciones tendientes al logro de todos los objetivos estratégicos de varias empresas pertenecientes a un grupo corporativo, con actividades diversificadas.')],
    },
    'PARAM_22': {
        "headers": ['VALOR TIPO COMPONENTE DE COSTO REPARACIÓN', 'DESCRIPCIÓN TIPO COMPONENTE DE COSTO REPARACIÓN'],
        "filas": [(1, 'Movimiento de Tierra (incluye excavación, rellenos, transporte y retiros de excedentes)'), (2, 'Reparación'), (3, 'Materiales e insumos'), (4, 'Rotura de Pavimento'), (5, 'Reposición de Pavimentos')],
    },
    'PARAM_23': {
        "headers": ['VALOR TIPO COMPONENTE INVERSIÓN TELECONTROL', 'DESCRIPCIÓN TIPO COMPONENTE INVERSIÓN TELECONTROL'],
        "filas": [(1, 'Estación de Telemetría - Sensor'), (2, 'Estación de Telemetría - Actuador'), (3, 'Estación de Telemetría - Controlador'), (4, 'Centro o Sala de Control - Nodo Central de Telemetría'), (5, 'Centro o Sala de Control - Plataforma Scada'), (6, 'Centro o Sala de Control - Sistema de Información Técnico'), (7, 'Centro o Sala de Control - Sistema de Proyección y Sonido')],
    },
    'PARAM_24': {
        "headers": ['VALOR TIPO COMPONENTE REMUNERACIONES', 'DESCRIPCIÓN TIPO COMPONENTE REMUNERACIONES', 'DESCRIPCIÓN'],
        "filas": [(1, 'Sueldo Base', 'Corresponde al sueldo contractual mensual recibido por el trabajador, antes de descuentos previsionales e impuestos.'), (2, 'Gratificación Legal Garantizada', 'Corresponde a la participación de utilidades que las empresas con fines de lucro deben distribuir entre sus trabajadores. Ésta se puede acoger a los siguientes itemes de distribución: 30% de utilidad líquida o 25% del sueldo base con tope de 4,75 IMM.'), (3, 'Gratificación Convencional Garantizada', 'Corresponde a una gratificación estipulada en contratos individuales o en instrumentos colectivos y que no puede ser menor a la gratificación legal.'), (4, 'Asignación de Zona', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por cumplir sus funciones en alguna región específica. Tiene por objeto cubrir los costos de vida que dicha región conlleva.'), (5, 'Asignación y Bonos por Condiciones de Trabajo (del Cargo)', 'Corresponde a la entrega de asignaciones mensuales relacionados al cargo o por desempeñar funciones en un área determinada (p.e. operaciones, mantención, construcción, administrativos, etc.).'), (6, 'Asignación de Título', 'Corresponde a un bono mensual entregado a los trabajadores que acrediten estar titulados, generalmente, de una carrera universitaria o técnica profesional.'), (7, 'Asignación de Responsabilidad', 'Suma de dinero que se agrega mensualmente a ciertos cargos específicos, por considerarse de mayor complejidad o exigencias técnicas.'), (8, 'Asignaciones y Bonos Personales', 'Suma de dinero que se otorga mensualmente por antecedentes de carácter personal de los trabajadores. Por ejemplo, bono por la cantidad de hijos.'), (9, 'Aguinaldo de Fiestas Patrias', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de Fiestas Patrias.'), (10, 'Aguinaldo de Navidad', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de Navidad.'), (11, 'Bonos de Vacaciones', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de vacaciones.'), (12, 'Asignación de Movilización', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por concepto de movilización al lugar de trabajo.'), (13, 'Asignación de Colación', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por concepto de colación (se excluyen los vales de almuerzo y servicios de casino sin costo para los colaboradores).'), (14, 'Asignación de Pérdida de Caja', 'Suma de dinero que se paga mensualmente a los trabajadores que cumplen funciones de cajeros o que con motivo de sus cargos, custodian o reciben dineros. Esta asignación tiene por objetivo cubrir a los colaboradores por eventuales pérdidas a las que pueden verse expuestos.'), (15, 'Asignación Zona Extrema', 'Suma de dinero que se entrega mensualmente (adicional a lo legal) a los trabajadores con domicilio particular y laboral en las regiones XV, I, XI y XII,o en las provincias de Chiloé y Palena.'), (16, 'Asignación Desgaste de Herramientas', 'Suma de dinero que se otorga mensualmente a los trabajadores que desempeñan sus labores con herramientas o instrumentos de su propiedad. Esta asignación tiene por objetivo compensar al colaborador por el desgaste de éstas.'), (17, 'Otros NO Imponibles y Tributables', 'Corresponde a otros ítems de remuneración no afectos a cotizaciones  e impuestos (p.e. viáticos).'), (18, 'Asignaciones de turno', 'Asignación mensual entregada a los trabajadores con una jornada laboral especial a la regular, por ejemplo nocturna. '), (19, 'Incentivos Individuales o Colectivos (Parte fija)', 'Corresponde a un porcentaje del bono de evaluación de desempeño que se entrega de forma garantizada a los trabajadores. '), (20, 'APV (Empresa)', 'Monto entregado mensualmente a los trabajadores como complemento de las cotizaciones obligatorias en la AFP. Tiene por objetivo mejorar los ingresos recibidos por los trabajadores durante su jubilación.'), (21, 'Salud (Empresa)', 'Suma de dinero que se entrega mensualmente a los trabajadores para aportar en la cobertura de gastos incurridos en salud física o psicológica de ellos y sus familias.'), (22, 'Incremento DL 3501', 'Suma de dinero entregada mensualmente con el objetivo de aumentar las cotizaciones previsionales de los trabajadores, manteniendo la renta líquida percibida (previo a la entrada en vigencia del decreto). \n'), (23, 'Depósitos Convenidos (Empresa)', 'Suma de dinero entregada por el empleador, tras un acuerdo con el trabajador, la cual es depositada en la cuenta de capitalización individual y sirve para incrementar el monto de las pensiones. Estos fondos no pueden ser retirados por el colaborador antes de pensionarse.'), (24, 'Asignación Habitacional (Monetaria)', 'Bono entregado mensualmente a los trabajador por concepto de vivienda, generalmente en casos de traslados de región.'), (25, 'Bono Energía y/o Bonificación Agua Potable', 'Suma de dinero que entregan las empresas mensualmente a los trabajadores con limitaciones de acceso o que paguen un alto costo por los servicios de energía, agua potable, servicio de alcantarillado y/o acumulación de aguas servidas cercanas al domicilio particular.'), (26, 'Bonos por Horas extras', 'Asignación que corresponde a los trabajadores que regularmente tienen una jornada laboral superior al máximo legal.'), (27, 'Horas Extras', 'Pago de las horas trabajadas superior a la jornada de trabajo regular.'), (28, 'Diferencias por Licencias Médicas', 'Pago del sueldo no percibido de los trabajadores que se encuentren de licencia médica.'), (29, 'Otros Imponibles y Tributables', 'Corresponde a las restantes sumas en dinero fijas que perciben los trabajadores, bajo cualquier otro concepto y que son imponibles y tributables.'), (30, 'Gratificación Variable (No Garantizada)', 'Corresponde a la gratificación estipulada en contratos individuales o en instrumentos colectivos y que está sujeta a las utilidades de la empresa en su ejercicio comercial.'), (31, 'Participación de Utilidades', 'Participación entregada por concepto de incentivos y con fines de retención.'), (32, 'Bono de Producción', 'Haberes no garantizados y sujetos a objetivos de producción o cumplimiento de indicadores de gestión del área o departamento.'), (33, 'Comisiones (Ventas, Cobranzas At. Clientes, etc.)', 'Haberes no garantizados y sujetos al cumplimiento de objetivos comerciales o de venta.'), (34, 'Bonos e Incentivos Generales (del negocio)', 'Haberes no garantizados y sujetos al cumplimiento de objetivos y desempeño de la empresa en su totalidad. Estos bonos se entregan anual o semestralmente.'), (35, 'Bonos e Incentivos Individuales', 'Bono entregado a los trabajadores con una evaluación de desempeño individual buena o sobresaliente, o sobre una nota, dentro de una escala definida por la empresa. Esta asignación, generalmente, se entrega anualmente.'), (36, 'Bonos e Incentivos de Mediano y Largo Plazo', 'Bonos entregados, principalmente a ejecutivos, con el objetivo de aumentar la probabilidad de retención en los cargos considerados clave para las operaciones de la empresa. ')],
    },
    'PARAM_25': {
        "headers": ['VALOR TIPO CONCEPTO COMPENSACIONES RRHH', 'DESCRIPCIÓN TIPO CONCEPTO COMPENSACIONES RRHH', 'DESCRIPCIÓN'],
        "filas": [(1, 'Sueldo Base', 'Corresponde al sueldo contractual mensual recibido por el trabajador, antes de descuentos previsionales e impuestos.'), (2, 'Gratificación Legal Garantizada', 'Corresponde a la participación de utilidades que las empresas con fines de lucro deben distribuir entre sus trabajadores. Ésta se puede acoger a los siguientes itemes de distribución: 30% de utilidad líquida o 25% del sueldo base con tope de 4,75 IMM.'), (3, 'Gratificación Convencional Garantizada', 'Corresponde a una gratificación estipulada en contratos individuales o en instrumentos colectivos y que no puede ser menor a la gratificación legal.'), (4, 'Asignación de Zona', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por cumplir sus funciones en alguna región específica. Tiene por objeto cubrir los costos de vida que dicha región conlleva.'), (5, 'Asignación y Bonos por Condiciones de Trabajo (del Cargo)', 'Corresponde a la entrega de asignaciones mensuales relacionados al cargo o por desempeñar funciones en un área determinada (p.e. operaciones, mantención, construcción, administrativos, etc.).'), (6, 'Asignación de Título', 'Corresponde a un bono mensual entregado a los trabajadores que acrediten estar titulados, generalmente, de una carrera universitaria o técnica profesional.'), (7, 'Asignación de Responsabilidad', 'Suma de dinero que se agrega mensualmente a ciertos cargos específicos, por considerarse de mayor complejidad o exigencias técnicas.'), (8, 'Asignaciones y Bonos Personales', 'Suma de dinero que se otorga mensualmente por antecedentes de carácter personal de los trabajadores. Por ejemplo, bono por la cantidad de hijos.'), (9, 'Aguinaldo de Fiestas Patrias', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de Fiestas Patrias.'), (10, 'Aguinaldo de Navidad', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de Navidad.'), (11, 'Bonos de Vacaciones', 'Corresponde a la suma de dinero percibida anualmente por el trabajador, por concepto de vacaciones.'), (12, 'Asignación de Movilización', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por concepto de movilización al lugar de trabajo.'), (13, 'Asignación de Colación', 'Corresponde a la suma de dinero percibida mensualmente por el trabajador, por concepto de colación (se excluyen los vales de almuerzo y servicios de casino sin costo para los colaboradores).'), (14, 'Asignación de Pérdida de Caja', 'Suma de dinero que se paga mensualmente a los trabajadores que cumplen funciones de cajeros o que con motivo de sus cargos, custodian o reciben dineros. Esta asignación tiene por objetivo cubrir a los colaboradores por eventuales pérdidas a las que pueden verse expuestos.'), (15, 'Asignación Zona Extrema', 'Suma de dinero que se entrega mensualmente (adicional a lo legal) a los trabajadores con domicilio particular y laboral en las regiones XV, I, XI y XII,o en las provincias de Chiloé y Palena.'), (16, 'Asignación Desgaste de Herramientas', 'Suma de dinero que se otorga mensualmente a los trabajadores que desempeñan sus labores con herramientas o instrumentos de su propiedad. Esta asignación tiene por objetivo compensar al colaborador por el desgaste de éstas.'), (17, 'Otros NO Imponibles y Tributables', 'Corresponde a otros ítems de remuneración no afectos a cotizaciones  e impuestos (p.e. viáticos).'), (18, 'Asignaciones de turno', 'Asignación mensual entregada a los trabajadores con una jornada laboral especial a la regular, por ejemplo nocturna. '), (19, 'Incentivos Individuales o Colectivos (Parte fija)', 'Corresponde a un porcentaje del bono de evaluación de desempeño que se entrega de forma garantizada a los trabajadores. '), (20, 'APV (Empresa)', 'Monto entregado mensualmente a los trabajadores como complemento de las cotizaciones obligatorias en la AFP. Tiene por objetivo mejorar los ingresos recibidos por los trabajadores durante su jubilación.'), (21, 'Salud (Empresa)', 'Suma de dinero que se entrega mensualmente a los trabajadores para aportar en la cobertura de gastos incurridos en salud física o psicológica de ellos y sus familias.'), (22, 'Incremento DL 3501', 'Suma de dinero entregada mensualmente con el objetivo de aumentar las cotizaciones previsionales de los trabajadores, manteniendo la renta líquida percibida (previo a la entrada en vigencia del decreto). \n'), (23, 'Depósitos Convenidos (Empresa)', 'Suma de dinero entregada por el empleador, tras un acuerdo con el trabajador, la cual es depositada en la cuenta de capitalización individual y sirve para incrementar el monto de las pensiones. Estos fondos no pueden ser retirados por el colaborador antes de pensionarse.'), (24, 'Asignación Habitacional (Monetaria)', 'Bono entregado mensualmente a los trabajador por concepto de vivienda, generalmente en casos de traslados de región.'), (25, 'Bono Energía y/o Bonificación Agua Potable', 'Suma de dinero que entregan las empresas mensualmente a los trabajadores con limitaciones de acceso o que paguen un alto costo por los servicios de energía, agua potable, servicio de alcantarillado y/o acumulación de aguas servidas cercanas al domicilio particular.'), (26, 'Bonos por Horas extras', 'Asignación que corresponde a los trabajadores que regularmente tienen una jornada laboral superior al máximo legal.'), (27, 'Horas Extras', 'Pago de las horas trabajadas superior a la jornada de trabajo regular.'), (28, 'Diferencias por Licencias Médicas', 'Pago del sueldo no percibido de los trabajadores que se encuentren de licencia médica.'), (29, 'Otros Imponibles y Tributables', 'Corresponde a las restantes sumas en dinero fijas que perciben los trabajadores, bajo cualquier otro concepto y que son imponibles y tributables.'), (30, 'Gratificación Variable (No Garantizada)', 'Corresponde a la gratificación estipulada en contratos individuales o en instrumentos colectivos y que está sujeta a las utilidades de la empresa en su ejercicio comercial.'), (31, 'Participación de Utilidades', 'Participación entregada por concepto de incentivos y con fines de retención.'), (32, 'Bono de Producción', 'Haberes no garantizados y sujetos a objetivos de producción o cumplimiento de indicadores de gestión del área o departamento.'), (33, 'Comisiones (Ventas, Cobranzas At. Clientes, etc.)', 'Haberes no garantizados y sujetos al cumplimiento de objetivos comerciales o de venta.'), (34, 'Bonos e Incentivos Generales (del negocio)', 'Haberes no garantizados y sujetos al cumplimiento de objetivos y desempeño de la empresa en su totalidad. Estos bonos se entregan anual o semestralmente.'), (35, 'Bonos e Incentivos Individuales', 'Bono entregado a los trabajadores con una evaluación de desempeño individual buena o sobresaliente, o sobre una nota, dentro de una escala definida por la empresa. Esta asignación, generalmente, se entrega anualmente.'), (36, 'Bonos e Incentivos de Mediano y Largo Plazo', 'Bonos entregados, principalmente a ejecutivos, con el objetivo de aumentar la probabilidad de retención en los cargos considerados clave para las operaciones de la empresa. '), (37, 'Bienestar: Bonificación, costo empresa', 'Aportes monetarios que realiza la empresa a un fondo común creado por sindicatos o del área de capital humano para incentivar o apoyar en el financiamiento de materias de bienestar para los trabajadores. Abarca a beneficios de diversa índole, tales como en temas de salud, educación, previsionales, etc.'), (38, 'Seguro Complementario Salud', 'Convenio que realiza el empleador con una empresa aseguradora, con el fin de mitigar el gasto que se genera en prestaciones médicas y de salud.'), (39, 'Seguro de Vida', 'Seguros contratados por las empresas para proteger el patrimonio de los trabajadores en caso de fallecimiento.'), (40, 'Aporte Plan de Isapre', 'Monto entregado mensualmente como aporte a la cotización legal a los afiliados en isapres, con la finalidad de que el trabajador contrate un plan de salud de mayor valor. '), (41, 'Beneficio Salud Empresa', 'Desarrollo de convenios específicos o servicios administrados por la empresa para el cuidado y prevención de la salud de sus trabajadores. Puede incluir a chequeos médicos, dentales, nutricionista, etc.'), (42, 'Otros Seguros (Costo prima empresa)', 'Entrega de otros seguros para los trabajadores, tales como seguro contra accidentes, seguro dental o similares.'), (43, 'Indemnizaciones por sobre lo legal (Provisiones)', 'Derecho laboral por despido que la empresa debe dar al colaborador, en proporción a la cantidad de años que haya trabajado, sin aplicar los topes legales de 90 UF y/o 11 años'), (44, 'Regalos por Diversos Conceptos: navidad, cumpleaños, etc.', 'Entrega de un producto o servicio por parte de la empresa, por motivo de cumpleaños del trabajador u otro acontecimiento.'), (45, 'Pago de paquetes vacacionales', 'Corresponde a la entrega de beneficios por motivo de vacaciones o descanso de los trabajadores, el cual puede incluir el traslado, alojamiento, alimentación y/o visitas programadas del lugar de destino.'), (46, 'Paquete de Fiestas Patrias', 'Corresponde a la entrega de una canasta de productos a los trabajadores por la celebración de Fiestas Patrias (p.e. carnes, cecinas, vino, verduras, etc.).'), (47, 'Fiesta de F. Patrias', 'Fiesta o celebración organizada por la empresa, para sus trabajadores, por motivo de Fiestas Patrias.'), (48, 'Paquete de Navidad', 'Corresponde a la entrega de una canasta de productos a los trabajadores por la celebración de Navidad (p.e. pavo, galletas, vino, verduras, etc.).'), (49, 'Regalos de Navidad (Personal, Niños, Etc.)', 'Entrega de un regalo (producto o servicio), por parte de la empresa, a los trabajadores por motivo de Navidad. '), (50, 'Fiesta de Navidad', 'Celebración realizada por la empresa para los trabajadores o también de sus familiares, por concepto de Navidad.'), (51, 'Fiesta de fin de año y/o Aniversario', 'Celebración realizada por la empresa para los trabajadores para conmemorar el aniversario de la organización o para festejar el fin del año calendario.'), (52, 'Consumos Habitacionales', 'Corresponde al pago total o parcial de las cuentas de servicios básicos (luz, agua, gas u otros) de los trabajadores.'), (53, 'Asigna Casa : (Valor estim. del beneficio)', 'Corresponde a la entrega de una vivienda, por un tiempo determinado, para el uso del trabajador o también de su familia.'), (54, 'Asigna Vehículo : (Valor estimado del beneficio)', 'Corresponde a la entrega de un vehículo para facilitar el traslado del trabajador a visitas de clientes u otras actividades relacionadas a sus funciones.'), (55, 'Reembolso uso vehículo propio (valor estim.)', 'Corresponde al pago de los costos incurridos debido al traslado del trabajador a visitas de clientes u otras actividades relacionadas a sus funciones (p.e. bencina, peajes, tags, etc.).'), (56, 'Pago de estacionamiento', 'Corresponde a un pago incurrido por la empresa en proveer de estacionamientos a sus colaboradores o una asignación en dinero para que el trabajador costee uno por cuenta propia.'), (57, 'Traslado, Vehículo valorizado y pasajes', 'Corresponde a la entrega o reembolso de los costos incurridos por movilización de viajes por motivos de trabajo. Se diferencia de otras asignaciones o reembolsos por traslados de menores distancias geográficas.'), (58, 'Salud, Vida, Previsional (si no están en Seguros)', 'Corresponde a la entrega de otros beneficios de salud o protección que no estén especificados anteriormente. '), (59, 'Personales: Escolaridades (si no están en Bonos Estudios)', 'Corresponde a la entrega de otros bonos como apoyo al financiamiento de la educación de los hijos de los trabajadores, tales como los bonos de excelencia académica o similares.'), (60, 'Beneficios Previsionales no Obligatorios (si no están en B. Previsionales)', 'Corresponde a la entrega de otros beneficios previsionales que no estén especificados anteriormente. '), (61, 'Vales, lo que no paga el trabajador (cheque R. / No constituye Renta)', 'Corresponde a la entrega de vales de colación o cheques de restaurant para que sean utilizados para la colación del trabajador.'), (62, 'Casino, lo que no paga el trabajador (No constituye Renta)', 'Servicio de alimentación ofrecido por la empresa o proveedor, sin costo para los trabajadores.'), (63, 'Bonos por Negociación Colectiva (Sólo BTN) Mensualizado', 'Monto entregado mensualmente por la empresa a los trabajadores para acordar la finalización de un proceso de negociación colectiva específico. '), (64, 'Otros Bonos por Negociación Colectiva', 'Corresponde a la entrega de otros bonos por acuerdos tras la finalización de negociaciones colectivas, tales como bono por práctica operacional o bono por renovación de jornada.'), (65, 'Bonos Escolares (Cargas)', 'Corresponde a la entrega de un bono para apoyar el financiamiento de la educación de los hijos de los trabajadores. Generalmente abarca a la educación básica y media (científico humanista o técnica profesional).'), (66, 'Becas de Estudio (Trabajador)', 'Financiamiento que entrega la empresa a los trabajadores que realizan cursos de perfeccionamiento (generalmente estudios de postgrado).'), (67, 'Entrega de Vestuario (No obligatorio)', 'Entrega de vestimenta de trabajo, superior a lo exigido por la ley. '), (68, 'Entrega de Vestuario Obligatorio (Arts. Seguridad y de Trabajo)', 'Entrega de vestimenta de trabajo para el desempeño y resguardo de la seguridad del trabajador durante la realización de las funciones propias de su cargo. '), (69, 'Bonos por matrimonio (valor total costo empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos en el caso de contraer matrimonio y/o acuerdo de unión civil del trabajador.'), (70, 'Bonos por nacimiento (valor total costo empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos por nacimiento por hijo del trabajador.'), (71, 'Cuota mortuoria pagada (costo Empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos por fallecimiento de familiares directos del trabajador. Los más comunes son en caso del fallecimiento de hijo o cónyuge, como también del padre o madre del colaborador.'), (72, 'Sala cuna (valor total costo Sala Cuna, Reembolso o Bono)', 'Asignación en dinero que entrega la empresa (adicional a lo legal) por concepto de sala cuna a sus trabajadores, por cada hijo menor a 24 meses de edad.'), (73, 'Pago de Club Privado u Otros', 'Corresponde a convenios suscritos por la empresa con entidades de recreación para los trabajadores.'), (74, 'Otros Beneficios Valorizables que no son Remuneración', 'Corresponde a cualquier otro beneficio no monetario que no haya sido especificado anteriormente.')],
    },
    'PARAM_26': {
        "headers": ['VALOR TIPO CONOCIMIENTOS', 'DESCRIPCIÓN TIPO CONOCIMIENTOS'],
        "filas": [(1, 'Enseñanza media completa clásica, comercial o técnica'), (2, 'Carreras intermedias de 2 años o estudios específicos, de más de un año de duración, que requieren título profesional'), (3, 'Carrera profesional completa, sin grado académico, duración 4 años'), (4, 'Carrera profesional completa, con grado de licenciado, duración mínima 5 años'), (5, 'Estudios de especialización o magister'), (6, 'Doctorado o estudios de valor e intensidad equivalentes')],
    },
    'PARAM_27': {
        "headers": ['VALOR TIPO CONSULTORÍA INVERSIÓN TIC', 'DESCRIPCIÓN TIPO CONSULTORÍA INVERSIÓN TIC'],
        "filas": [(1, ' Consultoría de Levantamiento de Procesos'), (2, ' Consultoría de Arquitectura'), (3, ' Consultoría de Selección de Solución'), (4, ' Consultoría de Contrato'), (5, ' Consultoría de Contraparte'), (6, ' Consultoría de Gestión del Cambio'), (7, ' Otra')],
    },
    'PARAM_28': {
        "headers": ['VALOR TIPO CONSUMO BÁSICO', 'DESCRIPCIÓN TIPO CONSUMO BÁSICO'],
        "filas": [(1, 'Agua potable y alcantarillado'), (2, 'Energía Eléctrica'), (3, 'Gas'), (4, 'Gastos Comunes')],
    },
    'PARAM_29': {
        "headers": ['VALOR TIPO CONTACTOS', 'DESCRIPCIÓN TIPO CONTACTOS'],
        "filas": [(1, 'Ocasional: Por lo menos una vez al mes'), (2, 'Frecuente: Por lo menos una vez a la semana'), (3, 'Continuo: Permanentemente')],
    },
    'PARAM_30': {
        "headers": ['VALOR TIPO CONTRATO', 'DESCRIPCIÓN TIPO CONTRATO'],
        "filas": [(1, 'Obras'), (2, 'Servicios prestados'), (3, 'Servicios recibidos'), (4, 'Suministros e insumos')],
    },
    'PARAM_31': {
        "headers": ['VALOR TIPO CONTRATO PERSONAS', 'DESCRIPCIÓN TIPO CONTRATO PERSONAS'],
        "filas": [(1, 'Honorarios'), (2, 'Jornada Parcial'), (3, 'Planta o permanente')],
    },
    'PARAM_32': {
        "headers": ['VALOR TIPO CONTROL DE CALIDAD AP', 'DESCRIPCIÓN TIPO CONTROL DE CALIDAD AP'],
        "filas": [(1, 'Nch 409'), (2, 'Parámetros Críticos'), (3, 'Control de Procesos'), (4, 'Control Extranormativo o Normativo Adicional')],
    },
    'PARAM_33': {
        "headers": ['VALOR TIPO CONTROL DE CALIDAD AS', 'DESCRIPCIÓN TIPO CONTROL DE CALIDAD AS'],
        "filas": [(1, 'Control Normativo-DS 90'), (2, 'Control Normativo-DS RCA'), (3, 'Control Normativo-DS 4'), (4, 'Control de Procesos'), (5, 'Control Extranormativo o Normativo adicional')],
    },
    'PARAM_34': {
        "headers": ['VALOR TIPO CUENTA LIBRO MAYOR', 'DESCRIPCIÓN TIPO CUENTA LIBRO MAYOR'],
        "filas": [(1, 'Ingresos'), (2, 'Costos y Gastos'), (3, 'Activos'), (4, 'Pasivos'), (5, 'Patrimonio')],
    },
    'PARAM_35': {
        "headers": ['VALOR TIPO DERECHOS SOBRE EL ACTIVO O BIEN', 'DESCRIPCIÓN TIPO DERECHOS SOBRE EL ACTIVO O BIEN'],
        "filas": [(1, 'Propio'), (2, 'Arrendado'), (3, 'Leasing'), (4, 'Comodato'), (5, 'Otro')],
    },
    'PARAM_36': {
        "headers": ['VALOR TIPO DOCUMENTO', 'DESCRIPCIÓN TIPO DOCUMENTO'],
        "filas": [(1, 'Boletas'), (2, 'Facturas'), (3, 'Avisos de Cortes Programados'), (4, 'Avisos de Inicio Período Sobreconsumo'), (5, 'Avisos de Término Período Sobreconsumo'), (6, 'Avisos de Notificación de Exceso de Consumo'), (7, 'Subsidio'), (8, 'Nueva Tarifa'), (9, 'Libros de Venta'), (10, 'Otros documentos')],
    },
    'PARAM_37': {
        "headers": ['VALOR TIPO DOCUMENTO LICITACIÓN', 'DESCRIPCIÓN TIPO DOCUMENTO LICITACIÓN'],
        "filas": [(1, 'Bases Técnicas'), (2, 'Oferta Económica'), (3, 'Oferta Técnica'), (4, 'Otros')],
    },
    'PARAM_38': {
        "headers": ['VALOR TIPO DOMINIO DE IDIOMA', 'DESCRIPCIÓN TIPO DOMINIO DE IDIOMA'],
        "filas": [(1, 'Básico'), (2, 'Medio'), (3, 'Avanzado - Bilingüe')],
    },
    'PARAM_39': {
        "headers": ['VALOR TIPO ZONAL', 'DESCRIPCIÓN TIPO ELEMENTO OBSTRUCTIVO'],
        "filas": [(1, 'Acumulación de arena'), (2, 'Intrusión de raíces'), (3, 'Depósitos de grasas'), (4, 'Piedras o rocas'), (5, 'Otros elementos sólidos')],
    },
    'PARAM_40': {
        "headers": ['VALOR TIPO EQUIPO', 'DESCRIPCIÓN TIPO EQUIPO'],
        "filas": [(1, 'Equipos de Lavado'), (2, 'Equipos de Desobstrucción'), (3, 'Equipos de Detección de Fugas'), (4, 'Equipos de Laboratorio'), (5, 'Otros Equipos')],
    },
    'PARAM_41': {
        "headers": ['VALOR TIPO ESPECIALIDAD', 'DESCRIPCIÓN TIPO ESPECIALIDAD'],
        "filas": [(1, 'Eléctrica'), (2, 'Electromecánica'), (3, 'Obras Civiles'), (4, 'Mecánica'), (5, 'Instrumentista')],
    },
    'PARAM_42': {
        "headers": ['VALOR TIPO ESTADO PROYECTO', 'DESCRIPCIÓN TIPO ESTADO PROYECTO'],
        "filas": [(1, 'Activo'), (2, 'Postergado. No hay actividades en ejecución, pero aún no se cierra'), (3, 'Cerrado. Culminaron todas las actividades esperadas'), (4, 'Cerrado de forma incompleta. No se logró ejecutar la totalidad de las actividades')],
    },
    'PARAM_43': {
        "headers": ['VALOR TIPO ESTUDIOS', 'DESCRIPCIÓN TIPO ESTUDIOS'],
        "filas": [(1, 'Análisis Financieros'), (2, 'Estudios de Prefactibilidad'), (3, 'Estudios Hidrológicos e Hidrogeológicos'), (4, 'Estudios de Diseño de Ingeniería (Topografía u otros)'), (5, 'Estudios de Servidumbres'), (6, 'Estudios y Declaraciones de Impacto Ambiental'), (7, 'Estudios Legales'), (8, 'Otros Estudios')],
    },
    'PARAM_44': {
        "headers": ['VALOR TIPO EXPERIENCIA', 'DESCRIPCIÓN TIPO EXPERIENCIA'],
        "filas": [(1, 'Hasta un año'), (2, 'Dos años'), (3, 'De 3 años a 5 años'), (4, 'De 6 años a 8 años'), (5, 'de 9 a 11 años'), (6, 'De 12 años a 14 años')],
    },
    'PARAM_45': {
        "headers": ['VALOR TIPO FACTURA', 'DESCRIPCIÓN TIPO FACTURA'],
        "filas": [(1, 'Obras'), (2, 'Servicios prestados'), (3, 'Servicios recibidos'), (4, 'Suministros e insumos')],
    },
    'PARAM_46': {
        "headers": ['VALOR TIPO FINANCIAMIENTO PROYECTO', 'DESCRIPCIÓN TIPO FINANCIAMIENTO PROYECTO'],
        "filas": [(1, 'Financiamiento Propio'), (2, 'Obras Aportadas por Terceros - Privados'), (3, 'Obras Aportadas por Terceros - Fisco (excluidas municipalidades)'), (4, 'Obras Aportadas por Terceros - Municipal'), (5, 'Empresa Relacionada'), (6, 'Otro Tipo de Financiamiento')],
    },
    'PARAM_47': {
        "headers": ['VALOR TIPO FUNCIÓN', 'DESCRIPCIÓN TIPO FUNCIÓN'],
        "filas": [(1, 'Sedimentación'), (2, 'Control de bulking'), (3, 'Control de algas'), (4, 'Desinfección'), (5, 'Floculación'), (6, 'Estabilización química e higienización de lodos'), (7, 'Tratamiento de aire y eliminación de olores'), (8, 'Fluoruración'), (9, 'Otra')],
    },
    'PARAM_48': {
        "headers": ['VALOR TIPO GRUPO COMPONENTE', 'DESCRIPCIÓN TIPO GRUPO COMPONENTE'],
        "filas": [(1, 'OOCC'), (2, 'TUBERIAS Y CONDUCCIONES'), (3, 'EQUIPOS'), (4, 'EQUIPOS ELÉCTRICOS'), (5, 'COSTOS INDIRECTOS')],
    },
    'PARAM_49': {
        "headers": ['VALOR TIPO HABILIDAD LIDERAZGO', 'DESCRIPCIÓN TIPO HABILIDAD LIDERAZGO'],
        "filas": [(1, 'Normal habilidad para prestar servicio y obtener cooperación (cuando se trata con asuntos rutinarios).'), (2, 'Más de la normal habilidad para prestar servicio y obtener cooperación.'), (3, 'Buena habilidad para negociar y obtener cooperación (entrevistas, discusiones en grupo, investigaciones de mercado, venta y compra, representando la empresa).'), (4, 'Muy buena habilidad para negociar y obtener cooperación.'), (5, 'Máxima habilidad en negociación y en obtener cooperación (negociaciones al más alto nivel acerca de los aspectos más vitales para la empresa con otras organizaciones y autoridades).')],
    },
    'PARAM_50': {
        "headers": ['VALOR TIPO HERRAMIENTA O EQUIPO', 'DESCRIPCIÓN TIPO HERRAMIENTA O EQUIPO'],
        "filas": [(1, 'Adaptadores de husillos'), (2, 'Alicate Caimán 10"'), (3, 'Alicate Caimán 8"'), (4, 'Alicate corriente'), (5, 'Alicate cortante 6"'), (6, 'Alicate crimpeadora'), (7, 'Alicate punta 7” con aislamiento'), (8, 'Alicate universal 12” con aislamiento'), (9, 'Alicate universal 9” con aislamiento'), (10, 'Alicates Seguros Segeer'), (11, 'Amperímetro de Tenaza'), (12, 'Aprieta Terminales'), (13, 'Aprieta Terminales Hidráulico'), (14, 'Atornillador de paleta de 5 mm'), (15, 'Atornillador de paleta de 6 mm'), (16, 'Atornillador de paleta de 8 mm'), (17, 'Atornillador phillips 5 x 100'), (18, 'Atornillador phillips 6 x 200'), (19, 'Caimán'), (20, 'Caja de herramientas PVC'), (21, 'Calibrador de Lazo'), (22, 'Cautín'), (23, 'Chuzo'), (24, 'Colorímetro'), (25, 'Combo de 4lb'), (26, 'Combo de 6lb'), (27, 'Combo de 8lb'), (28, 'Cuchillo pela cables'), (29, 'Diablitos'), (30, 'Equipo de Desobstrucción'), (31, 'Equipo de Detección de Fugas - Correlación Acústica'), (32, 'Equipo de Detección de Fugas - Gas Trazador'), (33, 'Equipo de Laboratorio'), (34, 'Equipo de Lavado de Redes'), (35, 'Esmeril Angular 4,5'), (36, 'Esmeril Angular a Batería'), (37, 'Huincha de medir de 7,5 metros'), (38, 'Indicador de Fases Eléctricas'), (39, 'Juego atornilladores'), (40, 'Juego de Cinceles'), (41, 'Juego de dados 6 a 20 mm'), (42, 'Juego de Limas'), (43, 'Juego de llaves punta corona ¼ a 1 ¼”'), (44, 'Juego de llaves punta corona 6-32 mm'), (45, 'Juego llaves Allen cardánicas 1/8 a ½”'), (46, 'Juego llaves Allen cardánicas 1-10 mm'), (47, 'Juego Sacabocados'), (48, 'Lámpara a Batería'), (49, 'Linterna recargable'), (50, 'Llave ajustable 6 - 8 - 10 - 12 "'), (51, 'Llave de cadena'), (52, 'Llave francesa de 12”'), (53, 'Llave francesa de 14"'), (54, 'Llave Saca Filtro'), (55, 'Llave stillson de 12”'), (56, 'Llave Stilson 18"'), (57, 'Llave tubo 10 "'), (58, 'Manilla especial para sacar tapa de grifo'), (59, 'Manillas de accionamiento de válvulas (distintos diámetros)'), (60, 'Manómetro de 1 a 100 mca'), (61, 'Marco de sierra'), (62, 'Martillo Antirrebote'), (63, 'Martillo de Peña'), (64, 'Medidor de Aislación'), (65, 'Medidor de mano para cloro residual'), (66, 'Medidor de mano para Ph'), (67, 'Medidor de mano para turbiedad'), (68, 'Multitester eléctrico'), (69, 'Pala'), (70, 'Pala punta huevo'), (71, 'Pistola Aire Caliente 1600W'), (72, 'Pistola Calafatera'), (73, 'Pistola Engrasadora Lagh 400'), (74, 'Soplador de aire frío'), (75, 'Taladro a Batería'), (76, 'Taladro eléctrico de ½” con percutor'), (77, 'Tenaza universal 1”'), (78, 'Tester de RED TCP'), (79, 'Vibrómetro'), (80, 'Otras herramientas o equipos')],
    },
    'PARAM_51': {
        "headers": ['VALOR TIPO IMPRESORAS Y PLOTTERS', 'DESCRIPCIÓN TIPO IMPRESORAS Y PLOTTERS'],
        "filas": [(1, 'Impresora Laser Monocromática'), (2, 'Impresora Laser Color'), (3, 'Impresora Inyección de Tinta Monocromática'), (4, 'Impresora Inyección de Tinta Color'), (5, 'Impresora 3D'), (6, 'Plotter Laser Monocromático'), (7, 'Plotter Laser Color'), (8, 'Plotter  Inyección de Tinta Monocromático'), (9, 'Plotter Inyección de Tinta Color')],
    },
    'PARAM_52': {
        "headers": ['VALOR TIPO INMUEBLE', 'DESCRIPCIÓN TIPO INMUEBLE'],
        "filas": [(1, 'Oficina Corporativa'), (2, 'Oficina Zonal'), (3, 'Oficinas de Atención de Público'), (4, 'Recinto Productivo'), (5, 'Bodega'), (6, 'Estacionamiento')],
    },
    'PARAM_53': {
        "headers": ['VALOR TIPO INSPECCIÓN TELEVISIVA', 'DESCRIPCIÓN TIPO INSPECCIÓN TELEVISIVA'],
        "filas": [(1, 'Tractor de oruga o ruedas con CCTV o HD'), (2, 'Balsa con CCTV o HD'), (3, 'Inspección Directa (Personal con cámara)')],
    },
    'PARAM_54': {
        "headers": ['VALOR TIPO INSTANCIA CORTE', 'DESCRIPCIÓN TIPO INSTANCIA CORTE'],
        "filas": [(1, 'Visita corte'), (2, 'Primera'), (3, 'Segunda Vereda sin Pavimento'), (4, 'Segunda Vereda con Pavimento'), (5, 'Segunda Calzada sin Pavimiento'), (6, 'Segunda Calzada con Pavimiento')],
    },
    'PARAM_55': {
        "headers": ['VALOR TIPO LLAMADA O CONTACTO', 'DESCRIPCIÓN TIPO LLAMADA O CONTACTO'],
        "filas": [(1, 'Llamadas Recibidas'), (2, 'Llamadas Efectuadas'), (3, 'Contactos vía Web Gestionados'), (4, 'Contactos vía Carta Gestionados'), (5, 'Contactos vía Correo Eléctrónicos'), (6, 'Contactos vía Redes Sociales'), (7, 'Contactos vía chat')],
    },
    'PARAM_56': {
        "headers": ['VALOR TIPO MANTENCIÓN GRIFOS', 'DESCRIPCIÓN TIPO MANTENCIÓN GRIFOS'],
        "filas": [(1, 'Inspección'), (2, 'Limpieza y Pintura'), (3, 'Reparación piezas especiales (ej. suministro y cambio de cabeza o de tapa gorro)'), (4, 'Renovación piezas especiales')],
    },
    'PARAM_57': {
        "headers": ['VALOR TIPO MANTENCIÓN RECINTOS', 'DESCRIPCIÓN TIPO MANTENCIÓN RECINTOS'],
        "filas": [(1, 'Desmalezado'), (2, 'Control de plagas'), (3, 'Pintura'), (4, 'Mantención de jardines')],
    },
    'PARAM_58': {
        "headers": ['VALOR TIPO MATERIAL E INSUMO DE LABORATORIO', 'DESCRIPCIÓN TIPO MATERIAL E INSUMO DE LABORATORIO'],
        "filas": [(1, 'Productos Quimicos'), (2, 'Repuestos de Equipos'), (3, 'Herramientas'), (4, 'Articulos de Escritorio'), (5, 'Ropa y artefactos de protección y seguridad')],
    },
    'PARAM_59': {
        "headers": ['VALOR TIPO MATERIAL E INSUMOS DE OFICINA Y BODEGA', 'DESCRIPCIÓN TIPO MATERIAL E INSUMOS DE OFICINA Y BODEGA'],
        "filas": [(1, 'Materiales de Oficina Varios'), (2, 'Materiales de Oficina Papelería'), (3, 'Insumos Computacionales'), (4, 'Materiales e insumos de bodega')],
    },
    'PARAM_60': {
        "headers": ['VALOR TIPO MATERIAL TUBERÍA', 'DESCRIPCIÓN TIPO MATERIAL TUBERÍA'],
        "filas": [(1, 'PVC'), (2, 'HDPE'), (3, 'Asbesto cemento'), (4, 'Acero'), (5, 'Cobre'), (6, 'Fierro Fundido'), (7, 'Cemento comprimido')],
    },
    'PARAM_61': {
        "headers": ['VALOR TIPO MOBILIARIO', 'DESCRIPCIÓN TIPO MOBILIARIO'],
        "filas": [(1, 'Mobiliario de Oficina'), (2, 'Mobiliario de Bodega'), (3, 'Mobiliario de Laboratorio'), (4, 'Mobiliario de Recintos Productivos'), (5, 'Otros Mobiliarios')],
    },
    'PARAM_62': {
        "headers": ['VALOR TIPO MODALIDAD PAGO', 'DESCRIPCIÓN TIPO MODALIDAD PAGO'],
        "filas": [(1, 'Remunerado'), (2, 'Honorarios'), (3, 'Servicio Externo')],
    },
    'PARAM_63': {
        "headers": ['VALOR TIPO MODALIDAD SERVICIO MANTENCIÓN', 'DESCRIPCIÓN TIPO MODALIDAD SERVICIO MANTENCIÓN'],
        "filas": [(1, 'Suministro de personal sin provisión de materiales'), (2, 'Completo, incluyendo provisión de materiales')],
    },
    'PARAM_64': {
        "headers": ['VALOR TIPO MONEDA', 'DESCRIPCIÓN TIPO MONEDA'],
        "filas": [(1, 'CLP (Pesos chilenos)'), (2, 'UF (Unidades de Fomento)'), (3, 'USD (Dólares americanos)'), (4, 'EURO'), (5, 'Otra')],
    },
    'PARAM_65': {
        "headers": ['VALOR TIPO MONITOREO', 'DESCRIPCIÓN TIPO MONITOREO'],
        "filas": [(1, 'Efluente'), (2, 'Cuerpo de agua'), (3, 'Sedimentos'), (4, 'Bentos submareal e intermareal')],
    },
    'PARAM_66': {
        "headers": ['VALOR TIPO MOTIVO EGRESO', 'DESCRIPCIÓN TIPO MOTIVO EGRESO'],
        "filas": [(1, 'Mutuo acuerdo de las partes'), (2, 'Renuncia del trabajador con aviso previo'), (3, 'Muerte del trabajador'), (4, 'Vencimiento del plazo convenido en el contrato'), (5, 'Conclusión del trabajo o servicio que dio origen al contrato'), (6, 'Caso fortuito o fuerza mayor'), (7, 'Conducta indebida de carácter grave (artículo 160 del Código del Trabajo)'), (8, 'Término del contrato por Necesidades de la Empresa'), (9, 'Empleador sometido mediante resolución judicial a procedimiento concursal de liquidación de sus bienes')],
    },
    'PARAM_67': {
        "headers": ['VALOR TIPO MUESTREO AS', 'DESCRIPCIÓN TIPO MUESTREO AS'],
        "filas": [(1, 'P'), (2, 'C-8'), (3, 'C-24')],
    },
    'PARAM_68': {
        "headers": ['VALOR TIPO NIVEL DE TERCERIZACIÓN', 'DESCRIPCIÓN TIPO NIVEL DE TERCERIZACIÓN'],
        "filas": [(1, 'Interno'), (2, 'Exterrno ')],
    },
    'PARAM_69': {
        "headers": ['VALOR TIPO OPERACIÓN DE EQUIPO', 'DESCRIPCIÓN TIPO OPERACIÓN DE EQUIPO'],
        "filas": [(1, 'Pruebas y mantención preventiva'), (2, 'Operación normal'), (3, 'Fallas o emergencias')],
    },
    'PARAM_70': {
        "headers": ['VALOR TIPO OPERACIÓN DE INFRAESTRUCTURA', 'DESCRIPCIÓN TIPO OPERACIÓN DE INFRAESTRUCTURA'],
        "filas": [(1, 'Normal'), (2, 'Emergencia')],
    },
    'PARAM_71': {
        "headers": ['VALOR TIPO ORDEN DE COMPRA', 'DESCRIPCIÓN TIPO ORDEN DE COMPRA'],
        "filas": [(1, 'Obras'), (2, 'Servicios prestados'), (3, 'Servicios recibidos'), (4, 'Suministros e insumos')],
    },
    'PARAM_72': {
        "headers": ['VALOR TIPO ORIGEN HORAS EXTRA', 'DESCRIPCIÓN TIPO ORIGEN HORAS EXTRA'],
        "filas": [(1, 'Emergencia'), (2, 'Turbiedad Extrema'), (3, 'Operación Normal'), (4, 'Otros orígenes')],
    },
    'PARAM_73': {
        "headers": ['VALOR TIPO OTROS BENEFICIOS ADICIONALES', 'DESCRIPCIÓN TIPO OTROS BENEFICIOS ADICIONALES', 'DESCRIPCIÓN'],
        "filas": [(1, 'Bienestar: Bonificación, costo empresa', 'Aportes monetarios que realiza la empresa a un fondo común creado por sindicatos o del área de capital humano para incentivar o apoyar en el financiamiento de materias de bienestar para los trabajadores. Abarca a beneficios de diversa índole, tales como en temas de salud, educación, previsionales, etc.'), (2, 'Seguro Complementario Salud', 'Convenio que realiza el empleador con una empresa aseguradora, con el fin de mitigar el gasto que se genera en prestaciones médicas y de salud.'), (3, 'Seguro de Vida', 'Seguros contratados por las empresas para proteger el patrimonio de los trabajadores en caso de fallecimiento.'), (4, 'Aporte Plan de Isapre', 'Monto entregado mensualmente como aporte a la cotización legal a los afiliados en isapres, con la finalidad de que el trabajador contrate un plan de salud de mayor valor. '), (5, 'Beneficio Salud Empresa', 'Desarrollo de convenios específicos o servicios administrados por la empresa para el cuidado y prevención de la salud de sus trabajadores. Puede incluir a chequeos médicos, dentales, nutricionista, etc.'), (6, 'Otros Seguros (Costo prima empresa)', 'Entrega de otros seguros para los trabajadores, tales como seguro contra accidentes, seguro dental o similares.'), (7, 'Indemnizaciones por sobre lo legal (Provisiones)', 'Derecho laboral por despido que la empresa debe dar al colaborador, en proporción a la cantidad de años que haya trabajado, sin aplicar los topes legales de 90 UF y/o 11 años'), (8, 'Regalos por Diversos Conceptos: navidad, cumpleaños, etc.', 'Entrega de un producto o servicio por parte de la empresa, por motivo de cumpleaños del trabajador u otro acontecimiento.'), (9, 'Pago de paquetes vacacionales', 'Corresponde a la entrega de beneficios por motivo de vacaciones o descanso de los trabajadores, el cual puede incluir el traslado, alojamiento, alimentación y/o visitas programadas del lugar de destino.'), (10, 'Paquete de Fiestas Patrias', 'Corresponde a la entrega de una canasta de productos a los trabajadores por la celebración de Fiestas Patrias (p.e. carnes, cecinas, vino, verduras, etc.).'), (11, 'Fiesta de F. Patrias', 'Fiesta o celebración organizada por la empresa, para sus trabajadores, por motivo de Fiestas Patrias.'), (12, 'Paquete de Navidad', 'Corresponde a la entrega de una canasta de productos a los trabajadores por la celebración de Navidad (p.e. pavo, galletas, vino, verduras, etc.).'), (13, 'Regalos de Navidad (Personal, Niños, Etc.)', 'Entrega de un regalo (producto o servicio), por parte de la empresa, a los trabajadores por motivo de Navidad. '), (14, 'Fiesta de Navidad', 'Celebración realizada por la empresa para los trabajadores o también de sus familiares, por concepto de Navidad.'), (15, 'Fiesta de fin de año y/o Aniversario', 'Celebración realizada por la empresa para los trabajadores para conmemorar el aniversario de la organización o para festejar el fin del año calendario.'), (16, 'Consumos Habitacionales', 'Corresponde al pago total o parcial de las cuentas de servicios básicos (luz, agua, gas u otros) de los trabajadores.'), (17, 'Asigna Casa : (Valor estim. del beneficio)', 'Corresponde a la entrega de una vivienda, por un tiempo determinado, para el uso del trabajador o también de su familia.'), (18, 'Asigna Vehículo : (Valor estimado del beneficio)', 'Corresponde a la entrega de un vehículo para facilitar el traslado del trabajador a visitas de clientes u otras actividades relacionadas a sus funciones.'), (19, 'Reembolso uso vehículo propio (valor estim.)', 'Corresponde al pago de los costos incurridos debido al traslado del trabajador a visitas de clientes u otras actividades relacionadas a sus funciones (p.e. bencina, peajes, tags, etc.).'), (20, 'Pago de estacionamiento', 'Corresponde a un pago incurrido por la empresa en proveer de estacionamientos a sus colaboradores o una asignación en dinero para que el trabajador costee uno por cuenta propia.'), (21, 'Traslado, Vehículo valorizado y pasajes', 'Corresponde a la entrega o reembolso de los costos incurridos por movilización de viajes por motivos de trabajo. Se diferencia de otras asignaciones o reembolsos por traslados de menores distancias geográficas.'), (22, 'Salud, Vida, Previsional (si no están en Seguros)', 'Corresponde a la entrega de otros beneficios de salud o protección que no estén especificados anteriormente. '), (23, 'Personales: Escolaridades (si no están en Bonos Estudios)', 'Corresponde a la entrega de otros bonos como apoyo al financiamiento de la educación de los hijos de los trabajadores, tales como los bonos de excelencia académica o similares.'), (24, 'Beneficios Previsionales no Obligatorios (si no están en B. Previsionales)', 'Corresponde a la entrega de otros beneficios previsionales que no estén especificados anteriormente. '), (25, 'Vales, lo que no paga el trabajador (cheque R. / No constituye Renta)', 'Corresponde a la entrega de vales de colación o cheques de restaurant para que sean utilizados para la colación del trabajador.'), (26, 'Casino, lo que no paga el trabajador (No constituye Renta)', 'Servicio de alimentación ofrecido por la empresa o proveedor, sin costo para los trabajadores.'), (27, 'Bonos por Negociación Colectiva (Sólo BTN) Mensualizado', 'Monto entregado mensualmente por la empresa a los trabajadores para acordar la finalización de un proceso de negociación colectiva específico. '), (28, 'Otros Bonos por Negociación Colectiva', 'Corresponde a la entrega de otros bonos por acuerdos tras la finalización de negociaciones colectivas, tales como bono por práctica operacional o bono por renovación de jornada.'), (29, 'Bonos Escolares (Cargas)', 'Corresponde a la entrega de un bono para apoyar el financiamiento de la educación de los hijos de los trabajadores. Generalmente abarca a la educación básica y media (científico humanista o técnica profesional).'), (30, 'Becas de Estudio (Trabajador)', 'Financiamiento que entrega la empresa a los trabajadores que realizan cursos de perfeccionamiento (generalmente estudios de postgrado).'), (31, 'Entrega de Vestuario (No obligatorio)', 'Entrega de vestimenta de trabajo, superior a lo exigido por la ley. '), (32, 'Entrega de Vestuario Obligatorio (Arts. Seguridad y de Trabajo)', 'Entrega de vestimenta de trabajo para el desempeño y resguardo de la seguridad del trabajador durante la realización de las funciones propias de su cargo. '), (33, 'Bonos por matrimonio (valor total costo empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos en el caso de contraer matrimonio y/o acuerdo de unión civil del trabajador.'), (34, 'Bonos por nacimiento (valor total costo empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos por nacimiento por hijo del trabajador.'), (35, 'Cuota mortuoria pagada (costo Empresa)', 'Ayuda monetaria entregada por la empresa para cubrir los gastos por fallecimiento de familiares directos del trabajador. Los más comunes son en caso del fallecimiento de hijo o cónyuge, como también del padre o madre del colaborador.'), (36, 'Sala cuna (valor total costo Sala Cuna, Reembolso o Bono)', 'Asignación en dinero que entrega la empresa (adicional a lo legal) por concepto de sala cuna a sus trabajadores, por cada hijo menor a 24 meses de edad.'), (37, 'Pago de Club Privado u Otros', 'Corresponde a convenios suscritos por la empresa con entidades de recreación para los trabajadores.'), (38, 'Otros Beneficios Valorizables que no son Remuneración', 'Corresponde a cualquier otro beneficio no monetario que no haya sido especificado anteriormente.')],
    },
    'PARAM_74': {
        "headers": ['VALOR TIPO PERIODICIDAD CONTRATO SERVICIOS', 'DESCRIPCIÓN TIPO PERIODICIDAD CONTRATO SERVICIOS'],
        "filas": [(1, 'Anual'), (2, 'Mensual'), (3, 'Eventos'), (4, 'Esporádico'), (5, 'Sin Contrato')],
    },
    'PARAM_75': {
        "headers": ['VALOR TIPO PERSONAL', 'DESCRIPCIÓN TIPO PERSONAL'],
        "filas": [(1, 'Permanente (Jornada completa)'), (2, 'Permanente (Jornada parcial)'), (3, 'Personal de Reemplazo'), (4, 'Alumno en Práctica')],
    },
    'PARAM_76': {
        "headers": ['VALOR TIPO PERSONAL CONTACTADO', 'DESCRIPCIÓN TIPO PERSONAL CONTACTADO'],
        "filas": [(1, 'Interno: Prima el contacto con personal de la empresa, no perteneciente a la misma dependencia.'), (2, 'Interno y externo: El Cargo se contacta primordialmente con personal interno pero también sus contactos externos son relevantes (o viceversa).'), (3, 'Externo: Prima el contacto con personal no vinculado laboralmente a la empresa.')],
    },
    'PARAM_77': {
        "headers": ['VALOR TIPO PÓLIZA', 'DESCRIPCIÓN TIPO PÓLIZA'],
        "filas": [(1, 'Incendio'), (2, 'Sismo'), (3, 'Robo'), (4, 'Responsabilidad Civil'), (5, 'Miscelánea (si la póliza incluye la cobertura para todos los siniestros antes mencionados)')],
    },
    'PARAM_78': {
        "headers": ['VALOR TIPO PONDERADOR SOLUCIÓN DE PROBLEMAS', 'DESCRIPCIÓN TIPO PONDERADOR SOLUCIÓN DE PROBLEMAS'],
        "filas": [(1, 'Bajo : Es un poco menos que lo definido en el ítem, pero no es el anterior.'), (2, 'Pleno o Medio: Es bastante equivalente a lo definido en el ítem.'), (3, 'Alto: Es un poco más que lo definido en el ítem, pero no es el siguiente.')],
    },
    'PARAM_79': {
        "headers": ['VALOR TIPO PRODUCTO QUÍMICO', 'DESCRIPCIÓN TIPO PRODUCTO QUÍMICO'],
        "filas": [(1, 'Ácido Fluorosilícico'), (2, 'Cal'), (3, 'Cat. Floc 8103 Plus'), (4, 'Cloro Líquido Cilindro'), (5, 'Cloro Líquido Cisterna'), (6, 'Cloro Líquido Contenedor'), (7, 'Cloruro Férrico'), (8, 'Flúor Sales'), (9, 'Flúor Fluorosilicato'), (10, 'Flúor Ácido'), (11, 'FO 4190 PWG'), (12, 'Cloro Gas Cilindro'), (13, 'Cloro Gas Contenedor'), (14, 'Hidróxido de Calcio'), (15, 'Hipoclorito de Calcio'), (16, 'Hipoclorito de Sodio'), (17, 'Permanganato de Potasio'), (18, 'Policloruro de aluminio'), (19, 'Polielectroli'), (20, 'Polímeros'), (21, 'Soda Cáustica'), (22, 'Sulfato de Aluminio'), (23, 'Sulfato de Cobre'), (24, 'Ácido Sulfúrico'), (25, 'AWC A – 102 Plus'), (26, 'Dolomita'), (27, 'EDTA – Na'), (28, 'Fluoruro de Sodio'), (29, 'Kit Arsénico'), (30, 'Reactivo DPD'), (31, 'Reactivo FerroVer'), (32, 'Reactivo Flúor'), (33, 'Otro')],
    },
    'PARAM_80': {
        "headers": ['VALOR TIPO PROVEEDOR O CLIENTE', 'DESCRIPCIÓN TIPO PROVEEDOR O CLIENTE'],
        "filas": [(1, 'Concesionaria Relacionada'), (2, 'Concesionaria No Relacionada'), (3, 'Otra Relacionada'), (4, 'Otra No Relacionada'), (5, 'Persona Natural')],
    },
    'PARAM_81': {
        "headers": ['VALOR TIPO PROYECTO', 'DESCRIPCIÓN TIPO PROYECTO'],
        "filas": [(1, 'Propio'), (2, 'Aportado por terceros'), (3, 'Mixto')],
    },
    'PARAM_82': {
        "headers": ['VALOR TIPO PLANTA', 'DESCRIPCIÓN TIPO PLANTA'],
        "filas": [(1, 'Lodo Activado'), (2, 'Laguna Aireada'), (3, 'Tratamiento Primario Químicamente Asistido'), (4, 'Filtros Biológicos (Biofiltros)'), (5, 'Lombrifiltros'), (6, 'Laguna Facultativa'), (7, 'Planta Pretratamiento AS de Emisario Submarino'), (8, 'Planta de Tramiento de Agua Potable (PTAP)'), (9, 'PEAS'), (10, 'PEAP')],
    },
    'PARAM_83': {
        "headers": ['VALOR TIPO RECAUDADOR', 'DESCRIPCIÓN TIPO RECAUDADOR'],
        "filas": [(1, 'PAC'), (2, 'Servipag'), (3, 'Sencillito'), (4, 'Transbank'), (5, 'Caja Vecina'), (6, 'Otro')],
    },
    'PARAM_84': {
        "headers": ['VALOR TIPO RED', 'DESCRIPCIÓN TIPO RED'],
        "filas": [(1, 'Red de Distribución '), (2, 'Conduccion de AP'), (3, 'Red de Recolección, Sólo Aguas Servidas'), (4, 'Red de Recolección, Red Unitaria 2'), (5, 'Red de Recolección, Sólo Aguas Lluvias'), (6, 'Conduccion de AS, Sólo Aguas Servidas'), (7, 'Conduccion de AS, Aguas Servidas y Aguas Lluvias'), (8, 'Conduccion de AS, Sólo Aguas Lluvias')],
    },
    'PARAM_85': {
        "headers": ['VALOR TIPO REMUNERACIÓN DIRECTOR', 'DESCRIPCIÓN TIPO REMUNERACIÓN DIRECTOR'],
        "filas": [(1, 'Remuneración Fija'), (2, 'Remuneración Variable (sesión)')],
    },
    'PARAM_86': {
        "headers": ['VALOR TIPO REPARACIÓN ARRANQUE', 'DESCRIPCIÓN TIPO REPARACIÓN ARRANQUE'],
        "filas": [(1, 'Reparación de Unión arranque-matriz o tubería de arranque'), (2, 'Cambio o reparación de llave de paso exterior'), (3, 'Cambio o reparación de llave de paso interior'), (4, 'Reparación entre llaves de paso')],
    },
    'PARAM_87': {
        "headers": ['VALOR TIPO RESIDUO', 'DESCRIPCIÓN TIPO RESIDUO'],
        "filas": [(1, 'Residuos sólidos'), (2, 'Arenas y grasas')],
    },
    'PARAM_88': {
        "headers": ['VALOR TIPO RESPALDO', 'DESCRIPCIÓN TIPO RESPALDO'],
        "filas": [(1, 'Contrato'), (2, 'Orden de Compra, de Trabajo, o de Servicio'), (3, 'Factura')],
    },
    'PARAM_89': {
        "headers": ['VALOR TIPO RESPONSABILIDAD EMERGENCIA', 'DESCRIPCIÓN TIPO RESPONSABILIDAD EMERGENCIA'],
        "filas": [(1, 'Empresa'), (2, 'Terceros'), (3, 'Eventos de la naturaleza')],
    },
    'PARAM_90': {
        "headers": ['VALOR TIPO RESPONSABILIDAD POR RESULTADOS', 'DESCRIPCIÓN TIPO RESPONSABILIDAD POR RESULTADOS'],
        "filas": [(1, 'Informativa. Sólo transmite información.'), (2, 'Indirecta. Da asesoría, apoyo e interpretación para que otros puedan realizar determinada labor.'), (3, 'Directa. El resultado de las decisiones es determinante en los logros finales.')],
    },
    'PARAM_91': {
        "headers": ['VALOR TIPO SERVICIO', 'DESCRIPCIÓN TIPO SERVICIO'],
        "filas": [(1, 'Agua Potable'), (2, 'Aguas Servidas'), (3, 'Agua Potable y Aguas Servidas')],
    },
    'PARAM_92': {
        "headers": ['VALOR TIPO SERVICIO DE DOCUMENTO', 'DESCRIPCIÓN TIPO SERVICIO DE DOCUMENTO'],
        "filas": [(1, 'Impresión'), (2, 'Suministro (insumos)'), (3, 'Mecanizado (ensobrado)')],
    },
    'PARAM_93': {
        "headers": ['VALOR TIPO SERVICIO DE MANTENCIÓN DE REDES', 'DESCRIPCIÓN TIPO SERVICIO DE MANTENCIÓN DE REDES'],
        "filas": [(1, 'Mantención de cámaras'), (2, 'Renovación de cámaras'), (3, 'Detección de fugas'), (4, 'Reparación de roturas'), (5, 'Mantención de grifos'), (6, 'Reparación de arranque'), (7, 'Renovación de arranque'), (8, 'Control de medidores'), (9, 'Recambio de medidores'), (10, 'Inspección televisiva de colectores'), (11, 'Desobstrucción de colectores'), (12, 'Lavado de red y/o colectores'), (13, 'Mantención de cámaras'), (14, 'Renovación de cámaras'), (15, 'Desobstrucción de uniones'), (16, 'Renovación de uniones')],
    },
    'PARAM_94': {
        "headers": ['VALOR TIPO SERVICIO INFORMÁTICO', 'DESCRIPCIÓN TIPO SERVICIO INFORMÁTICO'],
        "filas": [(1, 'Servicios Cloud'), (2, 'Servicio de Administración y Explotación'), (3, 'Monitoreo y Continuidad operacional'), (4, 'Operación del Centro de Procesamiento de Datos'), (5, 'Soporte y Mantención SW Macroinformática'), (6, 'Soporte Microinformática'), (7, 'Mantenimiento HW'), (8, 'Auditoría Informática'), (9, 'Seguridad Informática'), (10, 'Renovación de licencias'), (11, 'Help Desk'), (12, 'Hosting'), (13, 'Housing'), (14, 'Site Contingencia'), (15, 'Soporte en terreno'), (16, 'Ingeniería ante Contingencias'), (17, 'Custodia Medios Magnéticos')],
    },
    'PARAM_95': {
        "headers": ['VALOR TIPO SERVIDUMBRE', 'DESCRIPCIÓN TIPO SERVIDUMBRE'],
        "filas": [(1, 'Tránsito'), (2, 'Tránsito y Acueducto'), (3, 'Tránsito, Acueducto y Conducción Eléctrica'), (4, 'Tránsito y Conducción Eléctrica'), (5, 'Ocupación')],
    },
    'PARAM_96': {
        "headers": ['VALOR TIPO SISTEMA OPERATIVO', 'DESCRIPCIÓN TIPO SISTEMA OPERATIVO'],
        "filas": [(1, 'Windows'), (2, 'Mac'), (3, 'Linux'), (4, 'Unix'), (5, 'Otro')],
    },
    'PARAM_97': {
        "headers": ['VALOR TIPO SOFTWARE', 'DESCRIPCIÓN TIPO SOFTWARE'],
        "filas": [(1, 'Software World Class'), (2, 'Software Comercial Estándar de amplia disponibilidad'), (3, 'Desarrollo por Encargo'), (4, 'Desarrollo Propio'), (5, 'Software Libre'), (6, 'Otro')],
    },
    'PARAM_98': {
        "headers": ['VALOR TIPO SOLUCIÓN DE PROBLEMAS', 'DESCRIPCIÓN TIPO SOLUCIÓN DE PROBLEMAS'],
        "filas": [(1, 'Los problemas son mínimos, por cuanto el trabajo es rutinario, simple y repetitivo, con instrucciones muy específicas, en cuanto a la secuencia y el tiempo.'), (2, 'El trabajo es básicamente repetitivo, aunque implica la interpretación de instrucciones para afrontar pequeñas variaciones.'), (3, 'Los problemas son típicos y las soluciones son encontradas fácilmente en procedimientos bien definidos o con base en prácticas anteriores.'), (4, 'Los problemas tienen mayor variabilidad pero, en general, pueden manejarse con los procedimientos definidos o por comparación con soluciones o situaciones anteriores.'), (5, 'Los problemas son relativamente variados y complejos para lograr los objetivos específicos, requiriendo mejoras en métodos o el establecimiento de técnicas y estándares.'), (6, 'Los problemas tienen elementos poco usuales que no se habían presentado anteriormente, se exige análisis y evaluación de varias posibilidades. Las guías y políticas dan un marco general pero no la respuesta. Se necesita creatividad y criterio para encontrar la mejor solución.'), (7, 'Los problemas son altamente complejos y nuevos requiriendo mucha investigación y alta creatividad. Las políticas y planes estratégicos de la empresa pueden dar alguna guía, pero la solución puede estar en modificar alguno de ellos.'), (8, 'Los problemas son de un alto nivel de complejidad, debido tanto a la impredecibilidad del medio como a la necesidad de proyección a largo plazo. La solución puede estar en desarrollar nuevos objetivos y estrategias del negocio o redefinir la misión del mismo. Aquí se incluyen Cargos que implican la creación de modelos sofisticados de investigación científica')],
    },
    'PARAM_99': {
        "headers": ['VALOR TIPO SUPERFICIE ROTURA', 'DESCRIPCIÓN TIPO SUPERFICIE ROTURA'],
        "filas": [(1, 'Calzada Hormigón'), (2, 'Calzada Asfalto'), (3, 'Calzada Tierra'), (4, 'Calzada Otros'), (5, 'Vereda Pavimentada'), (6, 'Vereda Tierra'), (7, 'Bandejón Tierra'), (8, 'Bandejón Pavimentado')],
    },
    'PARAM_100': {
        "headers": ['VALOR TIPO SUPERFICIE TERRENO', 'DESCRIPCIÓN TIPO SUPERFICIE TERRENO'],
        "filas": [(1, 'Espacio Deportivo o Recreacional'), (2, 'Baldío o Eriazo'), (3, 'Espacio Libre')],
    },
    'PARAM_101': {
        "headers": ['VALOR TIPO SW MACROINFORMÁTICA', 'DESCRIPCIÓN TIPO SW MACROINFORMÁTICA'],
        "filas": [(1, 'Base de Datos'), (2, 'Modelos Matemáticos, Hidrológicos e Hidráulicos'), (3, 'Sistema de Laboratorio'), (4, 'GIS'), (5, 'Sistema Comercial o de Clientes'), (6, 'Call o Contact Center'), (7, 'ERP o Sistema Administrativo'), (8, 'Business Intelligence'), (9, 'E-learning'), (10, 'Sistema Documental'), (11, 'Sistema de Calidad'), (12, 'Servicio de Plataforma Web'), (13, 'Oficina Comercial Virtual'), (14, 'Sistema de Facturación Electrónica'), (15, 'Sistema de Correo Corporativo'), (16, 'Seguridad Informática'), (17, 'Helpdesk Informático'), (18, 'Otros')],
    },
    'PARAM_102': {
        "headers": ['VALOR TIPO TAMAÑO DE SUPERVISIÓN', 'DESCRIPCIÓN TIPO TAMAÑO DE SUPERVISIÓN'],
        "filas": [(1, 'Hasta 1 persona'), (2, '2 - 10 personas'), (3, '11- 40 personas'), (4, '41 - 100 personas'), (5, '101 - 300 personas'), (6, '301 - 1.000 personas'), (7, '1.001 - 4.000 personas'), (8, '4.001 - 7.000 personas'), (9, 'Más de 7.000 personas')],
    },
    'PARAM_103': {
        "headers": ['VALOR TIPO TARIFA ELÉCTRICA', 'DESCRIPCIÓN TIPO TARIFA ELÉCTRICA'],
        "filas": [(1, 'BT1a'), (2, 'BT1b'), (3, 'TRBT2'), (4, 'TRBT3'), (5, 'TRAT1'), (6, 'TRAT2'), (7, 'TRAT3'), (8, 'BT2'), (9, 'BT3'), (10, 'BT4.1'), (11, 'BT4.2'), (12, 'BT4.3'), (13, 'BT5'), (14, 'AT2'), (15, 'AT3'), (16, 'AT4.1'), (17, 'AT4.2'), (18, 'AT4,3'), (19, 'AT5'), (20, 'TFR')],
    },
    'PARAM_104': {
        "headers": ['VALOR TIPO TECNOLOGÍA DE OPERACIÓN', 'DESCRIPCIÓN TIPO TECNOLOGÍA DE OPERACIÓN'],
        "filas": [(1, 'Control Automático Local'), (2, 'Manual'), (3, 'Telecontrol'), (4, 'Telemedición')],
    },
    'PARAM_105': {
        "headers": ['VALOR TIPO ZONAL', 'DESCRIPCIÓN TIPO TECNOLOGÍA DETECCIÓN DE FUGAS'],
        "filas": [(1, 'Gas Trazador'), (2, 'Correlación Acústica'), (3, 'Otra')],
    },
    'PARAM_106': {
        "headers": ['VALOR TIPO UNIDAD ADQUISICIÓN TIC', 'DESCRIPCIÓN TIPO UNIDAD ADQUISICIÓN TIC'],
        "filas": [(1, 'Licencias'), (2, 'Servidores'), (3, 'Unidades Físicas'), (4, 'Horas Hombre de Desarrollo'), (5, 'Horas Hombre de Consultoría'), (6, 'Otro')],
    },
    'PARAM_107': {
        "headers": ['VALOR TIPO UNIDAD CONCENTRACIÓN', 'DESCRIPCIÓN TIPO UNIDAD CONCENTRACIÓN'],
        "filas": [(1, 'Molaridad: se expresa en moles de soluto por litro de disolvente (M).'), (2, 'Molalidad: se expresa en moles de soluto por kilogramo de disolvente (m).'), (3, 'Fracción molar: se expresa como la proporción de la cantidad de moles de soluto entre la cantidad total de moles presentes en la solución (X).'), (4, 'Porcentaje en masa: se expresa como la cantidad de soluto en gramos por cada 100 gramos de solución\xa0(%\xa0en\xa0masa).')],
    },
    'PARAM_108': {
        "headers": ['VALOR TIPO UNIDADES DE CAUDAL', 'DESCRIPCIÓN TIPO UNIDADES DE CAUDAL'],
        "filas": [(1, 'm3/s'), (2, 'lts/s'), (3, 'Acciones'), (4, 'Otra')],
    },
    'PARAM_109': {
        "headers": ['VALOR TIPO URBANIZACIÓN INMUEBLE', 'DESCRIPCIÓN TIPO URBANIZACIÓN INMUEBLE'],
        "filas": [(1, 'Áreas Verdes y Jardines'), (2, 'Cierre Perimetral'), (3, 'Puertas de Acceso'), (4, 'Vialidad (caminos de acceso y circulación interior)'), (5, 'Arborización'), (6, 'Otras Obras de Urbanización')],
    },
    'PARAM_110': {
        "headers": ['VALOR TIPO USO DE HERRAMIENTA O EQUIPO', 'DESCRIPCIÓN TIPO USO DE HERRAMIENTA O EQUIPO'],
        "filas": [(1, 'Oficina'), (2, 'Bodega'), (3, 'Laboratorio'), (4, 'Terreno y Recintos Productivos'), (5, 'Taller'), (6, 'Otro')],
    },
    'PARAM_111': {
        "headers": ['VALOR TIPO USO EDIFICACIÓN', 'DESCRIPCIÓN TIPO USO EDIFICACIÓN'],
        "filas": [(1, 'Edificación de Oficinas Administrativas'), (2, 'Edificación de Oficinas de Atención de Público'), (3, 'Edificación de Oficinas de Operación y Control'), (4, 'Bodegas'), (5, 'Laboratorio'), (6, 'Estacionamientos'), (7, 'Obras')],
    },
    'PARAM_112': {
        "headers": ['VALOR TIPO VEHÍCULO', 'DESCRIPCIÓN TIPO VEHÍCULO'],
        "filas": [(1, 'Automóvil'), (2, 'Furgón'), (3, 'Camioneta Cabina Simpla (CS)'), (4, 'Camioneta Doble Cabina (DC)'), (5, 'Camión'), (6, 'Camión Grúa'), (7, 'Camión Desobstructor'), (8, 'Cargador Frontal'), (9, 'Otro Vehículo')],
    },
    'PARAM_113': {
        "headers": ['VALOR TIPO VIAJE', 'DESCRIPCIÓN TIPO VIAJE'],
        "filas": [(1, 'Actividades propias del cargo'), (2, 'Capacitación'), (3, 'Eventos gremiales'), (4, 'Reuniones con SISS'), (5, 'Actividades recreacionales')],
    },
    'PARAM_114': {
        "headers": ['VALOR TIPO ZONAL', 'DESCRIPCIÓN TIPO ZONAL'],
        "filas": [(1, 'Operativa'), (2, 'Comercial'), (3, 'Territorial'), (4, 'Centralizada')],
    },
}


def identificar_tabla_generico(nombre_archivo, catalogo):
    """Extrae el nombre de tabla (ej. 'MEI_1', 'ST_12', 'GRH_8') de un nombre
    de archivo, tolerando texto adicional despues (fecha, version, etc.),
    sin importar mayusculas/minusculas. Se compara por el nombre MAS LARGO
    que calce primero.

    Si el nombre no calza con nada del catalogo explicito, se intenta un
    reconocimiento GENERICO: cualquier archivo con forma '{FAMILIA}_{NUMERO}'
    se reconoce igual SIEMPRE QUE ese código exista de verdad en
    TABLAS_SCR_OFICIAL (el catálogo oficial del SCR) — esto permite
    reconocer tablas del SCR que aún no están explícitamente en
    TABLAS_SIMPLES/ST_TABLES/GPA_TABLES sin abrir la puerta a "reconocer"
    códigos inventados que no existen (ej. GGI_8, que no es una tabla real).

    Devuelve None si no matchea ni el catalogo ni una tabla real del SCR."""
    base = re.sub(r"\.(XLSX|XLS)$", "", nombre_archivo.upper())
    candidatos = sorted(catalogo, key=len, reverse=True)
    for tabla in candidatos:
        tabla_up = tabla.upper()
        if base == tabla_up or base.startswith(tabla_up + "_") or base.startswith(tabla_up + "-") or base.startswith(tabla_up + " "):
            return tabla

    m = re.match(r"^([A-Z]{2,4})[_\-\s]+(\d{1,3})(?!\d)", base)
    if m:
        familia, numero = m.group(1), m.group(2)
        codigo_candidato = f"{familia}_{numero}"
        if codigo_candidato in TABLAS_SCR_OFICIAL:
            return codigo_candidato
    return None


def procesar_zip_carpeta(zip_bytes):
    """Abre el .zip subido, recorre TODOS los archivos .xlsx/.xls (sin
    importar en qué subcarpeta estén), y arma un diccionario
    {nombre_tabla: bytes}. Devuelve (archivos, avisos) donde avisos incluye
    archivos no reconocidos y duplicados (misma tabla encontrada 2+ veces)."""
    archivos = {}
    avisos = []
    duplicados = {}
    catalogo = _catalogo_completo()

    try:
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
    except zipfile.BadZipFile:
        return {}, ["El archivo subido no es un .zip válido (o está corrupto)."]

    no_reconocidos = []
    for info in zf.infolist():
        if info.is_dir():
            continue
        nombre_completo = info.filename
        nombre_base = os.path.basename(nombre_completo)
        if not nombre_base.lower().endswith((".xlsx", ".xls")):
            continue
        if nombre_base.startswith("~$"):  # archivo temporal de Excel abierto
            continue
        tabla = identificar_tabla_generico(nombre_base, catalogo)
        if tabla is None:
            no_reconocidos.append(nombre_completo)
            continue
        try:
            contenido = zf.read(info)
        except Exception as e:
            avisos.append(f"No se pudo leer '{nombre_completo}' del zip: {e}")
            continue
        if tabla in archivos:
            duplicados.setdefault(tabla, [nombre_completo]).append(nombre_completo)
            # Nos quedamos con el que se encontró primero; se avisa el duplicado.
            continue
        archivos[tabla] = contenido

    if no_reconocidos:
        avisos.append(
            f"{len(no_reconocidos)} archivo(s) dentro del zip no calzaron con ninguna tabla conocida "
            f"y se ignoraron: {', '.join(no_reconocidos[:10])}" + (", ..." if len(no_reconocidos) > 10 else "")
        )
    for tabla, nombres in duplicados.items():
        avisos.append(
            f"Se encontró más de un archivo para la tabla {tabla} dentro del zip "
            f"({', '.join(nombres)}); se usó solo el primero encontrado."
        )

    return archivos, avisos


def checklist_faltantes(archivos, tablas_requeridas):
    """Devuelve la lista de tablas de 'tablas_requeridas' que NO están en
    'archivos' (dict {tabla: bytes})."""
    return [t for t in tablas_requeridas if t not in archivos]


# ============================================================================
# UX: resúmenes KPI y comparación contra la generación anterior en la sesión
# ============================================================================
def mostrar_kpis_rep2(final_rows, avisos):
    reg = sum(r[7] for r in final_rows if r[5] in (11, 12))
    noreg = sum(r[7] for r in final_rows if r[5] in (21, 22))
    recursos = len(set(r[4] for r in final_rows))
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Gasto Regulado", f"${reg:,.0f}")
    c2.metric("Gasto No Regulado", f"${noreg:,.0f}")
    c3.metric("N° Recursos", recursos)
    c4.metric("Avisos", len(avisos))


def mostrar_kpis_rep3(final_rows3, avisos3):
    total = sum(r[8] for r in final_rows3)
    procesos = len(set(r[6] for r in final_rows3))
    recursos = len(set(r[4] for r in final_rows3))
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Gasto Anual (regulado)", f"${total:,.0f}")
    c2.metric("N° Procesos", procesos)
    c3.metric("N° Recursos", recursos)
    c4.metric("Avisos", len(avisos3))


def mostrar_kpis_cyg(cyg14, cyg8, cyg9, avisos_cyg):
    reg = sum(r[7] for r in cyg14.get("CYG_1", [])) + sum(r[7] for r in cyg14.get("CYG_2", []))
    noreg = sum(r[7] for r in cyg14.get("CYG_3", [])) + sum(r[7] for r in cyg14.get("CYG_4", []))
    total_filas = sum(len(v) for v in cyg14.values()) + len(cyg8) + len(cyg9)

    anterior = st.session_state.get("diff_prev_cyg_totales")
    delta_reg = delta_noreg = delta_filas = None
    if anterior is not None:
        delta_reg = reg - anterior["reg"]
        delta_noreg = noreg - anterior["noreg"]
        delta_filas = total_filas - anterior["filas"]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Regulado (CYG_1+CYG_2)", f"${reg:,.0f}", delta=f"{delta_reg:+,.0f}" if delta_reg is not None and abs(delta_reg) > 1 else None)
    c2.metric("No Regulado (CYG_3+CYG_4)", f"${noreg:,.0f}", delta=f"{delta_noreg:+,.0f}" if delta_noreg is not None and abs(delta_noreg) > 1 else None)
    c3.metric("Filas totales", f"{total_filas:,}", delta=delta_filas if delta_filas else None)
    c4.metric("Avisos", len(avisos_cyg))
    if anterior is not None and (delta_reg or delta_noreg or delta_filas):
        st.caption("↑↓ Comparado con la generación anterior en esta sesión.")
    st.session_state["diff_prev_cyg_totales"] = {"reg": reg, "noreg": noreg, "filas": total_filas}


def mostrar_comparacion_anterior(clave_diff, filas_nuevas, idx_agrupacion, idx_valor, etiqueta_grupo):
    """Compara 'filas_nuevas' contra lo guardado en session_state[clave_diff]
    (la generación anterior EN ESTA MISMA SESIÓN, si existe), agrupando por
    'idx_agrupacion' (un índice, o una tupla de índices para clave compuesta,
    ej. (recurso, familia)) y sumando 'idx_valor'. Muestra los cambios y
    luego actualiza session_state[clave_diff] para la próxima comparación."""
    def _clave(r):
        if isinstance(idx_agrupacion, (tuple, list)):
            return tuple(r[i] for i in idx_agrupacion)
        return r[idx_agrupacion]

    anterior = st.session_state.get(clave_diff)
    if anterior is not None:
        viejo = defaultdict(float)
        nuevo = defaultdict(float)
        for r in anterior:
            viejo[_clave(r)] += r[idx_valor]
        for r in filas_nuevas:
            nuevo[_clave(r)] += r[idx_valor]
        claves = set(viejo) | set(nuevo)
        cambios = []
        for k in claves:
            v_old, v_new = viejo.get(k, 0.0), nuevo.get(k, 0.0)
            if abs(v_new - v_old) > 1:
                cambios.append((k, v_old, v_new, v_new - v_old))
        if cambios:
            cambios.sort(key=lambda x: -abs(x[3]))
            with st.expander(f"🔄 Cambios respecto a la generación anterior en esta sesión ({len(cambios)})", expanded=True):
                df_cambios = pd.DataFrame(cambios, columns=[etiqueta_grupo, "Antes ($)", "Ahora ($)", "Diferencia ($)"])
                st.dataframe(
                    df_cambios.style.format({"Antes ($)": "{:,.0f}", "Ahora ($)": "{:,.0f}", "Diferencia ($)": "{:+,.0f}"}),
                    width="stretch",
                )
        else:
            st.caption("Sin cambios respecto a la generación anterior en esta sesión.")
    st.session_state[clave_diff] = filas_nuevas


# ---------------------------------------------------------------- UI --------
st.title("Generador REP_2 · Costos y Gastos por Familia de Servicios")
st.caption(
    "Consolida las familias GRH, GCP, GGV, GGI, GGM, OGG y MEI en la tabla "
    "REP_2 exigida por la SISS."
)

modo_jefatura = st.sidebar.toggle(
    "🔒 Vista Jefatura (solo lectura)",
    key="modo_jefatura",
    help="Activa una vista simplificada: sin carga de archivos ni botones de "
         "generación, solo el resumen y la validación de lo ya generado en esta sesión.",
)

sufijo_archivo = st.sidebar.text_input(
    "Sufijo para nombres de archivo (opcional)",
    value="",
    placeholder="_v09.07.26",
    help="Se agrega al final del nombre de cada archivo que descargues "
         "(ej. 'REP_2_v09.07.26.xlsx'). Déjalo vacío para usar el nombre por defecto.",
    key="sufijo_archivo",
)


def nombre_con_sufijo(nombre_base, extension):
    """Arma el nombre final de archivo para descarga, insertando el sufijo
    configurado por el usuario (si lo hay) ANTES de la extensión."""
    return f"{nombre_base}{sufijo_archivo}.{extension}"


if modo_jefatura:
    st.title("📊 Resumen Ejecutivo — REP_2 / REP_3 / CYG")
    st.caption(
        "Vista de solo lectura: muestra lo ya generado en esta sesión por el "
        "equipo operativo. Para cargar archivos o generar reportes, desactiva "
        "'Vista Jefatura' en la barra lateral."
    )

    hay_rep2 = "last_final_rows" in st.session_state
    hay_rep3 = "last_final_rows3" in st.session_state
    hay_cyg = "last_cyg14" in st.session_state

    if not (hay_rep2 or hay_rep3 or hay_cyg):
        st.info("Aún no se ha generado ningún reporte en esta sesión.")
    else:
        if hay_rep2:
            st.subheader("REP_2 — Costos y Gastos por Familia de Servicios")
            mostrar_kpis_rep2(st.session_state["last_final_rows"], st.session_state.get("last_avisos", []))
        if hay_rep3:
            st.subheader("REP_3 — Costos y Gastos No Activados por Proceso")
            mostrar_kpis_rep3(st.session_state["last_final_rows3"], st.session_state.get("last_avisos3", []))
        if hay_cyg:
            st.subheader("CYG — Costos y Gastos por Recurso")
            mostrar_kpis_cyg(
                st.session_state["last_cyg14"],
                st.session_state.get("last_cyg8", []),
                st.session_state.get("last_cyg9", []),
                st.session_state.get("last_avisos_cyg", []),
            )

        if hay_rep2 and hay_rep3 and hay_cyg:
            st.divider()
            st.subheader("✅ Validación cruzada")
            _r2 = st.session_state["last_final_rows"]
            _r3 = st.session_state["last_final_rows3"]
            _c14 = st.session_state["last_cyg14"]
            _reg2 = sum(r[7] for r in _r2 if r[5] in (11, 12))
            _noreg2 = sum(r[7] for r in _r2 if r[5] in (21, 22))
            _tot3 = sum(r[8] for r in _r3)
            _c12 = sum(r[7] for r in _c14.get("CYG_1", [])) + sum(r[7] for r in _c14.get("CYG_2", []))
            _c34 = sum(r[7] for r in _c14.get("CYG_3", [])) + sum(r[7] for r in _c14.get("CYG_4", []))
            _tol = 10
            if abs(_tot3 - _reg2) <= _tol and abs(_c12 - _reg2) <= _tol and abs(_c34 - _noreg2) <= _tol:
                st.success("Todo cuadra: REP_2, REP_3 y las tablas CYG son consistentes entre sí.")
            else:
                st.warning("Hay diferencias entre reportes fuera de tolerancia — revisar con el equipo operativo.")

        if "minuta_bytes" in st.session_state:
            st.divider()
            st.download_button(
                "📄 Descargar Minuta de Criterios de Asignación",
                data=st.session_state["minuta_bytes"],
                file_name=nombre_con_sufijo("Minuta_Parametrizacion", "docx"),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
    st.stop()

if not modo_jefatura:
    with st.sidebar:
        st.header("Archivos de entrada")
        st.caption(
            "Haz clic en 'Browse files' y selecciona directamente tu CARPETA "
            "completa (con sus subcarpetas, ej. una carpeta 'ST' con las tablas "
            "ST adentro) — no hace falta entrar a cada tabla una por una. "
            "También puedes arrastrar la carpeta, soltar varios .xlsx sueltos, "
            "o subir un .zip; lo que te resulte más cómodo. No importa la "
            "estructura de subcarpetas (se admite 1 nivel), el sistema busca "
            "cada tabla por su nombre en cualquier parte."
        )
        st.caption(
            "Los nombres pueden tener texto adicional (ej. 'MEI_1_2025.xlsx' o "
            "'GRH_8 (v2) 2025-07-01.xlsx') — se reconocen igual, mientras "
            "empiecen con el nombre de la tabla (ej. 'MEI_1', 'GRH_8'). Si vuelves "
            "a subir una tabla que ya tenías (por ejemplo porque la corregiste), "
            "se reemplaza automáticamente por la versión nueva."
        )
        try:
            # "directory" habilita seleccionar una carpeta completa con un clic
            # (requiere una versión de Streamlit razonablemente reciente).
            f_subidos = st.file_uploader(
                "Carpeta (clic para seleccionarla), archivos .xlsx sueltos, o .zip",
                type=["zip", "xlsx", "xls"],
                accept_multiple_files="directory",
                key="archivos_subidos",
            )
        except Exception:
            st.caption(
                "ℹ️ Tu versión de Streamlit no soporta seleccionar una carpeta con "
                "un clic (actualiza streamlit para esa opción); mientras tanto, "
                "puedes arrastrar la carpeta o subir varios archivos/.zip igual."
            )
            f_subidos = st.file_uploader(
                "Archivos .xlsx sueltos, o .zip (arrastra la carpeta aquí)",
                type=["zip", "xlsx", "xls"],
                accept_multiple_files=True,
                key="archivos_subidos_fallback",
            )

        class _ArchivoZip:
            """Envoltorio liviano para que el resto del código (que espera un
        objeto tipo UploadedFile de Streamlit, con .getvalue()/.name) siga
        funcionando igual, sin importar si el archivo vino de un uploader
        individual o de la carga unificada."""
            def __init__(self, name, data):
                self.name = name
                self._data = data
            def getvalue(self):
                return self._data

        def _procesar_archivos_subidos(lista_uploaded):
            """Recibe la lista de UploadedFile (mezcla posible de .zip sueltos,
        archivos .xlsx sueltos, o ambos -- como cuando se arrastra una
        carpeta completa, que el navegador entrega como muchos archivos
        individuales) y devuelve (archivos {tabla: bytes}, avisos).
        Si una misma tabla aparece más de una vez (ej. subiste una version
        corregida junto con el resto), se usa la ULTIMA que aparece en la
        lista subida -- asi una correccion siempre reemplaza a la anterior."""
            archivos = {}
            avisos = []
            catalogo = _catalogo_completo()
            duplicados = {}
            no_reconocidos = []
            for f in lista_uploaded:
                nombre = f.name
                datos = f.getvalue()
                if nombre.lower().endswith(".zip"):
                    archivos_zip_i, avisos_zip_i = procesar_zip_carpeta(datos)
                    for tabla, contenido in archivos_zip_i.items():
                        if tabla in archivos:
                            duplicados.setdefault(tabla, []).append(f"{nombre} (dentro del zip)")
                        archivos[tabla] = contenido
                    avisos.extend(avisos_zip_i)
                    continue
                nombre_base = os.path.basename(nombre)
                if nombre_base.startswith("~$"):
                    continue
                tabla = identificar_tabla_generico(nombre_base, catalogo)
                if tabla is None:
                    no_reconocidos.append(nombre)
                    continue
                if tabla in archivos:
                    duplicados.setdefault(tabla, []).append(nombre)
                archivos[tabla] = datos

            if no_reconocidos:
                avisos.append(
                    f"{len(no_reconocidos)} archivo(s) no calzaron con ninguna tabla conocida "
                    f"y se ignoraron: {', '.join(no_reconocidos[:10])}" + (", ..." if len(no_reconocidos) > 10 else "")
                )
            for tabla, nombres in duplicados.items():
                avisos.append(
                    f"Se encontró más de un archivo para la tabla {tabla} "
                    f"({', '.join(nombres)}); se usó la ÚLTIMA versión subida "
                    f"(útil si subiste una corrección)."
                )
            return archivos, avisos

        if f_subidos:
            firma_actual = tuple(sorted((f.name, f.size) for f in f_subidos))
            if st.session_state.get("_archivos_firma") != firma_actual:
                archivos_zip, avisos_zip = _procesar_archivos_subidos(f_subidos)
                st.session_state["archivos_zip"] = archivos_zip
                st.session_state["avisos_zip"] = avisos_zip
                st.session_state["_archivos_firma"] = firma_actual

        archivos_zip = st.session_state.get("archivos_zip", {})
        avisos_zip = st.session_state.get("avisos_zip", [])

        if archivos_zip:
            st.success(f"✅ {len(archivos_zip)} tabla(s) reconocida(s).")
            with st.expander("Ver tablas encontradas"):
                st.write(sorted(archivos_zip.keys()))
            if avisos_zip:
                with st.expander(f"⚠️ Avisos de carga ({len(avisos_zip)})", expanded=True):
                    for a in avisos_zip:
                        st.markdown(f"- {a}")

            st.markdown("**¿Qué falta para cada reporte?**")
            faltan_rep2 = checklist_faltantes(
                archivos_zip,
                ["GRH_8", "GRH_11", "GCP_4", "GCP_5", "GGV_4", "GGV_5", "GGI_5",
                 "GGM_1", "GGM_2", "GGM_3", "GGM_4", "GGM_5", "OGG_5",
                 "MEI_1", "MEI_2", "MEI_3", "MEI_4"] + ST_TABLES + GPA_TABLES,
            )
            faltan_rep3_extra = checklist_faltantes(archivos_zip, ["GRH_12", "GCP_6", "GGV_6", "GGI_6"])
            faltan_cyg9 = checklist_faltantes(archivos_zip, ["MCO_42", "ING_4"])

            if faltan_rep2:
                st.caption(f"Para REP_2 aún faltan: {', '.join(faltan_rep2)}")
            else:
                st.caption("✅ REP_2: todas las tablas base están.")
            if faltan_rep3_extra:
                st.caption(f"Para REP_3 (además de lo de REP_2) faltan: {', '.join(faltan_rep3_extra)}")
            else:
                st.caption("✅ REP_3: tablas de actividad completas.")
            if faltan_cyg9:
                st.caption(f"Para CYG_9 (clientes) faltan: {', '.join(faltan_cyg9)}")
            else:
                st.caption("✅ CYG_9: MCO_42 e ING_4 están.")

        st.markdown("**Plantilla (opcional)**")
        f_template = st.file_uploader("REP_2.xlsx (diccionario)", type="xlsx", key="template")

        def _get(nombre):
            return _ArchivoZip(f"{nombre}.xlsx", archivos_zip[nombre]) if nombre in archivos_zip else None

        f_grh8 = _get("GRH_8")
        f_grh11 = _get("GRH_11")
        f_gcp4 = _get("GCP_4")
        f_gcp5 = _get("GCP_5")
        f_ggv4 = _get("GGV_4")
        f_ggv5 = _get("GGV_5")
        f_ggi5 = _get("GGI_5")
        f_ggm1 = _get("GGM_1")
        f_ggm2 = _get("GGM_2")
        f_ggm3 = _get("GGM_3")
        f_ggm4 = _get("GGM_4")
        f_ggm5 = _get("GGM_5")
        f_ogg5 = _get("OGG_5")
        f_mei1 = _get("MEI_1")
        f_mei2 = _get("MEI_2")
        f_mei3 = _get("MEI_3")
        f_mei4 = _get("MEI_4")
        f_grh12 = _get("GRH_12")
        f_gcp6 = _get("GCP_6")
        f_ggv6 = _get("GGV_6")
        f_ggi6 = _get("GGI_6")
        f_mco42 = _get("MCO_42")
        f_ing4 = _get("ING_4")

        f_st_files = [_ArchivoZip(f"{t}.xlsx", archivos_zip[t]) for t in ST_TABLES if t in archivos_zip] or None
        f_gpa_files = [_ArchivoZip(f"{t}.xlsx", archivos_zip[t]) for t in GPA_TABLES if t in archivos_zip] or None
        f_gpa_files_rep3 = f_gpa_files

        archivos_regulares = [f_grh8, f_grh11, f_gcp4, f_gcp5, f_ggv4, f_ggv5, f_ggi5,
                               f_ggm1, f_ggm2, f_ggm3, f_ggm4, f_ggm5, f_ogg5,
                               f_mei1, f_mei2, f_mei3, f_mei4]
        hay_algo_cargado = any(archivos_regulares) or bool(f_st_files) or bool(f_gpa_files)

        st.divider()
        PARAM_SESSION_KEYS = [
            "ggm_overrides", "ogg_overrides", "mei_overrides", "st_overrides",
            "ggm_proceso_overrides", "ogg_proceso_overrides", "mei_proceso_overrides", "gpa_proceso_overrides",
        ]
        with st.expander("💾 Guardar / Cargar Parametrización"):
            st.caption(
                "La parametrización que configures (servicio y proceso) vive solo "
                "en esta sesión del navegador — se pierde si recargas la página. "
                "Descárgala para poder volver a cargarla la próxima vez, sin "
                "reconfigurar todo de nuevo."
            )
            datos_export = {k: st.session_state.get(k, []) for k in PARAM_SESSION_KEYS}
            hay_parametrizacion = any(datos_export.values())
            json_bytes = json.dumps(datos_export, ensure_ascii=False, indent=2).encode("utf-8")
            st.download_button(
                "Descargar parametrización (.json)",
                data=json_bytes,
                file_name="parametrizacion_suralis.json",
                mime="application/json",
                disabled=not hay_parametrizacion,
                help=None if hay_parametrizacion else "Aún no has configurado ninguna parametrización.",
            )
            f_param_json = st.file_uploader("Cargar parametrización (.json)", type="json", key="param_json_uploader")
            if f_param_json is not None and st.session_state.get("_param_json_procesado") != f_param_json.name:
                try:
                    datos_import = json.loads(f_param_json.getvalue().decode("utf-8"))
                    for k in PARAM_SESSION_KEYS:
                        if k in datos_import:
                            st.session_state[k] = [tuple(x) for x in datos_import[k]]
                    st.session_state["_param_json_procesado"] = f_param_json.name
                    st.success("Parametrización cargada correctamente.")
                    st.rerun()
                except Exception as e:
                    st.error(f"No se pudo leer el archivo de parametrización: {e}")


# ============================================================================
# NAVEGACIÓN PRINCIPAL: dos tarjetas grandes (Tablas REP / Tablas CYG)
# ============================================================================
st.markdown(
    """
    <style>
    div[data-testid="stButton"] button[kind="secondary"].tarjeta-nav {
        height: 110px;
        font-size: 20px;
        font-weight: 600;
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def _leer_parametrizacion(session_key):
    """Lee (sin renderizar ningún widget) la parametrización ya guardada en
    session_state para 'session_key' y la devuelve como {recurso: [(destino, pct), ...]}.
    Se usa para que CYG pueda usar la parametrización de proceso configurada
    en la tarjeta 'Tablas REP', incluso cuando esa tarjeta no está visible
    en este momento (los datos ya quedaron guardados en session_state)."""
    if session_key not in st.session_state:
        st.session_state[session_key] = []
    params = defaultdict(list)
    for cod_r, cod_destino, pct in st.session_state[session_key]:
        params[cod_r].append((cod_destino, pct))
    return dict(params)


# Se calculan siempre (sin mostrar widgets) para que estén disponibles sin
# importar qué tarjeta esté activa; la tarjeta "Tablas REP" además muestra
# los paneles interactivos para EDITAR estos mismos valores.
ggm_params = _leer_parametrizacion("ggm_overrides")
ogg_params = _leer_parametrizacion("ogg_overrides")
mei_params = _leer_parametrizacion("mei_overrides")
st_params = _leer_parametrizacion("st_overrides")
ggm_proceso_params = _leer_parametrizacion("ggm_proceso_overrides")
ogg_proceso_params = _leer_parametrizacion("ogg_proceso_overrides")
mei_proceso_params = _leer_parametrizacion("mei_proceso_overrides")
gpa_proceso_params = _leer_parametrizacion("gpa_proceso_overrides")

if "vista_activa" not in st.session_state:
    st.session_state["vista_activa"] = None

st.markdown("## ¿Qué quieres generar?")
col_rep, col_cyg = st.columns(2)
with col_rep:
    if st.button(
        "📊 TABLAS REP\n\nREP_2 + REP_3",
        key="btn_tarjeta_rep",
        width='stretch',
        type="primary" if st.session_state["vista_activa"] == "REP" else "secondary",
    ):
        st.session_state["vista_activa"] = "REP"
        st.rerun()
with col_cyg:
    if st.button(
        "📁 TABLAS CYG\n\nCYG_1 a CYG_9",
        key="btn_tarjeta_cyg",
        width='stretch',
        type="primary" if st.session_state["vista_activa"] == "CYG" else "secondary",
    ):
        st.session_state["vista_activa"] = "CYG"
        st.rerun()

vista_activa = st.session_state["vista_activa"]

if vista_activa is None:
    st.info(
        "👆 Elige qué quieres generar: **Tablas REP** (REP_2 y REP_3) o "
        "**Tablas CYG** (CYG_1 a CYG_9). Puedes cambiar entre ambas cuando quieras "
        "— lo ya cargado y parametrizado se mantiene."
    )

st.divider()

with st.expander("📖 Consultar Diccionario SISS (Maestros MAE / Tipificaciones)"):
    st.caption(
        "Consulta directa a las tablas oficiales del Maestro SISS embebidas en el "
        "sistema, para apoyar el trabajo diario y futuras funciones. MAE_7 (Cargos "
        "SISS, ~7.787 filas) y MAE_15 (Partidas SVI, ~14.134 filas) no están "
        "embebidas por su tamaño."
    )
    tipo_consulta = st.radio(
        "¿Qué quieres consultar?",
        ["Tablas MAE (maestros)", "Tipificaciones / Parámetros (PARAM)"],
        horizontal=True,
        key="tipo_consulta_diccionario",
    )

    if tipo_consulta == "Tablas MAE (maestros)":
        opciones = {f"{k} — {v['nombre']}": k for k, v in MAE_DATOS.items()}
        opciones_faltantes = {"MAE_7 — Cargos SISS (no embebida, muy grande)": None, "MAE_15 — Código Partidas SVI (no embebida, muy grande)": None}
        seleccion = st.selectbox("Tabla MAE", list(opciones.keys()) + list(opciones_faltantes.keys()), key="sel_tabla_mae")
        if seleccion in opciones_faltantes:
            st.warning("Esta tabla no está embebida por su tamaño. Si la necesitas, súbela directamente y avísame para agregarle soporte de consulta puntual.")
        else:
            tabla_key = opciones[seleccion]
            contenido = MAE_DATOS[tabla_key]
            df_mae = pd.DataFrame(contenido["filas"], columns=contenido["headers"]).astype(str).replace("None", "")
            busqueda = st.text_input("Buscar (en cualquier columna)", key="busqueda_mae")
            if busqueda:
                mask = df_mae.apply(lambda col: col.astype(str).str.contains(busqueda, case=False, na=False)).any(axis=1)
                df_mae = df_mae[mask]
            st.caption(f"{len(df_mae)} fila(s) mostradas.")
            st.dataframe(df_mae, width="stretch", height=350)

    else:
        opciones_param = {f"{k} — {TIPIFICACIONES_INDICE.get(k, '')}": k for k in TIPIFICACIONES_DATOS.keys()}
        seleccion_param = st.selectbox(
            "Tabla de Tipificación / Parámetro",
            sorted(opciones_param.keys(), key=lambda s: int(s.split(' — ')[0].split('_')[1])),
            key="sel_tabla_param",
        )
        tabla_key = opciones_param[seleccion_param]
        contenido = TIPIFICACIONES_DATOS[tabla_key]
        df_param = pd.DataFrame(contenido["filas"], columns=contenido["headers"]).astype(str).replace("None", "")
        busqueda_p = st.text_input("Buscar (en cualquier columna)", key="busqueda_param")
        if busqueda_p:
            mask = df_param.apply(lambda col: col.astype(str).str.contains(busqueda_p, case=False, na=False)).any(axis=1)
            df_param = df_param[mask]
        st.caption(f"{len(df_param)} fila(s) mostradas.")
        st.dataframe(df_param, width="stretch", height=350)

    st.markdown("**Búsqueda global de un código MCO**")
    busqueda_mco = st.text_input("Buscar en el catálogo de MCO (código o descripción)", key="busqueda_mco")
    if busqueda_mco:
        resultados_mco = [
            (k, v) for k, v in MCO_CATALOGO_COMPLETO.items()
            if busqueda_mco.upper() in k.upper() or busqueda_mco.upper() in v.upper()
        ]
        if resultados_mco:
            st.dataframe(pd.DataFrame(resultados_mco, columns=["MCO", "Descripción"]), width="stretch")
        else:
            st.caption("Sin resultados.")

if vista_activa == "REP":
    st.markdown(
        """
**Lógica aplicada**
- **GRH / GGV**: % de dedicación (por persona o por activo) aplicado a los montos de gasto.
- **GCP / GGI**: ya vienen abiertas por recurso y servicio.
- **GGM / OGG / MEI / ST**: sin apertura por servicio — por defecto 100% va al
  servicio regulado 1101; se puede parametrizar % a servicios no regulados
  en los paneles de abajo.
- **GPA**: 6 tablas, cada una con un servicio regulado FIJO (GPA_1→1201,
  GPA_2→1202, GPA_3→1203, GPA_4→1204, GPA_5→1205, GPA_6→1201). No requiere
  parametrización de servicios no regulados.
- **Generación parcial**: puedes generar REP_2 con solo algunas tablas
  cargadas. Las familias/recursos sin datos quedan en 0, y se muestra un
  aviso detallado de qué faltó o llegó vacío.
    """
    )

    st.subheader("Parametrización opcional — familias sin apertura por servicio")
    st.caption("Por defecto todo el gasto se asigna 100% al servicio regulado 1101.")

    with st.expander("Configurar parametrización GGM / OGG / MEI / ST", expanded=False):
        ggm_params = panel_parametrizacion("GGM (recursos 2401-2411)", RECURSOS_GGM, "ggm_overrides")
        st.divider()
        ogg_params = panel_parametrizacion("OGG (recursos 2501-2550)", RECURSOS_OGG, "ogg_overrides")
        st.divider()
        mei_params = panel_parametrizacion("MEI (recursos 4101-4106)", RECURSOS_MEI, "mei_overrides")
        st.divider()
        st.caption("ST: algunos códigos son compartidos por varias tablas (ej. 5110 en ST_22 a ST_30); la parametrización se aplica por código de recurso, no por tabla.")
        st_params = panel_parametrizacion("ST (Servicios Tercerizados)", RECURSOS_ST, "st_overrides")

    if "ggm_overrides" not in st.session_state:
        ggm_params = {}
    if "ogg_overrides" not in st.session_state:
        ogg_params = {}
    if "mei_overrides" not in st.session_state:
        mei_params = {}
    if "st_overrides" not in st.session_state:
        st_params = {}

    # Validar overflow > 100%
    def check_overflow(params):
        return {r: sum(p for _, p in lst) for r, lst in params.items() if sum(p for _, p in lst) > 1.0}

    overflow = {}
    overflow.update(check_overflow(ggm_params))
    overflow.update(check_overflow(ogg_params))
    overflow.update(check_overflow(mei_params))
    overflow.update(check_overflow(st_params))
    if overflow:
        st.error(f"La suma de % parametrizados supera 100% para el(los) recurso(s): {list(overflow.keys())}. Ajusta los valores.")

    run = st.button("Generar REP_2", type="primary", disabled=not hay_algo_cargado or bool(overflow))

    if not hay_algo_cargado:
        st.info("Sube al menos un archivo en el panel izquierdo para habilitar la generación (puede ser parcial).")

    if run:
        fb = {
            "grh8": f_grh8.getvalue() if f_grh8 else None,
            "grh11": f_grh11.getvalue() if f_grh11 else None,
            "gcp4": f_gcp4.getvalue() if f_gcp4 else None,
            "gcp5": f_gcp5.getvalue() if f_gcp5 else None,
            "ggv4": f_ggv4.getvalue() if f_ggv4 else None,
            "ggv5": f_ggv5.getvalue() if f_ggv5 else None,
            "ggi5": f_ggi5.getvalue() if f_ggi5 else None,
            "ggm1": f_ggm1.getvalue() if f_ggm1 else None,
            "ggm2": f_ggm2.getvalue() if f_ggm2 else None,
            "ggm3": f_ggm3.getvalue() if f_ggm3 else None,
            "ggm4": f_ggm4.getvalue() if f_ggm4 else None,
            "ggm5": f_ggm5.getvalue() if f_ggm5 else None,
            "ogg5": f_ogg5.getvalue() if f_ogg5 else None,
            "mei1": f_mei1.getvalue() if f_mei1 else None,
            "mei2": f_mei2.getvalue() if f_mei2 else None,
            "mei3": f_mei3.getvalue() if f_mei3 else None,
            "mei4": f_mei4.getvalue() if f_mei4 else None,
        }

        st_files = {}
        avisos_carga_archivos = []
        for f in (f_st_files or []):
            tabla = identificar_tabla_st(f.name)
            if tabla is None:
                avisos_carga_archivos.append(f"⚠️ El archivo **{f.name}** no coincide con ninguna tabla ST_3..ST_34 conocida; se ignora.")
                continue
            try:
                st_files[tabla] = leer_tabla_st(f.getvalue())
            except Exception as e:
                avisos_carga_archivos.append(f"⚠️ No se pudo leer **{f.name}** ({tabla}): {e}. Se excluye del cálculo.")

        gpa_files = {}
        for f in (f_gpa_files or []):
            tabla = identificar_tabla_gpa(f.name)
            if tabla is None:
                avisos_carga_archivos.append(f"⚠️ El archivo **{f.name}** no coincide con ninguna tabla GPA_1..GPA_6 conocida; se ignora.")
                continue
            try:
                gpa_files[tabla] = leer_tabla_st(f.getvalue())  # misma lectura robusta por encabezado
            except Exception as e:
                avisos_carga_archivos.append(f"⚠️ No se pudo leer **{f.name}** ({tabla}): {e}. Se excluye del cálculo.")

        try:
            final_rows, checks, familia_map, by_recurso_planas, gpa_detalle, avisos = build_rep2(
                fb, ggm_params, ogg_params, mei_params, st_files, st_params, gpa_files
            )
            avisos = avisos_carga_archivos + avisos
        except Exception as e:
            st.error(f"Error procesando los archivos: {e}")
            st.stop()

        df = pd.DataFrame(final_rows, columns=HEADERS)
        if len(df) == 0:
            st.error("No se generó ninguna fila. Revisa que al menos una tabla tenga datos válidos.")
            st.stop()
        st.success(f"REP_2 generado con {len(df)} filas.")
        mostrar_comparacion_anterior("diff_prev_rep2", final_rows, idx_agrupacion=(4, 5), idx_valor=7, etiqueta_grupo="(Recurso, Familia)")
        st.session_state['last_final_rows'] = final_rows
        st.session_state['last_familia_map'] = familia_map
        st.session_state['last_by_recurso_planas'] = by_recurso_planas
        st.session_state['last_gpa_detalle'] = gpa_detalle
        st.session_state['last_params_by_familia'] = {"GGM": ggm_params, "OGG": ogg_params, "MEI": mei_params, "ST": st_params}
        st.session_state['last_avisos'] = avisos
        st.session_state['last_checks'] = checks

    if 'last_final_rows' in st.session_state:
        final_rows = st.session_state['last_final_rows']
        familia_map = st.session_state['last_familia_map']
        by_recurso_planas = st.session_state['last_by_recurso_planas']
        gpa_detalle = st.session_state['last_gpa_detalle']
        avisos = st.session_state['last_avisos']
        checks = st.session_state['last_checks']
        df = pd.DataFrame(final_rows, columns=HEADERS)
        mostrar_kpis_rep2(final_rows, avisos)

        if avisos:
            with st.expander(f"⚠️ Avisos de carga ({len(avisos)})", expanded=True):
                for aviso in avisos:
                    st.markdown(f"- {aviso}")

        st.subheader("Validación de cuadratura por familia de gasto")
        st.caption("Solo se muestran las familias con datos cargados.")
        cols = st.columns(max(len(checks), 1))
        all_ok = True
        for col, (fam, chk) in zip(cols, checks.items()):
            with col:
                st.markdown(f"**{fam}**")
                st.metric("Δ GASTO", f"{chk['diff_gasto']:,.1f}")
                st.metric("Δ ACTIVADO", f"{chk['diff_act']:,.1f}")
                if abs(chk["diff_gasto"]) > 1 or abs(chk["diff_act"]) > 1:
                    all_ok = False
                    st.warning("Diff > $1")
                else:
                    st.info("OK")

        if not all_ok:
            st.warning("Alguna familia presenta una diferencia de cuadratura mayor a $1. Revisa los archivos fuente.")

        st.subheader("Detalle REP_2")
        st.dataframe(
            df.style.format({
                "% NO ACTIVADO ASIGNADO FAMILIA SERVICIOS": "{:.2f}%",
                "% ACTIVADO ASIGNADO FAMILIA SERVICIOS": "{:.2f}%",
                "GASTO ANUAL": "{:,.0f}",
                "MONTO ACTIVADO": "{:,.0f}",
            }),
            width='stretch',
        )

        params_by_familia = {"GGM": ggm_params, "OGG": ogg_params, "MEI": mei_params, "ST": st_params}
        excel_bytes = build_excel(
            final_rows, familia_map, by_recurso_planas, params_by_familia,
            gpa_detalle=gpa_detalle,
            avisos=avisos,
            template_bytes=f_template.getvalue() if f_template else None,
        )
        st.download_button(
            "Descargar REP_2.xlsx",
            data=excel_bytes,
            file_name=nombre_con_sufijo("REP_2", "xlsx"),

            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    # ============================================================================
    # PANEL DE VISUALIZACIÓN: evolución histórica y detección de anomalías
    # ============================================================================
    st.divider()
    st.header("📊 Panel de Visualización")
    st.caption(
        "Compara el año actual (generado arriba) contra un libro histórico "
        "consolidado (ej. REP_2_2020-2024.xlsx) para ver la evolución del gasto "
        "por Familia de Gasto, Código de Recurso y Familia de Servicio, además "
        "de detectar variaciones interanuales atípicas."
    )

    f_historico = st.file_uploader(
        "Libro histórico consolidado REP_2 (opcional, ej. REP_2_2020-2024.xlsx)",
        type="xlsx", key="historico_viz"
    )


    def cargar_historico(file_bytes):
        if file_bytes is None:
            return []
        wb_h = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
        ws_h = wb_h.active
        return [r for r in ws_h.iter_rows(min_row=2, values_only=True) if r[0] is not None]


    def gasto_total_fila(r):
        # Solo GASTO ANUAL (no activado). No se suma MONTO ACTIVADO.
        return r[7] or 0


    def detectar_anomalias(evol_dict, nombre_dim):
        anomalias = []
        for key, serie_dict in evol_dict.items():
            anios_serie = sorted(serie_dict.keys())
            if len(anios_serie) < 2:
                continue
            valores = [serie_dict[a] for a in anios_serie]
            cambios = []
            for i in range(1, len(valores)):
                prev, curr = valores[i - 1], valores[i]
                pct_cambio = None if prev == 0 else (curr - prev) / prev
                cambios.append((anios_serie[i], prev, curr, pct_cambio))
            pct_validos = [c[3] for c in cambios if c[3] is not None]
            if len(pct_validos) >= 3:
                media = statistics.mean(pct_validos)
                desv = statistics.stdev(pct_validos)
            else:
                media, desv = None, None
            for anio_actual, prev, curr, pct_cambio in cambios:
                if pct_cambio is None:
                    continue
                es_anomalia = False
                motivo = ""
                if desv and desv > 1e-9:
                    z = (pct_cambio - media) / desv
                    if abs(z) > 2:
                        es_anomalia = True
                        motivo = f"Variación atípica (z={z:.1f}) vs. patrón histórico"
                else:
                    if abs(pct_cambio) > 0.5:
                        es_anomalia = True
                        motivo = "Variación mayor a 50% (serie corta)"
                if es_anomalia:
                    anomalias.append({
                        "Dimensión": nombre_dim, "Grupo": key, "Año": anio_actual,
                        "Valor Año Anterior": round(prev, 0), "Valor Año Actual": round(curr, 0),
                        "% Variación": round(pct_cambio, 4), "Motivo": motivo,
                    })
        return anomalias


    final_rows_actual = st.session_state.get('last_final_rows', [])
    historico_rows = cargar_historico(f_historico.getvalue() if f_historico else None)
    combinado = list(historico_rows) + list(final_rows_actual)

    if not combinado:
        st.info("Genera el REP_2 del año actual arriba y/o sube el libro histórico para ver el panel de visualización.")
    else:
        anios_disponibles = sorted(set(r[2] for r in combinado))

        evol_familia = defaultdict(lambda: defaultdict(float))
        evol_recurso = defaultdict(lambda: defaultdict(float))
        evol_famserv = defaultdict(lambda: defaultdict(float))
        FAMSERV_NOMBRE = {11: '11 - Servicios Sanitarios (Agua/Alcantarillado)', 12: '12 - Servicios Sanitarios (Otros)', 22: '22 - Servicios No Regulados'}

        for r in combinado:
            fam = familia_de_recurso(r[4])
            evol_familia[fam][r[2]] += gasto_total_fila(r)
            evol_recurso[r[4]][r[2]] += gasto_total_fila(r)
            evol_famserv[r[5]][r[2]] += gasto_total_fila(r)

        nombre_recurso_map = {cod: f"{cod} - {nombre_recurso(cod)}" for cod in evol_recurso.keys()}
        famserv_map = {fs: FAMSERV_NOMBRE.get(fs, str(fs)) for fs in evol_famserv.keys()}

        vista = st.radio(
            "Ver evolución por:",
            ["Familia de Gasto", "Código de Recurso", "Familia de Servicio"],
            horizontal=True,
        )

        if vista == "Familia de Gasto":
            evol_dict, label_map = evol_familia, {k: k for k in evol_familia}
        elif vista == "Código de Recurso":
            evol_dict, label_map = evol_recurso, nombre_recurso_map
        else:
            evol_dict, label_map = evol_famserv, famserv_map

        tabla = pd.DataFrame({
            label_map[k]: [evol_dict[k].get(a, 0) for a in anios_disponibles]
            for k in evol_dict
        }, index=anios_disponibles)
        tabla.index.name = "Año"

        if vista == "Código de Recurso" and len(tabla.columns) > 15:
            top_cols = tabla.sum(axis=0).sort_values(ascending=False).head(15).index
            st.caption("Mostrando los 15 códigos de recurso con mayor gasto acumulado (usa la tabla completa para ver el resto).")
            st.line_chart(tabla[top_cols])
        else:
            st.line_chart(tabla)

        with st.expander("Ver tabla de datos"):
            st.dataframe(tabla.style.format("{:,.0f}"), width='stretch')

        # --- Detección de anomalías ---
        st.subheader("⚠️ Anomalías detectadas (variaciones interanuales atípicas)")
        anomalias_familia = detectar_anomalias(evol_familia, "Familia de Gasto")
        anomalias_recurso = detectar_anomalias({nombre_recurso_map[k]: v for k, v in evol_recurso.items()}, "Código de Recurso")
        anomalias_famserv = detectar_anomalias({famserv_map[k]: v for k, v in evol_famserv.items()}, "Familia de Servicio")
        todas_anomalias = anomalias_familia + anomalias_recurso + anomalias_famserv

        if todas_anomalias:
            df_anom = pd.DataFrame(todas_anomalias).sort_values("% Variación", key=lambda s: s.abs(), ascending=False)
            st.dataframe(
                df_anom.style.format({
                    "Valor Año Anterior": "{:,.0f}",
                    "Valor Año Actual": "{:,.0f}",
                    "% Variación": "{:+.1%}",
                }).background_gradient(subset=["% Variación"], cmap="RdYlGn_r"),
                width='stretch',
            )
        else:
            st.success("No se detectaron variaciones interanuales atípicas con los criterios actuales.")

        st.caption(
            "Criterio: para series con ≥3 variaciones interanuales se usa z-score "
            "(|z|>2) sobre el propio histórico de cada grupo; para series más "
            "cortas se marca si la variación supera 50%."
        )

        # --- Descarga del Excel completo (REP_2 + Panel de Visualización) ------
        st.divider()
        if 'last_final_rows' not in st.session_state:
            st.info("Genera el REP_2 del año actual arriba para poder descargar el Excel con el panel de visualización incluido.")
        else:
            if st.button("📥 Preparar Excel con Panel de Visualización (REP_2 + evolución histórica)"):
                excel_base = build_excel(
                    st.session_state['last_final_rows'],
                    st.session_state['last_familia_map'],
                    st.session_state['last_by_recurso_planas'],
                    st.session_state['last_params_by_familia'],
                    gpa_detalle=st.session_state.get('last_gpa_detalle'),
                    avisos=st.session_state.get('last_avisos'),
                    template_bytes=f_template.getvalue() if 'f_template' in dir() and f_template else None,
                )
                wb_completo = openpyxl.load_workbook(excel_base)

                header_fill = PatternFill("solid", start_color="1F4E78", end_color="1F4E78")
                header_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
                data_font = Font(name="Arial", size=10)
                thin = Side(style="thin", color="D9D9D9")
                border = Border(left=thin, right=thin, top=thin, bottom=thin)

                ws5 = wb_completo.create_sheet("Evolucion_Familia_Gasto")
                ws5.append(["Familia de Gasto"] + [str(a) for a in anios_disponibles])
                for c in range(1, len(anios_disponibles) + 2):
                    cell = ws5.cell(row=1, column=c)
                    cell.font = header_font; cell.fill = header_fill; cell.border = border
                    cell.alignment = Alignment(horizontal="center")
                for fam_k in sorted(evol_familia.keys()):
                    ws5.append([fam_k] + [round(evol_familia[fam_k].get(a, 0), 2) for a in anios_disponibles])
                for r in range(2, ws5.max_row + 1):
                    for c in range(1, len(anios_disponibles) + 2):
                        cell = ws5.cell(row=r, column=c)
                        cell.font = data_font; cell.border = border
                        if c > 1:
                            cell.number_format = '#,##0;(#,##0);"-"'
                ws5.column_dimensions["A"].width = 40
                for c in range(2, len(anios_disponibles) + 2):
                    ws5.column_dimensions[get_column_letter(c)].width = 16
                if ws5.max_row > 1 and len(anios_disponibles) > 1:
                    chart1 = LineChart()
                    chart1.title = "Evolución del Gasto por Familia"
                    chart1.y_axis.title = "Gasto Anual ($)"
                    chart1.x_axis.title = "Año"
                    data_ref = Reference(ws5, min_col=2, max_col=len(anios_disponibles) + 1, min_row=1, max_row=ws5.max_row)
                    cats_ref = Reference(ws5, min_col=1, min_row=2, max_row=ws5.max_row)
                    chart1.add_data(data_ref, titles_from_data=True)
                    chart1.set_categories(cats_ref)
                    chart1.width = 26; chart1.height = 12
                    ws5.add_chart(chart1, f"A{ws5.max_row + 3}")

                ws6 = wb_completo.create_sheet("Evolucion_Recurso")
                ws6.append(["Código Recurso", "Nombre Recurso"] + [str(a) for a in anios_disponibles])
                for c in range(1, len(anios_disponibles) + 3):
                    cell = ws6.cell(row=1, column=c)
                    cell.font = header_font; cell.fill = header_fill; cell.border = border
                    cell.alignment = Alignment(horizontal="center")
                for cod in sorted(evol_recurso.keys()):
                    ws6.append([cod, nombre_recurso(cod)] + [round(evol_recurso[cod].get(a, 0), 2) for a in anios_disponibles])
                for r in range(2, ws6.max_row + 1):
                    for c in range(1, len(anios_disponibles) + 3):
                        cell = ws6.cell(row=r, column=c)
                        cell.font = data_font; cell.border = border
                        if c > 2:
                            cell.number_format = '#,##0;(#,##0);"-"'
                ws6.column_dimensions["A"].width = 14
                ws6.column_dimensions["B"].width = 45
                for c in range(3, len(anios_disponibles) + 3):
                    ws6.column_dimensions[get_column_letter(c)].width = 16
                ws6.freeze_panes = "C2"
                if ws6.max_row > 1:
                    ultima_col = len(anios_disponibles) + 2
                    rango = f"{get_column_letter(ultima_col)}2:{get_column_letter(ultima_col)}{ws6.max_row}"
                    ws6.conditional_formatting.add(rango, ColorScaleRule(
                        start_type="min", start_color="63BE7B",
                        mid_type="percentile", mid_value=50, mid_color="FFEB84",
                        end_type="max", end_color="F8696B"))

                ws7 = wb_completo.create_sheet("Evolucion_Familia_Servicio")
                ws7.append(["Familia de Servicio"] + [str(a) for a in anios_disponibles])
                for c in range(1, len(anios_disponibles) + 2):
                    cell = ws7.cell(row=1, column=c)
                    cell.font = header_font; cell.fill = header_fill; cell.border = border
                    cell.alignment = Alignment(horizontal="center")
                for fs in sorted(evol_famserv.keys()):
                    ws7.append([famserv_map[fs]] + [round(evol_famserv[fs].get(a, 0), 2) for a in anios_disponibles])
                for r in range(2, ws7.max_row + 1):
                    for c in range(1, len(anios_disponibles) + 2):
                        cell = ws7.cell(row=r, column=c)
                        cell.font = data_font; cell.border = border
                        if c > 1:
                            cell.number_format = '#,##0;(#,##0);"-"'
                ws7.column_dimensions["A"].width = 42
                for c in range(2, len(anios_disponibles) + 2):
                    ws7.column_dimensions[get_column_letter(c)].width = 16
                if ws7.max_row > 1 and len(anios_disponibles) > 1:
                    chart2 = BarChart()
                    chart2.type = "col"; chart2.grouping = "clustered"
                    chart2.title = "Evolución del Gasto por Familia de Servicio"
                    chart2.y_axis.title = "Gasto Anual ($)"
                    chart2.x_axis.title = "Año"
                    data2 = Reference(ws7, min_col=2, max_col=len(anios_disponibles) + 1, min_row=1, max_row=ws7.max_row)
                    cats2 = Reference(ws7, min_col=1, min_row=2, max_row=ws7.max_row)
                    chart2.add_data(data2, titles_from_data=True)
                    chart2.set_categories(cats2)
                    chart2.width = 24; chart2.height = 11
                    ws7.add_chart(chart2, f"A{ws7.max_row + 3}")

                ws8 = wb_completo.create_sheet("Anomalias_Detectadas")
                ws8.append(["Dimensión", "Grupo", "Año con Anomalía", "Valor Año Anterior", "Valor Año Actual", "% Variación", "Motivo"])
                for c in range(1, 8):
                    cell = ws8.cell(row=1, column=c)
                    cell.font = header_font; cell.fill = header_fill; cell.border = border
                    cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                ws8.row_dimensions[1].height = 30
                if todas_anomalias:
                    for a_row in todas_anomalias:
                        ws8.append([a_row["Dimensión"], a_row["Grupo"], a_row["Año"], a_row["Valor Año Anterior"],
                                    a_row["Valor Año Actual"], a_row["% Variación"], a_row["Motivo"]])
                else:
                    ws8.append(["(Sin anomalías detectadas con los criterios actuales)", "", "", "", "", "", ""])
                for r in range(2, ws8.max_row + 1):
                    for c in range(1, 8):
                        cell = ws8.cell(row=r, column=c)
                        cell.font = data_font; cell.border = border
                        if c in (4, 5):
                            cell.number_format = '#,##0;(#,##0);"-"'
                        if c == 6 and isinstance(ws8.cell(row=r, column=6).value, float):
                            cell.number_format = "+0.00%;-0.00%"
                            val = ws8.cell(row=r, column=6).value
                            cell.fill = PatternFill("solid", start_color="FFC7CE" if val > 0 else "C6EFCE",
                                                     end_color="FFC7CE" if val > 0 else "C6EFCE")
                ws8.column_dimensions["A"].width = 16
                ws8.column_dimensions["B"].width = 42
                ws8.column_dimensions["C"].width = 14
                ws8.column_dimensions["D"].width = 18
                ws8.column_dimensions["E"].width = 18
                ws8.column_dimensions["F"].width = 12
                ws8.column_dimensions["G"].width = 55
                ws8.freeze_panes = "A2"

                out_completo = io.BytesIO()
                wb_completo.save(out_completo)
                out_completo.seek(0)
                st.session_state['excel_completo_bytes'] = out_completo.getvalue()
                st.success("Excel listo. Usa el botón de descarga que apareció abajo.")

            if 'excel_completo_bytes' in st.session_state:
                st.download_button(
                    "Descargar REP_2 + Panel de Visualización.xlsx",
                    data=st.session_state['excel_completo_bytes'],
                    file_name=nombre_con_sufijo("REP_2_con_evolucion", "xlsx"),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

    # ============================================================================
    # GENERADOR REP_3 - Costos y Gastos No Activados de Servicios Regulados - Proceso
    # ============================================================================
    st.divider()
    st.header("📋 Generar REP_3 (apertura por Código de Proceso)")
    st.caption(
        "REP_3 toma solo el gasto REGULADO de REP_2 (familias 11 y 12) y lo abre "
        "por Código de Proceso. El Código de Proceso corresponde a los primeros "
        "3 dígitos del Código Actividad, según MAE_1 del Maestro SISS."
    )

    HEADERS_REP3 = [
        "CÓDIGO EMPRESA", "PERÍODO INFORMACIÓN", "AÑO INFORMADO",
        "CÓDIGO SECTOR DECRETO TARIFARIO", "CÓDIGO RECURSO",
        "CÓDIGO FAMILIA SERVICIOS REGULADOS", "CÓDIGO PROCESO",
        "% ASIGNADO PROCESO", "GASTO ANUAL",
    ]


    def build_excel_rep3(final_rows3):
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "REP_3"

        header_fill = PatternFill("solid", start_color="1F4E78", end_color="1F4E78")
        header_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
        data_font = Font(name="Arial", size=10)
        thin = Side(style="thin", color="D9D9D9")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        ws.append(HEADERS_REP3)
        for c in range(1, len(HEADERS_REP3) + 1):
            cell = ws.cell(row=1, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        ws.row_dimensions[1].height = 30

        for row in final_rows3:
            ws.append(row)

        for r in range(2, ws.max_row + 1):
            for c in range(1, len(HEADERS_REP3) + 1):
                cell = ws.cell(row=r, column=c)
                cell.font = data_font
                cell.border = border
                if c == 8:
                    cell.number_format = "0.00"
                elif c == 9:
                    cell.number_format = '#,##0;(#,##0);"-"'
                else:
                    cell.alignment = Alignment(horizontal="center")

        widths = [14, 16, 14, 20, 14, 24, 14, 16, 18]
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A2"

        # Hoja de detalle por proceso (nombre descriptivo)
        ws2 = wb.create_sheet("Detalle_Procesos")
        ws2.append(["Código Proceso", "Nombre Proceso", "Gasto Anual Total"])
        for c in range(1, 4):
            cell = ws2.cell(row=1, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.border = border
        por_proceso = defaultdict(float)
        for row in final_rows3:
            por_proceso[row[6]] += row[8]
        for proceso in sorted(por_proceso.keys()):
            ws2.append([proceso, nombre_proceso(proceso), round(por_proceso[proceso], 2)])
        for r in range(2, ws2.max_row + 1):
            for c in range(1, 4):
                cell = ws2.cell(row=r, column=c)
                cell.font = data_font
                cell.border = border
                if c == 3:
                    cell.number_format = '#,##0;(#,##0);"-"'
        ws2.column_dimensions["A"].width = 14
        ws2.column_dimensions["B"].width = 55
        ws2.column_dimensions["C"].width = 20

        out = io.BytesIO()
        wb.save(out)
        out.seek(0)
        return out


    def panel_parametrizacion_proceso_generico(nombre, recursos_disponibles, session_key, defaults_map):
        if session_key not in st.session_state:
            st.session_state[session_key] = []
        st.markdown(f"**{nombre}**")
        col1, col2, col3, col4 = st.columns([1.2, 2.5, 1, 0.8])
        with col1:
            sel_recurso = st.selectbox("Código Recurso", recursos_disponibles, key=f"sel_recurso_proc_{session_key}")
        with col2:
            sel_proceso = st.selectbox(
                "Proceso destino",
                options=sorted(PROCESO_NOMBRE.keys()),
                format_func=lambda c: f"{c} - {PROCESO_NOMBRE[c]}",
                key=f"sel_proceso_{session_key}",
            )
        with col3:
            sel_pct = st.number_input(
                "% asignado", min_value=0.0, max_value=100.0, value=10.0,
                step=0.0000000001, format="%.10f", key=f"sel_pct_proc_{session_key}",
                help="Puedes ingresar hasta 10 decimales para minimizar la diferencia de redondeo en el monto resultante.",
            )
        with col4:
            st.write("")
            st.write("")
            if st.button("Agregar", key=f"add_proc_{session_key}"):
                st.session_state[session_key].append((sel_recurso, sel_proceso, sel_pct / 100.0))
        if st.session_state[session_key]:
            for i, (cod_r, cod_p, pct) in enumerate(st.session_state[session_key]):
                c1, c2, c3, c4 = st.columns([1.2, 2.5, 1, 0.8])
                c1.write(cod_r)
                c2.write(f"{cod_p} - {PROCESO_NOMBRE.get(cod_p, '')}")
                c3.write(f"{pct:.9%}")
                if c4.button("Quitar", key=f"del_proc_{session_key}_{i}"):
                    st.session_state[session_key].pop(i)
                    st.rerun()
        else:
            default_desc = ", ".join(f"{r}→{p}" for r, p in list(defaults_map.items())[:3])
            st.caption(f"Sin parametrización: se usa el default (ej. {default_desc}...).")
        params = defaultdict(list)
        for cod_r, cod_p, pct in st.session_state[session_key]:
            params[cod_r].append((cod_p, pct))
        return dict(params)


    with st.expander("📂 Parametrización de proceso para REP_3", expanded=False):
        st.caption(
            "REP_3 reutiliza automáticamente GRH_8/11/12, GCP_5/6, GGV_4/5/6, GGI_5/6, "
            "GGM, OGG, MEI, ST y GPA desde el mismo zip cargado arriba en la barra lateral "
            "— no hace falta volver a subir nada aquí."
        )
        faltan_para_rep3 = [t for t in ["GRH_12", "GCP_6", "GGV_6", "GGI_6"] if t not in archivos_zip]
        if faltan_para_rep3:
            st.warning(f"Aún faltan (para la apertura por actividad de REP_3): {', '.join(faltan_para_rep3)}")
        else:
            st.success("Todas las tablas de actividad (GRH_12, GCP_6, GGV_6, GGI_6) están cargadas.")

        st.markdown("**Parametrización de proceso — GGM y OGG (sin tabla de apertura)**")
        st.caption(
            "Por defecto: GGM → Informática (505) para recursos de TI/telecom, "
            "Abastecimiento y Servicios Generales (504) para materiales; "
            "OGG → Dirección Superior (501) para gastos de Directorio, "
            "Administración y Finanzas (502) para el resto. Editable abajo."
        )
        ggm_proceso_params = panel_parametrizacion_proceso_generico(
            "GGM (recursos 2401-2411)", RECURSOS_GGM, "ggm_proceso_overrides", DEFAULT_PROCESO_GGM
        )
        st.divider()
        ogg_proceso_params = panel_parametrizacion_proceso_generico(
            "OGG (recursos 2501-2550)", RECURSOS_OGG, "ogg_proceso_overrides", DEFAULT_PROCESO_OGG
        )
        st.divider()
        st.caption(
            "MEI_1 (Productos químicos, sin actividad): por defecto 70% al proceso "
            "103 (tratamiento AP) y 30% al proceso 204 (tratamiento AS), en base a "
            "la distribución real de MEI_2 (Energía Eléctrica) entre esos mismos "
            "procesos de tratamiento. MEI_2/3/4 ya traen su propio Código Actividad "
            "y no requieren parametrización."
        )
        mei_proceso_params = panel_parametrizacion_proceso_generico(
            "MEI_1 (recurso 4101)", [4101], "mei_proceso_overrides", DEFAULT_PROCESO_MEI1
        )
        st.divider()
        st.caption(
            "GPA: por defecto 100% al proceso 601 'Prestaciones Asociadas' (todas "
            "las tablas GPA corresponden a esa categoría según MAE_1)."
        )
        gpa_proceso_params = panel_parametrizacion_proceso_generico(
            "GPA (recursos 6101, 6201, 6301, 6401, 6501, 6601)", RECURSOS_GPA_REP3, "gpa_proceso_overrides", {r: 601 for r in [6101, 6201, 6301, 6401, 6501, 6601]}
        )

    run_rep3 = st.button("Generar REP_3", type="primary")

    if run_rep3:
        fb3 = {
            "grh8": f_grh8.getvalue() if f_grh8 else None,
            "grh11": f_grh11.getvalue() if f_grh11 else None,
            "grh12": f_grh12.getvalue() if f_grh12 else None,
            "gcp5": f_gcp5.getvalue() if f_gcp5 else None,
            "gcp6": f_gcp6.getvalue() if f_gcp6 else None,
            "ggv4": f_ggv4.getvalue() if f_ggv4 else None,
            "ggv5": f_ggv5.getvalue() if f_ggv5 else None,
            "ggv6": f_ggv6.getvalue() if f_ggv6 else None,
            "ggi5": f_ggi5.getvalue() if f_ggi5 else None,
            "ggi6": f_ggi6.getvalue() if f_ggi6 else None,
            "ggm1": f_ggm1.getvalue() if f_ggm1 else None,
            "ggm2": f_ggm2.getvalue() if f_ggm2 else None,
            "ggm3": f_ggm3.getvalue() if f_ggm3 else None,
            "ggm4": f_ggm4.getvalue() if f_ggm4 else None,
            "ggm5": f_ggm5.getvalue() if f_ggm5 else None,
            "ogg5": f_ogg5.getvalue() if f_ogg5 else None,
            "mei1": f_mei1.getvalue() if f_mei1 else None,
            "mei2": f_mei2.getvalue() if f_mei2 else None,
            "mei3": f_mei3.getvalue() if f_mei3 else None,
            "mei4": f_mei4.getvalue() if f_mei4 else None,
        }
        st_files_raw = {}
        for f in (f_st_files or []):
            tabla = identificar_tabla_st(f.name)
            if tabla:
                st_files_raw[tabla] = f.getvalue()

        gpa_files_raw_rep3 = {}
        for f in (f_gpa_files_rep3 or []):
            tabla = identificar_tabla_gpa(f.name)
            if tabla:
                gpa_files_raw_rep3[tabla] = f.getvalue()

        try:
            final_rows3, avisos3 = build_rep3(
                fb3, ggm_proceso_params, ogg_proceso_params, ggm_params, ogg_params, st_params, st_files_raw,
                mei_proceso_params=mei_proceso_params, gpa_proceso_params=gpa_proceso_params,
                mei_params=mei_params, gpa_files_raw=gpa_files_raw_rep3,
            )
        except Exception as e:
            st.error(f"Error generando REP_3: {e}")
            st.stop()

        if not final_rows3:
            st.error("No se generó ninguna fila de REP_3. Verifica que hayas cargado al menos GRH_8+GRH_11+GRH_12, o alguna otra familia con su tabla de proceso.")
        else:
            mostrar_comparacion_anterior("diff_prev_rep3", final_rows3, idx_agrupacion=(4, 5, 6), idx_valor=8, etiqueta_grupo="(Recurso, Familia, Proceso)")
            st.session_state['last_final_rows3'] = final_rows3
            st.session_state['last_avisos3'] = avisos3
            st.session_state['last_epas'] = tuple(final_rows3[0][:4])

    if 'last_final_rows3' in st.session_state:
        final_rows3 = st.session_state['last_final_rows3']
        avisos3 = st.session_state['last_avisos3']

        st.success(f"REP_3 generado con {len(final_rows3)} filas.")
        mostrar_kpis_rep3(final_rows3, avisos3)
        if avisos3:
            with st.expander(f"⚠️ Avisos REP_3 ({len(avisos3)})", expanded=True):
                for a in avisos3:
                    st.markdown(f"- {a}")

        df3 = pd.DataFrame(final_rows3, columns=HEADERS_REP3)
        st.dataframe(
            df3.style.format({"% ASIGNADO PROCESO": "{:.2f}%", "GASTO ANUAL": "{:,.0f}"}),
            width='stretch',
        )

        excel_rep3 = build_excel_rep3(final_rows3)
        st.download_button(
            "Descargar REP_3.xlsx",
            data=excel_rep3,
            file_name=nombre_con_sufijo("REP_3", "xlsx"),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    st.divider()
    st.header("📑 Listado de ID Respaldo → MCO (tablas de sustento)")
    st.caption(
        "Genera un Excel con todos los ID RESPALDO de MEI_1, MEI_2, MEI_4, "
        "GPA_1 a GPA_6, GGV_3 y OGG_5 (los que corresponda tener cargados), "
        "indicando de qué tabla vienen, su Tipo Respaldo, y a qué MCO se "
        "deben declarar según ese tipo."
    )
    with st.expander("Ver criterio de asignación MCO"):
        st.markdown(
            "**Tablas estándar** (GPA_1-6, GGV_3, OGG_5):\n"
            "- Tipo Respaldo 1 → MCO_4 (Contratos servicios recibidos)\n"
            "- Tipo Respaldo 2 → MCO_9 (Orden de compra servicio)\n"
            "- Tipo Respaldo 3 → MCO_12 (Facturas servicios)\n\n"
            "**Tablas MEI** (MEI_1, MEI_2, MEI_4):\n"
            "- Tipo Respaldo 1 → MCO_6 (Contrato Suministro insumos)\n"
            "- Tipo Respaldo 2 → MCO_9 (Orden de compra servicio)\n"
            "- Tipo Respaldo 3 → MCO_13 (Facturas de suministros)"
        )

    if st.button("Generar Listado de Respaldos", key="btn_generar_respaldos"):
        filas_resp, avisos_resp = armar_listado_respaldos(archivos_zip)
        st.session_state["last_filas_respaldos"] = filas_resp
        st.session_state["last_avisos_respaldos"] = avisos_resp

    if "last_filas_respaldos" in st.session_state:
        filas_resp = st.session_state["last_filas_respaldos"]
        avisos_resp = st.session_state["last_avisos_respaldos"]
        st.success(f"{len(filas_resp)} ID Respaldo único(s) encontrados.")
        sin_mco = sum(1 for f in filas_resp if f[3] is None)
        c1, c2 = st.columns(2)
        c1.metric("Total ID Respaldo", len(filas_resp))
        c2.metric("Sin MCO asignado", sin_mco)
        if avisos_resp:
            with st.expander(f"⚠️ Avisos ({len(avisos_resp)})", expanded=True):
                for a in avisos_resp:
                    st.markdown(f"- {a}")
        df_resp = pd.DataFrame(filas_resp, columns=["TABLA", "ID RESPALDO", "TIPO RESPALDO", "MCO", "DESCRIPCIÓN MCO"])
        df_resp["ID RESPALDO"] = df_resp["ID RESPALDO"].astype(str)
        st.dataframe(df_resp, width="stretch")
        excel_resp = build_excel_respaldos(filas_resp, avisos_resp)
        st.download_button(
            "Descargar Listado_Respaldos_MCO.xlsx",
            data=excel_resp,
            file_name=nombre_con_sufijo("Listado_Respaldos_MCO", "xlsx"),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    st.divider()
    st.header("🔗 Consolidado ID Servicio / Subservicio Tercerizado")
    st.caption(
        "Consolida el ID SERVICIO TERCERIZADO y el ID SUBSERVICIO TERCERIZADO "
        "de ST_3 a ST_34, GGI_1, GGI_2, GGM_1, GGM_2 y GGM_3 (las que tengas "
        "cargadas). ST_3-ST_10 solo traen subservicio (se deriva el servicio "
        "quitando la primera 'S'); GGI_1/2 y GGM_1-3 solo traen servicio (se "
        "deriva el subservicio agregando una 'S' al inicio); ST_11-ST_34 ya "
        "traen ambos."
    )
    with st.expander("Ver criterio de asignación MCO (hoja Agregado)"):
        st.markdown(
            "- ID Servicio empieza con 'CO-' o contiene '-CO-' → Tipo 1 → MCO_4 (Contrato Servicios)\n"
            "- ID Servicio empieza con 'OC-' o contiene '-OC-' → Tipo 2 → MCO_9 (Orden de Compra servicios)\n"
            "- ID Servicio empieza con 'FA-' o contiene '-FA-' → Tipo 3 → MCO_12 (Factura Servicios)\n"
            "- Si no calza con ninguno → Tipo 9999 → 'Asignar MCO' (requiere revisión manual)"
        )

    if st.button("Generar Consolidado Servicio Tercerizado", key="btn_generar_servicio_tercerizado"):
        trazab_serv, agregado_serv, avisos_serv = armar_trazabilidad_servicio_tercerizado(archivos_zip)
        st.session_state["last_trazab_servicio"] = trazab_serv
        st.session_state["last_agregado_servicio"] = agregado_serv
        st.session_state["last_avisos_servicio"] = avisos_serv

    if "last_trazab_servicio" in st.session_state:
        trazab_serv = st.session_state["last_trazab_servicio"]
        agregado_serv = st.session_state["last_agregado_servicio"]
        avisos_serv = st.session_state["last_avisos_servicio"]
        st.success(f"{len(trazab_serv)} fila(s) de trazabilidad · {len(agregado_serv)} ID Servicio único(s) en el agregado.")
        sin_mco_serv = sum(1 for f in agregado_serv if f[3] == 9999)
        c1, c2, c3 = st.columns(3)
        c1.metric("Filas trazabilidad", len(trazab_serv))
        c2.metric("ID Servicio únicos", len(agregado_serv))
        c3.metric("Sin MCO reconocido (9999)", sin_mco_serv)
        if avisos_serv:
            with st.expander(f"⚠️ Avisos ({len(avisos_serv)})", expanded=False):
                for a in avisos_serv:
                    st.markdown(f"- {a}")
        tab1, tab2 = st.tabs(["Trazabilidad", "Agregado"])
        with tab1:
            df_trazab = pd.DataFrame(trazab_serv, columns=["TABLA", "CÓDIGO RECURSO", "ID SERVICIO TERCERIZADO", "ID SUBSERVICIO TERCERIZADO", "GASTO ANUAL", "MONTO ACTIVADO"])
            st.dataframe(df_trazab, width="stretch")
        with tab2:
            df_agregado = pd.DataFrame(agregado_serv, columns=["ID SERVICIO TERCERIZADO", "GASTO ANUAL TOTAL", "MONTO ACTIVADO TOTAL", "TIPO RESPALDO", "MCO", "DESCRIPCIÓN MCO"])
            st.dataframe(df_agregado, width="stretch")
        excel_serv = build_excel_servicio_tercerizado(trazab_serv, agregado_serv, avisos_serv)
        st.download_button(
            "Descargar Consolidado_Servicio_Tercerizado.xlsx",
            data=excel_serv,
            file_name=nombre_con_sufijo("Consolidado_Servicio_Tercerizado", "xlsx"),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )


if vista_activa == "CYG":
    # ============================================================================
    # GENERADOR CYG - Costos y Gastos por Recurso (CYG_1-4, CYG_8, CYG_9)
    # ============================================================================
    st.divider()
    st.header("📊 Generar tablas CYG")
    st.caption(
        "Las tablas CYG (Costos y Gastos por Recurso) se construyen SIEMPRE desde "
        "las tablas fuente (no desde REP_2/REP_3, que ya colapsan el servicio "
        "específico y la actividad específica en familia/proceso). REP_2 y REP_3 "
        "se usan aquí solo como checkpoint de validación de cuadratura."
    )
    st.caption(
        "CYG_1: servicios 1101/1102 · CYG_2: servicios 1201-1205 · "
        "CYG_3: servicios no regulados 2101-2117 · CYG_4: servicios no regulados "
        "2201-2214 · CYG_8: apertura por actividad (solo gasto regulado) · "
        "CYG_9: servicios no regulados abiertos por ID Cliente (vía MCO_42 + ING_4)."
    )

    HEADERS_CYG_SERV = [
        "CÓDIGO EMPRESA", "PERÍODO INFORMACIÓN", "AÑO INFORMADO", "CÓDIGO SECTOR DECRETO TARIFARIO",
        "CÓDIGO SERVICIO", "CÓDIGO RECURSO",
        "% NO ACTIVADO ASIGNADO", "MONTO NO ACTIVADO", "% ACTIVADO ASIGNADO", "MONTO ACTIVADO",
    ]
    HEADERS_CYG_8 = [
        "CÓDIGO EMPRESA", "PERÍODO INFORMACIÓN", "AÑO INFORMADO", "CÓDIGO SECTOR DECRETO TARIFARIO",
        "CÓDIGO ACTIVIDAD", "CÓDIGO RECURSO",
        "% NO ACTIVADO ASIGNADO ACTIVIDAD", "MONTO NO ACTIVADO", "% ACTIVADO ASIGNADO ACTIVIDAD", "MONTO ACTIVADO",
    ]
    HEADERS_CYG_9 = [
        "CÓDIGO EMPRESA", "PERÍODO INFORMACIÓN", "AÑO INFORMADO", "CÓDIGO SECTOR DECRETO TARIFARIO",
        "ID CLIENTE", "CÓDIGO SERVICIO NO REGULADO", "CÓDIGO RECURSO",
        "% NO ACTIVADO ASIGNADO A CLIENTE", "MONTO NO ACTIVADO", "% ACTIVADO ASIGNADO A CLIENTE", "MONTO ACTIVADO",
    ]


    def _agregar_hoja_cyg(wb, nombre_hoja, headers, filas, pct_cols, monto_cols):
        ws = wb.create_sheet(nombre_hoja)
        header_fill = PatternFill("solid", start_color="1F4E78", end_color="1F4E78")
        header_font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
        data_font = Font(name="Arial", size=10)
        thin = Side(style="thin", color="D9D9D9")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        ws.append(headers)
        for c in range(1, len(headers) + 1):
            cell = ws.cell(row=1, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border
        ws.row_dimensions[1].height = 30

        for row in filas:
            ws.append(row)

        for r in range(2, ws.max_row + 1):
            for c in range(1, len(headers) + 1):
                cell = ws.cell(row=r, column=c)
                cell.font = data_font
                cell.border = border
                if c in pct_cols:
                    cell.number_format = "0.00"
                elif c in monto_cols:
                    cell.number_format = '#,##0;(#,##0);"-"'
                else:
                    cell.alignment = Alignment(horizontal="center")
        ws.freeze_panes = "A2"
        return ws


    def build_excel_cyg(cyg14, cyg8, cyg9, avisos_cyg):
        wb = openpyxl.Workbook()
        wb.remove(wb.active)

        for nombre in ["CYG_1", "CYG_2", "CYG_3", "CYG_4"]:
            ws = _agregar_hoja_cyg(wb, nombre, HEADERS_CYG_SERV, cyg14[nombre], pct_cols={7, 9}, monto_cols={8, 10})
            widths = [14, 16, 14, 20, 14, 14, 16, 18, 16, 18]
            for i, w in enumerate(widths, start=1):
                ws.column_dimensions[get_column_letter(i)].width = w

        ws8 = _agregar_hoja_cyg(wb, "CYG_8", HEADERS_CYG_8, cyg8, pct_cols={7, 9}, monto_cols={8, 10})
        widths8 = [14, 16, 14, 20, 14, 14, 20, 18, 20, 18]
        for i, w in enumerate(widths8, start=1):
            ws8.column_dimensions[get_column_letter(i)].width = w

        ws9 = _agregar_hoja_cyg(wb, "CYG_9", HEADERS_CYG_9, cyg9, pct_cols={8, 10}, monto_cols={9, 11})
        widths9 = [14, 16, 14, 20, 16, 20, 14, 18, 18, 18, 18]
        for i, w in enumerate(widths9, start=1):
            ws9.column_dimensions[get_column_letter(i)].width = w

        if avisos_cyg:
            wsa = wb.create_sheet("Avisos_CYG")
            wsa.append(["Aviso"])
            cell = wsa.cell(row=1, column=1)
            cell.font = Font(name="Arial", bold=True, color="FFFFFF", size=10)
            cell.fill = PatternFill("solid", start_color="1F4E78", end_color="1F4E78")
            for a in avisos_cyg:
                wsa.append([a.replace("**", "")])
            for r in range(2, wsa.max_row + 1):
                c = wsa.cell(row=r, column=1)
                c.font = Font(name="Arial", size=10)
                c.alignment = Alignment(wrap_text=True, vertical="top")
            wsa.column_dimensions["A"].width = 110

        out = io.BytesIO()
        wb.save(out)
        out.seek(0)
        return out


    with st.expander("📂 Estado de MCO_42 + ING_4 (CYG_9)", expanded=False):
        st.caption(
            "MCO_42 asigna cada servicio no regulado a un ID Ingreso. ING_4 abre "
            "cada ID Ingreso en ID Cliente + monto anual de ingreso. El gasto se "
            "reparte entre clientes proporcionalmente a su ingreso. Si un servicio "
            "no está cubierto, se declara con ID CLIENTE = '-1' (100%). Ambas se "
            "toman automáticamente del mismo zip cargado arriba en la barra lateral."
        )
        faltan_mco_ing = [t for t in ["MCO_42", "ING_4"] if t not in archivos_zip]
        if faltan_mco_ing:
            st.warning(f"Aún faltan: {', '.join(faltan_mco_ing)} (mientras tanto, CYG_9 declarará todo bajo ID CLIENTE '-1').")
        else:
            st.success("MCO_42 e ING_4 están cargadas.")

    def build_excel_cyg_individual(nombre_hoja, headers, filas, pct_cols, monto_cols, widths):
        """Arma un libro Excel con UNA sola tabla CYG (para descargar y
        enviar cada tabla por separado, en vez del libro combinado)."""
        wb_ind = openpyxl.Workbook()
        wb_ind.remove(wb_ind.active)
        ws_ind = _agregar_hoja_cyg(wb_ind, nombre_hoja, headers, filas, pct_cols=pct_cols, monto_cols=monto_cols)
        for i, w in enumerate(widths, start=1):
            ws_ind.column_dimensions[get_column_letter(i)].width = w
        out_ind = io.BytesIO()
        wb_ind.save(out_ind)
        out_ind.seek(0)
        return out_ind

    run_cyg = st.button("Generar tablas CYG", type="primary")

    if run_cyg:
        fb_cyg = {
            "grh8": f_grh8.getvalue() if f_grh8 else None,
            "grh11": f_grh11.getvalue() if f_grh11 else None,
            "grh12": f_grh12.getvalue() if f_grh12 else None,
            "gcp5": f_gcp5.getvalue() if f_gcp5 else None,
            "gcp6": f_gcp6.getvalue() if f_gcp6 else None,
            "ggv4": f_ggv4.getvalue() if f_ggv4 else None,
            "ggv5": f_ggv5.getvalue() if f_ggv5 else None,
            "ggv6": f_ggv6.getvalue() if f_ggv6 else None,
            "ggi5": f_ggi5.getvalue() if f_ggi5 else None,
            "ggi6": f_ggi6.getvalue() if f_ggi6 else None,
            "ggm1": f_ggm1.getvalue() if f_ggm1 else None,
            "ggm2": f_ggm2.getvalue() if f_ggm2 else None,
            "ggm3": f_ggm3.getvalue() if f_ggm3 else None,
            "ggm4": f_ggm4.getvalue() if f_ggm4 else None,
            "ggm5": f_ggm5.getvalue() if f_ggm5 else None,
            "ogg5": f_ogg5.getvalue() if f_ogg5 else None,
            "mei1": f_mei1.getvalue() if f_mei1 else None,
            "mei2": f_mei2.getvalue() if f_mei2 else None,
            "mei3": f_mei3.getvalue() if f_mei3 else None,
            "mei4": f_mei4.getvalue() if f_mei4 else None,
        }
        st_files_raw_cyg = {}
        for f in (f_st_files or []):
            tabla = identificar_tabla_st(f.name)
            if tabla:
                st_files_raw_cyg[tabla] = f.getvalue()
        gpa_files_raw_cyg = {}
        for f in (f_gpa_files_rep3 or []):
            tabla = identificar_tabla_gpa(f.name)
            if tabla:
                gpa_files_raw_cyg[tabla] = f.getvalue()

        try:
            agg_serv, agg_act, avisos_cyg, epas = build_cyg_core(
                fb_cyg, ggm_params, ogg_params, mei_params, st_params, st_files_raw_cyg,
                ggm_proceso_params, ogg_proceso_params, mei_proceso_params, gpa_proceso_params,
                gpa_files_raw_cyg,
            )
        except Exception as e:
            st.error(f"Error generando CYG: {e}")
            st.stop()

        cyg14 = build_cyg_1_a_4(agg_serv, epas)
        cyg8 = build_cyg_8(agg_act, epas)
        cyg9, avisos9 = build_cyg_9(
            agg_serv,
            f_mco42.getvalue() if f_mco42 else None,
            f_ing4.getvalue() if f_ing4 else None,
            epas,
        )
        avisos_cyg = avisos_cyg + avisos9

        st.session_state['last_cyg14'] = cyg14
        st.session_state['last_cyg8'] = cyg8
        st.session_state['last_cyg9'] = cyg9
        st.session_state['last_avisos_cyg'] = avisos_cyg

    if "last_cyg14" in st.session_state:
        cyg14 = st.session_state["last_cyg14"]
        cyg8 = st.session_state["last_cyg8"]
        cyg9 = st.session_state["last_cyg9"]
        avisos_cyg = st.session_state["last_avisos_cyg"]

        total_filas = sum(len(v) for v in cyg14.values()) + len(cyg8) + len(cyg9)
        st.success(f"Tablas CYG generadas: {total_filas} filas en total.")
        mostrar_kpis_cyg(cyg14, cyg8, cyg9, avisos_cyg)
        for nombre, filas in cyg14.items():
            st.write(f"**{nombre}**: {len(filas)} filas")
        st.write(f"**CYG_8**: {len(cyg8)} filas")
        st.write(f"**CYG_9**: {len(cyg9)} filas")

        if avisos_cyg:
            with st.expander(f"⚠️ Avisos CYG ({len(avisos_cyg)})", expanded=True):
                for a in avisos_cyg:
                    st.markdown(f"- {a}")


        st.markdown("**Descargar cada tabla por separado**")
        cols_cyg_dl = st.columns(3)
        widths_serv = [14, 16, 14, 20, 14, 14, 16, 18, 16, 18]
        widths8 = [14, 16, 14, 20, 14, 14, 20, 18, 20, 18]
        widths9 = [14, 16, 14, 20, 16, 20, 14, 18, 18, 18, 18]
        tablas_individuales = [
            ("CYG_1", HEADERS_CYG_SERV, cyg14["CYG_1"], {7, 9}, {8, 10}, widths_serv),
            ("CYG_2", HEADERS_CYG_SERV, cyg14["CYG_2"], {7, 9}, {8, 10}, widths_serv),
            ("CYG_3", HEADERS_CYG_SERV, cyg14["CYG_3"], {7, 9}, {8, 10}, widths_serv),
            ("CYG_4", HEADERS_CYG_SERV, cyg14["CYG_4"], {7, 9}, {8, 10}, widths_serv),
            ("CYG_8", HEADERS_CYG_8, cyg8, {7, 9}, {8, 10}, widths8),
            ("CYG_9", HEADERS_CYG_9, cyg9, {8, 10}, {9, 11}, widths9),
        ]
        for i, (nombre_tabla, headers_t, filas_t, pct_t, monto_t, widths_t) in enumerate(tablas_individuales):
            excel_individual = build_excel_cyg_individual(nombre_tabla, headers_t, filas_t, pct_t, monto_t, widths_t)
            with cols_cyg_dl[i % 3]:
                st.download_button(
                    f"📄 {nombre_tabla}.xlsx",
                    data=excel_individual,
                    file_name=nombre_con_sufijo(nombre_tabla, "xlsx"),
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    key=f"dl_individual_{nombre_tabla}",
                    width="stretch",
                )

        st.markdown("**O descargar todo junto en un solo libro**")
        excel_cyg = build_excel_cyg(cyg14, cyg8, cyg9, avisos_cyg)
        st.download_button(
            "Descargar CYG_completo.xlsx (las 6 tablas + avisos)",
            data=excel_cyg,
            file_name=nombre_con_sufijo("CYG_completo", "xlsx"),
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        )

    # ============================================================================
    # MINUTA DE CRITERIOS DE ASIGNACIÓN (Word, generada dinámicamente)
    # ============================================================================
    def build_minuta_docx(ggm_params, ogg_params, mei_params, st_params,
                           ggm_proceso_params, ogg_proceso_params, mei_proceso_params, gpa_proceso_params,
                           by_recurso_planas=None, epas=None):
        """Genera una minuta Word que documenta:
    1. Los criterios de asignación POR DEFECTO usados por el sistema (fijos,
       acordados con la jefatura).
    2. La parametrización ADICIONAL efectivamente ejecutada en ESTA
       generación (dinámica: solo aparece lo que el usuario configuró)."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        by_recurso_planas = by_recurso_planas or {}

        doc = Document()

        title = doc.add_heading("Minuta de Criterios de Asignación de Gasto", level=0)
        if epas:
            empresa, periodo, anio, sector = epas
            p = doc.add_paragraph()
            p.add_run(f"Empresa {empresa} · Período {periodo} · Año Informado {anio} · Sector {sector}").italic = True

        doc.add_paragraph(
            "Este documento resume, para la generación de REP_2, REP_3 y las tablas CYG "
            "correspondientes: (1) los criterios de asignación POR DEFECTO definidos por el "
            "sistema para las familias sin tabla de apertura propia, y (2) la parametrización "
            "ADICIONAL que el usuario configuró y ejecutó en esta generación específica, "
            "desviándose de esos defaults."
        )

        # ---------------- SECCIÓN 1: DEFAULTS ----------------
        doc.add_heading("1. Criterios de Asignación por Defecto", level=1)

        doc.add_heading("1.1 Asignación de Servicio (REP_2 / CYG_1-4)", level=2)
        doc.add_paragraph(
            "Las familias GGM, OGG, MEI y ST no tienen tabla propia de apertura por servicio. "
            "Por defecto, el 100% del Gasto Anual de cada recurso se asigna al servicio regulado "
            "1101. El Monto Activado SIEMPRE se mantiene 100% en el servicio 1101, "
            "independiente de cualquier parametrización de servicio no regulado."
        )
        doc.add_paragraph(
            "GPA tiene un servicio FIJO por tabla (no configurable): GPA_1→1201, GPA_2→1202, "
            "GPA_3→1203, GPA_4→1204, GPA_5→1205, GPA_6→1201."
        )

        doc.add_heading("1.2 Asignación de Proceso / Actividad (REP_3 / CYG_8)", level=2)
        tabla_defaults = doc.add_table(rows=1, cols=3)
        tabla_defaults.style = "Light Grid Accent 1"
        hdr = tabla_defaults.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Familia / Recurso", "Asignación por defecto", "Justificación"
        filas_default = [
            ("GGI (todos los recursos)", "Reparto equitativo entre las 21 actividades del proceso 504", "Sin tabla de actividad propia"),
            ("GGM 2401-2409 (TI/telecom)", "100% → proceso 505 Informática", "Naturaleza informática/telecom del gasto"),
            ("GGM 2410-2411 (materiales)", "100% → proceso 504 Abastecimiento", "Materiales de compra/abastecimiento general"),
            ("OGG 2501-2502 (Directorio)", "100% → proceso 501 Dirección Superior", "Gasto de gobierno corporativo"),
            ("OGG resto", "100% → proceso 502 Admin. y Finanzas", "Naturaleza administrativa/financiera general"),
            ("MEI_1 (4101, Q. Químicos)", "Proceso/actividad EXACTA vía tabla OBRA_TIPO_NBI", "Mapeo determinístico entregado por la jefatura"),
            ("GPA (todas las tablas)", "Reparto equitativo dentro del subproceso propio (proceso 601)", "Cada tabla = una prestación asociada específica"),
        ]
        for f, a, j in filas_default:
            row = tabla_defaults.add_row().cells
            row[0].text, row[1].text, row[2].text = f, a, j

        # ---------------- SECCIÓN 2: PARAMETRIZACIÓN EJECUTADA ----------------
        doc.add_heading("2. Parametrización Adicional Ejecutada en esta Generación", level=1)

        def _nombre_servicio(cod):
            return SERVICIOS_NO_REGULADOS.get(cod, f"Servicio {cod}")

        hubo_parametrizacion = False

        doc.add_heading("2.1 Reasignación de Servicio (desvío desde el servicio regulado 1101)", level=2)
        params_servicio = [("GGM", ggm_params), ("OGG", ogg_params), ("MEI", mei_params), ("ST", st_params)]
        filas_serv = []
        for nombre_familia, params in params_servicio:
            recurso_totales = by_recurso_planas.get(nombre_familia, {})
            for cod_recurso, overrides in params.items():
                if not overrides:
                    continue
                total_original = recurso_totales.get(cod_recurso, [0.0, 0.0])[0]
                for cod_serv_destino, pct in overrides:
                    monto = total_original * pct
                    filas_serv.append((nombre_familia, cod_recurso, cod_serv_destino, _nombre_servicio(cod_serv_destino), pct, monto))

        if filas_serv:
            hubo_parametrizacion = True
            tabla_serv = doc.add_table(rows=1, cols=6)
            tabla_serv.style = "Light Grid Accent 1"
            hdr = tabla_serv.rows[0].cells
            for i, h in enumerate(["Familia", "Recurso", "Servicio Destino", "Nombre Servicio", "%", "Monto Estimado ($)"]):
                hdr[i].text = h
            for fam, rec, serv, nombre_s, pct, monto in filas_serv:
                row = tabla_serv.add_row().cells
                row[0].text = fam
                row[1].text = str(rec)
                row[2].text = str(serv)
                row[3].text = nombre_s
                row[4].text = f"{pct*100:.6f}%"
                row[5].text = f"{monto:,.0f}"
        else:
            doc.add_paragraph("No se configuró ninguna reasignación de servicio en esta generación; se utilizó el default (100% servicio 1101) para todos los recursos de estas familias.")

        doc.add_heading("2.2 Reasignación de Proceso (desvío desde el proceso/actividad por defecto)", level=2)
        params_proceso = [("GGM", ggm_proceso_params), ("OGG", ogg_proceso_params), ("MEI_1", mei_proceso_params), ("GPA", gpa_proceso_params)]
        filas_proc = []
        for nombre_familia, params in params_proceso:
            for cod_recurso, overrides in params.items():
                if not overrides:
                    continue
                for cod_proceso_destino, pct in overrides:
                    filas_proc.append((nombre_familia, cod_recurso, cod_proceso_destino, nombre_proceso(cod_proceso_destino), pct))

        if filas_proc:
            hubo_parametrizacion = True
            tabla_proc = doc.add_table(rows=1, cols=5)
            tabla_proc.style = "Light Grid Accent 1"
            hdr = tabla_proc.rows[0].cells
            for i, h in enumerate(["Familia", "Recurso", "Proceso Destino", "Nombre Proceso", "%"]):
                hdr[i].text = h
            for fam, rec, proc, nombre_p, pct in filas_proc:
                row = tabla_proc.add_row().cells
                row[0].text = fam
                row[1].text = str(rec)
                row[2].text = str(proc)
                row[3].text = nombre_p
                row[4].text = f"{pct*100:.6f}%"
        else:
            doc.add_paragraph("No se configuró ninguna reasignación de proceso en esta generación; se utilizaron los defaults de la sección 1.2 para todos los recursos de estas familias.")

        doc.add_paragraph()
        nota = doc.add_paragraph()
        nota.add_run(
            "Nota: esta minuta refleja el estado de la parametrización EN EL MOMENTO de generar "
            "este documento. Si se modifica la parametrización y se vuelve a generar REP_2/REP_3/CYG, "
            "debe volver a generarse esta minuta para que quede actualizada."
        ).italic = True

        # Corrige un detalle de settings.xml (atributo 'percent' del zoom) que
        # python-docx a veces omite; Word lo tolera pero validadores XSD lo marcan.
        try:
            settings = doc.settings.element
            for zoom in settings.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}zoom"):
                if zoom.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}percent") is None:
                    zoom.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}percent", "100")
        except Exception:
            pass

        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out


st.divider()
st.header("📄 Minuta de Criterios de Asignación")
st.caption(
    "Genera un documento Word que explica los criterios de asignación por defecto "
    "y detalla la parametrización adicional (servicio y proceso) efectivamente "
    "configurada en esta sesión — útil para respaldo documental ante la jefatura o auditoría."
)
if st.button("Generar Minuta (Word)"):
    try:
        minuta_bytes = build_minuta_docx(
            ggm_params, ogg_params, mei_params, st_params,
            ggm_proceso_params, ogg_proceso_params, mei_proceso_params, gpa_proceso_params,
            by_recurso_planas=st.session_state.get('last_by_recurso_planas'),
            epas=st.session_state.get('last_epas'),
        )
        st.session_state['minuta_bytes'] = minuta_bytes.getvalue()
        st.success("Minuta generada.")
    except Exception as e:
        st.error(f"Error generando la minuta: {e}")

if 'minuta_bytes' in st.session_state:
    st.download_button(
        "Descargar Minuta_Parametrizacion.docx",
        data=st.session_state['minuta_bytes'],
        file_name=nombre_con_sufijo("Minuta_Parametrizacion", "docx"),
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# ============================================================================
# PANEL DE VALIDACIÓN CRUZADA — REP_2 vs REP_3 vs CYG (gráfico de barras)
# ============================================================================
st.divider()
st.header("📊 Panel de Validación Cruzada REP_2 / REP_3 / CYG")
st.caption(
    "Compara visualmente el Gasto Anual total (regulado y no regulado) de REP_2 "
    "contra REP_3, CYG_1+CYG_2 (regulado) y CYG_3+CYG_4 (no regulado), para "
    "verificar de un vistazo que todo cuadre. Requiere haber generado REP_2, "
    "REP_3 y las tablas CYG en esta misma sesión."
)

faltan_para_panel = []
if 'last_final_rows' not in st.session_state:
    faltan_para_panel.append("REP_2")
if 'last_final_rows3' not in st.session_state:
    faltan_para_panel.append("REP_3")
if 'last_cyg14' not in st.session_state:
    faltan_para_panel.append("tablas CYG")

if faltan_para_panel:
    st.info(f"Falta generar: {', '.join(faltan_para_panel)}. Genera los reportes arriba para habilitar este panel.")
else:
    final_rows_p2 = st.session_state['last_final_rows']
    final_rows_p3 = st.session_state['last_final_rows3']
    cyg14_p = st.session_state['last_cyg14']

    rep2_regulado = sum(r[7] for r in final_rows_p2 if r[5] in (11, 12))
    rep2_no_regulado = sum(r[7] for r in final_rows_p2 if r[5] in (21, 22))
    rep2_total = rep2_regulado + rep2_no_regulado

    rep3_total = sum(r[8] for r in final_rows_p3)

    cyg1_total = sum(r[7] for r in cyg14_p.get("CYG_1", []))
    cyg2_total = sum(r[7] for r in cyg14_p.get("CYG_2", []))
    cyg3_total = sum(r[7] for r in cyg14_p.get("CYG_3", []))
    cyg4_total = sum(r[7] for r in cyg14_p.get("CYG_4", []))
    cyg1_2_total = cyg1_total + cyg2_total
    cyg3_4_total = cyg3_total + cyg4_total

    col1, col2 = st.columns(2)
    with col1:
        st.metric("REP_2 — Regulado (11+12)", f"${rep2_regulado:,.0f}")
        st.metric("REP_3 (total)", f"${rep3_total:,.0f}", delta=f"{rep3_total - rep2_regulado:,.0f}")
        st.metric("CYG_1 + CYG_2", f"${cyg1_2_total:,.0f}", delta=f"{cyg1_2_total - rep2_regulado:,.0f}")
    with col2:
        st.metric("REP_2 — No Regulado (21+22)", f"${rep2_no_regulado:,.0f}")
        st.metric("CYG_3 + CYG_4", f"${cyg3_4_total:,.0f}", delta=f"{cyg3_4_total - rep2_no_regulado:,.0f}")
        st.metric("REP_2 — Total", f"${rep2_total:,.0f}")

    fig, ax = plt.subplots(figsize=(9, 5.5))

    grupo_regulado = ["REP_2\n(Regulado)", "REP_3", "CYG_1+CYG_2"]
    valores_regulado = [rep2_regulado, rep3_total, cyg1_2_total]
    grupo_no_regulado = ["REP_2\n(No Regulado)", "CYG_3+CYG_4"]
    valores_no_regulado = [rep2_no_regulado, cyg3_4_total]

    x_reg = list(range(len(grupo_regulado)))
    x_noreg = list(range(len(grupo_regulado) + 1, len(grupo_regulado) + 1 + len(grupo_no_regulado)))

    barras_reg = ax.bar(x_reg, valores_regulado, color="#1F4E78", width=0.6, label="Regulado (11+12)")
    barras_noreg = ax.bar(x_noreg, valores_no_regulado, color="#C00000", width=0.6, label="No Regulado (21+22)")

    for x, v in zip(x_reg, valores_regulado):
        ax.text(x, v, f"${v:,.0f}", ha="center", va="bottom", fontsize=8)
    for x, v in zip(x_noreg, valores_no_regulado):
        ax.text(x, v, f"${v:,.0f}", ha="center", va="bottom", fontsize=8)

    ax.set_xticks(x_reg + x_noreg)
    ax.set_xticklabels(grupo_regulado + grupo_no_regulado, fontsize=9)
    ax.set_ylabel("Gasto Anual ($)")
    ax.set_title("Validación Cruzada: REP_2 vs REP_3 vs CYG")
    ax.legend(loc="upper right")
    ax.ticklabel_format(style="plain", axis="y")
    fig.tight_layout()

    st.pyplot(fig)

    diff_reg_rep3 = rep3_total - rep2_regulado
    diff_reg_cyg = cyg1_2_total - rep2_regulado
    diff_noreg_cyg = cyg3_4_total - rep2_no_regulado

    tolerancia = 10  # pesos, por redondeo
    ok_rep3 = abs(diff_reg_rep3) <= tolerancia
    ok_cyg_reg = abs(diff_reg_cyg) <= tolerancia
    ok_cyg_noreg = abs(diff_noreg_cyg) <= tolerancia

    if ok_rep3 and ok_cyg_reg and ok_cyg_noreg:
        st.success("✅ Todo cuadra: REP_2, REP_3 y las tablas CYG son consistentes entre sí (diferencias dentro de tolerancia de redondeo).")
    else:
        if not ok_rep3:
            st.warning(f"⚠️ REP_3 difiere de REP_2 (regulado) en ${diff_reg_rep3:,.2f}.")
        if not ok_cyg_reg:
            st.warning(f"⚠️ CYG_1+CYG_2 difiere de REP_2 (regulado) en ${diff_reg_cyg:,.2f}.")
        if not ok_cyg_noreg:
            st.warning(f"⚠️ CYG_3+CYG_4 difiere de REP_2 (no regulado) en ${diff_noreg_cyg:,.2f}.")
